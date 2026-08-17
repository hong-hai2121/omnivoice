# -*- coding: utf-8 -*-
"""
Nhận diện audio TIẾNG TRUNG (Chinese speech-to-text).

Dựa trên cách làm của project GetTextFromLink (faster-whisper + ffmpeg),
nhưng đây là một script ĐỘC LẬP: chỉ cần truyền vào 1 file audio hoặc video,
script sẽ trích xuất âm thanh và nhận diện thành văn bản tiếng Trung.

Cách dùng:
    python nhandien_giongnoi.py "duong_dan/file.mp3"
    python nhandien_giongnoi.py "video.mp4" --model large-v3 --out ketqua.txt
    python nhandien_giongnoi.py "audio.wav" --pinyin

Yêu cầu cài đặt:
    pip install faster-whisper
    (khuyến nghị) pip install torch        # để chạy GPU (CUDA)
    (tùy chọn)   pip install pypinyin      # nếu muốn xuất kèm phiên âm pinyin
    ffmpeg + ffprobe phải có trong PATH.
"""

import argparse
import gc
import glob
import os
import shutil
import subprocess
import sys

# Bảo đảm console Windows in được tiếng Trung/Việt
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

FFMPEG_PATH = shutil.which("ffmpeg")
FFPROBE_PATH = shutil.which("ffprobe")

# Thư mục lưu cache model whisper (tải 1 lần, dùng lại các lần sau)
WHISPER_MODEL_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "whisper_cache")

# Các thư mục cache model có sẵn khác để TÌM model đã tải (không tải lại).
# Tìm thấy model ở đâu thì nạp THẲNG từ đó, không hề gọi mạng.
EXTRA_CACHE_DIRS = [
    r"D:\Python\GetLinktoText\whisper_cache",
]


def _find_local_model(model_name):
    """Tìm thư mục model faster-whisper đã tải sẵn (định dạng CTranslate2) cho model_name.

    Trả về đường dẫn thư mục snapshot (chứa model.bin) nếu có, ngược lại None.
    Quét cache của script trước, rồi tới EXTRA_CACHE_DIRS.

    Chủ repo để "*": phần lớn model do Systran chuyển đổi, riêng large-v3-turbo là
    mobiuslabsgmbh (faster-whisper tự trỏ tới đó) — dò cứng "Systran" sẽ không thấy
    turbo trong cache và lại đi tải về qua mạng.
    """
    repo = f"models--*--faster-whisper-{model_name}"
    for root in [WHISPER_MODEL_PATH, *EXTRA_CACHE_DIRS]:
        # Bố cục HF: <cache>/<repo>/snapshots/<hash>/model.bin
        for model_bin in glob.glob(os.path.join(root, repo, "snapshots", "*", "model.bin")):
            return os.path.dirname(model_bin)
        # Trường hợp model nằm thẳng trong 1 thư mục (đã giải nén sẵn)
        if os.path.isfile(os.path.join(root, model_name, "model.bin")):
            return os.path.join(root, model_name)
    return None

# Phần mở rộng được coi là audio sẵn (không cần trích xuất lại)
AUDIO_EXTS = {".wav", ".mp3", ".m4a", ".aac", ".flac", ".ogg", ".opus", ".wma"}

# Tham số cố định cho faster-whisper khi nhận diện TIẾNG TRUNG.
# - language="zh": ép model nhận diện tiếng Trung (phổ thông / Mandarin).
# - initial_prompt: gợi ý bằng tiếng Trung giúp model thêm dấu câu chuẩn hơn.
# - condition_on_previous_text=False: không để lỗi của câu trước lan sang câu sau.
# - vad_filter: tự bỏ đoạn im lặng/nhạc nền nên không cần cắt file thủ công.
#
# Chống Whisper kẹt vòng lặp ("啊啊啊啊..." nuốt trọn cửa sổ 30s) — 3 lớp:
#   1. temperature là DÃY, không phải số đơn. Whisper giải mã ở nhiệt độ 0, thấy
#      compression_ratio/log_prob vượt ngưỡng thì giải mã LẠI ở nhiệt độ cao hơn.
#      Đặt temperature=0 là tắt hẳn cơ chế này: ngưỡng vẫn tính nhưng không có gì
#      để lùi về, nên đoạn lặp bị giữ nguyên.
#   2. no_repeat_ngram_size=3: cấm lặp lại bất kỳ 3-gram nào NGAY LÚC giải mã.
#   3. repetition_penalty=1.1: phạt token đã sinh, làm nguội vòng lặp từ sớm.
# Lưu ý: chế độ batched (BatchedInferencePipeline) chỉ đọc temperatures[0] và
# KHÔNG có vòng fallback — ở đó lớp 1 vô hiệu. Bù lại bằng _repair_looped_segments():
# đoạn nào vẫn kẹt lặp thì giải mã lại bằng model thường (có đủ lớp 1).
_TRANSCRIBE_OPTS = dict(
    language="zh",
    condition_on_previous_text=False,
    temperature=(0.0, 0.2, 0.4, 0.6, 0.8, 1.0),
    beam_size=5,
    initial_prompt="以下是普通话的句子，请用简体中文加上标点符号。",
    no_speech_threshold=0.6,
    compression_ratio_threshold=2.4,
    log_prob_threshold=-1.0,
    no_repeat_ngram_size=3,
    repetition_penalty=1.1,
    vad_filter=True,
    vad_parameters=dict(min_silence_duration_ms=500),
)


def extract_audio(media_path, output_wav, tempo=1.0):
    """Trích xuất âm thanh mono 16kHz bằng ffmpeg (định dạng whisper ưa thích).

    tempo < 1.0 → làm CHẬM giọng đọc (vd 0.7 = 70% tốc độ gốc) để Whisper bắt chữ
    dễ hơn với audio đọc quá nhanh. atempo chỉ nhận 0.5–2.0; ngoài khoảng đó bỏ qua.
    """
    if not FFMPEG_PATH:
        print("❌ Không tìm thấy ffmpeg trong PATH. Hãy cài ffmpeg trước.")
        return None
    af = []
    if tempo and abs(tempo - 1.0) > 1e-3 and 0.5 <= tempo <= 2.0:
        af = ["-filter:a", f"atempo={tempo}"]
        print(f"🐢 Làm chậm audio còn {tempo:.2f}x để nhận diện chuẩn hơn...")
    command = [FFMPEG_PATH, "-i", media_path, "-ac", "1", "-ar", "16000", *af, "-y", "-vn", output_wav]
    try:
        subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        print(f"✅ Đã trích xuất âm thanh: {output_wav}")
        return output_wav
    except subprocess.CalledProcessError as e:
        print(f"❌ Lỗi khi trích xuất âm thanh: {e}")
        return None


def get_audio_duration(audio_path):
    """Lấy thời lượng (giây) bằng ffprobe — không nạp cả file vào RAM."""
    if FFPROBE_PATH:
        try:
            result = subprocess.run(
                [
                    FFPROBE_PATH, "-v", "error",
                    "-show_entries", "format=duration",
                    "-of", "default=noprint_wrappers=1:nokey=1",
                    audio_path,
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True,
            )
            return float(result.stdout.strip())
        except Exception as e:
            print(f"⚠️ Không lấy được thời lượng audio: {e}")
            return None
    return None


# Model được giữ lại trong bộ nhớ giữa các lần chạy (đỡ phải nạp lại 1.5GB mỗi file)
_CACHE = {"name": None, "model": None, "batched": None, "device": None}


def free_model():
    """Giải phóng model khỏi RAM/VRAM (gọi khi đóng app nếu muốn)."""
    _CACHE["batched"] = None
    _CACHE["model"] = None
    _CACHE["name"] = None
    gc.collect()
    try:
        import torch
        torch.cuda.empty_cache()
    except Exception:
        pass


def get_model(model_name, use_batched=True):
    """Lấy model faster-whisper (có cache). Trả về (model, batched_pipeline|None, device).

    - Tự dùng GPU (CUDA, float16) nếu có, không thì CPU (int8).
    - batched_pipeline: chạy nhiều cửa sổ 30s song song → nhanh hơn 2-4× cho file dài.
    """
    if _CACHE["name"] == model_name and _CACHE["model"] is not None:
        return _CACHE["model"], _CACHE["batched"], _CACHE["device"]

    free_model()  # đổi model → bỏ model cũ trước
    try:
        from faster_whisper import WhisperModel, BatchedInferencePipeline
    except ImportError:
        print("❌ Chưa cài faster-whisper. Chạy: pip install faster-whisper")
        sys.exit(1)

    device, compute_type = "cpu", "int8"
    try:
        import torch
        if torch.cuda.is_available():
            device, compute_type = "cuda", "float16"
    except Exception:
        pass

    local_dir = _find_local_model(model_name)
    if local_dir:
        # Đã có sẵn trên đĩa → nạp thẳng từ thư mục, KHÔNG gọi mạng.
        print(f"📂 Nạp model '{model_name}' từ cache local ({device.upper()}, {compute_type}):\n   {local_dir}")
        model = WhisperModel(local_dir, device=device, compute_type=compute_type,
                             local_files_only=True)
    else:
        # Chưa có → tải về 1 lần vào WHISPER_MODEL_PATH, các lần sau sẽ nạp local.
        print(f"🔄 Chưa có local, đang TẢI model '{model_name}' trên {device.upper()} ({compute_type})...")
        os.makedirs(WHISPER_MODEL_PATH, exist_ok=True)
        model = WhisperModel(model_name, device=device, compute_type=compute_type,
                             download_root=WHISPER_MODEL_PATH)
    batched = None
    if use_batched:
        try:
            batched = BatchedInferencePipeline(model=model)
        except Exception as e:
            print(f"⚠️ Không bật được chế độ batched, dùng thường: {e}")
    _CACHE.update(name=model_name, model=model, batched=batched, device=device)
    print("✅ Model đã sẵn sàng.")
    return model, batched, device


def _max_repeat_run(text, max_unit=4):
    """Số lần lặp liên tiếp NHIỀU NHẤT của một cụm ngắn (1-4 ký tự) trong text.

    "啊啊啊啊啊啊" -> 6 (kẹt vòng lặp). "看看" -> 2, "高高兴兴" -> 2, "讨论讨论" -> 2:
    điệp từ là ngữ pháp bình thường của tiếng Trung nên ngưỡng phải để cao, đừng cắt.
    """
    s = "".join(text.split())
    best = 1
    for n in range(1, max_unit + 1):
        i = 0
        while i + n <= len(s):
            unit = s[i:i + n]
            k = 1
            while s[i + k * n:i + (k + 1) * n] == unit:
                k += 1
            if k > best:
                best = k
            i += k * n if k > 1 else 1
    return best


# Một câu bị nghi kẹt vòng lặp khi: gzip nén được quá tốt (= quá lặp), HOẶC có cụm
# ngắn lặp liên tiếp >= 6 lần (điệp từ thật của tiếng Trung không bao giờ tới 6).
_LOOP_CR_THRESHOLD = _TRANSCRIBE_OPTS["compression_ratio_threshold"]
_LOOP_RUN_THRESHOLD = 6


def _looks_looped(text, compression_ratio=None):
    if compression_ratio and compression_ratio > _LOOP_CR_THRESHOLD:
        return True
    return _max_repeat_run(text) >= _LOOP_RUN_THRESHOLD


def _repair_looped_segments(model, audio_path, records, suspects):
    """Giải mã LẠI các đoạn kẹt lặp bằng model thường (sửa records tại chỗ).

    Chế độ batched chạy nhanh nhưng chỉ dùng temperatures[0] nên không lùi nhiệt độ
    được: đoạn nào kẹt lặp là kẹt luôn, nuốt trọn cửa sổ 30s tiếng nói thật. Ở đây gom
    đúng những mốc giờ bị hỏng rồi chạy lại MỘT lượt bằng model thường — có đủ vòng
    temperature fallback. Chỉ vài chục giây audio nên gần như không ảnh hưởng tốc độ.

    Trả về số câu đã sửa được.
    """
    # Gộp các câu nghi ngờ nằm sát nhau thành từng cửa sổ liền mạch.
    windows = []  # [start, end, [chỉ số record]]
    for i in suspects:
        start, end = records[i][0], records[i][1]
        if end - start < 0.2:
            continue
        if windows and start - windows[-1][1] < 0.5:
            windows[-1][1] = max(windows[-1][1], end)
            windows[-1][2].append(i)
        else:
            windows.append([start, end, [i]])
    if not windows:
        return 0

    tong = sum(w[1] - w[0] for w in windows)
    print(f"🔁 Nghi {len(suspects)} câu kẹt lặp ({tong:.0f}s) → giải mã lại có fallback...")

    opts = dict(_TRANSCRIBE_OPTS)
    # clip_timestamps làm faster-whisper bỏ qua VAD; nói rõ ra cho khỏi hiểu nhầm.
    opts.pop("vad_parameters", None)
    opts["vad_filter"] = False
    clip = ",".join(f"{w[0]:.3f},{w[1]:.3f}" for w in windows)
    try:
        segments, _ = model.transcribe(audio_path, clip_timestamps=clip, **opts)
        fixed_text = [[] for _ in windows]
        for seg in segments:
            piece = seg.text.strip()
            if not piece:
                continue
            # Xếp câu vừa nghe lại vào cửa sổ chứa nó (theo điểm giữa).
            mid = (seg.start + seg.end) / 2
            for wi, w in enumerate(windows):
                if w[0] - 0.5 <= mid <= w[1] + 0.5:
                    fixed_text[wi].append(piece)
                    break
    except Exception as e:
        print(f"⚠️ Không giải mã lại được, giữ nguyên bản cũ: {e}")
        return 0

    n_fixed = 0
    for w, texts in zip(windows, fixed_text):
        new = "".join(texts).strip()
        # Vẫn lặp sau khi chạy lại -> fallback bó tay, giữ nguyên còn hơn mất chữ.
        if not new or _looks_looped(new):
            continue
        idxs = w[2]
        records[idxs[0]][2] = new
        for i in idxs[1:]:
            records[i][2] = ""
        n_fixed += len(idxs)
    print(f"✅ Đã sửa {n_fixed}/{len(suspects)} câu kẹt lặp." if n_fixed
          else "⚠️ Chạy lại vẫn lặp — giữ nguyên bản cũ.")
    return n_fixed


def transcribe_chinese(media_path, model_name="medium", batch_size=8,
                       use_batched=True, on_segment=None, on_progress=None,
                       partial_path=None, speed=0.7):
    """Nhận diện 1 file audio/video thành văn bản tiếng Trung.

    Cải tiến: chạy MỘT lần cho cả file rồi đọc từng câu (segment) theo kiểu stream —
    không cắt cứng giữa câu, báo tiến độ, và lưu dần ra file tạm để crash không mất hết.

    on_segment(text, fraction): gọi mỗi khi có 1 câu mới (fraction = % đã xong, có thể None).
    on_progress(fraction):      gọi để cập nhật tiến độ 0..1.
    partial_path:               nếu có, ghi dần kết quả vào đây; xoá khi xong thành công.
    speed:                      tốc độ phát audio khi nhận diện (1.0 = gốc, 0.7 = chậm lại
                                để bắt chữ tốt hơn với giọng đọc nhanh). Văn bản không đổi,
                                chỉ thay đổi tốc độ → độ chính xác.
    """
    if not os.path.isfile(media_path):
        print(f"❌ Không tìm thấy file: {media_path}")
        return None

    work_dir = os.path.dirname(os.path.abspath(media_path))
    base = os.path.splitext(os.path.basename(media_path))[0]
    ext = os.path.splitext(media_path)[1].lower()

    # Cần xử lý qua ffmpeg khi: là video (tách tiếng) HOẶC cần làm chậm tốc độ.
    # Audio đúng định dạng + tốc độ gốc thì dùng thẳng, khỏi tạo file tạm.
    need_slow = bool(speed) and abs(speed - 1.0) > 1e-3
    temp_wav = None
    if ext in AUDIO_EXTS and not need_slow:
        audio_path = media_path
    else:
        temp_wav = os.path.join(work_dir, f"{base}_temp_audio.wav")
        if not extract_audio(media_path, temp_wav, tempo=speed):
            return None
        audio_path = temp_wav

    model, batched, device = get_model(model_name, use_batched=use_batched)

    duration = get_audio_duration(audio_path) or 0
    if duration:
        print(f"⏳ Thời lượng: {duration / 60:.1f} phút")

    print("📝 Đang nhận diện giọng nói tiếng Trung...")
    if batched is not None:
        segments, info = batched.transcribe(audio_path, batch_size=batch_size, **_TRANSCRIBE_OPTS)
    else:
        segments, info = model.transcribe(audio_path, **_TRANSCRIBE_OPTS)

    total = duration or float(getattr(info, "duration", 0) or 0)
    records = []   # [start, end, text] từng câu — giữ mốc giờ để còn chạy lại đoạn hỏng
    suspects = []  # chỉ số các câu nghi kẹt vòng lặp
    pf = open(partial_path, "w", encoding="utf-8") if partial_path else None
    try:
        for seg in segments:
            piece = seg.text.strip()
            if not piece:
                continue
            if _looks_looped(piece, getattr(seg, "compression_ratio", None)):
                suspects.append(len(records))
            records.append([seg.start, seg.end, piece])
            if pf:
                pf.write(piece)
                pf.flush()
            frac = min(seg.end / total, 1.0) if total else None
            if on_segment:
                on_segment(piece, frac)
            if on_progress and frac is not None:
                on_progress(frac)

        # Batched không có temperature fallback -> vá lại ở đây (xem _TRANSCRIBE_OPTS).
        # Bản thường đã tự lùi nhiệt độ rồi, chạy lại y hệt cũng vô ích nên bỏ qua.
        if suspects and batched is not None:
            _repair_looped_segments(model, audio_path, records, suspects)
    finally:
        if pf:
            pf.close()
        if temp_wav and os.path.exists(temp_wav):
            try:
                os.remove(temp_wav)
            except Exception:
                pass
        gc.collect()
        if device == "cuda":
            try:
                import torch
                torch.cuda.empty_cache()
            except Exception:
                pass

    if on_progress:
        on_progress(1.0)
    transcript = "".join(r[2] for r in records).strip()
    # Xong xuôi -> bỏ file tạm
    if partial_path and os.path.exists(partial_path):
        try:
            os.remove(partial_path)
        except Exception:
            pass
    return transcript


# Dấu kết thúc câu tiếng Trung (và một số dấu phương Tây phòng khi lẫn).
_SENT_ENDERS = "。！？!?…\n"
# Dấu đóng (ngoặc kép/ngoặc đơn) có thể đứng NGAY SAU dấu kết câu — nuốt theo để
# không tách lệch, ví dụ: 他说：“走吧。” phải cắt sau dấu ” chứ không sau 。
_SENT_CLOSERS = "”’」』）)\"'"


def split_sentences(text):
    """Tách văn bản tiếng Trung thành danh sách câu, GIỮ nguyên dấu câu.

    Cắt sau mỗi cụm dấu kết câu (。！？…) kèm các dấu đóng ngoặc/ngoặc kép theo sau.
    """
    sentences, buf = [], []
    i, n = 0, len(text)
    while i < n:
        ch = text[i]
        buf.append(ch)
        if ch in _SENT_ENDERS:
            j = i + 1
            while j < n and (text[j] in _SENT_ENDERS or text[j] in _SENT_CLOSERS):
                buf.append(text[j])
                j += 1
            sentences.append("".join(buf))
            buf = []
            i = j
        else:
            i += 1
    if buf:
        sentences.append("".join(buf))
    return [s for s in (x.strip() for x in sentences) if s]


def split_into_chunks(text, min_len=1000, max_len=1500):
    """Gộp câu thành các đoạn ~min_len..max_len ký tự, KHÔNG cắt giữa câu.

    Tham lam: nối câu đến khi vượt max_len thì chốt đoạn (luôn cắt ở cuối câu).
    Đoạn cuối quá ngắn thì gộp vào đoạn trước cho gọn.
    """
    chunks, cur = [], ""
    for s in split_sentences(text):
        if cur and len(cur) + len(s) > max_len:
            chunks.append(cur)
            cur = s
        else:
            cur += s
    if cur:
        # Đoạn cuối hơi ngắn thì gộp vào đoạn trước — nhưng chỉ khi không vượt max_len.
        if chunks and len(cur) < min_len and len(chunks[-1]) + len(cur) <= max_len:
            chunks[-1] += cur
        else:
            chunks.append(cur)
    return chunks


def save_docx(transcript, out_path, title=None, pinyin_text=None,
              chunk=True, min_chars=1000, max_chars=1500):
    """Lưu kết quả ra file Word (.docx), set font Microsoft YaHei để chữ Trung hiển thị đúng.

    chunk=True: tách văn bản thành các đoạn ~min_chars..max_chars ký tự (cắt ở cuối câu),
    mỗi đoạn có tiêu đề "ĐOẠN k" để tiện dịch/sao chép từng phần.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    def _set_cjk(run, font="Microsoft YaHei", size=13):
        run.font.name = font
        run.font.size = Pt(size)
        # Bắt buộc gán riêng font cho chữ Đông Á (CJK) thì Word mới hiển thị đúng
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)

    doc = Document()
    if title:
        h = doc.add_heading(level=1)
        _set_cjk(h.add_run(title), size=16)

    if chunk:
        chunks = split_into_chunks(transcript, min_len=min_chars, max_len=max_chars)
        print(f"✂️  Đã tách thành {len(chunks)} đoạn (mỗi đoạn ~{min_chars}–{max_chars} ký tự, cắt ở cuối câu).")
        for idx, ck in enumerate(chunks, 1):
            doc.add_heading(f"ĐOẠN {idx} ({len(ck)} ký tự)", level=2)
            para = doc.add_paragraph()
            _set_cjk(para.add_run(ck))
    else:
        for line in transcript.split("\n"):
            para = doc.add_paragraph()
            _set_cjk(para.add_run(line))

    if pinyin_text:
        doc.add_paragraph()
        doc.add_heading("Pinyin", level=2)
        doc.add_paragraph(pinyin_text)

    doc.save(out_path)
    return out_path


def to_pinyin(text):
    """Chuyển văn bản tiếng Trung sang pinyin (nếu đã cài pypinyin)."""
    try:
        from pypinyin import pinyin, Style
    except ImportError:
        print("⚠️ Chưa cài pypinyin, bỏ qua phần phiên âm. (pip install pypinyin)")
        return None
    syllables = pinyin(text, style=Style.TONE)
    return " ".join(item[0] for item in syllables)


def main():
    # Không truyền đường dẫn → mở giao diện (GUI) cho tiện thay vì báo lỗi
    if len(sys.argv) == 1:
        gui_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "nhandien_gui.py")
        if os.path.exists(gui_path):
            print("ℹ️ Không có đường dẫn file → mở giao diện (GUI)...")
            subprocess.run([sys.executable, gui_path])
            return
        print("❌ Hãy truyền đường dẫn file, ví dụ:\n"
              "   python nhandien_giongnoi.py \"duong_dan/file.mp3\"")
        return

    parser = argparse.ArgumentParser(
        description="Nhận diện audio tiếng Trung bằng faster-whisper."
    )
    parser.add_argument("media", help="Đường dẫn file audio hoặc video cần nhận diện.")
    parser.add_argument(
        "--model", default="medium",
        help="Tên model whisper: tiny/base/small/medium/large-v3 (mặc định: medium).",
    )
    parser.add_argument(
        "--out", default=None,
        help="File .docx để lưu kết quả (mặc định: cùng tên file gốc + _zh.docx).",
    )
    parser.add_argument(
        "--pinyin", action="store_true",
        help="Xuất kèm phiên âm pinyin (cần cài pypinyin).",
    )
    parser.add_argument(
        "--batch-size", type=int, default=8,
        help="Số cửa sổ 30s chạy song song (lớn hơn = nhanh hơn nhưng tốn VRAM hơn).",
    )
    parser.add_argument(
        "--speed", type=float, default=0.7,
        help="Tốc độ audio khi nhận diện (1.0 = gốc, 0.7 = chậm lại cho dễ bắt chữ).",
    )
    args = parser.parse_args()

    def _progress(frac):
        print(f"\r⏳ Tiến độ: {frac * 100:5.1f}%", end="", flush=True)

    partial = os.path.splitext(args.media)[0] + "_zh.partial.txt"
    transcript = transcribe_chinese(
        args.media, model_name=args.model, batch_size=args.batch_size,
        on_progress=_progress, partial_path=partial, speed=args.speed,
    )
    print()  # xuống dòng sau thanh tiến độ
    if transcript is None:
        sys.exit(1)

    print("\n================ KẾT QUẢ (中文) ================")
    print(transcript)

    py = to_pinyin(transcript) if args.pinyin else None
    if py:
        print("\n--- Pinyin ---")
        print(py)

    out_path = args.out or os.path.splitext(args.media)[0] + "_zh.docx"
    save_docx(transcript, out_path, title=os.path.basename(args.media), pinyin_text=py)
    print(f"\n💾 Đã lưu kết quả vào: {out_path}")


if __name__ == "__main__":
    main()
