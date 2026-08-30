"""
Giao diện desktop cho Voice Cloning — chạy: python taogiong_gui.py
"""

import sys, os
# Gốc repo OmniVoice (chứa package omnivoice + venv) — lùi 2 cấp từ myvoice/scripts/
_REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, os.pardir))
_VENV_SCRIPTS = os.path.join(_REPO_ROOT, "venv", "Scripts")
_VENV_PYTHON  = os.path.join(_VENV_SCRIPTS, "python.exe")
_VENV_PYTHONW = os.path.join(_VENV_SCRIPTS, "pythonw.exe")   # bản KHÔNG mở cửa sổ console đen
def _samepath(a, b):
    return os.path.normcase(os.path.abspath(a)) == os.path.normcase(os.path.abspath(b))
# Nếu chưa chạy bằng interpreter của venv → chạy lại bằng pythonw.exe (ẩn cửa sổ console).
# Đã ở venv (python.exe hoặc pythonw.exe) thì không khởi động lại nữa.
if __name__ == "__main__" and (os.path.exists(_VENV_PYTHON) or os.path.exists(_VENV_PYTHONW)) \
        and not (_samepath(sys.executable, _VENV_PYTHON) or _samepath(sys.executable, _VENV_PYTHONW)):
    import subprocess
    _launcher = _VENV_PYTHONW if os.path.exists(_VENV_PYTHONW) else _VENV_PYTHON
    subprocess.Popen([_launcher] + sys.argv)
    sys.exit()
# Để import được package omnivoice ở gốc repo dù chạy từ thư mục con
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)
# Để import được video_khung.py nằm cùng thư mục scripts/
_SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPTS_DIR not in sys.path:
    sys.path.insert(0, _SCRIPTS_DIR)

# Bộ cấp phát VRAM của PyTorch giữ lại khối đã cắt để tái dùng; mỗi đoạn TTS lại
# có độ dài chuỗi khác nhau nên qua vài trăm đoạn là phần GIỮ LẠI phình to hơn
# hẳn phần thực dùng (đo được: thực dùng đỉnh 2.1 GB mà tiến trình chiếm 6.5 GB).
# Trên Windows driver KHÔNG báo hết VRAM mà lặng lẽ đẩy tensor sang RAM chạy qua
# PCIe, chậm 20-50× — đúng cái làm tập 08 trôi từ 50s/đoạn xuống 285s/đoạn.
# Ngưỡng này bảo allocator tự nhả bớt khối rảnh khi đã giữ tới 80% card.
# (expandable_segments hợp lý hơn nhưng torch 2.8 KHÔNG hỗ trợ trên Windows — đã
# thử, nó chỉ cảnh báo rồi bỏ qua.) PHẢI đặt TRƯỚC lần import torch đầu tiên.
os.environ.setdefault("PYTORCH_CUDA_ALLOC_CONF", "garbage_collection_threshold:0.8")

import re
import hashlib
import threading
import logging
import queue
import numpy as np
import soundfile as sf
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox
from pathlib import Path

# GUI chạy bằng pythonw.exe (không console) nên mỗi lần gọi ffmpeg/ffprobe Windows
# lại bật một cửa sổ console mới → nhấp nháy. Cờ này cho chạy hẳn không console.
CREATE_NO_WINDOW = 0x08000000 if sys.platform == "win32" else 0

BASE_DIR   = Path(__file__).resolve().parent.parent   # myvoice/
VOICE_DIR  = BASE_DIR / "voice"
SCRIPT_DIR = BASE_DIR / "kịch_bản"
OUTPUT_DIR = SCRIPT_DIR / "output"                    # nơi gom mọi kết quả (wav + video + chunks)
VIDEODOC_DIR   = BASE_DIR / "videodoc"                # kho clip DỌC (có thể chia thư mục con theo chủ đề)
VIDEODOC_INPUT = VIDEODOC_DIR / "input.txt"           # ghi TÊN 1 thư mục con → dựng video dọc Facebook từ đó
VIDEONGANG_DIR = BASE_DIR / "videongang"              # kho clip NGANG (có thể chia thư mục con theo chủ đề)
NGANG_SOURCE_ALL = "(tất cả)"                         # nhãn combobox = dùng cả kho videongang/
GEMINI_DOCX = SCRIPT_DIR / "gemini_result.docx"       # kết quả dịch Gemini → nguồn nội dung TTS
SEO_DOCX   = SCRIPT_DIR / "seoYoutube.docx"           # SEO YouTube (Gemini) — chạy sau bước dịch
CHINESE_DOCX = SCRIPT_DIR / "tiengTrung.docx"         # văn bản tiếng Trung (nguồn để dịch Gemini)
YOUTUBE_DIR = BASE_DIR / "YOUTUBE"                    # nơi chứa seo_youtube_gemini.py
DOWNLOAD_DIR = Path(__file__).resolve().parent / "downloads_zh"  # mp3 tải từ link
DRIVE_SCRIPT_FOLDER_ID = "1cDUrHiQmzyIK7a8rqY3pFspHsizaceei"  # thư mục Drive "kịch bản"
PREFIX_FILE = Path(__file__).resolve().parent / "copy_prefix.txt"  # câu mở đầu dịch (chèn đoạn 1)
FAV_FILE   = BASE_DIR / "voice_favorites.json"        # danh sách giọng mẫu yêu thích
EFFECT_FAV_FILE = BASE_DIR / "effect_favorites.json"  # danh sách hiệu ứng yêu thích (★)
PIPE_FILE  = BASE_DIR / "taogiong_pipeline.json"      # cài đặt quy trình tạo kịch bản (auto + model/tốc độ)
OPTS_FILE  = BASE_DIR / "taogiong_options.json"        # cài đặt mục "Cài đặt" (nhớ lần chạy trước)
# Mặc định quy trình: ① tự chạy ②, ② tự chạy ③, ③ tự chạy tạo giọng (OmniVoice).
# upload=False: ĐĂNG YOUTUBE mặc định TẮT — đăng là việc hướng ra ngoài, chỉ chạy
# khi bạn chủ động tick.
PIPE_DEFAULTS = dict(auto2=True, auto3=True, auto_tts=True, seo=True, model="medium", speed="0.7",
                     shutdown=False, sleep=False, upload=False)
# Số phút chờ trước khi tắt máy khi bật ô "Xong thì tắt máy" (huỷ bằng: shutdown /a).
SHUTDOWN_DELAY_MIN = 5
# Số phút chờ trước khi CHO MÁY NGỦ khi bật ô "🌙 Xong thì cho máy ngủ". Ngắn hơn
# tắt máy vì ngủ dễ huỷ hơn nhiều: bỏ tick là xong, mà lỡ ngủ rồi thì chạm chuột
# là dậy — không mất gì cả.
SLEEP_DELAY_MIN = 3
AUDIO_EXTS = {".mp3", ".wav", ".MP3", ".WAV", ".flac", ".FLAC"}
STAR       = "★ "                                     # tiền tố hiển thị cho giọng yêu thích

# Kho NHẠC NỀN (myvoice/Music) — chèn vào video TikTok, mix nhỏ hơn giọng.
MUSIC_DIR  = BASE_DIR / "Music"
MUSIC_EXTS = {".mp3", ".wav", ".flac", ".m4a", ".aac", ".ogg", ".opus", ".wma"}

# Kho hiệu ứng phủ lên video (scripts/hieuung/) — thường là .mov có alpha
EFFECTS_DIR = Path(__file__).resolve().parent / "hieuung"
EFFECT_EXTS = {".mov", ".mp4", ".webm", ".mkv", ".avi", ".gif"}
EFFECT_NONE = "Không (mặc định)"                       # mục "không thêm hiệu ứng"
DEFAULT_EFFECT = "bubbles_overlay_6.mov"               # hiệu ứng chọn sẵn nếu có trong hieuung/

# Kiểu phụ đề cho video ngang (YouTube) — xem video_gansub.py.
#   • .srt rời: chỉ xuất file, KHÔNG đụng video → tải lên YouTube Studio, người xem
#     bật/tắt được. Nhanh (chỉ tốn thời gian Whisper nghe audio).
#   • vẽ cứng : đốt chữ vào khung hình → phải mã hoá lại cả video, lâu hơn nhiều và
#     tốn thêm dung lượng đĩa bằng đúng một bản video nữa trong lúc ghi đè.
SUB_MODE_SRT  = "file .srt rời"
SUB_MODE_BURN = "vẽ cứng vào hình"


def suspend_computer() -> str:
    """Đưa máy vào chế độ NGỦ ngay. Trả "" nếu gọi được, hoặc mô tả lỗi.

    Hàm chỉ trả về khi máy đã DẬY lại (SetSuspendState chặn tại đó).

    Tham số (bHibernate=0, bForce=1, bWakeupEventsDisabled=0): số 0 đầu là xin NGỦ
    chứ không ngủ đông. Máy nào đang bật sẵn ngủ đông vẫn có thể vào hibernate thay
    vì sleep — muốn ngủ thật thì tắt nó một lần bằng:  powercfg -h off

    NGUỒN DUY NHẤT cho cả GUI lẫn bảng web (myvoice/web/power.py gọi lại hàm này).
    """
    import ctypes
    import subprocess
    try:
        if ctypes.windll.powrprof.SetSuspendState(0, 1, 0):
            return ""
    except Exception as e:
        logging.warning(f"⚠️ Gọi SetSuspendState không được ({e}) — thử lại bằng rundll32.")
    # Dự phòng: đúng lệnh ngủ Windows vẫn dùng, chạy ở tiến trình riêng nên không
    # dính chuyện quyền/hàm nạp hỏng của tiến trình đang chạy.
    try:
        subprocess.run(["rundll32.exe", "powrprof.dll,SetSuspendState", "0,1,0"],
                       check=False, creationflags=CREATE_NO_WINDOW)
        return ""
    except Exception as e:
        return str(e)


def _ensure_youtube_path():
    """Cho phép import các module trong myvoice/YOUTUBE (seo_docx_parser, đăng video...)."""
    youtube_dir = str(YOUTUBE_DIR)
    if youtube_dir not in sys.path:
        sys.path.insert(0, youtube_dir)


def upload_slots_text(default: str = "08:00 & 18:00") -> str:
    """Mô tả các khung giờ đăng, lấy từ dang_video_youtube.UPLOAD_SLOTS (nguồn DUY
    NHẤT quy định giờ) để dòng gợi ý trên giao diện không bị lệch với lúc chạy thật."""
    try:
        _ensure_youtube_path()
        import dang_video_youtube as yt
        return " & ".join(f"{h:02d}:{m:02d}" for h, m in yt.UPLOAD_SLOTS)
    except Exception:
        return default


def video_gansub_max_chars_doc(default: int = 27) -> int:
    """Số ký tự/dòng phụ đề của khung DỌC 1080x1920 (lấy từ video_gansub).

    Nạp muộn để dựng giao diện không phải kéo cả module vào; thiếu module thì lấy
    số mặc định, chỉ ảnh hưởng dòng gợi ý trên UI chứ không đổi kết quả.
    """
    try:
        import video_gansub as gs
        return gs.SUB_MAX_CHARS_DOC
    except Exception:
        return default

# Mặc định mục "Cài đặt" (dùng khi chưa có taogiong_options.json) — sau đó được
# ghi đè bằng giá trị của LẦN CHẠY TRƯỚC để mỗi lần mở giữ lại lựa chọn cũ.
OPTS_DEFAULTS = dict(
    from_gemini=True, chunk=300,
    make_video=True, ngang_speed="1.0", ngang_source=NGANG_SOURCE_ALL, effect=DEFAULT_EFFECT,
    make_video_doc=True, doc_speed="1.0", doc_percent=100,
    doc_from_ngang=False,
    doc_from_subfolder=False,
    doc_no_effect=False, make_tiktok=False, tiktok_speed="1.0",
    tiktok_percent=50,
    # YouTube Short: cắt ≤2:50 từ chính video TikTok rồi đăng tự động sau bản chính
    # 1 giờ. Bật sẵn — cắt bằng `-c copy` nên gần như không tốn thêm thời gian dựng.
    make_short=True,
    tiktok_no_effect=False, tiktok_caption_pos=40,
    tiktok_music=False, tiktok_music_db=-12, bring_front=True,
    make_sub=False, sub_mode=SUB_MODE_SRT, sub_model="large-v3-turbo", sub_max_chars=50,
    # Kiểu phụ đề khi vẽ cứng — kho scripts/kieusub_mau/*.json (dùng chung với
    # myvideo): hopbo = hộp bo góc cũ, ngoài ra karaoke/hormozi/anton/neon…
    sub_kieu="hopbo",
    # Font chữ đè lên font của kiểu (rỗng = theo kiểu; kho myvoice/fonts +
    # vài font Windows — xem kieusub.danh_sach_font).
    sub_font="",
    # Màu chữ đè lên màu của kiểu (RGB hex vd "FFD700"; rỗng = theo kiểu).
    sub_mau="",
    # Màu VIỀN quanh chữ, đè viền của kiểu (rỗng = theo kiểu). Kiểu không có
    # viền (hộp bo góc, CapCut) thì đổi màu này không thấy gì.
    sub_mau_vien="",
    # Vị trí chữ: % chiều cao tính từ đáy (rỗng = mặc định 173px ngang/380px dọc).
    sub_vitri="",
    # Cỡ chữ: % phóng to/thu nhỏ so với cỡ gốc của kiểu (rỗng/100 = giữ nguyên).
    # Trần ký tự/dòng tự rút theo tỉ lệ ngược nên chữ to không tràn mép.
    sub_cochu="",
    # Bề ngang dòng chữ: % so với bề ngang chuẩn (rỗng/100 = như cũ; 20–150).
    # Hẹp lại thì xuống dòng sớm — chữ gom về giữa; quá 100% dòng dài ra,
    # kéo quá tay là chạm mép khung (xem kieusub.ap_bengang).
    sub_bengang="",
    # Số DÒNG mỗi lần hiện chữ: 2 = gom hai dòng như ảnh mẫu của kho kiểu
    # (chữ ở lại lâu gấp đôi), 1 = mỗi lần một dòng như nếp cũ.
    sub_dong=2,
    # Phụ đề cho VIDEO DỌC (facebook.mp4) — dùng chung kiểu/model/chế độ với video
    # ngang, chỉ khác khung hình (1080x1920) và số ký tự/dòng (tự rút xuống 23).
    make_sub_doc=False,
)


def resolve_videodoc_subfolder(log=print):
    """Đọc videodoc/input.txt → TÊN thư mục con → trả về Path thư mục con hợp lệ.

    File input: dòng đầu KHÁC rỗng và không bắt đầu bằng '#' được coi là tên thư mục
    con trong videodoc/. Trả về None (→ dùng cả kho videodoc/) khi file thiếu/rỗng,
    thư mục con không tồn tại, hoặc không có clip .mp4 nào bên trong.
    """
    try:
        if not VIDEODOC_INPUT.exists():
            log(f"[video dọc] Không thấy {VIDEODOC_INPUT} → dùng cả kho videodoc/.")
            return None
        name = ""
        for line in VIDEODOC_INPUT.read_text(encoding="utf-8", errors="replace").splitlines():
            s = line.strip()
            if s and not s.startswith("#"):
                name = s
                break
        if not name:
            log("[video dọc] input.txt rỗng → dùng cả kho videodoc/.")
            return None
        sub = VIDEODOC_DIR / name
        if not sub.is_dir():
            log(f"[video dọc] Thư mục con '{name}' không có trong videodoc/ → dùng cả kho.")
            return None
        if not any(sub.glob("*.mp4")):
            log(f"[video dọc] Thư mục con '{name}' không có .mp4 → dùng cả kho videodoc/.")
            return None
        return sub
    except Exception as e:
        log(f"[video dọc] Lỗi đọc input.txt ({e}) → dùng cả kho videodoc/.")
        return None


def list_ngang_sources() -> list[str]:
    """Tên các thư mục con của videongang/ CÓ ít nhất 1 clip .mp4 (chọn nguồn theo chủ đề)."""
    try:
        return sorted(
            d.name for d in VIDEONGANG_DIR.iterdir()
            if d.is_dir() and any(d.glob("*.mp4"))
        )
    except Exception:
        return []


def resolve_ngang_source(name: str, log=print):
    """TÊN thư mục con videongang/ → Path hợp lệ (có .mp4); None → dùng cả kho videongang/.

    name rỗng hoặc = NGANG_SOURCE_ALL → None. Thư mục con không tồn tại / không có
    clip .mp4 cũng trả None (fallback dùng cả kho) để không chặn việc dựng video.
    """
    name = (name or "").strip()
    if not name or name == NGANG_SOURCE_ALL:
        return None
    sub = VIDEONGANG_DIR / name
    if not sub.is_dir():
        log(f"[video ngang] Thư mục con '{name}' không có trong videongang/ → dùng cả kho.")
        return None
    if not any(sub.glob("*.mp4")):
        log(f"[video ngang] Thư mục con '{name}' không có .mp4 → dùng cả kho videongang/.")
        return None
    return sub

# ── BẢNG MÀU GIAO DIỆN (nền trắng) ───────────────────────────────────────────
UI = dict(
    bg="#ffffff", card="#ffffff", border="#e4e7ec", field="#ffffff",
    fg="#1f2430", muted="#7b828f",
    accent="#e84393", accent_dk="#c92f7b", accent_soft="#f4c4dc",
    track="#edeff2", hover="#f1f3f6", press="#e6e9ee",
    log_bg="#fbfbfc", log_info="#475063", log_warn="#b07400", log_err="#d62828",
)

SCRIPT_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Tự chuyển input.txt cũ từ voice/ sang kịch_bản/ nếu chưa có
_old_input = VOICE_DIR / "input.txt"
_new_input = SCRIPT_DIR / "input.txt"
if not _new_input.exists():
    if _old_input.exists():
        _new_input.write_bytes(_old_input.read_bytes())
    else:
        _new_input.write_text("", encoding="utf-8")


def list_voice_files():
    if not VOICE_DIR.exists():
        return []
    return sorted(f.name for f in VOICE_DIR.iterdir() if f.suffix in AUDIO_EXTS)


def list_effect_files():
    """Danh sách file hiệu ứng trong scripts/hieuung/ (chỉ tên file)."""
    if not EFFECTS_DIR.exists():
        return []
    return sorted(f.name for f in EFFECTS_DIR.iterdir()
                  if f.is_file() and f.suffix.lower() in EFFECT_EXTS)


def strip_star(label: str) -> str:
    """Bỏ tiền tố ★ để lấy lại tên file thật từ chuỗi hiển thị trong combobox."""
    return label[len(STAR):] if label.startswith(STAR) else label


def load_favorites() -> set:
    try:
        import json
        return set(json.loads(FAV_FILE.read_text(encoding="utf-8")))
    except Exception:
        return set()


def save_favorites(favorites: set):
    import json
    try:
        FAV_FILE.write_text(
            json.dumps(sorted(favorites), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as e:
        logging.warning(f"Không lưu được danh sách yêu thích: {e}")


def load_effect_favorites() -> set:
    try:
        import json
        return set(json.loads(EFFECT_FAV_FILE.read_text(encoding="utf-8")))
    except Exception:
        return set()


def save_effect_favorites(favorites: set):
    import json
    try:
        EFFECT_FAV_FILE.write_text(
            json.dumps(sorted(favorites), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception as e:
        logging.warning(f"Không lưu được hiệu ứng yêu thích: {e}")


def load_pipe_settings() -> dict:
    """Cài đặt quy trình tạo kịch bản đã lưu (auto + model/tốc độ); thiếu thì dùng mặc định."""
    data = dict(PIPE_DEFAULTS)
    try:
        import json
        data.update(json.loads(PIPE_FILE.read_text(encoding="utf-8")))
    except Exception:
        pass
    return data


def save_pipe_settings(data: dict):
    import json
    try:
        PIPE_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2),
                             encoding="utf-8")
    except Exception as e:
        logging.warning(f"Không lưu được cài đặt quy trình: {e}")


def load_opt_settings() -> dict:
    """Cài đặt mục 'Cài đặt' đã lưu (mặc định dựa vào lần chạy trước); thiếu thì dùng mặc định."""
    data = dict(OPTS_DEFAULTS)
    try:
        import json
        data.update(json.loads(OPTS_FILE.read_text(encoding="utf-8")))
    except Exception:
        pass
    return data


def save_opt_settings(data: dict):
    import json
    try:
        OPTS_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2),
                             encoding="utf-8")
    except Exception as e:
        logging.warning(f"Không lưu được cài đặt: {e}")


def load_prefix() -> str:
    """Câu mở đầu dịch (copy_prefix.txt, dùng chung với GUI nhận diện); chưa có thì rỗng."""
    try:
        if PREFIX_FILE.exists():
            return PREFIX_FILE.read_text(encoding="utf-8").strip()
    except Exception:
        pass
    return ""


def read_chinese_docx_chunks(path) -> list[str]:
    """Đọc nội dung tiếng Trung (bỏ Heading/Title), tách đoạn như khi nhận diện."""
    from docx import Document
    import nhandien_giongnoi as recog
    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs
             if p.text.strip() and not (p.style.name or "").startswith(("Heading", "Title"))]
    return recog.split_into_chunks("".join(parts))


def read_zh_docx_chunks(path) -> list[str]:
    """Đọc lại các ĐOẠN từ file *_zh.docx (do recog.save_docx tạo: mỗi đoạn 1 đoạn
    văn dưới tiêu đề 'ĐOẠN k'). Mỗi đoạn văn = 1 chunk → tái dùng để TIẾP TỤC dịch
    mà KHÔNG cần nhận diện lại. Trả [] nếu đọc lỗi/không có đoạn nào."""
    try:
        from docx import Document
        doc = Document(str(path))
    except Exception:
        return []
    chunks = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if (p.style.name or "").startswith(("Heading", "Title")):
            continue   # bỏ tiêu đề 'ĐOẠN k' và tiêu đề file
        chunks.append(p.text.strip())
    return chunks


# Tên CỐ ĐỊNH cho bản nhận diện tiếng Trung trong mỗi thư mục tập (thay cho tên dài
# theo tiêu đề video như "…_bilibili_zh.docx"). Trùng tên với CHINESE_DOCX của bản 1 link.
ZH_DOCX_NAME = "tiengTrung.docx"


def find_zh_docx(folder):
    """Tìm bản nhận diện tiếng Trung trong 1 thư mục tập: ưu tiên tiengTrung.docx
    (tên mới, cố định); lùi về *_zh.docx cho các bản CŨ đặt theo tên video."""
    folder = Path(folder)
    fixed = folder / ZH_DOCX_NAME
    if fixed.is_file():
        return fixed
    return next(iter(sorted(folder.glob("*_zh.docx"))), None)


def kiem_ban_dich_folder(folder):
    """Đoạn HỎNG trong bản dịch của 1 thư mục tập — chốt dùng chung cho MỌI cửa
    (tạo input / tạo giọng+video / đăng YouTube): BẤT KỲ đoạn nào hỏng là BỎ CẢ
    TẬP, không làm tiếp, không đăng (tập 85/87 từng lọt tới tận YouTube).

    Trả về list (số_đoạn, lý_do) theo dich_gemini.bad_chunks — [] là đủ và lành.
    Thiếu dữ kiện (chưa có bản nhận diện / bản dịch, đọc lỗi) → None: không đủ cơ
    sở kết luận, bên gọi cho qua để không chặn oan tập cũ thiếu file trung gian."""
    folder = Path(folder)
    zh = find_zh_docx(folder)
    gem = folder / "gemini_result.docx"
    if not zh or not gem.exists():
        return None
    try:
        import dich_gemini as g
        chunks = read_zh_docx_chunks(zh)
        if not chunks:
            return None
        return g.bad_chunks(chunks, g.read_results_docx(gem, len(chunks)))
    except Exception as e:
        logging.warning(f"⚠️ {folder.name}: không kiểm được bản dịch ({e}) — cho qua.")
        return None


def download_audio_mp3(url: str, out_dir: Path):
    """Tải audio từ link video (yt-dlp) → trả về đường dẫn .mp3 (None nếu lỗi).

    Dùng cho bước ① nhận diện khi đầu vào là LINK thay vì file có sẵn.
    """
    try:
        import yt_dlp
    except ImportError:
        logging.error("❌ Chưa cài yt-dlp. Chạy: pip install yt-dlp")
        return None
    import nhandien_giongnoi as recog
    ffmpeg_dir = os.path.dirname(recog.FFMPEG_PATH) if getattr(recog, "FFMPEG_PATH", None) else None
    out_dir.mkdir(parents=True, exist_ok=True)

    def hook(d):
        if d.get("status") == "downloading":
            pct = d.get("_percent_str", "").strip()
            if pct:
                logging.info(f"⬇️  Tải... {pct} {d.get('_speed_str', '').strip()}")
        elif d.get("status") == "finished":
            logging.info("✅ Tải xong, đang chuyển sang MP3...")

    ydl_opts = {
        "format": "bestaudio/best",
        # %(id)s tránh tên file có ký tự đặc biệt / tiếng Trung
        "outtmpl": os.path.join(str(out_dir), "%(id)s.%(ext)s"),
        "postprocessors": [
            {"key": "FFmpegExtractAudio", "preferredcodec": "mp3", "preferredquality": "192"}
        ],
        "quiet": True, "no_warnings": True, "noprogress": True, "progress_hooks": [hook],
    }
    if ffmpeg_dir:
        ydl_opts["ffmpeg_location"] = ffmpeg_dir
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            logging.info(f"🎬 Tiêu đề: {info.get('title', '')}")
            mp3_path = os.path.splitext(ydl.prepare_filename(info))[0] + ".mp3"
            if os.path.exists(mp3_path):
                return mp3_path
    except Exception as e:
        logging.error(f"❌ Lỗi khi tải video: {e}")
        return None
    # Phòng khi prepare_filename không khớp: lấy mp3 mới nhất trong thư mục.
    mp3s = sorted(out_dir.glob("*.mp3"), key=lambda p: p.stat().st_mtime, reverse=True)
    return str(mp3s[0]) if mp3s else None


# Số tập ĐÃ TẠO gần nhất — dùng chung với GUI thumbnail (thumbnail_gui_state.json).
THUMB_STATE_FILE = YOUTUBE_DIR / "thumbnail_gui_state.json"


def load_episode_number() -> int:
    """Số tập đã tạo gần nhất (thumbnail_gui_state.json); chưa có → 0."""
    try:
        import json
        d = json.loads(THUMB_STATE_FILE.read_text(encoding="utf-8"))
        n = str(d.get("episode_number", "")).strip()
        if n.isdecimal():
            return int(n)
    except Exception:
        pass
    return 0


def save_episode_number(n: int) -> None:
    """Lưu số tập vừa tạo để lần/ link sau tăng tiếp (đồng bộ với GUI thumbnail)."""
    try:
        import json
        THUMB_STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
        THUMB_STATE_FILE.write_text(
            json.dumps({"episode_number": str(n).zfill(2)}, ensure_ascii=False, indent=2),
            encoding="utf-8")
    except Exception as e:
        # Nuốt lỗi ở đây là lần chạy sau ĐÁNH SỐ TẬP SAI mà không ai biết → phải log.
        logging.warning(f"⚠️ Không lưu được số tập ({n}) vào {THUMB_STATE_FILE.name}: {e}")


# ── SỐ TẬP BỎ QUA (đặt trước cho tương lai) ──────────────────────────────────
# Danh sách số tập KHÔNG được cấp cho tập mới. Ví dụ đang ở tập 31, đặt trước 33
# và 34 → sau 32 sẽ nhảy thẳng sang 35 (thường vì 33/34 để dành cho video khác).
# Chỉ chặn lúc CẤP SỐ MỚI; tập đã tạo rồi thì giữ nguyên số của nó.
SKIP_EPISODES_FILE = BASE_DIR / "taogiong_skip_episodes.json"


def load_skip_episodes() -> set:
    """Tập hợp số tập cần bỏ qua; file thiếu/hỏng → rỗng."""
    try:
        import json
        data = json.loads(SKIP_EPISODES_FILE.read_text(encoding="utf-8"))
        return {int(n) for n in data if str(n).strip().isdecimal() and int(n) > 0}
    except Exception:
        return set()


def save_skip_episodes(nums) -> None:
    try:
        import json
        SKIP_EPISODES_FILE.write_text(
            json.dumps(sorted({int(n) for n in nums if int(n) > 0}), indent=2),
            encoding="utf-8")
    except Exception as e:
        logging.warning(f"Không lưu được danh sách tập bỏ qua: {e}")


def next_episode_number(after: int, skip=None) -> int:
    """Số tập kế tiếp sau `after`, nhảy qua mọi số nằm trong danh sách bỏ qua."""
    skip = load_skip_episodes() if skip is None else skip
    n = int(after) + 1
    while n in skip:
        n += 1
    return n


# ── MANIFEST: nhớ link nào ↔ thư mục tập nào + tiến độ (để chạy tiếp/báo cáo) ──
# File tổng đặt trong kịch_bản/ (đã gitignore). Mỗi lần chạy ghi nguồn (link/file)
# kèm số tập + bước đã xong. Lần sau chạy CÙNG link → tái dùng ĐÚNG thư mục cũ và
# bỏ qua phần đã làm; dù NHẬP KHÁC THỨ TỰ / thiếu link vẫn đúng tập.
MANIFEST_FILE = SCRIPT_DIR / "batch_manifest.json"

# Bỏ qua việc GỬI LẠI Gemini khi gemini_result.docx đã TỒN TẠI: không dò từng đoạn
# is_translation_done để quyết định gửi lại, vì check đó hay báo nhầm "chưa xong" →
# gửi LẠI đoạn đã dịch lên Gemini. Đổi về False để dò từng đoạn và dịch tiếp phần thiếu.
#
# ⚠️ CỜ NÀY CHỈ ẢNH HƯỞNG VIỆC GỬI LẠI — KHÔNG bỏ qua chốt chặn _translation_complete.
# Trước đây nó gán thẳng translation_ok = True, khiến tập dịch dở (vd tập 42: Gemini
# chết ở đoạn 4/7, đoạn 4-7 còn "(chưa dịch)") vẫn chạy tiếp ra audio/video chỉ 45%
# thời lượng. Nay _translation_complete LUÔN chạy trước khi tạo input.txt.
SKIP_TRANSLATE_DETAIL_CHECK = True


def norm_source(src: str) -> str:
    """Chuẩn hoá chuỗi nguồn để làm khoá manifest ổn định.

    - Bỏ khoảng trắng + nháy bao quanh.
    - FILE LOCAL: đưa về đường dẫn TUYỆT ĐỐI + normcase (trên Windows: hạ hoa/thường,
      đổi '/'→'\\'). Nhờ vậy cùng một file gõ khác kiểu — hoa/thường ổ đĩa (D:\\ vs d:\\),
      gạch chéo (\\ vs /), tương đối vs tuyệt đối — vẫn ra CÙNG khoá → không nhận diện lại.
    - URL (http/https): giữ NGUYÊN (đường dẫn mạng phân biệt hoa/thường).
    Chỉ ảnh hưởng KHOÁ manifest; thao tác đọc file vẫn dùng chuỗi gốc.
    """
    s = (src or "").strip().strip('"').strip("'").strip()
    if not s:
        return ""
    if s.lower().startswith(("http://", "https://")):
        return s
    try:
        return os.path.normcase(os.path.abspath(s))
    except Exception:
        return s


# ── THƯ MỤC TẬP: "01" (kiểu cũ) hoặc "01 - <tên nguồn>" (kiểu mới) ───────────
# Thư mục tập mang thêm TÊN NGUỒN để nhìn là biết tập đó làm từ link/file nào:
# "01 - 95", "07 - 陈家有女初长成". SỐ TẬP luôn là phần ĐẦU tên thư mục nên thư
# mục cũ (tên thuần số) vẫn chạy bình thường — mọi nơi tra tập đều đi qua các hàm
# dưới đây THAY CHO việc so tên thư mục bằng .isdecimal().
_EP_DIR_RE = re.compile(r'^(\d+)\s*(?:-\s*(.*))?$')
# Ký tự Windows cấm đặt trong tên thư mục (+ ký tự điều khiển).
_BAD_FS_CHARS = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
MAX_SOURCE_LABEL = 60      # cắt bớt tên nguồn quá dài cho tên thư mục gọn


def episode_of(name: str):
    """Số tập (chuỗi 2 chữ số) lấy từ TÊN thư mục tập; không phải thư mục tập → None.

    "1" → "01" · "01" → "01" · "01 - 95" → "01" · "output"/"downloads_zh" → None.
    """
    m = _EP_DIR_RE.match((name or "").strip())
    return m.group(1).zfill(2) if m else None


def episode_dirs() -> list:
    """Mọi thư mục tập trong kịch_bản/ (cả tên cũ "01" lẫn mới "01 - 95"), theo SỐ TẬP."""
    if not SCRIPT_DIR.exists():
        return []
    out = [p for p in SCRIPT_DIR.iterdir() if p.is_dir() and episode_of(p.name)]
    return sorted(out, key=lambda p: (int(episode_of(p.name)), p.name))


def find_episode_dir(episode):
    """Thư mục CÓ SẴN của tập (khớp số tập, bất kể có kèm tên nguồn hay không)."""
    ep = str(episode).strip().zfill(2)
    for p in episode_dirs():
        if episode_of(p.name) == ep:
            return p
    return None


def safe_folder_name(name: str) -> str:
    """Rút gọn chuỗi thành phần tên thư mục hợp lệ trên Windows."""
    s = _BAD_FS_CHARS.sub(" ", (name or "")).strip()
    s = re.sub(r"\s+", " ", s)
    if len(s) > MAX_SOURCE_LABEL:
        s = s[:MAX_SOURCE_LABEL].rstrip()
    return s.rstrip(" .")          # Windows cấm tên kết thúc bằng '.' hoặc ' '


def source_label(source: str) -> str:
    """TÊN NGUỒN ngắn để gắn vào tên thư mục tập.

    • File local → tên file bỏ đuôi ("C:\\Users\\PC\\Downloads\\95.mp4" → "95").
    • Link video → tiêu đề video (yt-dlp, CHỈ đọc metadata, không tải); lấy không
      được thì lùi về đoạn cuối của URL.
    Không lấy được gì → "" (thư mục giữ tên thuần số như cũ).
    """
    s = (source or "").strip().strip('"').strip("'")
    if not s:
        return ""
    if not s.lower().startswith(("http://", "https://")):
        return safe_folder_name(Path(s).stem)
    try:
        import yt_dlp
        with yt_dlp.YoutubeDL({"quiet": True, "no_warnings": True,
                               "skip_download": True, "noplaylist": True}) as ydl:
            info = ydl.extract_info(s, download=False) or {}
        label = safe_folder_name(info.get("title") or info.get("id") or "")
        if label:
            return label
    except Exception as e:
        logging.info(f"ℹ️ Không lấy được tiêu đề video ({e}) — đặt tên theo link.")
    tail = s.split("?")[0].rstrip("/").rsplit("/", 1)[-1]
    return safe_folder_name(tail)


def episode_dir_for(episode, source=None) -> Path:
    """Đường dẫn thư mục của tập: DÙNG LẠI thư mục đã có (không đổi tên thư mục cũ);
    chưa có thì đặt tên mới "<số tập> - <tên nguồn>" (thiếu tên nguồn → "<số tập>")."""
    existing = find_episode_dir(episode)
    if existing is not None:
        return existing
    ep = str(episode).strip().zfill(2)
    label = source_label(source) if source else ""
    return SCRIPT_DIR / (f"{ep} - {label}" if label else ep)


def load_manifest() -> dict:
    """Đọc manifest (nguồn→{episode, steps, done, updated}); lỗi/thiếu → {}."""
    try:
        import json
        d = json.loads(MANIFEST_FILE.read_text(encoding="utf-8"))
        return d if isinstance(d, dict) else {}
    except Exception:
        return {}


def save_manifest(data: dict) -> None:
    try:
        import json
        MANIFEST_FILE.parent.mkdir(parents=True, exist_ok=True)
        MANIFEST_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2),
                                 encoding="utf-8")
    except Exception as e:
        logging.warning(f"Không lưu được manifest: {e}")


# ── QUEUE ĐỂ TRUYỀN LOG TỪ THREAD VỀ GUI ───────────────────────────────────
log_queue = queue.Queue()


class QueueHandler(logging.Handler):
    def emit(self, record):
        log_queue.put((record.levelno, self.format(record)))


# ── LOGIC CLONE ──────────────────────────────────────────────────────────────
# ── PHÁT HIỆN SPIKE ──────────────────────────────────────────────────────────
_SPIKE_RATIO  = 5.0   # khung vượt N× median RMS → spike
_SILENT_RMS   = 0.005 # median quá thấp = gần im lặng
_FRAME_MS     = 50    # độ dài mỗi khung phân tích (ms)

# Im lặng đệm vào CUỐI audio đã ghép (giây) — xem chỗ ghép chunk.
TAIL_SILENCE_SEC = 0.7
# Thời gian fade-out NHẠC NỀN ở cuối bản trộn (giây).
BGM_FADE_OUT_SEC = 3.0

# ── CHUẨN HOÁ ÂM LƯỢNG ───────────────────────────────────────────────────────
# Giọng clone nghe "lúc to lúc nhỏ". Đo trên 6 tập đã dựng thì thủ phạm KHÔNG
# phải chuyện đoạn nọ to hơn đoạn kia (các chunk chỉ lệch nhau 1,6–2,9 dB) mà là
# CÂU nọ to hơn CÂU kia trong cùng một đoạn: mức từng câu trải 3,7–3,9 dB
# (p5–p95), cực trị chênh nhau tới 11 dB. Nên xử lý 3 lớp:
#   ① CÂN ĐOẠN (numpy, lúc ghép chunk) — kéo mọi đoạn về cùng mức rồi đặt cả bài
#      vào một mốc RMS cố định, để ngưỡng của lớp ② luôn ăn đúng chỗ.
#   ② NÉN GIỮA CÁC CÂU (ffmpeg acompressor) — lớp trị đúng bệnh: câu to bị ghìm
#      lại, câu nhỏ nổi lên.
#   ③ CHUẨN R128 (ffmpeg loudnorm 2 lượt) — đưa cả file về mức phát hành chuẩn.
# Cả giọng lẫn bản trộn đều để -14 LUFS: đó đúng là mốc YouTube/TikTok/Facebook
# kéo mọi video về. Nộp sẵn đúng mốc thì nền tảng không hạ mà cũng không phải
# nâng, video này so với video khác nghe đều nhau.
# (-16 LUFS là mốc của podcast/Spotify — để mức đó thì lên YouTube nghe nhỏ hơn
# video người khác 2 dB, nên KHÔNG dùng ở đây.)
# Đặt None để TẮT lớp ③.
VOICE_LUFS   = -14.0   # audio giọng (output.wav và mọi bản cắt/tăng tốc từ nó)
MIX_LUFS     = -14.0   # bản trộn giọng + nhạc nền
TRUE_PEAK_DB = -1.5    # chừa 1.5dB đỉnh thật, tránh méo khi nén AAC lúc dựng video
LOUDNORM_LRA = 11.0    # dải động cho phép (mức khuyến nghị của EBU cho lời nói)
# Trần kéo mỗi đoạn khi cân ở lớp ① (dB). Đoạn lệch hơn mức này gần như luôn là
# đoạn generate lỗi — kéo hết cỡ chỉ tổ khuếch đại tiếng ồn nền.
CHUNK_GAIN_MAX_DB = 12.0
# Mốc RMS (dBFS) của phần CÓ TIẾNG sau khi cân đoạn. Ngưỡng bộ nén ở lớp ② tính
# theo mức tuyệt đối, nên phải đặt bài nào cũng vào đúng một mốc thì nén mới ăn
# đều; -18 dBFS còn chừa dư đỉnh cho những chữ bật mạnh.
SPEECH_RMS_DBFS = -18.0
# Lớp ②, ba khâu nối tiếp:
#   • acompressor — nén CHẬM và NHẸ: ngưỡng thấp hơn mốc trên 2dB, tỉ lệ 2,5:1,
#     ra/vào chậm (50ms/400ms) nên nó bám theo mức của CẢ CÂU chứ không ghìm
#     từng âm tiết. Đo trên 5 phút giữa tập 38: dải mức giữa các câu 3,9 → 2,1 dB,
#     max-min 6,3 → 3,6 dB, mà nhấn nhá TRONG câu gần như nguyên vẹn (3,7 → 3,6).
#     (Đã thử dynaudnorm và speechnorm: dynaudnorm bám theo ĐỈNH nên còn làm dải
#     mức giữa các câu RỘNG THÊM 3,9 → 4,4 dB; speechnorm chỉ xuống 3,2 dB.)
#   • volume=9dB — nén xong thì cả bài tụt xuống ≈ -23 LUFS, nâng lại đúng 9dB
#     là về sát -14. Con số cố định được vì đầu vào luôn ở mốc SPEECH_RMS_DBFS.
#   • alimiter — attack 50ms của bộ nén cố tình cho tiếng bật (p, t, k) lọt qua
#     để giọng không mất độ nảy, nhưng chính mấy tiếng đó đội ĐỈNH lên. Bộ chặn
#     gọt riêng chúng, nhờ vậy loudnorm ở lớp ③ đủ chỗ để chỉ nâng MỘT mức gain
#     (chế độ Linear). Không có khâu này, loudnorm phải bung chế độ Dynamic —
#     vẫn ra -14 LUFS nhưng gain chạy qua lại trong bài.
# Đặt "" để tắt lớp ②.
SPEECH_COMPRESSOR = ("acompressor=threshold=-20dB:ratio=2.5:attack=50:release=400,"
                     "volume=9dB,"
                     "alimiter=limit=-2dB:attack=5:release=50:level=disabled")


def detect_spike(path: Path, sr: int) -> list[float]:
    """Trả về danh sách thời điểm (giây) bị spike, rỗng nếu OK."""
    data, _ = sf.read(str(path), dtype="float32", always_2d=False)
    if data.ndim > 1:
        data = data.mean(axis=1)
    frame = int(sr * _FRAME_MS / 1000)
    frames = [data[i:i+frame] for i in range(0, len(data) - frame, frame)]
    if not frames:
        return [0.0]
    rms = np.array([np.sqrt(np.mean(f**2)) for f in frames])
    median = float(np.median(rms))
    if median < _SILENT_RMS:
        return [0.0]
    threshold = _SPIKE_RATIO * max(median, 1e-4)
    bad = np.where(rms > threshold)[0]
    return [round(float(t) * _FRAME_MS / 1000, 2) for t in bad]


def _speech_rms(data: np.ndarray, sr: int) -> float:
    """Mức to của phần CÓ TIẾNG trong `data` (RMS, bỏ khoảng lặng). 0.0 nếu câm.

    Lấy RMS trung bình cả đoạn thì đoạn nào nhiều khoảng nghỉ sẽ bị chấm là "nhỏ"
    rồi kéo to lên oan. Nên chỉ tính trên các khung vượt 20% khung to nhất — đủ
    để loại khoảng nghỉ giữa câu mà vẫn giữ được chữ nói nhẹ.
    """
    if data.ndim > 1:
        data = data.mean(axis=1)
    frame = max(1, int(sr * _FRAME_MS / 1000))
    n = len(data) // frame
    if n < 1:
        return float(np.sqrt(np.mean(data ** 2))) if len(data) else 0.0
    rms = np.sqrt(np.mean(data[:n * frame].reshape(n, frame).astype("float64") ** 2, axis=1))
    gate = float(np.percentile(rms, 95)) * 0.2
    speech = rms[rms >= gate]
    if speech.size == 0:
        return 0.0
    return float(np.sqrt(np.mean(speech ** 2)))


def _balance_chunk_levels(parts: list, sr: int) -> list:
    """Kéo mọi đoạn về CÙNG mức to trước khi ghép — hết cảnh đoạn to đoạn nhỏ.

    Mốc là TRUNG VỊ mức tiếng nói của cả bài (không phải trung bình): vài đoạn
    lỗi to/nhỏ bất thường sẽ không kéo lệch mốc của toàn bài. Mỗi đoạn chỉ được
    chỉnh tối đa ±CHUNK_GAIN_MAX_DB, và không bao giờ kéo tới mức vỡ tiếng.

    Xong thì kéo CẢ BÀI về mốc SPEECH_RMS_DBFS (một mức gain chung cho tất cả,
    không đổi tương quan giữa các đoạn) — bước này để bộ nén ở sau luôn nhận
    được đầu vào cùng một mức, bài nào cũng nén y như bài nào.

    Chỉ đổi ÂM LƯỢNG tổng của từng đoạn, không đụng gì bên trong đoạn → nhấn nhá
    trong câu giữ nguyên, không có tiếng "bơm" như khi nén động.
    """
    levels = [_speech_rms(p, sr) for p in parts]
    ok = [lv for lv in levels if lv > 1e-4]
    if len(ok) < 2:
        return parts
    target = float(np.median(ok))
    lim = 10 ** (CHUNK_GAIN_MAX_DB / 20)
    out, gains = [], []
    for p, lv in zip(parts, levels):
        if lv <= 1e-4:                       # đoạn câm: để nguyên, kéo chỉ ra tiếng ồn
            out.append(p)
            continue
        g = min(max(target / lv, 1 / lim), lim)
        peak = float(np.max(np.abs(p))) if p.size else 0.0
        if peak * g > 0.99:                  # không để vỡ tiếng sau khi kéo
            g = 0.99 / peak
        out.append((p * g).astype("float32", copy=False))
        gains.append(20 * float(np.log10(max(g, 1e-6))))
    if gains:
        logging.info(f"🔊 Cân âm lượng {len(gains)} đoạn về cùng mức "
                     f"(chỉnh {min(gains):+.1f} … {max(gains):+.1f} dB).")

    # Đưa cả bài về mốc cố định, nhưng không bao giờ để chạm trần 0dBFS.
    g_all = 10 ** (SPEECH_RMS_DBFS / 20) / max(target, 1e-6)
    peak_all = max((float(np.max(np.abs(p))) for p in out if p.size), default=0.0)
    if peak_all * g_all > 0.97:
        g_all = 0.97 / peak_all
    if abs(20 * np.log10(max(g_all, 1e-6))) > 0.05:
        out = [(p * g_all).astype("float32", copy=False) for p in out]
        logging.info(f"🔊 Đặt cả bài về mốc {SPEECH_RMS_DBFS:g} dBFS "
                     f"({20 * np.log10(g_all):+.1f} dB).")
    return out


# ── ĐỘ DÀI MỖI ĐOẠN, TÍNH THEO ÂM TIẾT TIẾNG VIỆT ────────────────────────────
# OmniVoice sinh song song: nó CHỐT trước số token audio rồi mới nhồi chữ vào
# đúng khung đó. Khung này để mặc định thì do RuleDurationEstimator đoán theo
# TRỌNG SỐ KÝ TỰ (chữ cái latin 1.0, dấu câu 0.5, chữ số 3.5...). Nhưng thời
# gian nói tiếng Việt phụ thuộc SỐ ÂM TIẾT chứ không phải số chữ cái: "nghiêng"
# 7 chữ mà chỉ 1 âm tiết, "ừ" 1 chữ cũng 1 âm tiết. Nên đoạn nhiều từ dài bị
# cho khung quá rộng (đọc lê thê), đoạn nhiều từ ngắn bị khung quá hẹp (đọc gấp)
# → giọng lúc nhanh lúc chậm giữa các đoạn.
#
# Đo trên tập 38 (10 đoạn rải đều, cùng ngochuyen.mp3):
#   để model tự đoán : 3.76–4.39 âm tiết/giây, lệch 17%
#   tự tính bên dưới : 4.05–4.30 âm tiết/giây, lệch 6%
# Ba hằng số đã hiệu chỉnh để nhịp đọc TRUNG BÌNH giữ nguyên như trước (~4.2
# âm tiết/giây) — chỉ hết dao động, không làm audio dài/ngắn hơn tổng thể.
# Muốn đọc chậm lại thì giảm VN_RATE, đọc nhanh hơn thì tăng.
VN_RATE      = 4.65   # âm tiết/giây khi đang nói liên tục
VN_PAUSE_MID = 0.112  # giây nghỉ thêm cho mỗi dấu , ; :
VN_PAUSE_END = 0.262  # giây nghỉ thêm cho mỗi dấu . ! ?

_VN_TOKEN    = re.compile(r'[0-9A-Za-zÀ-ỹà-ỹ]+')
_VN_MID_RE   = re.compile(r'[,;:]')
_VN_END_RE   = re.compile(r'[.!?]')
# "26" đọc là "hai mươi sáu" — 2 chữ số thành 3 âm tiết. Đếm mỗi chữ số 1.5 âm
# tiết để đoạn có số không bị khung quá hẹp rồi đọc vội/nuốt chữ.
_VN_SYL_PER_DIGIT = 1.5
# Chữ Hán sót qua bước dịch: model vẫn đọc (~1 âm tiết/chữ) mà _VN_TOKEN không
# bắt được → khung hụt giờ nặng. Đếm riêng từng chữ.
_VN_CJK      = re.compile(r'[一-鿿㐀-䶿豈-﫿]')
# Cụm nguyên âm liền nhau (kể cả có dấu). Mỗi TIẾNG Việt đúng 1 cụm ("nghiêng",
# "khuya", "quyết" → 1); từ ngoại lai nhiều cụm ("piano" 2, "camera" 3,
# "audio" 2, "elise" 3) — trước đây đều bị đếm 1 âm tiết nên đoạn có nhạc cụ /
# đồ tây / tên riêng tây bị hụt giờ, đọc gấp. Đếm mỗi cụm 1 âm tiết.
_VN_VOWELS   = re.compile(r'[aeiouyàáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩị'
                          r'òóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵ]+')


def _vn_syl_token(tok: str) -> float:
    """Số âm tiết đọc ra của MỘT token (chunk đã .lower() trước khi vào đây)."""
    if tok.isdigit():
        return len(tok) * _VN_SYL_PER_DIGIT
    if any(c.isdigit() for c in tok):
        # Chữ–số dính nhau ("12a8" đọc "mười hai a tám"): đếm rời từng ký tự.
        return sum(_VN_SYL_PER_DIGIT if c.isdigit() else 1 for c in tok)
    return max(1, len(_VN_VOWELS.findall(tok)))


def vn_duration(text: str):
    """Số giây nên dành cho đoạn `text`, tính theo số âm tiết + số dấu câu.

    Trả None khi đoạn không có âm tiết nào (chỉ dấu câu) — lúc đó để model tự
    đoán như cũ, vì ép duration=0 sẽ lỗi.
    """
    syllables = sum(_vn_syl_token(tok) for tok in _VN_TOKEN.findall(text))
    syllables += len(_VN_CJK.findall(text))
    if not syllables:
        return None
    return (syllables / VN_RATE
            + len(_VN_MID_RE.findall(text)) * VN_PAUSE_MID
            + len(_VN_END_RE.findall(text)) * VN_PAUSE_END)


# ── SỐ BƯỚC GIẢI MÃ MỖI ĐOẠN ─────────────────────────────────────────────────
# OmniVoice là masked-diffusion, KHÔNG có KV cache: mỗi bước chạy lại TOÀN BỘ
# chuỗi (prompt giọng mẫu + văn bản + khung audio), nên thời gian sinh một đoạn
# tỉ lệ THẲNG với số này. Đo trên RTX 4060: 32 bước 4.1s/đoạn (~12 phút cả bài
# 191 đoạn), 16 bước 2.0s/đoạn (~6 phút).
#
# GIỮ 32 — là mặc định tác giả model đã kiểm. Có thử 16: tai nghe không phân
# biệt được, detect_spike chấm 0/60 đoạn lỗi ở CẢ HAI mức. Nhưng detect_spike
# chỉ bắt lỗi thô (nổ, rít, câm), không thấy sai thanh điệu hay lệch nhịp — nên
# đó là "chưa thấy lỗi", không phải "tương đương". Ít bước thì mỗi bước chốt
# nhiều token hơn mà chúng không nhìn thấy nhau, rủi ro nằm ở đuôi phân phối,
# 60 đoạn chưa đủ để loại. Đổi lấy 6 phút trên một việc chạy tự động thì không
# đáng — nút thắt thật là VRAM, đã xử lý rồi.
#
# Số này nằm trong chữ ký cấu hình: đổi là các đoạn cũ tự sinh lại, không bao
# giờ trộn hai mức chất lượng vào cùng một bài.
NUM_STEP = 32

# KHÔNG gom nhiều đoạn vào một lượt generate. Nghe thì hợp lý, nhưng đo trên
# RTX 4060 (4 đoạn 142-268 ký tự) thì gom lô 4 CHẬM HƠN: 1.9s/đoạn so với 1.5s.
# Lý do: card đã chạy hết công suất ngay từ 1 đoạn, mà gom lô thì mọi đoạn phải
# đệm cho bằng đoạn dài nhất — ở đây thừa 26% phép tính. Card mạnh hơn (chưa bão
# hoà ở lô 1) thì gom lô mới có lãi.

# ── VUỐT TẮT CUỐI MỖI ĐOẠN ───────────────────────────────────────────────────
# Package OmniVoice mặc định fade_duration=0.1: vuốt 100ms CUỐI của đoạn về 0 để
# chống tiếng "cạch" ở mối nối. Chỗ đó đáng lẽ là im lặng — nhưng model chốt trước
# số frame rồi nhồi chữ vào cho vừa, nên đoạn nào cũng có xu hướng kết thúc SÁT MÉP,
# không còn im lặng để fade nhai. Đo trên 12 đoạn của bài 59 (sinh với fade tắt rồi
# tự áp fade lên CÙNG mảng audio, loại yếu tố may rủi của lần sinh):
#   fade 100ms → mất trung bình 19% năng lượng ở 100ms cuối (nặng nhất 38%)
#   fade  10ms → mất 0%
# Đây chính là "đuôi từ cuối bị cắt mất một chút" nghe thấy. 10ms vẫn thừa sức
# chống cạch (chỗ nối đã có pad_duration=0.1 im lặng hai đầu).
#
# ĐÃ THỬ VÀ LOẠI: nới khung thời lượng (hạ VN_RATE 4.65→4.4→4.2). Không ăn thua —
# model không dùng chỗ dư để im lặng, nó tính ra tỉ lệ tốc độ rồi đọc CHẬM lại cho
# vừa khung, vẫn kết thúc sát mép (xem chú thích ở vn_duration).
FADE_OUT_SEC = 0.01

# Cứ ngần này đoạn thì trả các khối VRAM đã giữ mà không dùng lại về cho driver.
VRAM_CLEAN_EVERY = 20


def _tra_vram(ly_do=""):
    """Trả các khối VRAM mà torch đang GIỮ LẠI về cho driver, rồi ghi mức trống.

    Bộ cấp phát của torch giữ lại khối đã cắt để tái dùng: bỏ tham chiếu tới
    model thì tensor chỉ về tay allocator, còn driver (và mọi tiến trình khác:
    NVENC, faster-whisper, Chrome) vẫn thấy card đầy. Chỉ empty_cache() mới trả
    thật.

    ĐIỀU KIỆN DÙNG: gọi khi đã KHÔNG CÒN tham chiếu nào tới model. Gọi lúc model
    vẫn còn sống thì hàm này gần như không trả được gì — mà dòng log lại in ra
    một con số đẹp, nên rất dễ tưởng đã dọn xong.
    """
    import gc
    import torch      # nạp trong hàm như chỗ khác trong file — GUI mở cho nhẹ
    gc.collect()
    try:
        torch.cuda.empty_cache()
        free_b, total_b = torch.cuda.mem_get_info()
        logging.info(f"🧹 Đã trả VRAM{f' ({ly_do})' if ly_do else ''} — còn trống "
                     f"{free_b/2**30:.1f}/{total_b/2**30:.1f} GB.")
    except Exception as e:
        logging.warning(f"Không dọn được VRAM ({ly_do or 'không rõ bước'}): {e}")


def _free_asr(model):
    """Xả model Whisper mà OmniVoice tự nạp để phiên âm giọng mẫu.

    create_voice_clone_prompt() không kèm ref_text thì OmniVoice nạp
    whisper-large-v3-turbo (~1.6 GB fp16) để nghe file giọng mẫu — dùng ĐÚNG MỘT
    LẦN rồi nằm lì trong VRAM tới hết bài. Trên card 8GB chính nó là phần đẩy
    tiến trình sang vùng tràn ra RAM.

    Ở đây xả được TRIỆT ĐỂ vì tham chiếu DUY NHẤT tới pipeline nằm trong chính
    thuộc tính này: gán None là refcount về 0 ngay, nên _tra_vram() gọi liền sau
    đó trả được thật (khác hẳn trường hợp _do_omnivoice, xem chú thích ở đó).

    Đây là Whisper của RIÊNG OmniVoice (transformers pipeline). Whisper của bước
    nhận diện tiếng Trung và bước gắn phụ đề là faster-whisper, đã có
    nhandien_giongnoi.free_model() lo — không đụng gì tới nhau.
    """
    if getattr(model, "_asr_pipe", None) is None:
        return
    model._asr_pipe = None
    _tra_vram("xả Whisper phiên âm giọng mẫu, ~1.6 GB")


def _do_omnivoice(model, ly_do="nhường chỗ cho ASR kiểm tra mất chữ"):
    """Bỏ tham chiếu tới OmniVoice. Trả None để bên gọi gán lại vào biến `model`.

    ⚠️ PHẢI viết ĐÚNG HAI DÒNG, ĐÚNG THỨ TỰ NÀY:

        model = _do_omnivoice(model, "lý do")   # ① bên gọi bỏ tham chiếu của mình
        _tra_vram()                             # ② giờ mới trả VRAM cho driver

    Vì sao phải tách đôi: `del model` bên trong hàm chỉ bỏ tham chiếu CỦA HÀM,
    biến `model` bên gọi vẫn trỏ vào đúng model đó nên trọng số chưa chết. Bản
    cũ gộp cả empty_cache() vào trong hàm này → lúc dọn thì model VẪN CÒN SỐNG,
    dọn xong chẳng trả được gì, mà dòng log lại in mức trống ĐO TRƯỚC KHI dỡ nên
    nhìn cứ tưởng đã xong. Hậu quả thật trên card 8GB: taogiong_kiemtra_matchu
    đo mem_get_info() thấy card còn đầy nên lùi ASR về CPU (chậm hơn nhiều lần),
    và NVENC dựng video cũng thiếu chỗ nên rớt về libx264.

    Nạp lại sau đó chỉ tốn thời gian đọc trọng số — prompt giọng mẫu đã lưu ra
    .prompt.pt nên không phải chạy lại Whisper phiên âm (xem _voice_prompt).
    """
    if model is not None:
        _free_asr(model)   # xả Whisper con của OmniVoice trước, nếu còn
        logging.info(f"📤 Dỡ OmniVoice khỏi VRAM ({ly_do})...")
    return None


def _voice_prompt(model, ref_audio):
    """VoiceClonePrompt dựng 1 lần rồi dùng lại cho mọi đoạn — và mọi lần chạy sau.

    Gọi model.generate(ref_audio=...) sẽ dựng LẠI prompt cho TỪNG đoạn: đọc
    mp3, cắt lặng, chạy Whisper large-v3-turbo phiên âm, mã hoá token. Bài 159
    đoạn tức là chạy Whisper 159 lần cho đúng một file — đo được ~7s thừa mỗi
    đoạn.

    Prompt còn được LƯU RA ĐĨA cạnh file giọng mẫu, nên từ lần chạy sau nạp
    thẳng file .pt và Whisper không phải vào VRAM lần nào nữa. Sửa file giọng
    mẫu (mtime mới hơn) thì prompt tự dựng lại.
    """
    cache = getattr(model, "_vn_prompt_cache", None)
    if cache is None:
        cache = model._vn_prompt_cache = {}
    if ref_audio in cache:
        return cache[ref_audio]

    from omnivoice.models.omnivoice import VoiceClonePrompt

    ref_path = Path(str(ref_audio))
    # Đuôi .pt nên file này KHÔNG lọt vào danh sách giọng của GUI (lọc AUDIO_EXTS).
    pt_path = ref_path.with_name(ref_path.name + ".prompt.pt")
    prompt = None

    try:
        if pt_path.exists() and pt_path.stat().st_mtime >= ref_path.stat().st_mtime:
            prompt = VoiceClonePrompt.load(str(pt_path))
            logging.info(f"🎙️  Nạp prompt giọng mẫu đã lưu: {pt_path.name} "
                         "— khỏi phải chạy Whisper.")
    except Exception as e:
        logging.warning(f"Prompt giọng mẫu đã lưu không đọc được ({e}) → dựng lại.")
        prompt = None

    if prompt is None:
        prompt = model.create_voice_clone_prompt(ref_audio=ref_audio)
        logging.info(f"🎙️  Đã dựng prompt giọng mẫu (dùng lại cho mọi đoạn): "
                     f"{ref_path.name}")
        try:
            prompt.save(str(pt_path))
            logging.info(f"💾 Đã lưu prompt giọng mẫu → {pt_path.name}")
        except Exception as e:
            logging.warning(f"Không lưu được prompt giọng mẫu: {e}")

    # Nạp từ đĩa hay vừa dựng thì tới đây Whisper cũng hết việc cho cả bài.
    _free_asr(model)

    cache[ref_audio] = prompt
    return prompt


# ── TÁCH LOGIC GENERATE 1 CHUNK ───────────────────────────────────────────────
def _generate_chunk(model, mode, voice_param, chunk, dur_scale=1.0):
    # language="vi": cả pipeline này chỉ sinh tiếng Việt, khai báo rõ giúp model
    # phát âm chuẩn hơn là để chế độ đoán ngôn ngữ.
    # dur_scale: nới khung thời lượng (render lại đoạn bị nuốt chữ lần 2 dùng
    # 1.06 — thêm ~6% chỗ để model nhét lại câu đã nuốt; lần đầu chỉ re-roll).
    dur = vn_duration(chunk)
    if dur is not None:
        dur *= dur_scale
    kw = {"language": "vi", "duration": dur, "num_step": NUM_STEP,
          "fade_duration": FADE_OUT_SEC}
    if mode == "clone":
        return model.generate(text=chunk,
                              voice_clone_prompt=_voice_prompt(model, voice_param), **kw)
    elif mode == "design":
        return model.generate(text=chunk, instruct=voice_param, **kw)
    return model.generate(text=chunk, **kw)


SPLIT_CHARS = re.compile(r'(?<=[.!?。！？\n])\s*')

# Ký tự thay bằng khoảng trắng (tránh ghép từ)
_REPLACE_WITH_SPACE = re.compile(r'[—–\-]+')
# Ký tự xóa hoàn toàn
_REMOVE = re.compile(r'["""\'\'\'`~@#$%^&*_+=|\\<>\[\]{}]')
# Dấu ba chấm → dấu chấm
_ELLIPSIS = re.compile(r'…+|\.{2,}')
# Nhiều khoảng trắng → 1
_SPACES = re.compile(r'[ \t]+')


def clean_text(text: str) -> str:
    text = _ELLIPSIS.sub('.', text)
    text = _REPLACE_WITH_SPACE.sub(' ', text)
    text = _REMOVE.sub('', text)
    text = _SPACES.sub(' ', text)
    return text.strip()


def split_chunks(text: str, max_len: int):
    parts = SPLIT_CHARS.split(text)
    chunks, current = [], ""
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if len(current) + len(part) + 1 <= max_len:
            current = (current + " " + part).strip()
        else:
            if current:
                chunks.append(current)
            current = part
    if current:
        chunks.append(current)
    return chunks


# ── DỌN & CHÈN LẠI CÂU QUẢNG BÁ KÊNH ──────────────────────────────────────────
# Bản dịch có sẵn rất nhiều câu quảng bá rải rác (xem _PROMO_PATTERNS). Ta xóa hết
# rồi chèn lại 3 câu dưới đây ở đúng vị trí: mở đầu / thân bài / kết bài.
# SỬA 3 CÂU NÀY nếu muốn đổi lời.
PROMO_OPENING = "Lại là Mimi Audio đây, mời bạn nghe câu chuyện hôm nay."
PROMO_OPENING_ENABLED = False
PROMO_BODY    = "Nếu thấy hay, bạn thả tim ủng hộ mình nhé."
# KHÔNG dùng "..." ở đây: clean_text đổi "..." → "." nên "không có thật,..." biến
# thành "không có thật,." (phẩy dính chấm) → TTS đọc cụt và tách chunk sai chỗ.
PROMO_ENDING  = ("Cảm ơn bạn đã lắng nghe. "
                 "Đây là câu chuyện không có thật, mọi tình tiết chỉ là hư cấu. "
                 "Nếu thấy hay, nhớ thích và theo dõi Mimi Audio nhé. "
                 "Hẹn gặp lại ở câu chuyện tiếp theo.")

# Số câu PROMO_BODY chèn giữa bài (rải đều). 0 = không chèn câu nào ở thân bài.
PROMO_BODY_COUNT = 2

# Nhận diện câu quảng bá/kết bài để XÓA SẠCH trước khi chèn lại. Gemini dịch theo
# lô nên mỗi lô tự thêm một cụm quảng bá → bản dịch có hàng chục câu rải rác, có
# chỗ 4–5 câu liên tiếp. Mỗi mẫu phải bám sát nguyên văn câu quảng bá, KHÔNG bắt
# từ khóa trần ("theo dõi", "ủng hộ", "hẹn gặp", "lắng nghe" đều xuất hiện bình
# thường trong truyện: "bỏ theo dõi", "tỷ lệ ủng hộ 87%", "hẹn gặp ở luyện võ
# trường", "ngoan ngoãn lắng nghe"…).
_PROMO_PATTERNS = re.compile(
    r'mimi\s*(?:audio|truyện|chuyện)'                       # tên kênh
    r'|cảm ơn\s+(?:bạn|các bạn|quý vị|mọi người)\s+đã\s+(?:lắng\s+)?nghe'
    r'|hẹn\s+gặp\s+lại\s+(?:bạn\s+|các bạn\s+|mọi người\s+)?'
    r'(?:ở|trong|vào|tại)\s+(?:câu\s+chuyện|truyện|video|tập|clip)'
    r'|câu\s+chuyện\s+(?:này\s+)?(?:hoàn toàn\s+)?không\s+có\s+thật'
    r'|mọi\s+tình\s+tiết\s+(?:chỉ\s+)?là\s+hư\s+cấu'
    r'|thả\s+tim\s+ủng\s+hộ'
    r'|(?:thích|like)\s+và\s+theo\s+dõi'
    r'|đăng\s+ký\s+kênh',
    re.IGNORECASE)
# Câu quảng bá luôn ngắn. Chặn trên để không nuốt nhầm câu truyện dài lỡ chứa cụm.
_PROMO_MAX_LEN = 150
# Mảnh chỉ gồm dấu câu/khoảng trắng. Câu "…không có thật,..." bị tách thành
# "…không có thật,." + "." + "." → xóa câu quảng bá xong phải dọn 2 dấu chấm lẻ.
_PUNCT_ONLY = re.compile(r'[\s.,;:!?…]+')
# Tách câu nhưng GIỮ dấu kết câu/xuống dòng ở cuối mỗi mảnh để ghép lại nguyên trạng.
_PROMO_SENT_SPLIT = re.compile(r'(?<=[.!?。！？\n])')


def _is_promo_sentence(s: str) -> bool:
    """Câu này có phải câu quảng bá kênh / kết bài không?"""
    core = s.strip()
    return bool(core) and len(core) <= _PROMO_MAX_LEN and bool(_PROMO_PATTERNS.search(core))


def _line_break_near(pieces: list[str], idx: int, window: int = 80) -> int:
    """Vị trí chèn gần idx nhất mà rơi đúng sau một lần xuống dòng (hết đoạn).
    Không tìm được trong bán kính window thì chèn thẳng tại idx."""
    for d in range(window):
        for j in (idx + d, idx - d):
            if 0 <= j < len(pieces) and pieces[j].endswith("\n"):
                return j + 1
    return idx


def replace_channel_promo(text: str) -> tuple[str, int, int]:
    """XÓA SẠCH mọi câu quảng bá kênh/kết bài trong bản dịch, rồi CHÈN LẠI đúng
    số câu ở đúng vị trí mình kiểm soát:

        • PROMO_OPENING  → đầu bài
        • PROMO_BODY     → rải đều PROMO_BODY_COUNT lần ở thân bài
        • PROMO_ENDING   → cuối bài

    Gemini dịch theo lô, mỗi lô tự thêm một cụm quảng bá, nên bản dịch thô có
    hàng chục câu rải rác — có chỗ 4–5 câu đọc liên tiếp trong 10 giây, và cụm
    kết bài có khi rơi vào giữa truyện (rồi cuối bài lại không có câu kết nào).
    Vì vậy phải DỌN TRƯỚC, CHÈN SAU thay vì thay 1-đổi-1 tại chỗ.

    Trả về (text_đã_xử_lý, số_câu_đã_xóa, số_câu_đã_chèn). Chạy lại trên văn bản
    đã xử lý KHÔNG cộng dồn: 3 câu vừa chèn cũng khớp mẫu nên bị xóa rồi chèn lại,
    kết quả luôn đúng 1 mở đầu + PROMO_BODY_COUNT thân bài + 1 kết bài (vị trí
    thân bài có thể xê dịch một chỗ xuống dòng ở lần chạy thứ hai rồi đứng yên).
    """
    kept, removed = [], []
    just_removed = False
    for s in _PROMO_SENT_SPLIT.split(text):
        drop = _is_promo_sentence(s)
        if drop:
            removed.append(s.strip())
        elif just_removed and _PUNCT_ONLY.fullmatch(s):
            drop = True                 # dấu chấm lẻ sót lại của câu vừa xóa
        if drop:
            just_removed = True
            if s.endswith("\n"):        # giữ xuống dòng để không dính 2 đoạn vào nhau
                kept.append("\n")
        else:
            just_removed = False
            kept.append(s)

    # ── CHÈN LẠI câu thân bài, rải đều, cắt đúng chỗ hết đoạn ─────────────────
    inserted = 0
    if PROMO_BODY_COUNT > 0 and kept:
        spots = sorted(
            {_line_break_near(kept, len(kept) * k // (PROMO_BODY_COUNT + 1))
             for k in range(1, PROMO_BODY_COUNT + 1)},
            reverse=True)
        for pos in spots:                          # chèn từ cuối lên để giữ chỉ số
            kept.insert(pos, PROMO_BODY + "\n")
            inserted += 1

    out = re.sub(r'\n{3,}', '\n\n', "".join(kept).strip())
    if PROMO_OPENING_ENABLED and PROMO_OPENING.strip():
        out = PROMO_OPENING + "\n" + out
        inserted += 1
    if PROMO_ENDING.strip():
        out = out + "\n" + PROMO_ENDING
        inserted += 1

    if removed:
        forms = {}
        for s in removed:
            forms[s] = forms.get(s, 0) + 1
        logging.info(
            f"🧹 Đã xóa {len(removed)} câu quảng bá/kết bài rải rác trong bản dịch: "
            + " | ".join(f"{n}× {s[:50]!r}" for s, n in
                         sorted(forms.items(), key=lambda kv: -kv[1])[:8]))
    return out, len(removed), inserted


# ── SỬA TỪ TIẾNG ANH BỊ SÓT → tiếng Việt ─────────────────────────────────────
# Gemini đôi khi để sót nguyên từ tiếng Anh giữa câu Việt: "But/but" (但),
# "If thấy hay…", "Twenty năm sau", "Sit vào trong xe", "Hand ... run lên".
# Bảng thay + danh sách giữ nguyên nằm ở scripts/tienganh_map.tsv (dich_tienganh).
def replace_leaked_english(text: str) -> tuple[str, int, list]:
    """Thay từ tiếng Anh Gemini sót → tiếng Việt, và dò từ lạ còn lại.

    Trả về (text_đã_thay, số_từ_đã_thay, danh_sách_từ_nghi_ngờ). Từ nghi ngờ chỉ
    để CẢNH BÁO (tên riêng/từ mượn mới), không tự sửa. Thiếu module → no-op."""
    try:
        import dich_tienganh as en
    except Exception as e:
        logging.warning(f"⚠️ Bỏ qua kiểm tra tiếng Anh: {e}")
        return text, 0, []
    out, n, suspects = en.fix_english(text, on_log=logging.info)
    if n:
        logging.info(f"🔁 Đã thay {n} từ tiếng Anh sót → tiếng Việt.")
    if suspects:
        logging.warning("⚠️ Từ lạ (không giống tiếng Việt) — tự xem lại, app KHÔNG "
                        "tự sửa: " + ", ".join(suspects[:25])
                        + (f" … (+{len(suspects) - 25})" if len(suspects) > 25 else "")
                        + f"\n   → Nếu là tên riêng/từ mượn, thêm vào {en._MAP_PATH.name} "
                          "với cột 2 là dấu '=' để hết cảnh báo.")
    return out, n, suspects


def chunks_dir_for(output_path: Path) -> Path:
    """Thư mục chunks dùng chung cho mọi bản đánh số (output, output1, output2…).

    Bỏ phần số đuôi của tên file để các lần chạy ghi vào CÙNG một thư mục
    (output_chunks), tránh tạo output1_chunks, output2_chunks… mỗi lần và
    giữ được khả năng tái dùng/“resume” chunk đã tạo.
    """
    stem = output_path.stem
    base = re.match(r"^(.*?)(\d*)$", stem).group(1) or stem
    return output_path.parent / (base + "_chunks")


def unique_path(path: Path) -> Path:
    """Nếu file đã tồn tại, trả về tên mới tăng số: output.wav → output1.wav → output2.wav…"""
    if not path.exists():
        return path
    m = re.match(r"^(.*?)(\d*)$", path.stem)
    base = m.group(1)
    n = int(m.group(2)) + 1 if m.group(2) else 1
    while True:
        cand = path.with_name(f"{base}{n}{path.suffix}")
        if not cand.exists():
            return cand
        n += 1


def find_cover_doc(folder, log=print):
    """Tìm thumbnail DỌC của tập (thumbnail<NN>_dọc.png, 1080×1920) để đè lên FRAME
    ĐẦU video dọc/TikTok — TikTok & Facebook lấy frame đầu làm ảnh bìa mặc định.

    • Thư mục tập (kịch_bản/NN) chỉ có 1 file → lấy luôn.
    • Thư mục chung (kịch_bản/) chứa thumbnail của nhiều tập → lấy đúng SỐ TẬP đang
      đặt ở tab Thumbnail, không có thì lấy file mới nhất.
    Trả về None nếu chưa có thumbnail dọc (khi đó video dựng như cũ)."""
    folder = Path(folder)
    cands = list(folder.glob("thumbnail*_dọc.png"))
    if not cands:
        return None
    if len(cands) == 1:
        return cands[0]
    ep = load_episode_number()
    exact = folder / f"thumbnail{ep:02d}_dọc.png"
    if exact.is_file():
        return exact
    newest = max(cands, key=lambda p: p.stat().st_mtime)
    log(f"[ảnh bìa] {folder.name} có {len(cands)} thumbnail dọc, không khớp số tập "
        f"{ep:02d} → dùng bản mới nhất: {newest.name}")
    return newest


def _speedup_audio_for_doc(src, factor):
    """Tăng tốc audio cho VIDEO DỌC bằng ffmpeg atempo (giữ cao độ, không bị 'chipmunk').

    factor nằm trong khoảng atempo cho phép (0.5–2.0). Trả về file mới *_spedNNN.wav.
    """
    import shutil
    import subprocess
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        raise RuntimeError("Không tìm thấy ffmpeg trong PATH.")
    src = Path(src)
    tag = f"{factor:.2f}".replace(".", "")          # 1.07 -> "107"
    out = src.with_name(f"{src.stem}_sped{tag}{src.suffix}")
    cmd = [ffmpeg, "-y", "-i", str(src), "-filter:a", f"atempo={factor:.4f}",
           "-vn", str(out)]
    r = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8",
                       errors="replace", creationflags=CREATE_NO_WINDOW)
    if r.returncode != 0 or not out.exists():
        raise RuntimeError(f"ffmpeg atempo lỗi: {(r.stderr or '')[-300:] or r.returncode}")
    return out


def list_music_files() -> list[str]:
    """Danh sách file nhạc nền trong myvoice/Music (chỉ tên file)."""
    if not MUSIC_DIR.exists():
        return []
    return sorted(f.name for f in MUSIC_DIR.iterdir()
                  if f.is_file() and f.suffix.lower() in MUSIC_EXTS)


def _detect_peak_db(path: Path, ffmpeg: str, seconds: int = 150):
    """Đọc đỉnh (max_volume, dBFS) của audio bằng volumedetect (tối đa `seconds`
    giây đầu cho nhanh). Trả về float dB hoặc None nếu không đo được."""
    import subprocess
    r = subprocess.run(
        [ffmpeg, "-hide_banner", "-t", str(seconds), "-i", str(path),
         "-af", "volumedetect", "-f", "null", "-"],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        creationflags=CREATE_NO_WINDOW)
    m = re.search(r"max_volume:\s*(-?\d+(?:\.\d+)?)\s*dB", r.stderr or "")
    return float(m.group(1)) if m else None


def _measure_loudness(path: Path, ffmpeg: str, seconds: int | None = None,
                      target: float = -16.0, pre_filter: str = ""):
    """Đo độ to cảm nhận theo EBU R128 (lượt 1 của loudnorm).

    Trả dict {input_i, input_tp, input_lra, input_thresh, target_offset} (đơn vị
    LUFS/dBTP, kiểu float) hoặc None nếu không đo được. `seconds` giới hạn số giây
    đầu đem đo — đủ dùng cho nhạc nền, còn giọng thì đo cả bài cho chuẩn.

    `target` là mức LUFS mà lượt 2 sẽ nhắm tới: chỉ ảnh hưởng target_offset (số
    hiệu chỉnh cho lượt 2), nên phải truyền đúng mức sẽ dùng ở lượt sau.

    `pre_filter` là chuỗi filter chạy TRƯỚC loudnorm (vd bộ nén). Phải đo trên
    tín hiệu ĐÃ qua filter đó, vì chính nó làm đổi độ to — đo bản chưa nén rồi
    đem số đó chỉnh bản đã nén là lệch mức.

    Vì sao LUFS mà không phải đỉnh (max_volume): đỉnh chỉ là một mẫu to nhất,
    một tiếng "p" bật hơi cũng đội đỉnh lên trong khi cả bài vẫn nhỏ. LUFS đo
    đúng cái tai nghe thấy.
    """
    import json
    import subprocess
    cmd = [ffmpeg, "-hide_banner", "-nostats"]
    if seconds:
        cmd += ["-t", str(seconds)]
    cmd += ["-i", str(path), "-af",
            (pre_filter + "," if pre_filter else "") +
            f"loudnorm=I={target:g}:TP={TRUE_PEAK_DB:g}:LRA={LOUDNORM_LRA:g}:"
            f"print_format=json", "-f", "null", "-"]
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8",
                           errors="replace", creationflags=CREATE_NO_WINDOW)
    except OSError:
        return None
    # JSON nằm ở CUỐI stderr; lấy khối { } cuối cùng có chứa input_i.
    blocks = re.findall(r"\{[^{}]*\"input_i\"[^{}]*\}", r.stderr or "", re.S)
    if not blocks:
        return None
    try:
        raw = json.loads(blocks[-1])
        d = {k: float(raw[k]) for k in
             ("input_i", "input_tp", "input_lra", "input_thresh", "target_offset")}
    except (ValueError, KeyError, TypeError):
        return None                       # file câm → "-inf"/"nan", float() ném lỗi
    if any(not np.isfinite(v) for v in d.values()):
        return None
    return d


def _loudnorm_file(src: Path, target_lufs, out: Path | None = None,
                   label: str = "", pre_filter: str = "") -> Path:
    """Đưa file audio về đúng mức to chuẩn `target_lufs` (LUFS) — loudnorm 2 lượt.

    2 lượt = đo cả file trước rồi mới chỉnh, nên đúng mức ngay lần đầu (chạy 1
    lượt thì loudnorm vừa nghe vừa dò, đầu file luôn lệch). Thêm linear=true để
    cả file chỉ ăn MỘT mức gain duy nhất: giữ nguyên nhấn nhá, chỉ đổi to/nhỏ
    tổng thể — khác hẳn nén động (dynaudnorm) vốn hay làm giọng bị "bơm".

    `pre_filter` (vd SPEECH_COMPRESSOR) chạy trước loudnorm ở CẢ hai lượt, nên
    san bằng câu và chuẩn mức gộp chung một lần chạy, không tốn thêm lượt nào.

    out=None → ghi đè chính file nguồn. `target_lufs=None` hoặc đo/chỉnh lỗi →
    giữ nguyên file cũ (không làm hỏng bước sau), trả về đường dẫn kết quả.
    """
    import shutil
    import subprocess
    src = Path(src)
    out = Path(out) if out else src
    if target_lufs is None:
        return src
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        logging.warning("🔊 Không thấy ffmpeg → bỏ qua chuẩn hoá âm lượng.")
        return src
    name = label or src.name

    d = _measure_loudness(src, ffmpeg, target=target_lufs, pre_filter=pre_filter)
    if not d:
        logging.warning(f"🔊 Không đo được độ to ({name}) → giữ nguyên âm lượng.")
        return src

    try:
        sr = sf.info(str(src)).samplerate
    except Exception:
        sr = 44100
    # loudnorm chạy nội bộ ở 192kHz; không ép -ar là file ra bị đổi tần số lấy mẫu.
    filt = ((pre_filter + "," if pre_filter else "") +
            f"loudnorm=I={target_lufs}:TP={TRUE_PEAK_DB}:LRA={LOUDNORM_LRA}:"
            f"measured_I={d['input_i']:.2f}:measured_TP={d['input_tp']:.2f}:"
            f"measured_LRA={d['input_lra']:.2f}:measured_thresh={d['input_thresh']:.2f}:"
            f"offset={d['target_offset']:.2f}:linear=true:print_format=summary")
    tmp = out.with_name(out.stem + "_chuanam.tmp.wav")
    cmd = [ffmpeg, "-y", "-hide_banner", "-nostats", "-i", str(src),
           "-af", filt, "-ar", str(sr), "-c:a", "pcm_s16le", str(tmp)]
    r = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8",
                       errors="replace", creationflags=CREATE_NO_WINDOW)
    if r.returncode != 0 or not tmp.exists() or tmp.stat().st_size <= 4096:
        tmp.unlink(missing_ok=True)
        logging.warning(f"🔊 Chuẩn hoá âm lượng lỗi ({name}), giữ nguyên bản cũ: "
                        f"{(r.stderr or '')[-200:]}")
        return src
    os.replace(tmp, out)
    logging.info(f"🔊 Chuẩn âm {name}: {d['input_i']:.1f} → {target_lufs:.1f} LUFS "
                 f"(đỉnh {TRUE_PEAK_DB:g} dBTP)")
    return out


def _chuan_am_neu_can(wav: Path, status_var=None) -> bool:
    """Chuẩn hoá `wav` NẾU nó chưa đạt mức đích — dùng cho luồng ♻ dùng lại audio.
    Trả True nếu có sửa file.

    Không cần cờ đánh dấu, cứ ĐO rồi quyết: file đã qua chuẩn hoá thì độ to nằm
    sát VOICE_LUFS (sai số < 0.7 LUFS) → bỏ qua. Nhờ vậy chạy lại bao nhiêu lần
    cũng không nén chồng lên nhau (nén hai lần là giọng bẹt), mà tập cũ dựng từ
    trước khi có tính năng này vẫn được chuẩn lại.
    """
    import shutil
    wav = Path(wav)
    if VOICE_LUFS is None or not wav.exists():
        return False
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        return False
    d = _measure_loudness(wav, ffmpeg, seconds=300, target=VOICE_LUFS)
    if d and abs(d["input_i"] - VOICE_LUFS) <= 0.7:
        logging.info(f"🔊 {wav.name} đã đúng mức chuẩn ({d['input_i']:.1f} LUFS) "
                     "→ không chỉnh lại.")
        return False
    if status_var is not None:
        status_var.set("Đang chuẩn hoá âm lượng bản cũ...")
    logging.info("🔊 Audio cũ chưa chuẩn âm lượng → chuẩn lại (nén câu + R128).")
    return _loudnorm_file(wav, VOICE_LUFS, pre_filter=SPEECH_COMPRESSOR) == wav


def _probe_duration(path: Path) -> float | None:
    """Độ dài file (giây) bằng ffprobe. None nếu không đo được."""
    import shutil
    import subprocess
    ffprobe = shutil.which("ffprobe")
    if not ffprobe:
        return None
    try:
        r = subprocess.run(
            [ffprobe, "-v", "error", "-show_entries", "format=duration",
             "-of", "default=nw=1:nk=1", str(path)],
            capture_output=True, text=True, encoding="utf-8", errors="replace",
            creationflags=CREATE_NO_WINDOW)
        return float((r.stdout or "").strip())
    except (ValueError, OSError):
        return None


def make_youtube_sub(video_path: Path, script_path: Path, mode: str,
                     model: str = "medium", max_chars: int = 50,
                     progress=None, doc: bool = False,
                     kieu: str = "hopbo", font: str = "",
                     mau: str = "", vitri="", cochu="",
                     dong: int = 1, mau_vien: str = "",
                     bengang="") -> Path | None:
    """Tạo phụ đề cho 1 video. Trả về file .srt, None nếu lỗi.

    doc=False (mặc định): khung NGANG 1920x1080 — YOUTUBE.mp4.
    doc=True            : khung DỌC  1080x1920 — facebook.mp4 / tiktok.mp4. Đổi
                          PlayRes + MarginV của file .ass, và ép max_chars xuống
                          SUB_MAX_CHARS_DOC nếu bên gọi để số của khung ngang
                          (50 ký tự đặt vào khung 1080 là khung nền tràn ra ngoài).

    Tái dùng nguyên module video_gansub: Whisper CHỈ để lấy mốc giờ, còn CHỮ lấy
    từ input.txt gốc nên không bao giờ sai chính tả. Chạy ngay trong tiến trình
    này (không bung subprocess) để dùng chung log + thanh tiến độ.

    cochu = % phóng to/thu nhỏ chữ của kiểu (rỗng/100 = cỡ gốc), dong = số dòng
    tối đa mỗi lần hiện chữ (1 như cũ, 2 = gom hai dòng giống ảnh mẫu kho kiểu),
    mau_vien = màu viền quanh chữ (rỗng = viền gốc của kiểu),
    bengang = % BỀ NGANG dòng chữ so với bề ngang chuẩn (rỗng/100 = như cũ) —
    chỉ đổi chỗ xuống dòng, không đổi cỡ chữ (xem kieusub.ap_bengang).

    mode SUB_MODE_BURN thì đốt chữ vào hình rồi GHI ĐÈ lên video gốc — burn ra file
    tạm trước, xong mới thay thế, để video cũ không mất nếu ffmpeg chết giữa chừng.

    Whisper được giải phóng khỏi VRAM ở cuối: bước sau còn dựng video bằng NVENC,
    thiếu VRAM là rớt về CPU (libx264) chậm hẳn.
    """
    video_path, script_path = Path(video_path), Path(script_path)
    if not video_path.exists():
        logging.warning(f"📝 Chưa có video để gắn phụ đề: {video_path.name}")
        return None
    if not script_path.is_file() or script_path.stat().st_size == 0:
        logging.warning(f"📝 Không có kịch bản ({script_path.name}) → bỏ qua phụ đề.")
        return None

    tmp_wav = None
    try:
        import kieusub
        import video_gansub as gs
        from nhandien_giongnoi import extract_audio, get_audio_duration

        # Trần ký tự/dòng theo KIỂU phụ đề (font to → trần nhỏ) + khung hình:
        # khung DỌC hẹp hơn phải ngắt dòng ngắn lại, không thì chữ tràn ra mép
        # (kiểu mặc định hopbo ra đúng 23 như SUB_MAX_CHARS_DOC cũ).
        # `font`/`mau` khác rỗng = đè font/màu chữ của kiểu (ap_font thay cả
        # fontfile đo chữ; ap_mau chỉ đổi mau_chu, viền/nền giữ nguyên).
        # `cochu` (% cỡ chữ) phóng to/thu nhỏ chữ của kiểu và rút trần ký tự/dòng
        # theo tỉ lệ ngược — phải áp TRƯỚC chon_max_chars mới cắt dòng đúng.
        # `bengang` (% bề ngang) áp SAU cochu: chỉ ghi ti_ngang cho
        # chon_max_chars nhân vào trần — hẹp thì xuống dòng sớm, nới thì dài ra.
        kieu_dict = kieusub.ap_bengang(kieusub.ap_cochu(
            kieusub.ap_mau_vien(
                kieusub.ap_mau(kieusub.ap_font(kieusub.lay(kieu), font), mau),
                mau_vien), cochu), bengang)
        mc = kieusub.chon_max_chars(kieu_dict, max_chars, doc=doc)
        if mode == SUB_MODE_BURN and mc != max_chars:
            logging.info(f"📝 Kiểu '{kieu_dict['ten']}'"
                         + (" + khung dọc" if doc else "")
                         + f": đổi {max_chars} → {mc} ký tự/dòng theo "
                         "khung + bề ngang.")
            max_chars = mc
        elif doc and max_chars > gs.SUB_MAX_CHARS_DOC:
            # Chế độ .srt rời không có kiểu — giữ phép rút cũ của khung dọc.
            max_chars = gs.SUB_MAX_CHARS_DOC

        # dong=2: mỗi lần hiện chữ gom TỚI 2 DÒNG (như ảnh mẫu kho kiểu) thay vì
        # 1 dòng — chữ ở lại lâu gấp đôi, đọc đỡ hụt.
        cues = gs.build_cues(script_path.read_text(encoding="utf-8"),
                             max_chars, dong)
        if not cues:
            logging.warning("📝 Kịch bản trống → không có gì làm phụ đề.")
            return None
        logging.info(f"📝 Cắt {len(cues)} dòng phụ đề từ {script_path.name}, "
                     f"đang nghe audio bằng Whisper '{model}'...")

        # Nghe TIẾNG CỦA CHÍNH VIDEO (không phải output.wav) — video có thể đã tăng
        # tốc x1.1, lấy nhầm audio gốc là phụ đề trôi dần.
        tmp_wav = str(video_path.with_name(video_path.stem + "_sub16k.wav"))
        if not extract_audio(str(video_path), tmp_wav):
            return None
        audio_dur = get_audio_duration(tmp_wav) or 0

        words, starts, ends = gs.transcribe_words(tmp_wav, model, on_progress=progress)
        if not words:
            logging.warning("📝 Whisper không nghe được từ nào → phụ đề rải đều theo thời lượng.")
        cue_times = gs.align_cues(cues, words, starts, ends, audio_dur)

        srt_path = video_path.with_name(video_path.stem + ".srt")
        gs.write_srt(cues, cue_times, srt_path)
        logging.info(f"📝 Đã ghi phụ đề → {srt_path.name} ({len(cues)} dòng)")

        if mode == SUB_MODE_BURN:
            geo = ((gs.ASS_PLAY_W_DOC, gs.ASS_PLAY_H_DOC, gs.SUB_MARGIN_V_DOC) if doc
                   else (gs.ASS_PLAY_W, gs.ASS_PLAY_H, gs.SUB_MARGIN_V))
            # `vitri` (% chiều cao từ đáy) đè MarginV mặc định — áp cho cả hai
            # khung theo cùng tỉ lệ %, rỗng thì giữ 173/380 như cũ.
            geo = (geo[0], geo[1], kieusub.vitri_margin(vitri, geo[1], geo[2]))
            ass_path = kieusub.write_ass_kieu(cues, cue_times,
                                              srt_path.with_suffix(".ass"),
                                              kieu_dict, *geo)
            tmp_mp4 = video_path.with_name(video_path.stem + "_subtmp.mp4")
            logging.info("📝 Đang vẽ CỨNG phụ đề vào hình (mã hoá lại, hơi lâu)...")
            if gs.burn_subs(video_path, ass_path, tmp_mp4):
                os.replace(tmp_mp4, video_path)
                logging.info(f"📝 Đã đốt phụ đề vào {video_path.name}")
            else:
                tmp_mp4.unlink(missing_ok=True)
                logging.error("📝 Vẽ cứng phụ đề LỖI → giữ nguyên video, chỉ còn file .srt.")
        return srt_path
    except Exception as e:
        logging.error(f"📝 Lỗi tạo phụ đề: {e}")
        return None
    finally:
        if tmp_wav and os.path.exists(tmp_wav):
            try:
                os.remove(tmp_wav)
            except Exception:
                pass
        try:
            from nhandien_giongnoi import free_model
            free_model()                      # trả VRAM lại cho NVENC dựng video sau
        except Exception:
            pass


def _mix_bg_music(voice_wav: Path, music_file: Path, below_db: float,
                  out_wav: Path) -> Path:
    """Trộn NHẠC NỀN vào giọng rồi CHUẨN HOÁ cả bản trộn về mức phát hành.

    Nhạc được hạ xuống thấp hơn GIỌNG đúng |below_db| LUFS (vd giọng -16 LUFS,
    below=-12 → nhạc ≈ -28 LUFS). Nhạc LẶP cho đủ dài, fade-in 1s + fade-out cuối
    bài, cắt bằng độ dài giọng. Trả về out_wav (giữ nguyên độ dài giọng).

    Cân theo LUFS chứ không theo đỉnh: nhạc có đỉnh cao mà nghe rất nhỏ là chuyện
    thường (trống, tiếng gảy dây), lấy đỉnh làm mốc thì bài nào cũng ra một mức
    nhạc khác nhau. Đo lỗi thì lùi về cách cũ (so đỉnh) chứ không bỏ nhạc.

    Trộn ra file tạm 32-bit float (không có trần 0dBFS nên cộng hai nguồn không
    thể vỡ tiếng), xong mới chuẩn về MIX_LUFS rồi ghi ra out_wav 16-bit.

    Fade-out CHỈ áp cho nhạc, KHÔNG áp cho giọng — nếu fade cả bản trộn thì câu
    kết bị nhỏ dần đi.
    """
    import shutil
    import subprocess
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        raise RuntimeError("Không tìm thấy ffmpeg trong PATH.")
    voice_wav, music_file, out_wav = Path(voice_wav), Path(music_file), Path(out_wav)

    # Mốc so sánh: độ to cảm nhận (LUFS). Nhạc chỉ cần đo 5 phút đầu là đủ đại diện.
    voice_l = _measure_loudness(voice_wav, ffmpeg)
    music_l = _measure_loudness(music_file, ffmpeg, seconds=300)
    if voice_l and music_l:
        gain_db = (voice_l["input_i"] + below_db) - music_l["input_i"]
        gain_db = max(-45.0, min(gain_db, 12.0))   # chặn hai đầu, tránh số vô lý
        logging.info(f"🎼 Nhạc {music_l['input_i']:.1f} LUFS, giọng "
                     f"{voice_l['input_i']:.1f} LUFS → chỉnh nhạc {gain_db:+.1f} dB")
    else:
        voice_peak = _detect_peak_db(voice_wav, ffmpeg)
        music_peak = _detect_peak_db(music_file, ffmpeg)
        if voice_peak is None:
            voice_peak = -6.0                     # giọng OmniVoice chuẩn hoá ≈ -6dBFS
        target_db = voice_peak + below_db         # đỉnh nhạc mong muốn (dưới giọng)
        gain_db = (target_db - music_peak) if music_peak is not None else target_db
        gain_db = min(gain_db, 0.0)               # không khuếch đại nhạc vượt gốc
        logging.warning(f"🎼 Không đo được LUFS → cân nhạc theo đỉnh ({gain_db:+.1f} dB).")

    # Độ dài giọng = độ dài bản trộn (amix duration=first) → biết mốc bắt đầu fade.
    voice_sec = _probe_duration(voice_wav)
    fade_out = ""
    if voice_sec and voice_sec > BGM_FADE_OUT_SEC * 2:
        st = voice_sec - BGM_FADE_OUT_SEC
        fade_out = f"afade=t=out:st={st:.3f}:d={BGM_FADE_OUT_SEC:g},"
    else:
        logging.warning("⚠️ Không đo được độ dài giọng → bỏ fade-out nhạc nền.")

    filt = (f"[1:a]volume={gain_db:.2f}dB,afade=t=in:d=1.0,{fade_out}"
            f"aresample=44100[bg];"
            f"[0:a][bg]amix=inputs=2:duration=first:normalize=0[a]")
    raw_mix = out_wav.with_name(out_wav.stem + "_tronthô.tmp.wav")
    cmd = [ffmpeg, "-y", "-i", str(voice_wav),
           "-stream_loop", "-1", "-i", str(music_file),   # lặp nhạc cho đủ dài
           "-filter_complex", filt, "-map", "[a]",
           "-c:a", "pcm_f32le", str(raw_mix)]
    r = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8",
                       errors="replace", creationflags=CREATE_NO_WINDOW)
    if r.returncode != 0 or not raw_mix.exists():
        raw_mix.unlink(missing_ok=True)
        raise RuntimeError(f"ffmpeg mix nhạc lỗi: {(r.stderr or '')[-300:] or r.returncode}")

    try:
        # Bản trộn to hơn giọng trần (cộng thêm nhạc) → chuẩn lại lần cuối cho
        # đúng mức nền tảng, đây mới là audio đi vào video.
        done = _loudnorm_file(raw_mix, MIX_LUFS, out=out_wav, label=out_wav.name)
        if done != out_wav:
            os.replace(raw_mix, out_wav)          # chuẩn hoá lỗi → dùng bản trộn thô
    finally:
        raw_mix.unlink(missing_ok=True)
    if not out_wav.exists():
        raise RuntimeError("Không tạo được file trộn nhạc nền.")
    return out_wav


def _render_tiktok_caption_png(text: str, out_png: Path, canvas=(1080, 1920),
                               y_ratio: float = 0.40) -> Path | None:
    """Vẽ 'Mimi audio Số ..' ĐẸP lên PNG TRONG SUỐT đúng khung dọc, TÂM ở ~y_ratio
    chiều cao (0.40 = 40%). Trả về out_png, hoặc None nếu lỗi.

    Thiết kế: nền pill bo góc bán trong suốt + VIỀN vàng + bóng đổ mềm; chữ tô
    GRADIENT vàng→cam có VIỀN tối. Dùng PIL (tiếng Việt có dấu tốt) → tránh
    drawtext/escaping của ffmpeg.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont, ImageFilter
    except Exception as e:
        logging.warning(f"Không tạo được chữ TikTok (thiếu PIL): {e}")
        return None
    W, H = canvas
    # Font đậm hỗ trợ tiếng Việt (thử vài font Windows quen thuộc).
    font = None
    for fp, sz in [("C:/Windows/Fonts/segoeuib.ttf", 78),
                   ("C:/Windows/Fonts/arialbd.ttf", 78),
                   ("C:/Windows/Fonts/tahomabd.ttf", 74),
                   ("C:/Windows/Fonts/arial.ttf", 78)]:
        try:
            font = ImageFont.truetype(fp, size=sz)
            break
        except Exception:
            continue
    if font is None:
        font = ImageFont.load_default()

    base = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    measure = ImageDraw.Draw(base)
    stroke = 4
    l, t, r, b = measure.textbbox((0, 0), text, font=font, stroke_width=stroke)
    tw, th = r - l, b - t
    cx, cy = W / 2, H * y_ratio

    # ── Nền pill bo góc + viền vàng + bóng đổ mềm ──
    pad_x, pad_y = 54, 30
    pw, ph = tw + pad_x * 2, th + pad_y * 2
    x0, y0, x1, y1 = cx - pw / 2, cy - ph / 2, cx + pw / 2, cy + ph / 2
    radius = int(ph / 2)
    shadow = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    ImageDraw.Draw(shadow).rounded_rectangle([x0, y0 + 8, x1, y1 + 8], radius=radius,
                                             fill=(0, 0, 0, 130))
    base = Image.alpha_composite(base, shadow.filter(ImageFilter.GaussianBlur(12)))
    panel = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    ImageDraw.Draw(panel).rounded_rectangle([x0, y0, x1, y1], radius=radius,
                                            fill=(25, 12, 35, 190),
                                            outline=(255, 205, 60, 255), width=6)
    base = Image.alpha_composite(base, panel)

    # ── Chữ: viền tối vẽ trước, rồi tô gradient vàng→cam qua mask ──
    tx = cx - tw / 2 - l
    ty = cy - th / 2 - t
    outline = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    ImageDraw.Draw(outline).text((tx, ty), text, font=font, fill=(0, 0, 0, 0),
                                 stroke_width=stroke, stroke_fill=(70, 25, 0, 255))
    base = Image.alpha_composite(base, outline)
    grad = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    gd = ImageDraw.Draw(grad)
    top_c, bot_c = (255, 244, 170), (255, 138, 20)
    gy0, gy1 = int(cy - th / 2), int(cy + th / 2)
    for yy in range(gy0, gy1 + 1):
        f = (yy - gy0) / max(1, gy1 - gy0)
        col = tuple(int(top_c[i] + (bot_c[i] - top_c[i]) * f) for i in range(3))
        gd.line([(int(x0), yy), (int(x1), yy)], fill=col + (255,))
    mask = Image.new("L", (W, H), 0)
    ImageDraw.Draw(mask).text((tx, ty), text, font=font, fill=255)
    base.paste(grad, (0, 0), mask)

    try:
        out_png.parent.mkdir(parents=True, exist_ok=True)
        base.save(str(out_png))
        return out_png
    except Exception as e:
        logging.warning(f"Không lưu được ảnh chữ TikTok: {e}")
        return None


def _play_done_sound(success: bool = True) -> None:
    """Phát âm báo khi chạy xong (async, không chặn; bỏ qua nếu không phát được)."""
    try:
        import winsound
        # SystemAsterisk = báo nhẹ khi hoàn tất; SystemHand = âm báo lỗi.
        alias = "SystemAsterisk" if success else "SystemHand"
        winsound.PlaySound(alias, winsound.SND_ALIAS | winsound.SND_ASYNC)
    except Exception:
        pass


class _NullWidget:
    """Nút giả cho run_tts khi chạy ở chế độ batch (không có nút thật để bật/tắt)."""
    def config(self, *args, **kwargs):
        pass
    configure = config


def run_tts(mode, voice_param, chunks, output, progress_var, status_var, btn_run, btn_pause, btn_preview, pause_event, make_video=False, effect=None, make_video_doc=False, doc_speed=1.0, doc_percent=100, ngang_speed=1.0, reuse=False, doc_from_ngang=False, doc_no_effect=False, doc_from_subfolder=False, ngang_out=None, doc_out=None, make_tiktok=False, tiktok_out=None, tiktok_speed=1.0, tiktok_no_effect=False, tiktok_caption=None, tiktok_caption_pos=40, tiktok_music=False, tiktok_music_db=-12.0, video_only=False, ngang_source=None, tiktok_percent=50, make_sub=False, sub_mode=SUB_MODE_SRT, sub_model="medium", sub_max_chars=50, sub_kieu="hopbo", sub_font="", sub_mau="", sub_vitri="", sub_cochu="", sub_bengang="", sub_dong=1, sub_mau_vien="", make_sub_doc=False, make_short=False, short_out=None):
    import torch
    from omnivoice.models.omnivoice import OmniVoice
    from omnivoice.utils.common import get_best_device

    failed = False
    model = None        # khai báo sớm: khối finally dưới cùng luôn dọn được VRAM

    # Bước dựng video báo tiến trình qua THANH (progress_var) + dòng trạng thái,
    # KHÔNG spam % ra nhật ký. label đứng trước, % chạy trên thanh.
    def _video_progress(label):
        def _cb(pct, cur, total, speed):
            progress_var.set(int(pct))
            status_var.set(f"{label} {pct:.0f}%  ({cur:.0f}/{total:.0f}s · {speed})")
        return _cb

    try:
        total = len(chunks)
        output_path = Path(output)
        tmp_dir = chunks_dir_for(output_path)
        tmp_dir.mkdir(parents=True, exist_ok=True)

        audio_ready = output_path.exists() and output_path.stat().st_size > 4096

        # video_only: BỎ QUA tạo giọng, chỉ DỰNG VIDEO từ audio có sẵn (nút "Dựng
        # lại" riêng của từng mục). Cần audio đã tồn tại.
        if video_only:
            if not audio_ready:
                logging.error(f"🎬 Chưa có audio để dựng video: {output_path}")
                status_var.set("🎬 Chưa có audio — hãy tạo giọng trước.")
                return
            sig = old_sig = None
            reuse_audio = True
        else:
            # Chữ ký cấu hình (giọng/chế độ/văn bản) — quyết định có thể DÙNG LẠI
            # audio cũ hay phải tạo lại (text/giọng đổi → chữ ký khác → tạo lại).
            # NUM_STEP nằm trong chữ ký: đổi số bước là đổi chất lượng giọng, mà
            # ghép đoạn 32 bước với đoạn 16 bước trong cùng một bài thì nghe cứ
            # lệch lệch mà chẳng biết tại đâu. Cho vào đây là đổi số → sinh lại hết.
            sig = hashlib.sha1("|".join([mode, str(voice_param), f"step{NUM_STEP}",
                                         f"fade{FADE_OUT_SEC}", *chunks])
                               .encode("utf-8")).hexdigest()
            sig_file = tmp_dir / "_signature.txt"
            old_sig = sig_file.read_text(encoding="utf-8").strip() if sig_file.exists() else None
            # ♻ Dùng lại: chỉ tái dùng khi audio đã có VÀ chữ ký khớp (cùng văn bản/giọng).
            # Nếu văn bản/giọng đã đổi thì vẫn tạo lại để không ghép nhầm bản cũ.
            reuse_audio = reuse and audio_ready and old_sig == sig

        # Dựng video: chế độ dùng-lại BỎ QUA video đã có (chỉ dựng phần thiếu). Nhưng
        # nút "Dựng lại" (video_only) thì LUÔN dựng lại video dù đã tồn tại.
        skip_video = reuse_audio and not video_only
        # Audio DẪN XUẤT (bản %/tiktok/nhạc nền): luồng ♻ thường thì tái dùng cho
        # nhanh; nhưng nút "Dựng lại" (video_only) TẠO LẠI để áp cài đặt GUI mới nhất
        # (tốc độ, dB nhạc…). Audio GỐC (output.wav) vẫn dùng lại (không chạy lại TTS).
        reuse_derived = reuse_audio and not video_only

        if reuse_audio:
            logging.info(f"♻ Dùng lại audio đã có (bỏ qua tạo giọng): {output_path.name}")
            status_var.set(f"♻ Dùng audio đã có → {output_path.name}")
            # Audio của các tập dựng bằng bản cũ chưa qua chuẩn hoá → chuẩn nốt,
            # không thì tập cũ dựng lại vẫn to nhỏ thất thường. Chỉ đụng vào file
            # còn LỆCH mức đích: file đã chuẩn rồi mà nén thêm lần nữa là bẹt giọng.
            if _chuan_am_neu_can(output_path, status_var):
                # Audio gốc vừa đổi mức → bản %/tăng tốc/trộn nhạc làm từ bản
                # CŨ không còn khớp nữa, phải dựng lại từ bản vừa chuẩn.
                reuse_derived = False
            progress_var.set(100)
            btn_preview.config(state="normal")
        else:
            if reuse and not audio_ready:
                logging.info("♻ Chưa có audio để dùng lại → tạo giọng từ đầu.")
            elif reuse and old_sig != sig:
                logging.info("♻ Văn bản/giọng đã đổi → tạo lại giọng (không dùng bản cũ).")

            # Chữ ký khác lần trước → xóa chunk cũ để tạo lại, tránh ghép nhầm.
            if old_sig != sig:
                stale = list(tmp_dir.glob("*.wav"))
                for w in stale:
                    w.unlink()
                if stale:
                    logging.warning(
                        f"Cấu hình đổi (giọng/chế độ/văn bản) → xóa {len(stale)} "
                        "chunk cũ, tạo lại từ đầu."
                    )
                sig_file.write_text(sig, encoding="utf-8")

            logging.info(f"Tổng {total} đoạn — tải model...")
            status_var.set("Đang tải model...")

            device = get_best_device()
            model = OmniVoice.from_pretrained(
                "k2-fsa/OmniVoice", device_map=device, dtype=torch.float16
            )
            sr = model.sampling_rate

            for i, chunk in enumerate(chunks):
                # Chờ nếu đang tạm dừng
                if not pause_event.is_set():
                    status_var.set(f"⏸  Đã tạm dừng (đoạn {i+1}/{total})")
                    logging.info("Tạm dừng...")
                    pause_event.wait()
                    status_var.set(f"Đoạn {i+1}/{total}...")
                    logging.info("Tiếp tục.")

                tmp_file = tmp_dir / f"{i:04d}.wav"
                pct = int((i / total) * 100)
                progress_var.set(pct)
                status_var.set(f"Đoạn {i+1}/{total}...")

                if tmp_file.exists() and tmp_file.stat().st_size > 4096:
                    logging.info(f"[{i+1}/{total}] bỏ qua (đã có)")
                    continue
                elif tmp_file.exists():
                    logging.warning(f"[{i+1}/{total}] file cũ bị lỗi/cụt → generate lại")
                    tmp_file.unlink()

                logging.info(f"[{i+1}/{total}] {chunk[:60]!r}")
                result = _generate_chunk(model, mode, voice_param, chunk)
                sf.write(str(tmp_file), result[0], sr)

                # Mỗi đoạn một độ dài chuỗi khác nhau → allocator giữ lại đủ kiểu
                # khối rời. Dọn định kỳ để phần GIỮ LẠI không phình dần tới mức
                # tràn sang RAM (xem chú thích PYTORCH_CUDA_ALLOC_CONF đầu file).
                if (i + 1) % VRAM_CLEAN_EVERY == 0:
                    torch.cuda.empty_cache()
                    try:
                        free_b, total_b = torch.cuda.mem_get_info()
                        logging.info(f"🧹 Dọn VRAM sau {i+1}/{total} đoạn — còn trống "
                                     f"{free_b/2**30:.1f}/{total_b/2**30:.1f} GB.")
                    except Exception:
                        pass

            # ── KIỂM TRA SPIKE SAU KHI GENERATE XONG ────────────────────────
            status_var.set("Kiểm tra chất lượng audio...")
            logging.info("Kiểm tra spike toàn bộ chunks...")
            bad_chunks = []
            for i in range(total):
                f = tmp_dir / f"{i:04d}.wav"
                spikes = detect_spike(f, sr)
                if spikes:
                    bad_chunks.append(i)
                    logging.warning(f"  [SPIKE] {f.name} tại {spikes[:3]}s → render lại")

            if bad_chunks:
                logging.info(f"Render lại {len(bad_chunks)} chunk lỗi: {bad_chunks}")
                for idx, i in enumerate(bad_chunks):
                    status_var.set(f"Render lại chunk lỗi {idx+1}/{len(bad_chunks)}...")
                    tmp_file = tmp_dir / f"{i:04d}.wav"
                    tmp_file.unlink(missing_ok=True)
                    result = _generate_chunk(model, mode, voice_param, chunks[i])
                    sf.write(str(tmp_file), result[0], sr)
                    logging.info(f"  [{i:04d}] render lại xong")
            else:
                logging.info("Không phát hiện spike — tất cả chunk OK.")

            # ── KIỂM TRA MẤT CHỮ + TIẾNG LẠ (ASR đối chiếu từng đoạn) ───────
            # OmniVoice thi thoảng NUỐT nguyên câu ngắn (đo tập 49/53: 7–12% số
            # đoạn — hay dính nhất là câu thuật thoại "tôi hỏi"/"tôi gật đầu"
            # và cụm lặp "không chia. không chia."). detect_spike không thấy
            # được vì audio vẫn sạch, chỉ thiếu lời. Đối chiếu ASR ↔ văn bản để
            # bắt rồi render lại; lượt 2 nới khung thời lượng +6% cho model đủ
            # chỗ nhét lại câu đã nuốt.
            #
            # Cùng lượt ASR đó còn soi TIẾNG LẠ (rè/rít/ù không phải giọng
            # đọc — detect_spike bó tay khi nó ở mức âm lượng ngang lời nói):
            # nghe thừa chữ, độ tin cậy whisper tụt, và vùng ngoài lời nói mà
            # RMS vẫn cao (xem taogiong_kiemtra_matchu). Đoạn dính cờ nào cũng
            # về chung dict `miss` nên vòng render-lại dưới đây xử lý cả hai
            # loại lỗi như nhau; với tiếng lạ, render lại là re-roll — nới
            # khung +6% ở lượt 2 vô hại.
            #
            # VRAM: bước nghe lại chạy large-v3-turbo trên GPU (nhanh hơn hẳn
            # small/medium trên CPU của bản cũ), mà card 8GB không chứa nổi cả
            # nó lẫn OmniVoice → luân phiên nhau: dỡ OmniVoice → quét → có đoạn
            # thiếu thì trả VRAM của turbo, nạp lại OmniVoice render lại, rồi dỡ
            # ra quét tiếp. Nạp lại chỉ tốn thời gian trọng số: prompt giọng mẫu
            # đã lưu sẵn ra .prompt.pt nên KHÔNG phải chạy lại Whisper phiên âm
            # giọng mẫu (xem _voice_prompt).
            try:
                import taogiong_kiemtra_matchu as _mc
            except Exception as e:
                _mc = None
                logging.warning(f"⚠️ Bỏ qua kiểm tra mất chữ + tiếng lạ (thiếu module: {e})")
            if _mc is not None:
                model = _do_omnivoice(model)
                _tra_vram()          # dỡ xong mới thật sự còn chỗ cho turbo
                status_var.set("Kiểm tra mất chữ & tiếng lạ (ASR đối chiếu)...")
                logging.info("Kiểm tra mất chữ + tiếng lạ toàn bộ chunks (ASR đối chiếu văn bản)...")
                miss = _mc.quet_va_xac_minh(chunks, tmp_dir, on_log=logging.info,
                                            status=status_var.set)
                for attempt in (1, 2):
                    if not miss:
                        break
                    logging.warning(
                        f"📢 {len(miss)} đoạn nuốt chữ / có tiếng lạ → render lại "
                        f"(lượt {attempt}): {sorted(miss)}")
                    # Trả VRAM của ASR rồi mới nạp OmniVoice — hai model không
                    # cùng lúc nằm trên card 8GB được.
                    _mc.giai_phong()
                    status_var.set("Nạp lại OmniVoice để render lại...")
                    logging.info("↩ Nạp lại OmniVoice để render các đoạn lỗi...")
                    model = OmniVoice.from_pretrained(
                        "k2-fsa/OmniVoice", device_map=device, dtype=torch.float16
                    )
                    for i in sorted(miss):
                        logging.info(f"  [{i:04d}] lỗi: {' | '.join(miss[i])}")
                        status_var.set(f"Render lại đoạn lỗi {i+1}/{total} "
                                       f"(lượt {attempt})...")
                        tmp_file = tmp_dir / f"{i:04d}.wav"
                        tmp_file.unlink(missing_ok=True)
                        result = _generate_chunk(
                            model, mode, voice_param, chunks[i],
                            dur_scale=1.0 if attempt == 1 else 1.06)
                        sf.write(str(tmp_file), result[0], sr)
                        spikes = detect_spike(tmp_file, sr)
                        if spikes:
                            logging.warning(f"  [{i:04d}] bản render lại có spike "
                                            f"tại {spikes[:3]}s — nghe kiểm tra tay.")
                    model = _do_omnivoice(model)   # trả VRAM lại cho lượt quét sau
                    _tra_vram()
                    miss = _mc.quet_va_xac_minh(chunks, tmp_dir,
                                                on_log=logging.info,
                                                chi_cac_doan=set(miss),
                                                status=status_var.set)
                if miss:
                    for i in sorted(miss):
                        logging.error(
                            f"⛔ Đoạn {i:04d} VẪN lỗi sau 2 lần render lại "
                            f"({' | '.join(miss[i])}) — nghe kiểm tra tay: "
                            f"{chunks[i][:60]!r}")
                else:
                    logging.info("✅ Kiểm tra mất chữ + tiếng lạ: tất cả đoạn đạt.")
                _mc.giai_phong()

            status_var.set("Đang ghép file...")
            logging.info("Ghép tất cả đoạn...")
            parts = [
                sf.read(str(tmp_dir / f"{i:04d}.wav"), dtype="float32")[0]
                for i in range(total)
            ]

            # Cân âm lượng GIỮA các đoạn trước khi ghép (xem _balance_chunk_levels).
            # Phải làm ở đây chứ không phải trên file đã ghép: sau khi ghép thì
            # không còn biết ranh giới đoạn nào với đoạn nào nữa.
            status_var.set("Đang cân âm lượng các đoạn...")
            parts = _balance_chunk_levels(parts, sr)

            # Crossfade ngắn giữa các chunk để tránh click/vấp tại ranh giới
            fade = min(256, min(len(p) for p in parts) // 2)
            fade_in  = np.linspace(0, 1, fade, dtype="float32")
            fade_out = np.linspace(1, 0, fade, dtype="float32")
            merged = parts[0].copy()
            merged[-fade:] *= fade_out
            for p in parts[1:]:
                p = p.copy()
                p[:fade] *= fade_in
                merged[-fade:] += p[:fade]
                merged = np.concatenate([merged, p[fade:]])

            # OmniVoice kết thúc chunk cuối chỉ ~0.13s sau từ cuối cùng → audio
            # (và video) dừng phựt. Đệm thêm im lặng cho có chỗ thở.
            merged = np.concatenate(
                [merged, np.zeros(int(TAIL_SILENCE_SEC * sr), dtype="float32")])

            sf.write(output, merged, sr)

            # San bằng mức GIỮA CÁC CÂU + chuẩn độ to cả bài về mức phát hành.
            # Làm NGAY tại đây vì mọi bản dẫn xuất (cắt TikTok, tăng tốc, trộn
            # nhạc) đều lấy từ output.wav — chuẩn một lần ở gốc là cả dây
            # chuyền chuẩn theo.
            status_var.set("Đang chuẩn hoá âm lượng...")
            _loudnorm_file(output_path, VOICE_LUFS, pre_filter=SPEECH_COMPRESSOR)

            progress_var.set(100)
            status_var.set(f"Xong!  →  {output}")
            logging.info(f"Đã lưu → {output}")
            btn_preview.config(state="normal")   # cho phép nghe thử kết quả

            # ── GIẢI PHÓNG OMNIVOICE KHỎI VRAM ─────────────────────────────
            # Audio đã tạo + ghép xong; các bước còn lại (cắt, dựng video bằng
            # h264_nvenc) KHÔNG dùng tới OmniVoice. Model nạp mới mỗi lần chạy
            # (không cache), nên xóa ngay để trả ~vài GB VRAM cho NVENC — tránh
            # thiếu VRAM khiến dựng video rớt về CPU (libx264) chậm.
            # Bước kiểm tra mất chữ ở trên thường đã dỡ sẵn rồi (model = None) →
            # lúc đó hàm này không làm gì.
            model = _do_omnivoice(model, "trước khi dựng video")
            _tra_vram()

        # ── TỰ DỰNG VIDEO NGANG TỪ AUDIO FULL (nếu bật) ────────────────────
        ngang_video_path = None   # video ngang vừa dựng (để video dọc dùng lại nếu bật)
        if make_video:
            # Tăng tốc audio (giữ cao độ) trước khi dựng — nếu chọn mức > 1.0
            ngang_audio = output_path
            if ngang_speed and ngang_speed > 1.001:
                status_var.set(f"Đang tăng tốc audio x{ngang_speed:.2f} cho video ngang...")
                try:
                    ngang_audio = _speedup_audio_for_doc(output_path, ngang_speed)
                    logging.info(f"⏩ Tăng tốc audio video ngang x{ngang_speed:.2f} → {ngang_audio.name}")
                except Exception as e:
                    logging.warning(f"Không tăng tốc được audio video ngang (giữ tốc độ gốc): {e}")
                    ngang_audio = output_path
            status_var.set("Đang dựng video...")
            logging.info("Bắt đầu dựng video từ audio vừa tạo...")
            try:
                from video_khung import build_video
                progress_var.set(0)
                ngang_src_dir = resolve_ngang_source(ngang_source, log=logging.info)
                video_out = build_video(ngang_audio, log=logging.info, effect=effect,
                                        progress=_video_progress("🎬 Dựng video ngang..."),
                                        skip_existing=skip_video, output=ngang_out,
                                        source_dir=ngang_src_dir)
                progress_var.set(100)
                ngang_video_path = video_out
                status_var.set(f"Xong! Video → {video_out}")
                logging.info(f"Đã tạo video → {video_out}")
            except Exception as e:
                logging.error(f"Lỗi dựng video: {e}")
                status_var.set(f"Audio xong, lỗi dựng video: {e}")

        # ── (TÙY CHỌN) PHỤ ĐỀ CHO VIDEO NGANG (YouTube) ────────────────────
        # Làm SAU khi video ngang xong (cần nghe tiếng của chính video đó) và TRƯỚC
        # video dọc/TikTok — vì kiểu "vẽ cứng" ghi đè lên YOUTUBE.mp4, mà video dọc
        # có thể dùng lại chính file này làm nguồn hình.
        if make_sub and ngang_video_path and Path(ngang_video_path).exists():
            def _sub_progress(frac):
                pct = int(frac * 100)
                progress_var.set(pct)
                status_var.set(f"📝 Nghe audio làm phụ đề... {pct}%")
            status_var.set("📝 Đang tạo phụ đề cho video YouTube...")
            progress_var.set(0)
            make_youtube_sub(Path(ngang_video_path),
                             Path(ngang_video_path).parent / "input.txt",
                             sub_mode, sub_model, sub_max_chars,
                             progress=_sub_progress, kieu=sub_kieu,
                             font=sub_font, mau=sub_mau, vitri=sub_vitri,
                             cochu=sub_cochu, dong=sub_dong,
                             mau_vien=sub_mau_vien, bengang=sub_bengang)
            progress_var.set(100)

        # ── (TÙY CHỌN) DỰNG VIDEO DỌC (1080x1920, KHÔNG khung) ─────────────
        # Lấy AUDIO FULL; muốn ngắn hơn thì đặt % bên dưới (doc_percent).
        if make_video_doc:
            doc_audio = output_path
            # (TÙY CHỌN) Cắt ~doc_percent% ĐẦU của audio video dọc (cắt ở CUỐI CÂU) —
            # dùng CHUNG cơ chế với TikTok. doc_percent ≥ 100 → giữ nguyên (không cắt),
            # nên tương thích ngược. Cắt TRƯỚC bước tăng tốc.
            dpct = max(10, min(int(doc_percent or 100), 100))
            if dpct < 99:
                try:
                    from video_timclip import (cut_audio_at_sentence_end,
                                               probe_audio_duration)
                    d_total = probe_audio_duration(doc_audio)
                    d_target = d_total * dpct / 100.0
                    # Chỉ cắt khi audio đủ dài và mốc % nằm trong khoảng hợp lý.
                    if d_total > 0 and d_target >= 15 and d_target < d_total - 5:
                        dp_wav = doc_audio.with_name(
                            doc_audio.stem + f"_docpct{dpct}" + doc_audio.suffix)
                        if reuse_derived and dp_wav.exists() and dp_wav.stat().st_size > 4096:
                            doc_audio = dp_wav
                            logging.info(f"♻ Dùng lại audio video dọc {dpct}% đã có: {dp_wav.name}")
                        else:
                            status_var.set(f"Đang cắt ~{dpct}% audio cho video dọc (cắt cuối câu)...")
                            t_min = d_target / 60.0
                            margin = max(0.75, t_min * 0.08)   # cửa sổ ±max(0.75', 8%)
                            cut_seconds, _ = cut_audio_at_sentence_end(
                                doc_audio, dp_wav, target_minutes=t_min,
                                min_minutes=max(0.05, t_min - margin),
                                max_minutes=min(d_total / 60.0, t_min + margin),
                                silence_db=-35.0, min_silence=0.5)
                            doc_audio = dp_wav
                            m, s = divmod(cut_seconds, 60)
                            real_pct = cut_seconds / d_total * 100 if d_total else dpct
                            logging.info(f"✂ Audio video dọc ~{dpct}%: cắt tại {int(m)}:{s:05.2f} "
                                         f"(≈{real_pct:.0f}% thực) → {dp_wav.name}")
                    else:
                        logging.info(f"📱 Video dọc lấy {dpct}% → audio ngắn/mốc sát biên, dùng nguyên audio.")
                except Exception as e:
                    logging.warning(f"Không cắt được ~{dpct}% cho video dọc (dùng nguyên audio): {e}")
            # Tăng tốc audio (giữ cao độ) trước khi dựng — nếu chọn mức > 1.0
            if doc_speed and doc_speed > 1.001:
                status_var.set(f"Đang tăng tốc audio x{doc_speed:.2f} cho video dọc...")
                try:
                    doc_audio = _speedup_audio_for_doc(doc_audio, doc_speed)
                    logging.info(f"⏩ Tăng tốc audio video dọc x{doc_speed:.2f} → {doc_audio.name}")
                except Exception as e:
                    logging.warning(f"Không tăng tốc được audio (giữ tốc độ gốc): {e}")
            # NGUỒN video dọc — ưu tiên theo thứ tự:
            #   1) 📁 Thư mục con videodoc/ (theo videodoc/input.txt) nếu bật + hợp lệ
            #   2) ♻ Dùng lại VIDEO NGANG đã dựng (phóng to khớp chiều cao dọc)
            #   3) Ghép random cả kho videodoc/ (mặc định)
            # Ở mọi trường hợp, audio áp vào là doc_audio (audio video dọc/Facebook).
            ngang_src = None
            doc_source_dir = None
            if doc_from_subfolder:
                doc_source_dir = resolve_videodoc_subfolder(log=logging.info)
                if doc_source_dir is not None:
                    logging.info(f"📁 Video dọc ghép từ thư mục con: videodoc/{doc_source_dir.name}")
            # Chỉ dùng lại video ngang khi KHÔNG dùng (được) thư mục con.
            if doc_from_ngang and doc_source_dir is None:
                if ngang_video_path and ngang_video_path.exists():
                    ngang_src = ngang_video_path
                elif ngang_out and Path(ngang_out).exists():
                    ngang_src = Path(ngang_out)   # bản tự động: YOUTUBE.mp4 đã có
                else:
                    cands = sorted(
                        output_path.parent.glob(output_path.stem + "*_videodone.mp4"),
                        key=lambda p: p.stat().st_mtime, reverse=True)
                    ngang_src = cands[0] if cands else None
                if ngang_src is None:
                    logging.warning("Bật 'dùng lại video ngang' nhưng chưa có video "
                                    "ngang → dựng video dọc bình thường.")
                else:
                    logging.info(f"♻ Video dọc dùng lại video ngang: {ngang_src.name}")

            status_var.set("Đang dựng video dọc...")
            logging.info(f"Bắt đầu dựng video dọc từ {doc_audio.name}...")
            try:
                from video_doc import build_video_doc
                progress_var.set(0)
                # "Không áp hiệu ứng cho video dọc" → bỏ effect ở mọi trường hợp.
                doc_effect = None if doc_no_effect else effect
                # Ảnh bìa = thumbnail dọc của tập, đè lên ĐÚNG 1 frame đầu (xem
                # find_cover_doc). Chưa có thumbnail → None → dựng như cũ.
                vdoc_out = build_video_doc(doc_audio, log=logging.info, effect=doc_effect,
                                           progress=_video_progress("📱 Dựng video dọc..."),
                                           skip_existing=skip_video,
                                           source_video=ngang_src, source_dir=doc_source_dir,
                                           output=doc_out,
                                           cover_png=find_cover_doc(output_path.parent,
                                                                    log=logging.info))
                progress_var.set(100)
                status_var.set(f"Xong! Video dọc → {vdoc_out.name}")
                logging.info(f"Đã tạo video dọc → {vdoc_out.name}")
            except Exception as e:
                logging.error(f"Lỗi dựng video dọc: {e}")
                status_var.set(f"Lỗi video dọc: {e}")

        # ── (TÙY CHỌN) VIDEO TIKTOK — ~tiktok_percent% ĐẦU của audio. Nguồn hình:
        #    • Bật "📁 Ghép từ thư mục con videodoc" → DÙNG LẠI video Facebook đã dựng.
        #    • Ngược lại → ghép NGUYÊN video trong videodoc/ như cũ. File riêng để đăng TikTok.
        if make_tiktok:
            from video_doc import build_video_doc
            from video_timclip import (cut_audio_at_sentence_end,
                                       probe_audio_duration)
            # 1) Audio: lấy ~tiktok_percent% ĐẦU của audio, cắt ở CUỐI CÂU (khoảng lặng
            #    gần mốc % nhất) nên độ dài chỉ xấp xỉ %. ≥99% / audio quá ngắn / cắt
            #    lỗi → dùng nguyên audio.
            tk_audio = output_path
            try:
                total_sec = probe_audio_duration(output_path)
            except Exception:
                total_sec = 0.0
            pct = max(10, min(int(tiktok_percent or 50), 100))
            target_sec = total_sec * pct / 100.0
            if pct < 99 and total_sec > 0 and target_sec >= 15 and target_sec < total_sec - 5:
                tk_wav = output_path.with_name(output_path.stem + "_tiktok" + output_path.suffix)
                if reuse_derived and tk_wav.exists() and tk_wav.stat().st_size > 4096:
                    tk_audio = tk_wav
                    logging.info(f"♻ Dùng lại audio TikTok đã có: {tk_wav.name}")
                else:
                    status_var.set(f"Đang cắt ~{pct}% audio cho TikTok (cắt cuối câu)...")
                    try:
                        # Dò khoảng lặng cuối câu quanh mốc %: cửa sổ ± max(0.75 phút,
                        # 8% mốc), kẹp trong (0, tổng]. Chọn khoảng lặng gần mốc % nhất.
                        target_min = target_sec / 60.0
                        margin_min = max(0.75, target_min * 0.08)
                        min_min = max(0.05, target_min - margin_min)
                        max_min = min(total_sec / 60.0, target_min + margin_min)
                        cut_seconds, _ = cut_audio_at_sentence_end(
                            output_path, tk_wav,
                            target_minutes=target_min, min_minutes=min_min, max_minutes=max_min,
                            silence_db=-35.0, min_silence=0.5)
                        tk_audio = tk_wav
                        m, s = divmod(cut_seconds, 60)
                        real_pct = cut_seconds / total_sec * 100 if total_sec else pct
                        logging.info(f"✂ Audio TikTok ~{pct}%: cắt tại {int(m)}:{s:05.2f} "
                                     f"(≈{real_pct:.0f}% thực) → {tk_wav.name}")
                    except Exception as e:
                        logging.warning(f"Không cắt được ~{pct}% cho TikTok (dùng audio đầy đủ): {e}")
                        tk_audio = output_path
            else:
                logging.info(f"🎵 TikTok lấy {pct}% (≈ nguyên audio {total_sec/60:.1f} phút) → dùng nguyên audio.")

            # 1b) Tăng tốc audio (giữ cao độ) nếu chọn mức > 1.0 — như video ngang/dọc.
            if tiktok_speed and tiktok_speed > 1.001:
                status_var.set(f"Đang tăng tốc audio x{tiktok_speed:.2f} cho TikTok...")
                try:
                    tk_audio = _speedup_audio_for_doc(tk_audio, tiktok_speed)
                    logging.info(f"⏩ Tăng tốc audio TikTok x{tiktok_speed:.2f} → {tk_audio.name}")
                except Exception as e:
                    logging.warning(f"Không tăng tốc được audio TikTok (giữ tốc độ gốc): {e}")

            # Giữ lại bản GIỌNG TRẦN (trước khi trộn nhạc) để bước cắt Short còn dò
            # được khoảng lặng cuối câu — trộn nhạc xong thì nhạc chạy suốt, không
            # còn chỗ nào im lặng để nhận ra ranh giới câu. Xem video_short.
            tk_voice = tk_audio

            # 1c) Chèn NHẠC NỀN (từ Music/), mix nhỏ hơn giọng |tiktok_music_db| dB.
            if tiktok_music:
                musics = list_music_files()
                if not musics:
                    logging.warning(f"🎼 Bật nhạc nền nhưng {MUSIC_DIR} trống → bỏ qua nhạc.")
                else:
                    import random as _rnd
                    music_file = MUSIC_DIR / _rnd.choice(musics)
                    mix_out = (Path(tiktok_out).with_name(Path(tiktok_out).stem + "_bgm.wav")
                               if tiktok_out
                               else tk_audio.with_name(tk_audio.stem + "_bgm.wav"))
                    if reuse_derived and mix_out.exists() and mix_out.stat().st_size > 4096:
                        tk_audio = mix_out
                        logging.info(f"♻ Dùng lại audio TikTok + nhạc đã có: {mix_out.name}")
                    else:
                        status_var.set("Đang chèn nhạc nền cho TikTok...")
                        try:
                            tk_audio = _mix_bg_music(tk_audio, music_file,
                                                     float(tiktok_music_db), mix_out)
                            logging.info(f"🎼 Nhạc nền: {music_file.name} (≈{tiktok_music_db:.0f}dB "
                                         f"dưới giọng) → {mix_out.name}")
                        except Exception as e:
                            logging.warning(f"Không chèn được nhạc nền (giữ giọng gốc): {e}")

            # 2) Nguồn hình TikTok phụ thuộc "📁 Ghép từ thư mục con videodoc":
            #    - BẬT  → DÙNG LẠI video Facebook (facebook.mp4) đã dựng, KHÔNG ghép lại.
            #    - TẮT  → ghép NGUYÊN video trong kho videodoc/ như cũ (giữ nguyên chức năng).
            #    Tên riêng để KHÔNG đè video dọc (facebook/output_doc).
            tk_video_out = (Path(tiktok_out) if tiktok_out
                            else output_path.with_name(output_path.stem + "_tiktok.mp4"))
            tk_source = None
            if doc_from_subfolder:
                fb_video = (Path(doc_out) if doc_out
                            else output_path.with_name(output_path.stem + "_doc.mp4"))
                if fb_video.exists() and fb_video.stat().st_size > 0:
                    tk_source = fb_video
                    logging.info(f"🎵 TikTok dùng lại video Facebook: {fb_video.name}")
                else:
                    logging.warning("Bật ghép thư mục con nhưng chưa có video Facebook "
                                    "→ TikTok ghép từ kho videodoc/.")
            # Chữ overlay 'Mimi audio Số …' ở vị trí % chiều cao do người dùng chọn.
            cap_png = None
            if tiktok_caption:
                y_ratio = max(0.0, min(tiktok_caption_pos / 100.0, 1.0))
                cap_tmp = tk_video_out.with_name(tk_video_out.stem + "_caption.png")
                cap_png = _render_tiktok_caption_png(tiktok_caption, cap_tmp, y_ratio=y_ratio)
                if cap_png:
                    logging.info(f"🔤 Chữ TikTok (≈{tiktok_caption_pos}% cao): {tiktok_caption!r}")
            src_label = "video Facebook" if tk_source else "kho videodoc/"
            status_var.set("Đang dựng video TikTok...")
            logging.info(f"Bắt đầu dựng video TikTok từ {tk_audio.name} ({src_label})...")
            try:
                progress_var.set(0)
                tk_effect = None if tiktok_no_effect else effect
                # Ảnh bìa nằm TRÊN chữ caption, nhưng chỉ ở frame đầu nên chữ
                # 'Mimi audio Số …' vẫn hiện bình thường từ frame thứ 2 trở đi.
                tk_out = build_video_doc(tk_audio, log=logging.info, effect=tk_effect,
                                         progress=_video_progress("🎵 Dựng video TikTok..."),
                                         skip_existing=skip_video,
                                         source_video=tk_source, output=tk_video_out,
                                         caption_png=cap_png,
                                         cover_png=find_cover_doc(output_path.parent,
                                                                  log=logging.info))
                progress_var.set(100)
                status_var.set(f"Xong! Video TikTok → {tk_out.name}")
                logging.info(f"Đã tạo video TikTok → {tk_out.name}")
            except Exception as e:
                logging.error(f"Lỗi dựng video TikTok: {e}")
                status_var.set(f"Lỗi video TikTok: {e}")
            finally:
                # Dọn ảnh chữ tạm (đã nung vào video). Bỏ qua nếu Windows còn khóa.
                if cap_png:
                    try:
                        Path(cap_png).unlink(missing_ok=True)
                    except OSError:
                        pass

            # ── (TÙY CHỌN) YOUTUBE SHORT — cắt ≤2:50 từ CHÍNH video TikTok ────
            # Hình `-c:v copy` nên gần như tức thì và giữ nguyên chất lượng; short
            # thừa hưởng luôn khung dọc, khung tiêu đề, chữ 'Mimi audio Số N', ảnh
            # bìa frame đầu. Riêng TIẾNG thay bằng tk_voice (giọng trần) để không
            # đẩy nhạc nền TikTok lên YouTube — xem video_short.
            # Bước ⑥ đăng YouTube sẽ tự đăng file này sau bản chính 1 giờ.
            if make_short and tk_video_out.exists():
                from video_short import build_short, SHORT_NAME
                sh_out = Path(short_out) if short_out else tk_video_out.with_name(SHORT_NAME)
                status_var.set("Đang cắt video YouTube Short...")
                try:
                    build_short(tk_video_out, sh_out, voice_audio=tk_voice,
                                skip_existing=skip_video, log=logging.info)
                    status_var.set(f"Xong! YouTube Short → {sh_out.name}")
                except Exception as e:
                    # Short là video PHỤ — hỏng thì bỏ, không kéo đổ cả mẻ video.
                    logging.error(f"Lỗi cắt video Short: {e}")

        # ── (TÙY CHỌN) PHỤ ĐỀ CHO VIDEO DỌC (Facebook) ─────────────────────
        # Đặt SAU cùng, sau cả TikTok — CỐ Ý: khi bật "ghép từ thư mục con", video
        # TikTok lấy HÌNH của chính facebook.mp4 nhưng dùng AUDIO ngắn hơn (~50%).
        # Burn sub vào facebook trước là TikTok thừa hưởng phụ đề lệch hẳn nhịp.
        # Làm ở đây thì facebook.mp4 có phụ đề, tiktok.mp4 giữ hình sạch.
        if make_sub_doc:
            doc_video = (Path(doc_out) if doc_out
                         else output_path.with_name(output_path.stem + "_doc.mp4"))
            if doc_video.exists():
                def _sub_doc_progress(frac):
                    pct = int(frac * 100)
                    progress_var.set(pct)
                    status_var.set(f"📝 Nghe audio làm phụ đề video dọc... {pct}%")
                status_var.set("📝 Đang tạo phụ đề cho video dọc...")
                progress_var.set(0)
                make_youtube_sub(doc_video, output_path.parent / "input.txt",
                                 sub_mode, sub_model, sub_max_chars,
                                 progress=_sub_doc_progress, doc=True,
                                 kieu=sub_kieu, font=sub_font, mau=sub_mau,
                                 vitri=sub_vitri, cochu=sub_cochu,
                                 dong=sub_dong, mau_vien=sub_mau_vien,
                                 bengang=sub_bengang)
                progress_var.set(100)
            else:
                logging.warning(f"📝 Chưa có video dọc ({doc_video.name}) → bỏ qua phụ đề dọc.")

    except Exception as e:
        failed = True
        logging.error(f"Lỗi: {e}")
        status_var.set(f"Lỗi: {e}")
    finally:
        # Lỗi giữa chừng (hết VRAM, ffmpeg chết, người dùng đóng tab…) thì model
        # vẫn nằm nguyên trong VRAM — mà GUI thì sống mãi, nên card kẹt tới lần
        # chạy sau. Chạy trót lọt thì model đã là None, đoạn này không tốn gì.
        if model is not None:
            model = _do_omnivoice(model, "dọn sau khi lỗi")
            _tra_vram()
        # Lỗi ngay trong lượt quét mất chữ thì large-v3-turbo cũng còn nằm đó.
        # giai_phong() gọi lại nhiều lần vô hại nên cứ gọi cho chắc.
        try:
            import taogiong_kiemtra_matchu as _mc_clean
            _mc_clean.giai_phong()
        except Exception:
            pass
        pause_event.set()
        btn_run.config(state="normal")
        btn_pause.config(state="disabled", text="⏸  Tạm dừng")
        _play_done_sound(success=not failed)   # âm báo khi chạy xong (hoặc lỗi)


# ── GIAO DIỆN ────────────────────────────────────────────────────────────────
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("OmniVoice TTS")
        self.resizable(True, True)
        self.configure(bg=UI["bg"])
        self._apply_theme()
        self._pause_event = threading.Event()
        self._pause_event.set()
        self._playing = False
        self._preview_after = None
        self._last_output = None
        self._pipe_busy = False
        self._pipe_settings = load_pipe_settings()   # auto-chain + model/tốc độ đã lưu
        # ⏻ Xong thì tắt máy / 🌙 Xong thì cho máy ngủ — mỗi ô DÙNG CHUNG cho 2 chỗ
        # tick (tab Giọng nói + bước ③), nên tạo ở đây, TRƯỚC khi dựng panel: tick ở
        # đâu cũng là một cài đặt duy nhất. Hai ô loại trừ nhau (xem _pick_power_mode).
        self.var_shutdown = tk.BooleanVar(value=self._pipe_settings.get("shutdown", False))
        self.var_sleep = tk.BooleanVar(value=self._pipe_settings.get("sleep", False))
        # Đếm ngược giờ ngủ bằng threading.Timer chứ không phải after(): cờ bận được
        # hạ ở LUỒNG CHẠY VIỆC (finally của các _pipe_set_busy(False)), mà gọi after()
        # từ luồng khác là chuyện Tk không hứa.
        self._sleep_timer = None
        self._opt_settings = load_opt_settings()      # mục "Cài đặt" của lần chạy trước
        self._favorites = load_favorites()
        self._effect_favorites = load_effect_favorites()
        self._log_boxes = []          # mọi ô nhật ký (panel video + tab kịch bản) — cùng nhận log
        # Nút Tạm dừng / Dừng của batch có ở NHIỀU tab (Tạo kịch bản + Nhận diện) →
        # gom vào list để bật/tắt/đổi chữ ĐỒNG BỘ mọi nơi (cùng điều khiển 1 batch).
        self._batch_pause_widgets = []
        self._batch_stop_widgets = []
        # ĐĂNG YOUTUBE chạy trong luồng RIÊNG, song song với việc dựng video tập kế:
        # tải lên tốn MẠNG, dựng video tốn GPU — bắt tập sau đợi tập trước tải xong
        # là phí cả tiếng đồng hồ mỗi mẻ. Hàng đợi để một luồng đăng lần lượt (không
        # tải 2 video cùng lúc cho khỏi chia nhỏ băng thông).
        self._upload_q = queue.Queue()
        self._upload_thread = None
        self._upload_lock = threading.Lock()   # giữ nhịp put ↔ luồng đăng tự kết thúc
        self._upload_done = 0                  # số tập đã đăng xong trong mẻ hiện tại
        self._setup_logging()
        self._build_ui()
        self._poll_log()
        self._start_web_log_mirror()   # nhật ký đăng bên bảng web → ô Nhật ký ở đây
        self.update_idletasks()
        self.minsize(1280, 680)
        self._center(1560, 720)   # đủ rộng/cao cho Home + tab Thumbnail (các nút không bị che)

    def _apply_theme(self):
        """Theme nền trắng, phẳng, hiện đại (dựa trên 'clam' để tùy biến màu)."""
        st = ttk.Style(self)
        try:
            st.theme_use("clam")
        except tk.TclError:
            pass
        C = UI
        base_font  = ("Segoe UI", 10)
        small_font = ("Segoe UI", 8)

        st.configure(".", background=C["bg"], foreground=C["fg"],
                     font=base_font, bordercolor=C["border"],
                     focuscolor=C["bg"], troughcolor=C["track"])
        st.configure("TFrame", background=C["bg"])
        st.configure("TLabel", background=C["bg"], foreground=C["fg"])
        st.configure("Header.TLabel", font=("Segoe UI", 19, "bold"),
                     foreground=C["fg"])
        st.configure("Brand.TLabel", font=("Segoe UI", 19, "bold"),
                     foreground=C["accent"])
        st.configure("Sub.TLabel", font=("Segoe UI", 9), foreground=C["muted"])
        st.configure("Hint.TLabel", font=small_font, foreground=C["muted"])

        # Khung "thẻ" có viền nhẹ
        st.configure("TLabelframe", background=C["card"], bordercolor=C["border"],
                     relief="solid", borderwidth=1, padding=10)
        st.configure("TLabelframe.Label", background=C["card"],
                     foreground=C["accent"], font=("Segoe UI", 10, "bold"))

        # Nhập liệu
        for w in ("TEntry", "TSpinbox"):
            st.configure(w, fieldbackground=C["field"], background=C["field"],
                         bordercolor=C["border"], lightcolor=C["border"],
                         darkcolor=C["border"], insertcolor=C["fg"], padding=5)
            st.map(w, bordercolor=[("focus", C["accent"])],
                   lightcolor=[("focus", C["accent"])])
        st.configure("TSpinbox", arrowcolor=C["muted"])

        st.configure("TCombobox", fieldbackground=C["field"], background=C["field"],
                     bordercolor=C["border"], lightcolor=C["border"],
                     darkcolor=C["border"], arrowcolor=C["muted"], padding=5)
        st.map("TCombobox",
               fieldbackground=[("readonly", C["field"])],
               foreground=[("readonly", C["fg"])],
               selectbackground=[("readonly", C["field"])],
               selectforeground=[("readonly", C["fg"])],
               bordercolor=[("focus", C["accent"])],
               lightcolor=[("focus", C["accent"])])
        # Danh sách xổ xuống của combobox
        self.option_add("*TCombobox*Listbox.background", C["field"])
        self.option_add("*TCombobox*Listbox.foreground", C["fg"])
        self.option_add("*TCombobox*Listbox.selectBackground", C["accent"])
        self.option_add("*TCombobox*Listbox.selectForeground", "#ffffff")
        self.option_add("*TCombobox*Listbox.font", base_font)

        # Radio
        st.configure("TRadiobutton", background=C["card"], foreground=C["fg"])
        st.map("TRadiobutton",
               background=[("active", C["card"])],
               foreground=[("active", C["accent"]), ("selected", C["accent"])],
               indicatorcolor=[("selected", C["accent"]), ("!selected", "#cfd3da")])

        # Checkbox
        st.configure("TCheckbutton", background=C["card"], foreground=C["fg"])
        st.map("TCheckbutton",
               background=[("active", C["card"])],
               foreground=[("active", C["accent"]), ("selected", C["accent"])],
               indicatorcolor=[("selected", C["accent"]), ("!selected", "#cfd3da")])

        # Nút phụ (xám nhạt) + nút chính (nhấn hồng)
        st.configure("TButton", background="#eef0f3", foreground=C["fg"],
                     bordercolor=C["border"], relief="flat",
                     focusthickness=0, padding=(14, 8), font=base_font)
        st.map("TButton",
               background=[("active", C["hover"]), ("pressed", C["press"]),
                           ("disabled", "#f4f5f7")],
               foreground=[("disabled", "#aeb4be")])
        st.configure("Accent.TButton", background=C["accent"], foreground="#ffffff",
                     padding=(20, 9), font=("Segoe UI", 10, "bold"))
        st.map("Accent.TButton",
               background=[("active", C["accent_dk"]), ("pressed", C["accent_dk"]),
                           ("disabled", C["accent_soft"])],
               foreground=[("disabled", "#ffffff")])

        # Thanh tiến trình
        st.configure("TProgressbar", background=C["accent"],
                     troughcolor=C["track"], bordercolor=C["track"],
                     lightcolor=C["accent"], darkcolor=C["accent"], thickness=12)

    def _center(self, w, h):
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        x, y = (sw - w) // 2, (sh - h) // 3
        self.geometry(f"{w}x{h}+{max(x,0)}+{max(y,0)}")

    def _bring_to_front(self):
        """Đưa cửa sổ GUI lên trên cùng + lấy focus (vd khi đang làm việc ở VS Code).

        Dùng để bật lên ở bước tạo giọng. Gọi được TỪ THREAD batch: thực thi qua
        self.after(0, ...) cho chạy trên main thread (Tk không an toàn đa luồng).
        Mẹo topmost True→False để cửa sổ bật lên trước mà KHÔNG kẹt 'luôn trên cùng'.
        """
        def _do():
            try:
                if self.state() == "iconic":
                    self.deiconify()             # phòng khi đang thu nhỏ
                self.lift()
                self.attributes("-topmost", True)
                self.update_idletasks()
                self.attributes("-topmost", False)
                self.focus_force()
            except Exception:
                pass
        try:
            self.after(0, _do)
        except Exception:
            pass

    def _setup_logging(self):
        handler = QueueHandler()
        handler.setFormatter(logging.Formatter("%(asctime)s  %(message)s", "%H:%M:%S"))
        logging.getLogger().addHandler(handler)
        logging.getLogger().setLevel(logging.INFO)

    def _build_ui(self):
        root = ttk.Frame(self, padding=18)
        root.grid(sticky="nsew")
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        # 2 vùng: [nút nhỏ chuyển view] · [nội dung]  (Nhật ký nằm trong panel video)
        root.columnconfigure(0, weight=0)   # sidebar nút nhỏ
        root.columnconfigure(1, weight=1)   # nội dung view
        root.rowconfigure(0, weight=1)

        # ── Sidebar trái: 3 nút nhỏ — Home (đầy đủ) · Tạo kịch bản · Giọng nói ──
        side = ttk.Frame(root)
        side.grid(row=0, column=0, sticky="n", padx=(0, 14))
        self._nav_buttons = {}
        for key, label in [("home",   "🏠  Home\n(đầy đủ)"),
                           ("script", "🛠  Tạo\nkịch bản"),
                           ("voice",  "🎧  Giọng\nnói"),
                           ("thumb",  "🖼  Thumb\nnail"),
                           ("copy",   "📑  Copy\nSEO"),
                           ("report", "📋  Tiến\nđộ"),
                           ("recog",  "🎧  Nhận\ndiện")]:
            b = ttk.Button(side, text=label, width=11,
                           command=lambda k=key: self._show_view(k))
            b.pack(fill="x", pady=(0, 8))
            self._nav_buttons[key] = b

        # Bảng điều khiển WEB — không phải một view của cửa sổ này mà là bản chạy
        # trong trình duyệt (xem myvoice/web/), nên tách khỏi nhóm nút trên.
        ttk.Separator(side, orient="horizontal").pack(fill="x", pady=(4, 10))
        self.btn_web = ttk.Button(side, text="🌐  Bảng\nweb", width=11,
                                  command=self._open_web_panel)
        self.btn_web.pack(fill="x")

        # ── Vùng nội dung: 3 panel (pipeline · TTS · video) HIỆN/ẨN theo view ──
        # Mỗi panel chỉ dựng MỘT lần (không trùng widget); nút chỉ bật/tắt hiển thị
        # bằng grid()/grid_remove(). Home = hiện cả 3 → đúng giao diện gốc đầy đủ.
        content = ttk.Frame(root)
        content.grid(row=0, column=1, sticky="nsew", padx=(0, 16))
        self._content = content
        content.rowconfigure(0, weight=1)
        content.columnconfigure(0, weight=0)   # pipeline
        content.columnconfigure(1, weight=0)   # TTS
        content.columnconfigure(2, weight=1)   # video/hành động (giãn)

        frame_pipeline = ttk.Frame(content)
        frame_pipeline.grid(row=0, column=0, sticky="nsew", padx=(0, 16))
        frame_pipeline.columnconfigure(0, weight=1)
        frame_pipeline.rowconfigure(0, weight=1)

        frame_tts = ttk.Frame(content)
        frame_tts.grid(row=0, column=1, sticky="nsew", padx=(0, 16))
        frame_tts.columnconfigure(0, weight=1)
        frame_tts.rowconfigure(0, weight=1)

        frame_video = ttk.Frame(content)
        frame_video.grid(row=0, column=2, sticky="nsew")
        frame_video.columnconfigure(0, weight=1)
        frame_video.rowconfigure(0, weight=1)

        self._panels = {"pipeline": frame_pipeline, "tts": frame_tts,
                        "video": frame_video}

        # Panel Thumbnail — nhúng YOUTUBE/thumbnail_gui.py; phủ cả vùng nội dung khi chọn.
        frame_thumb = ttk.Frame(content)
        frame_thumb.grid(row=0, column=0, columnspan=3, sticky="nsew")
        frame_thumb.rowconfigure(0, weight=1)
        frame_thumb.columnconfigure(0, weight=1)
        self._panels["thumbnail"] = frame_thumb
        self._build_thumbnail_panel(frame_thumb)

        # Panel Tiến độ — bảng các link đã gửi (manifest) + đã làm tới đâu.
        frame_report = ttk.Frame(content)
        frame_report.grid(row=0, column=0, columnspan=3, sticky="nsew")
        frame_report.rowconfigure(0, weight=1)
        frame_report.columnconfigure(0, weight=1)
        self._panels["report"] = frame_report
        self._build_report_panel(frame_report)

        # Panel Copy SEO — danh sách tập (thư mục số trong kịch_bản) + copy tiêu đề/mô tả/thẻ tag.
        frame_copyseo = ttk.Frame(content)
        frame_copyseo.grid(row=0, column=0, columnspan=3, sticky="nsew")
        frame_copyseo.rowconfigure(0, weight=1)
        frame_copyseo.columnconfigure(0, weight=1)
        self._panels["copyseo"] = frame_copyseo
        self._build_copyseo_panel(frame_copyseo)

        # Panel Nhận diện — luồng RIÊNG: chỉ nhận diện các link rồi ngưng (không dịch/SEO/video).
        # Build SAU pipeline (bên dưới) vì dùng chung pipe_var_model/pipe_var_speed.
        frame_recog = ttk.Frame(content)
        frame_recog.grid(row=0, column=0, columnspan=3, sticky="nsew")
        frame_recog.rowconfigure(0, weight=1)
        frame_recog.columnconfigure(0, weight=1)
        self._panels["recog"] = frame_recog
        self._recog_frame = frame_recog

        # ════════════════════════════════════════════════
        # PANEL pipeline — Quy trình tạo kịch bản (nhận diện → Gemini → SEO → input.txt)
        # ════════════════════════════════════════════════
        self._build_pipeline_column(frame_pipeline, 0)
        # Tab '🎧 Nhận diện' build SAU pipeline vì dùng chung pipe_var_model/speed.
        self._build_recog_panel(self._recog_frame)

        # ════════════════════════════════════════════════
        # PANEL TTS — toàn bộ điều khiển TTS
        # ════════════════════════════════════════════════
        left = ttk.Frame(frame_tts)
        left.grid(row=0, column=0, sticky="nsew")
        left.columnconfigure(0, weight=1)

        # (Header '🎧 OmniVoice TTS' + dòng mô tả đã bỏ cho tiết kiệm chỗ dọc.
        #  Row 0 để trống — grid không cấp chiều cao cho hàng rỗng.)

        # ── Chế độ ── (bỏ chữ tiêu đề group, chỉ giữ viền)
        sec_mode = ttk.LabelFrame(left, text="")
        sec_mode.grid(row=1, column=0, sticky="ew", pady=(0, 12))
        sec_mode.columnconfigure(0, weight=1)

        self.var_mode = tk.StringVar(value="clone")
        mode_row = ttk.Frame(sec_mode)
        mode_row.grid(row=0, column=0, sticky="w")
        for label, val, desc in [
            ("🎙  Clone",   "clone",   "Nhái giọng mẫu"),
            ("🎨  Thiết kế", "design",  "Mô tả giọng"),
            ("🔊  Mặc định", "default", "Model tự chọn"),
        ]:
            f = ttk.Frame(mode_row)
            f.pack(side="left", padx=(0, 18))
            ttk.Radiobutton(f, text=label, variable=self.var_mode,
                            value=val, command=self._on_mode_change).pack(anchor="w")
            ttk.Label(f, text=desc, style="Hint.TLabel").pack(anchor="w", padx=22)

        ttk.Separator(sec_mode, orient="horizontal").grid(
            row=1, column=0, sticky="ew", pady=10)

        self.voice_frame = ttk.Frame(sec_mode)
        self.voice_frame.grid(row=2, column=0, sticky="ew")

        # Clone
        self.frm_clone = ttk.Frame(self.voice_frame)
        ttk.Label(self.frm_clone, text="Giọng mẫu:", width=11, anchor="w").pack(side="left", padx=(0, 6))
        self.var_ref = tk.StringVar()
        self.cb_ref = ttk.Combobox(self.frm_clone, textvariable=self.var_ref,
                                   values=[], width=30, state="readonly")
        self.cb_ref.pack(side="left")
        self.cb_ref.bind("<<ComboboxSelected>>", lambda e: self._update_fav_button())
        self.btn_fav = ttk.Button(self.frm_clone, text="☆", width=3,
                                  command=self._toggle_favorite)
        self.btn_fav.pack(side="left", padx=(6, 0))
        ttk.Button(self.frm_clone, text="↻", width=3,
                   command=self._refresh_voices).pack(side="left", padx=(6, 0))
        self._reload_voice_combo()   # nạp danh sách (yêu thích ★ lên đầu) + chọn mục đầu

        # Design
        self.frm_design = ttk.Frame(self.voice_frame)
        lang_row = ttk.Frame(self.frm_design)
        lang_row.pack(anchor="w", pady=(0, 4))
        ttk.Label(lang_row, text="Ngôn ngữ:").pack(side="left", padx=(0, 8))
        self.var_lang = tk.StringVar(value="en")
        ttk.Radiobutton(lang_row, text="English", variable=self.var_lang,
                        value="en", command=self._on_lang_change).pack(side="left", padx=4)
        ttk.Radiobutton(lang_row, text="中文", variable=self.var_lang,
                        value="zh", command=self._on_lang_change).pack(side="left", padx=4)
        self.design_attr_frame = ttk.Frame(self.frm_design)
        self.design_attr_frame.pack(anchor="w", fill="x")
        res_row = ttk.Frame(self.frm_design)
        res_row.pack(anchor="w", fill="x", pady=(6, 0))
        ttk.Label(res_row, text="Lệnh model:").pack(side="left", padx=(0, 8))
        self.var_instruct = tk.StringVar(value="female, young adult")
        ttk.Entry(res_row, textvariable=self.var_instruct, width=36).pack(side="left", fill="x", expand=True)

        self._design_vars: list[tk.StringVar] = []
        self._design_sep = ", "
        self._build_design_dropdowns()

        # Default
        self.frm_default = ttk.Frame(self.voice_frame)
        ttk.Label(self.frm_default, text="Model tự động chọn giọng phù hợp với văn bản",
                  style="Sub.TLabel").pack(side="left", padx=2)

        self._on_mode_change()

        # ── Tệp ── (bỏ chữ tiêu đề group, chỉ giữ viền)
        sec_file = ttk.LabelFrame(left, text="")
        sec_file.grid(row=2, column=0, sticky="ew", pady=(0, 12))
        sec_file.columnconfigure(1, weight=1)

        for r, (lbl, attr, default, is_save) in enumerate([
            ("Văn bản (.txt):", "var_txt", str(SCRIPT_DIR / "input.txt"),  False),
            ("Kết quả (.wav):", "var_out", str(OUTPUT_DIR / "output.wav"), True),
        ]):
            ttk.Label(sec_file, text=lbl, width=14, anchor="w").grid(
                row=r, column=0, sticky="w", padx=(0, 8), pady=4)
            var = tk.StringVar(value=default)
            setattr(self, attr, var)
            ttk.Entry(sec_file, textvariable=var).grid(
                row=r, column=1, sticky="ew", pady=4)
            cmd = (lambda v=var: self._pick_save(v, [("WAV", "*.wav")])) if is_save \
                else (lambda v=var: self._pick_file(v, [("Text", "*.txt")]))
            ttk.Button(sec_file, text="Chọn…", width=8, command=cmd).grid(
                row=r, column=2, padx=(8, 0), pady=4)

        # ── Cài đặt ── (bỏ chữ tiêu đề group, chỉ giữ viền)
        sec_opt = ttk.LabelFrame(left, text="")
        sec_opt.grid(row=3, column=0, sticky="ew", pady=(0, 12))

        # Nguồn nội dung: lấy từ Gemini (gemini_result.docx) + kiểm tra trước khi tạo
        gem_row = ttk.Frame(sec_opt)
        gem_row.pack(anchor="w", fill="x", pady=(0, 8))
        self.var_from_gemini = tk.BooleanVar(value=self._opt_settings["from_gemini"])
        ttk.Checkbutton(gem_row,
                        text="🌐  Lấy nội dung từ Gemini + kiểm tra trước khi tạo",
                        variable=self.var_from_gemini).pack(side="left")
        ttk.Label(gem_row, text="(gemini_result.docx → input.txt)",
                  style="Hint.TLabel").pack(side="left", padx=8)

        chunk_row = ttk.Frame(sec_opt)
        chunk_row.pack(anchor="w", fill="x")
        ttk.Label(chunk_row, text="Độ dài đoạn (ký tự):").pack(side="left", padx=(0, 8))
        self.var_chunk = tk.IntVar(value=self._opt_settings["chunk"])
        ttk.Spinbox(chunk_row, from_=100, to=1000, increment=50,
                    textvariable=self.var_chunk, width=7).pack(side="left")
        ttk.Label(chunk_row, text="(nhỏ = nhẹ GPU)",
                  style="Hint.TLabel").pack(side="left", padx=8)

        # Nguồn clip VIDEO NGANG theo chủ đề = 1 thư mục con của videongang/
        # (vd thiennhien, nauan). NGANG_SOURCE_ALL = ghép random từ CẢ kho videongang/.
        # Dồn về PHẢI của hàng cho gọn.
        self.var_ngang_source = tk.StringVar()
        _ngang_sources = [NGANG_SOURCE_ALL] + list_ngang_sources()
        _saved_src = self._opt_settings.get("ngang_source", NGANG_SOURCE_ALL)
        self.var_ngang_source.set(_saved_src if _saved_src in _ngang_sources else NGANG_SOURCE_ALL)
        self.cb_ngang_source = ttk.Combobox(chunk_row, textvariable=self.var_ngang_source,
                                            values=_ngang_sources, width=14, state="readonly")
        self.cb_ngang_source.pack(side="right")
        ttk.Label(chunk_row, text="Nguồn ngang:").pack(side="right", padx=(0, 6))

        video_row = ttk.Frame(sec_opt)
        video_row.pack(anchor="w", fill="x", pady=(8, 0))
        self.var_make_video = tk.BooleanVar(value=self._opt_settings["make_video"])
        ttk.Checkbutton(video_row, text="🎬  Tự dựng video (ngang)",
                        variable=self.var_make_video).pack(side="left")
        ttk.Label(video_row, text="Tăng tốc:").pack(side="left", padx=(12, 2))
        self.var_ngang_speed = tk.StringVar(value=self._opt_settings["ngang_speed"])
        ttk.Combobox(video_row, textvariable=self.var_ngang_speed, width=6,
                     values=["1.0", "1.05", "1.1", "1.15", "1.2", "1.25"]).pack(side="left")
        ttk.Label(video_row, text="x (giữ cao độ)",
                  style="Hint.TLabel").pack(side="left", padx=(4, 0))
        # Nút DỰNG LẠI (chỉ dựng video ngang từ audio có sẵn) — dồn về PHẢI của hàng
        # để nằm ở góc trên-phải mà không phá bố cục hàng.
        self.btn_run_ngang = ttk.Button(video_row, text="▶ Dựng lại", width=11,
                                        command=lambda: self._rebuild_video("ngang"))
        self.btn_run_ngang.pack(side="right")

        # ── Phụ đề cho video ngang (YouTube) — BẬT/TẮT được ──────────────────
        # Dùng lại video_gansub.py: Whisper nghe audio CỦA CHÍNH VIDEO để lấy mốc
        # giờ, còn chữ lấy từ input.txt gốc nên không sai chính tả.
        sub_row = ttk.Frame(sec_opt)
        sub_row.pack(anchor="w", fill="x", pady=(8, 0))
        self.var_make_sub = tk.BooleanVar(value=self._opt_settings["make_sub"])
        ttk.Checkbutton(sub_row, text="📝  Phụ đề cho video YouTube",
                        variable=self.var_make_sub).pack(side="left")
        ttk.Label(sub_row, text="Kiểu:").pack(side="left", padx=(12, 2))
        self.var_sub_mode = tk.StringVar(value=self._opt_settings["sub_mode"])
        ttk.Combobox(sub_row, textvariable=self.var_sub_mode, width=15, state="readonly",
                     values=[SUB_MODE_SRT, SUB_MODE_BURN]).pack(side="left")
        ttk.Label(sub_row, text="Model:").pack(side="left", padx=(10, 2))
        self.var_sub_model = tk.StringVar(value=self._opt_settings["sub_model"])
        ttk.Combobox(sub_row, textvariable=self.var_sub_model, width=14, state="readonly",
                     values=["small", "medium", "large-v3-turbo", "large-v3"]).pack(side="left")

        sub_row2 = ttk.Frame(sec_opt)
        sub_row2.pack(anchor="w", fill="x", pady=(4, 0))
        ttk.Label(sub_row2, text="Dài mỗi dòng:").pack(side="left", padx=(26, 2))
        self.var_sub_max_chars = tk.IntVar(value=self._opt_settings["sub_max_chars"])
        ttk.Spinbox(sub_row2, from_=20, to=90, increment=1, width=5,
                    textvariable=self.var_sub_max_chars).pack(side="left")
        # Kiểu TRÌNH BÀY khi vẽ cứng — kho scripts/kieusub_mau (dùng chung myvideo);
        # web có ảnh xem trước từng kiểu, GUI chỉ chọn theo tên.
        ttk.Label(sub_row2, text="Kiểu chữ:").pack(side="left", padx=(10, 2))
        self.var_sub_kieu = tk.StringVar(
            value=self._opt_settings.get("sub_kieu", "hopbo"))
        try:
            import kieusub
            _kieu_ids = [k["id"] for k in kieusub.danh_sach()] or ["hopbo"]
            _font_ten = [""] + [f["ten"] for f in kieusub.danh_sach_font()]
        except Exception:
            _kieu_ids, _font_ten = ["hopbo"], [""]
        ttk.Combobox(sub_row2, textvariable=self.var_sub_kieu, width=12,
                     state="readonly", values=_kieu_ids).pack(side="left")
        # Font đè lên font của kiểu — rỗng = dùng font gốc của kiểu.
        ttk.Label(sub_row2, text="Font:").pack(side="left", padx=(10, 2))
        self.var_sub_font = tk.StringVar(
            value=self._opt_settings.get("sub_font", ""))
        ttk.Combobox(sub_row2, textvariable=self.var_sub_font, width=16,
                     state="readonly", values=_font_ten).pack(side="left")
        # Màu CHỮ + màu VIỀN đè lên màu của kiểu — nút mở bảng chọn màu Windows,
        # × = trả về màu gốc của kiểu.
        self.var_sub_mau = tk.StringVar(
            value=self._opt_settings.get("sub_mau", ""))
        self.btn_sub_mau = self._o_mau_sub(
            sub_row2, "Màu:", self.var_sub_mau,
            "Màu chữ phụ đề (đè màu của kiểu)")
        self.var_sub_mau_vien = tk.StringVar(
            value=self._opt_settings.get("sub_mau_vien", ""))
        self.btn_sub_mau_vien = self._o_mau_sub(
            sub_row2, "Viền:", self.var_sub_mau_vien,
            "Màu viền quanh chữ phụ đề (đè viền của kiểu)")

        # Hàng riêng cho vị trí / cỡ chữ / số dòng — hàng trên đã kín chỗ.
        sub_row2b = ttk.Frame(sec_opt)
        sub_row2b.pack(anchor="w", fill="x", pady=(4, 0))
        # Vị trí chữ theo % chiều cao từ đáy — để TRỐNG là mặc định (173/380px).
        ttk.Label(sub_row2b, text="Vị trí % đáy:").pack(side="left", padx=(26, 2))
        self.var_sub_vitri = tk.StringVar(
            value=str(self._opt_settings.get("sub_vitri", "")))
        # 0–100: hết chiều cao khung, cùng khoảng với thanh kéo bên web (chung
        # taogiong_options.json) — để hẹp hơn thì mũi tên sẽ kéo tụt giá trị web.
        ttk.Spinbox(sub_row2b, from_=0, to=100, increment=1, width=4,
                    textvariable=self.var_sub_vitri).pack(side="left")
        # Cỡ chữ theo % cỡ gốc của kiểu — để TRỐNG (hoặc 100) là giữ nguyên.
        ttk.Label(sub_row2b, text="Cỡ chữ %:").pack(side="left", padx=(10, 2))
        self.var_sub_cochu = tk.StringVar(
            value=str(self._opt_settings.get("sub_cochu", "")))
        ttk.Spinbox(sub_row2b, from_=50, to=200, increment=5, width=4,
                    textvariable=self.var_sub_cochu).pack(side="left")
        # BỀ NGANG dòng chữ theo % bề ngang chuẩn — TRỐNG (hoặc 100) là như cũ;
        # hẹp lại thì xuống dòng sớm (chữ gom về giữa), quá 100% dòng dài ra —
        # cùng khoảng 20–150 với thanh kéo bên web (chung taogiong_options.json).
        ttk.Label(sub_row2b, text="Ngang %:").pack(side="left", padx=(10, 2))
        self.var_sub_bengang = tk.StringVar(
            value=str(self._opt_settings.get("sub_bengang", "")))
        ttk.Spinbox(sub_row2b, from_=20, to=150, increment=5, width=4,
                    textvariable=self.var_sub_bengang).pack(side="left")
        # Số dòng mỗi lần hiện chữ (2 = như ảnh mẫu của kho kiểu).
        ttk.Label(sub_row2b, text="Số dòng:").pack(side="left", padx=(10, 2))
        self.var_sub_dong = tk.IntVar(
            value=int(self._opt_settings.get("sub_dong", 2) or 2))
        ttk.Combobox(sub_row2b, textvariable=self.var_sub_dong, width=3,
                     state="readonly", values=[1, 2]).pack(side="left")

        # Phụ đề cho VIDEO DỌC (facebook.mp4) — tick riêng, dùng chung Kiểu/Model ở trên.
        # Khung 1080x1920 nên tự rút số ký tự/dòng xuống 27 cho khỏi tràn mép.
        sub_row3 = ttk.Frame(sec_opt)
        sub_row3.pack(anchor="w", fill="x", pady=(4, 0))
        self.var_make_sub_doc = tk.BooleanVar(value=self._opt_settings["make_sub_doc"])
        ttk.Checkbutton(sub_row3, text="📝  Phụ đề cho video DỌC (Facebook)",
                        variable=self.var_make_sub_doc).pack(side="left")

        # Hiệu ứng phủ lên toàn bộ video (từ đầu đến cuối) — lấy từ scripts/hieuung/
        fx_row = ttk.Frame(sec_opt)
        fx_row.pack(anchor="w", fill="x", pady=(8, 0))
        ttk.Label(fx_row, text="✨  Hiệu ứng:").pack(side="left", padx=(0, 8))
        self.var_effect = tk.StringVar(value=EFFECT_NONE)
        self.cb_effect = ttk.Combobox(fx_row, textvariable=self.var_effect,
                                      values=[EFFECT_NONE], width=24, state="readonly")
        self.cb_effect.pack(side="left")
        self.cb_effect.bind("<<ComboboxSelected>>",
                            lambda e: self._update_effect_fav_button())
        self.btn_effect_fav = ttk.Button(fx_row, text="☆", width=3,
                                         command=self._toggle_effect_favorite)
        self.btn_effect_fav.pack(side="left", padx=(6, 0))
        ttk.Button(fx_row, text="↻", width=3,
                   command=self._refresh_effects).pack(side="left", padx=(6, 0))
        ttk.Label(fx_row, text="(phủ lên toàn video)",
                  style="Hint.TLabel").pack(side="left", padx=8)
        # Nạp danh sách (yêu thích ★ lên đầu) + chọn lại hiệu ứng của lần chạy trước
        self._reload_effect_combo(keep=self._opt_settings.get("effect", EFFECT_NONE))

        # (Khối "Video dọc" đã chuyển sang CỘT 3 bên phải cho đỡ chật.)

        # ════════════════════════════════════════════════
        # PANEL video — Video dọc + Hành động + tiến trình
        # ════════════════════════════════════════════════
        right = ttk.Frame(frame_video)
        right.grid(row=0, column=0, sticky="nsew")
        right.columnconfigure(0, weight=1)
        right.rowconfigure(2, weight=1)   # hàng NHẬT KÝ (dưới Video dọc + TikTok) giãn — nhật ký lên cao

        # ── Video dọc ──
        vdoc = ttk.LabelFrame(right, text="  📱  Video dọc (1080×1920)  ")
        vdoc.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        self.var_make_video_doc = tk.BooleanVar(value=self._opt_settings["make_video_doc"])
        ttk.Checkbutton(vdoc, text="Dựng video dọc",
                        variable=self.var_make_video_doc).pack(anchor="w")
        vdoc_opts = ttk.Frame(vdoc)
        vdoc_opts.pack(anchor="w", fill="x", pady=(6, 0))
        ttk.Label(vdoc_opts, text="Tăng tốc:").pack(side="left", padx=(0, 2))
        self.var_doc_speed = tk.StringVar(value=self._opt_settings["doc_speed"])
        # Combobox để sửa tay được (vd 1.07), kèm các mức gợi ý sẵn.
        ttk.Combobox(vdoc_opts, textvariable=self.var_doc_speed, width=6,
                     values=["1.0", "1.05", "1.1", "1.15", "1.2", "1.25"]).pack(side="left")
        ttk.Label(vdoc_opts, text="x (giữ cao độ)",
                  style="Hint.TLabel").pack(side="left", padx=(4, 0))

        # Lấy ~N% thời lượng audio (cắt ở CUỐI CÂU) — giống TikTok. 100 = cả bài
        # (giữ nguyên hành vi cũ). Kết hợp được với "dùng audio không cắt" ở trên.
        vdoc_optsp = ttk.Frame(vdoc)
        vdoc_optsp.pack(anchor="w", fill="x", pady=(6, 0))
        ttk.Label(vdoc_optsp, text="Lấy khoảng:").pack(side="left", padx=(0, 2))
        self.var_doc_percent = tk.StringVar(value=str(self._opt_settings["doc_percent"]))
        ttk.Combobox(vdoc_optsp, textvariable=self.var_doc_percent, width=5,
                     values=["25", "30", "40", "50", "60", "70", "75", "80", "90", "100"]
                     ).pack(side="left")
        ttk.Label(vdoc_optsp, text="% thời lượng (cắt cuối câu · 100 = cả bài)",
                  style="Hint.TLabel").pack(side="left", padx=(4, 0))

        # Dùng lại video ngang đã dựng (phóng to khớp chiều cao dọc + cắt giữa),
        # thay âm bằng audio video dọc. Tắt → dựng từ kho videodoc/ như cũ.
        vdoc_opts2 = ttk.Frame(vdoc)
        vdoc_opts2.pack(anchor="w", fill="x", pady=(6, 0))
        self.var_doc_from_ngang = tk.BooleanVar(value=self._opt_settings["doc_from_ngang"])
        ttk.Checkbutton(vdoc_opts2,
                        text="♻  Dùng lại video ngang (phóng to khớp chiều cao)",
                        variable=self.var_doc_from_ngang,
                        command=lambda: self._exclusive_doc_source("ngang")).pack(side="left")

        # 📁 Ghép video dọc từ 1 THƯ MỤC CON của videodoc/ (tên ghi trong
        # videodoc/input.txt). LOẠI TRỪ với "dùng lại video ngang": bật cái này
        # tự tắt cái kia (xem _exclusive_doc_source).
        vdoc_opts2b = ttk.Frame(vdoc)
        vdoc_opts2b.pack(anchor="w", fill="x", pady=(6, 0))
        self.var_doc_from_subfolder = tk.BooleanVar(value=self._opt_settings["doc_from_subfolder"])
        # Settings cũ lỡ bật cả hai → ưu tiên thư mục con, tắt dùng-lại-ngang.
        if self.var_doc_from_subfolder.get() and self.var_doc_from_ngang.get():
            self.var_doc_from_ngang.set(False)
        ttk.Checkbutton(vdoc_opts2b,
                        text="📁  Ghép từ thư mục con videodoc (theo videodoc/input.txt)",
                        variable=self.var_doc_from_subfolder,
                        command=lambda: self._exclusive_doc_source("subfolder")).pack(side="left")

        # Không phủ hiệu ứng lên video dọc (mọi trường hợp) — video dọc sạch hiệu ứng.
        vdoc_opts3 = ttk.Frame(vdoc)
        vdoc_opts3.pack(anchor="w", fill="x", pady=(6, 0))
        self.var_doc_no_effect = tk.BooleanVar(value=self._opt_settings["doc_no_effect"])
        ttk.Checkbutton(vdoc_opts3, text="🚫  Không áp hiệu ứng cho video dọc",
                        variable=self.var_doc_no_effect).pack(side="left")

        # Nút DỰNG LẠI ở GÓC TRÊN-PHẢI: dùng place() nên KHÔNG chiếm chỗ trong luồng
        # pack → không đẩy/vỡ layout của mục. (Dựng lại video dọc từ audio có sẵn.)
        self.btn_run_doc = ttk.Button(vdoc, text="▶ Dựng lại",
                                      command=lambda: self._rebuild_video("doc"))
        self.btn_run_doc.place(relx=1.0, x=-6, y=2, anchor="ne")

        # ── Video TikTok (group riêng, DƯỚI 'Video dọc', TRÊN nút Chạy) ──
        # Lấy ~N% thời lượng audio (cắt ở CUỐI CÂU) + dùng lại video Facebook (khi
        # bật 📁) hoặc ghép từ videodoc/ → 1 video dọc riêng để đăng TikTok.
        tiktok = ttk.LabelFrame(right, text="  🎵  Video TikTok (cắt theo %)  ")
        tiktok.grid(row=1, column=0, sticky="ew", pady=(0, 10))
        self.var_make_tiktok = tk.BooleanVar(value=self._opt_settings["make_tiktok"])
        ttk.Checkbutton(tiktok, text="Tạo video TikTok (dùng lại Facebook khi bật 📁, "
                        "không thì ghép từ videodoc)",
                        variable=self.var_make_tiktok).pack(anchor="w")
        # Short cắt TỪ video TikTok nên chỉ có tác dụng khi ô trên được bật.
        self.var_make_short = tk.BooleanVar(value=self._opt_settings["make_short"])
        ttk.Checkbutton(tiktok, text="📱  Tạo + đăng YouTube Short (cắt ≤2:50 từ video TikTok, "
                        "hẹn sau bản chính 1 giờ)",
                        variable=self.var_make_short).pack(anchor="w")
        # Lấy khoảng bao nhiêu % thời lượng (cắt ở cuối câu → độ dài chỉ xấp xỉ %).
        tiktok_opts0 = ttk.Frame(tiktok)
        tiktok_opts0.pack(anchor="w", fill="x", pady=(6, 0))
        ttk.Label(tiktok_opts0, text="Lấy khoảng:").pack(side="left", padx=(0, 2))
        self.var_tiktok_percent = tk.StringVar(value=str(self._opt_settings["tiktok_percent"]))
        ttk.Combobox(tiktok_opts0, textvariable=self.var_tiktok_percent, width=5,
                     values=["25", "30", "40", "50", "60", "70", "75", "80", "90", "100"]
                     ).pack(side="left")
        ttk.Label(tiktok_opts0, text="% thời lượng (cắt cuối câu · 100 = cả bài)",
                  style="Hint.TLabel").pack(side="left", padx=(4, 0))
        tiktok_opts = ttk.Frame(tiktok)
        tiktok_opts.pack(anchor="w", fill="x", pady=(6, 0))
        ttk.Label(tiktok_opts, text="Tăng tốc:").pack(side="left", padx=(0, 2))
        self.var_tiktok_speed = tk.StringVar(value=self._opt_settings["tiktok_speed"])
        ttk.Combobox(tiktok_opts, textvariable=self.var_tiktok_speed, width=6,
                     values=["1.0", "1.05", "1.1", "1.15", "1.2", "1.25"]).pack(side="left")
        ttk.Label(tiktok_opts, text="x (giữ cao độ)",
                  style="Hint.TLabel").pack(side="left", padx=(4, 0))
        # Số tập cho chữ TikTok — MẶC ĐỊNH = số ở tab Thumbnail + 1 (TikTok của tập
        # đang dựng đi TRƯỚC thumbnail 1 số). Hiện ở đây để xem/sửa trước khi dựng;
        # chữ trên video sẽ là 'Mimi audio Số <số này>'. Tự cập nhật khi đổi số thumbnail.
        ttk.Label(tiktok_opts, text="Số tập:").pack(side="left", padx=(16, 2))
        self.var_tiktok_episode = tk.StringVar(value=self._default_tiktok_episode())
        ttk.Entry(tiktok_opts, textvariable=self.var_tiktok_episode, width=5,
                  justify="center").pack(side="left")
        ttk.Label(tiktok_opts, text="(= thumbnail + 1)",
                  style="Hint.TLabel").pack(side="left", padx=(4, 0))
        # Không phủ hiệu ứng lên video TikTok (video sạch hiệu ứng).
        tiktok_opts2 = ttk.Frame(tiktok)
        tiktok_opts2.pack(anchor="w", fill="x", pady=(6, 0))
        self.var_tiktok_no_effect = tk.BooleanVar(value=self._opt_settings["tiktok_no_effect"])
        ttk.Checkbutton(tiktok_opts2, text="🚫  Không áp hiệu ứng cho TikTok",
                        variable=self.var_tiktok_no_effect).pack(side="left")
        # Danh sách số tập ĐẶT TRƯỚC để bỏ qua khi cấp số tập mới (xem next_episode_number).
        ttk.Button(tiktok_opts2, text="⏭  Tập bỏ qua…",
                   command=self._edit_skip_episodes).pack(side="left", padx=(18, 0))
        self.lbl_skip_eps = ttk.Label(tiktok_opts2, text="", style="Hint.TLabel")
        self.lbl_skip_eps.pack(side="left", padx=(8, 0))
        self._refresh_skip_episodes_label()
        # Vị trí chữ theo chiều cao (0 = trên cùng, 100 = dưới cùng).
        tiktok_opts3 = ttk.Frame(tiktok)
        tiktok_opts3.pack(anchor="w", fill="x", pady=(6, 0))
        ttk.Label(tiktok_opts3, text="Vị trí chữ (% cao):").pack(side="left", padx=(0, 6))
        self.var_tiktok_caption_pos = tk.IntVar(value=self._opt_settings["tiktok_caption_pos"])
        ttk.Spinbox(tiktok_opts3, from_=5, to=95, increment=5,
                    textvariable=self.var_tiktok_caption_pos, width=5).pack(side="left")
        ttk.Label(tiktok_opts3, text="(0 = trên, 100 = dưới)",
                  style="Hint.TLabel").pack(side="left", padx=(6, 0))
        # Nhạc nền (từ Music/) + mức nhỏ hơn giọng (dB). Giọng ≈ -6dB → -12 = nhạc ≈ -18dB.
        tiktok_opts4 = ttk.Frame(tiktok)
        tiktok_opts4.pack(anchor="w", fill="x", pady=(6, 0))
        self.var_tiktok_music = tk.BooleanVar(value=self._opt_settings["tiktok_music"])
        ttk.Checkbutton(tiktok_opts4, text="🎼  Chèn nhạc nền (Music/)",
                        variable=self.var_tiktok_music).pack(side="left")
        ttk.Label(tiktok_opts4, text="nhỏ hơn giọng:").pack(side="left", padx=(10, 2))
        self.var_tiktok_music_db = tk.IntVar(value=self._opt_settings["tiktok_music_db"])
        ttk.Spinbox(tiktok_opts4, from_=-30, to=0, increment=1,
                    textvariable=self.var_tiktok_music_db, width=5).pack(side="left")
        ttk.Label(tiktok_opts4, text="dB", style="Hint.TLabel").pack(side="left", padx=(4, 0))

        # Nút DỰNG LẠI ở GÓC TRÊN-PHẢI (place() → không phá layout). Dựng lại video
        # TikTok từ audio có sẵn.
        self.btn_run_tiktok = ttk.Button(tiktok, text="▶ Dựng lại",
                                         command=lambda: self._rebuild_video("tiktok"))
        self.btn_run_tiktok.place(relx=1.0, x=-6, y=2, anchor="ne")

        # ── Nhóm hành động (Chạy / Tạm dừng / Nghe thử / Xóa output) — đặt trong CỘT
        #    TTS, NGAY DƯỚI group box "Cài đặt" (NGOÀI group box). left rows: 4 và 5. ──
        act = ttk.Frame(left)
        act.grid(row=4, column=0, sticky="ew", pady=(8, 8))
        self.btn_run = ttk.Button(act, text="▶  Chạy", command=self._start,
                                  style="Accent.TButton")
        self.btn_run.pack(side="left", padx=(0, 8))
        self.btn_pause = ttk.Button(act, text="⏸  Tạm dừng", command=self._toggle_pause,
                                    state="disabled")
        self.btn_pause.pack(side="left", padx=(0, 8))
        self.btn_preview = ttk.Button(act, text="🔊  Nghe thử", command=self._toggle_preview,
                                      state="disabled")
        self.btn_preview.pack(side="left", padx=(0, 8))
        ttk.Button(act, text="🗑  Xóa output", command=self._clear_output).pack(side="left")

        # ── 3 ô tick GỘP 1 HÀNG (chữ rút gọn cho vừa bề ngang cột giữa) ───────────
        #  ♻ dùng lại : bỏ qua tạo giọng nếu output.wav còn đúng văn bản/giọng và bỏ
        #               qua video đã dựng — chỉ dựng phần còn thiếu. Mặc định TẮT
        #               (mỗi lần mở app) để tránh vô tình dùng lại bản cũ.
        #  ⬆️ hiện cửa sổ: tới bước clone giọng thì GUI tự bật lên trên (mỗi link 1 lần).
        #  ⏻ tắt máy : hẹn tắt sau SHUTDOWN_DELAY_MIN phút khi xong (huỷ: shutdown /a);
        #  🌙 ngủ    : cho máy ngủ sau SLEEP_DELAY_MIN phút khi xong (huỷ: bỏ tick).
        #               Cả hai dùng CHUNG biến với ô tick ở bước ③.
        # Cột giữa có weight=0 (co theo nội dung) nên hàng này phải ngắn, kẻo cột
        # giữa phình ra lấn sang cột video bên phải.
        act2 = ttk.Frame(left)
        act2.grid(row=5, column=0, sticky="ew", pady=(0, 10))
        self.var_reuse = tk.BooleanVar(value=False)
        ttk.Checkbutton(act2, text="♻  Dùng lại audio/video",
                        variable=self.var_reuse).pack(side="left")
        self.var_bring_front = tk.BooleanVar(value=self._opt_settings["bring_front"])
        ttk.Checkbutton(act2, text="⬆️  Hiện cửa sổ khi tạo giọng",
                        variable=self.var_bring_front).pack(side="left", padx=(14, 0))
        ttk.Checkbutton(act2, text="⏻  Xong thì tắt máy", variable=self.var_shutdown,
                        command=lambda: self._pick_power_mode("shutdown")
                        ).pack(side="left", padx=(14, 0))
        ttk.Checkbutton(act2, text="🌙  Xong thì ngủ", variable=self.var_sleep,
                        command=lambda: self._pick_power_mode("sleep")
                        ).pack(side="left", padx=(14, 0))

        # ── Tiến trình ── GỘP LÀM 1 với thanh dưới bước 3: self.progress/self.status
        # DÙNG CHUNG biến với pipe_progress/pipe_status → chỉ một tiến trình duy nhất
        # (tạo giọng + dựng video chạy trên cùng thanh đó). Thanh trong khối này chỉ
        # HIỆN ở view 'Giọng nói' (khi cột quy trình bị ẩn); các view khác đã thấy
        # thanh dưới bước 3 nên _show_view ẩn thanh này để không lặp 2 thanh.
        prog_frame = ttk.Frame(left)
        prog_frame.grid(row=6, column=0, sticky="ew", pady=(0, 10))
        prog_frame.columnconfigure(0, weight=1)
        self._voice_prog_frame = prog_frame
        self.progress = self.pipe_progress        # cùng biến với thanh dưới bước 3
        self.status = self.pipe_status            # cùng dòng trạng thái
        self.progress_bar = ttk.Progressbar(prog_frame, variable=self.progress,
                                            maximum=100, mode="determinate")
        self.progress_bar.grid(row=0, column=0, sticky="ew")
        ttk.Label(prog_frame, textvariable=self.status,
                  style="Sub.TLabel").grid(row=1, column=0, sticky="w", pady=(4, 0))

        # ── Nhật ký: ngay dưới Video dọc + TikTok (đẩy lên cao) ──
        self._build_log_panel(right, 2)
        self._show_view("home")   # mặc định mở Home (đầy đủ như giao diện gốc)

    def _show_view(self, key):
        """Hiện/ẩn 3 panel theo view và làm nổi nút đang chọn:
        • home   → cả 3 panel (giống giao diện gốc đầy đủ)
        • script → chỉ panel quy trình tạo kịch bản
        • voice  → panel TTS + panel video/hành động
        Ô nhật ký luôn hiển thị (nằm ngoài 3 panel)."""
        show = {
            "home":   ("pipeline", "tts", "video"),
            "script": ("pipeline",),
            "voice":  ("tts", "video"),
            "thumb":  ("thumbnail",),
            "copy":   ("copyseo",),
            "report": ("report",),
            "recog":  ("recog",),
        }[key]
        for name, fr in self._panels.items():
            if name in show:
                fr.grid()
            else:
                fr.grid_remove()
        # Chỉ 1 thanh tiến trình: thanh trong khối 'Giọng nói' CHỈ hiện ở view voice
        # (lúc cột quy trình bị ẩn). Các view khác đã có thanh dưới bước 3 → ẩn đi để
        # không hiện 2 thanh trùng nhau (cả hai vốn dùng chung 1 biến).
        if hasattr(self, "_voice_prog_frame"):
            if key == "voice":
                self._voice_prog_frame.grid()
            else:
                self._voice_prog_frame.grid_remove()
        if key == "report":
            self._refresh_report()   # cập nhật bảng tiến độ mỗi khi mở tab
        elif key == "copy":
            self._refresh_copyseo()  # quét lại danh sách tập mỗi khi mở tab
        elif key == "recog":
            self._recog_refresh_table()  # cập nhật bảng trạng thái các tập
        # Riêng tab "Tạo kịch bản": dàn 3 bước theo NGANG + hiện nhật ký, lấp đầy
        # chiều rộng. Ở Home/Giọng nói: pipeline là cột DỌC hẹp (như giao diện gốc).
        is_script = (key == "script")
        self._pipeline_set_layout(horizontal=is_script)
        self._content.columnconfigure(0, weight=1 if is_script else 0)
        for k, b in self._nav_buttons.items():
            b.configure(style="Accent.TButton" if k == key else "TButton")

    def _pipeline_set_layout(self, horizontal: bool):
        """Sắp lại các bước ①②③ của 'Tạo kịch bản':
        • horizontal=True  → 3 bước dàn ngang, nhật ký bên phải (tab xem riêng).
        • horizontal=False → cột dọc hẹp, ẩn nhật ký (Home/Giọng nói)."""
        w = self._pipe_wrap
        s1, s2, s3 = self._pipe_steps
        # Chặn bề ngang 3 dòng trạng thái theo kiểu dàn trang: chữ dài thì XUỐNG DÒNG
        # chứ không được kéo cột rộng ra (xem ghi chú ở _build_pipeline_column).
        for lb in getattr(self, "_pipe_status_labels", ()):
            lb.configure(wraplength=900 if horizontal else 430)
        if horizontal:
            for c, wt in ((0, 0), (1, 0), (2, 0), (3, 1)):
                w.columnconfigure(c, weight=wt)
            w.rowconfigure(1, weight=1)
            self._pipe_hdr.grid_configure(row=0, column=0, columnspan=3, sticky="w")
            s1.grid_configure(row=1, column=0, columnspan=1, sticky="new", padx=(0, 12))
            s2.grid_configure(row=1, column=1, columnspan=1, sticky="new", padx=(0, 12))
            s3.grid_configure(row=1, column=2, columnspan=1, sticky="new", padx=(0, 12))
            self._pipe_pf.grid_configure(row=2, column=0, columnspan=3, sticky="ew")
            self._pipe_btn_open.grid_configure(row=3, column=0, columnspan=3, sticky="w")
            self._pipe_btn_reset.grid_configure(row=4, column=0, columnspan=3, sticky="w")
            self._pipe_log_frame.grid()
            self._batch_ctrl_frame.grid()        # nút Tạm dừng/Dừng: hiện ở tab Tạo kịch bản
        else:
            for c, wt in ((0, 1), (1, 0), (2, 0), (3, 0)):
                w.columnconfigure(c, weight=wt)
            w.rowconfigure(1, weight=0)
            self._pipe_hdr.grid_configure(row=0, column=0, columnspan=1, sticky="ew")
            s1.grid_configure(row=1, column=0, columnspan=1, sticky="ew", padx=0)
            s2.grid_configure(row=2, column=0, columnspan=1, sticky="ew", padx=0)
            s3.grid_configure(row=3, column=0, columnspan=1, sticky="ew", padx=0)
            self._pipe_pf.grid_configure(row=4, column=0, columnspan=1, sticky="ew")
            self._pipe_btn_open.grid_configure(row=5, column=0, columnspan=1, sticky="w")
            self._pipe_btn_reset.grid_configure(row=6, column=0, columnspan=1, sticky="w")
            self._pipe_log_frame.grid_remove()
            self._batch_ctrl_frame.grid_remove()  # ẩn nút Tạm dừng/Dừng ở Home/Giọng nói

    # ── TAB TIẾN ĐỘ: bảng các link đã gửi (manifest) + đã làm tới đâu ────────────
    # Thứ tự bước hiển thị trong cột "Tiến độ" (nhãn ngắn + khoá trong steps dict).
    _REPORT_STEPS = (("Dịch", "translate"), ("SEO", "seo"), ("Thumb", "thumbnail"),
                     ("Giọng", "audio"), ("Vid ngang", "video_ngang"),
                     ("Vid dọc", "video_doc"), ("Đăng", "upload"))

    def _build_report_panel(self, parent):
        wrap = ttk.Frame(parent, padding=4)
        wrap.grid(row=0, column=0, sticky="nsew")
        wrap.columnconfigure(0, weight=1)
        wrap.rowconfigure(2, weight=1)

        hdr = ttk.Frame(wrap)
        hdr.grid(row=0, column=0, sticky="ew")
        ttk.Label(hdr, text="📋  Tiến độ các link đã gửi", style="Header.TLabel").pack(side="left")
        ttk.Label(wrap, text="Mỗi link nhớ ĐÚNG số tập của nó. Chạy lại (chưa xóa output) "
                  "sẽ tiếp tục đúng tập còn dở. ✅ = xong · 🟡 = đang dở · 🔴 = chưa làm.",
                  style="Sub.TLabel").grid(row=1, column=0, sticky="w", pady=(2, 8))

        bar = ttk.Frame(wrap)
        bar.grid(row=2, column=0, sticky="ew")
        bar.columnconfigure(0, weight=1)
        # Bảng
        table = ttk.Frame(bar)
        table.grid(row=0, column=0, sticky="nsew")
        bar.rowconfigure(0, weight=1)
        table.columnconfigure(0, weight=1)
        table.rowconfigure(0, weight=1)
        cols = ("st", "episode", "progress", "source", "updated")
        tv = ttk.Treeview(table, columns=cols, show="headings")
        for c, t, w, anchor in (("st", "", 34, "center"), ("episode", "Tập", 50, "center"),
                                ("progress", "Tiến độ", 360, "w"),
                                ("source", "Nguồn (link/file)", 380, "w"),
                                ("updated", "Cập nhật", 140, "w")):
            tv.heading(c, text=t)
            tv.column(c, width=w, anchor=anchor, stretch=(c == "source"))
        tv.grid(row=0, column=0, sticky="nsew")
        vs = ttk.Scrollbar(table, orient="vertical", command=tv.yview)
        vs.grid(row=0, column=1, sticky="ns")
        tv.configure(yscrollcommand=vs.set)
        self._report_tv = tv

        foot = ttk.Frame(wrap)
        foot.grid(row=3, column=0, sticky="ew", pady=(8, 0))
        ttk.Button(foot, text="🔄  Làm mới", command=self._refresh_report).pack(side="left")
        ttk.Button(foot, text="▶  Chạy tiếp tập đang chọn", style="Accent.TButton",
                   command=self._resume_selected_folder).pack(side="left", padx=(12, 0))
        self.report_summary = tk.StringVar(value="")
        ttk.Label(foot, textvariable=self.report_summary,
                  style="Sub.TLabel").pack(side="left", padx=(12, 0))
        # Nhấp đúp 1 dòng = chạy tiếp tập đó.
        tv.bind("<Double-1>", lambda e: self._resume_selected_folder())

    def _refresh_report(self):
        """Đổ lại bảng tiến độ từ manifest + quét cả thư mục tập chưa có trong manifest."""
        tv = getattr(self, "_report_tv", None)
        if tv is None:
            return
        tv.delete(*tv.get_children())
        m = load_manifest()
        # Gộp: mục manifest (theo episode) + thư mục tập rời chưa có trong manifest.
        rows = {}   # episode(int) -> (source, entry|None)
        for src, e in m.items():
            ep = e.get("episode", "")
            if str(ep).isdecimal():
                rows[int(ep)] = (e.get("source", src), e)
        for p in episode_dirs():
            n = int(episode_of(p.name))
            if n not in rows:
                rows[n] = ("(không rõ — tạo trước khi có manifest)", None)

        n_done = 0
        for ep in sorted(rows):
            src, e = rows[ep]
            episode = str(ep).zfill(2)
            folder = find_episode_dir(episode) or (SCRIPT_DIR / episode)
            steps = (e or {}).get("steps") or self._folder_steps(folder, episode)
            # "Hoàn tất" vẫn tính theo việc DỰNG xong (video dọc và đăng YouTube là
            # tuỳ chọn) — thêm 2 bước đó vào điều kiện sẽ làm mọi tập cũ hoá dở dang.
            done = bool((e or {}).get("done")) if e else all(
                steps.get(k) for _, k in self._REPORT_STEPS
                if k not in ("video_doc", "upload"))
            prog = "  ".join(("✅" if steps.get(k) else "⬜") + lbl
                             for lbl, k in self._REPORT_STEPS)
            mark = "✅" if done else ("🟡" if any(steps.values()) else "🔴")
            if done:
                n_done += 1
            updated = (e or {}).get("updated", "")
            tv.insert("", "end", values=(mark, episode, prog, src[:90], updated))
        self.report_summary.set(
            f"{len(rows)} tập • {n_done} hoàn tất • manifest: {MANIFEST_FILE}")

    def _resume_selected_folder(self):
        """Chạy TIẾP 1 tập đang chọn trong bảng — KHÔNG cần link gốc nếu đã có
        bản nhận diện (*_zh.docx). Chỉ làm các bước còn thiếu (vd chỉ còn video)."""
        if self._pipe_busy:
            messagebox.showinfo("Đang bận", "Đang có tác vụ chạy — đợi xong đã nhé.")
            return
        tv = getattr(self, "_report_tv", None)
        sel = tv.selection() if tv else None
        if not sel:
            messagebox.showinfo("Chọn tập", "Hãy chọn 1 tập trong bảng rồi bấm 'Chạy tiếp'.")
            return
        episode = str(tv.item(sel[0], "values")[1]).strip().zfill(2)
        folder = find_episode_dir(episode) or (SCRIPT_DIR / episode)
        if not folder.exists():
            messagebox.showwarning("Không có thư mục", f"Không thấy thư mục tập {episode}.")
            return
        # Không có link gốc vẫn chạy tiếp được MIỄN LÀ đã nhận diện (tiengTrung.docx/*_zh.docx).
        if find_zh_docx(folder) is None:
            messagebox.showwarning(
                "Thiếu bản nhận diện",
                f"Tập {episode} chưa có *_zh.docx nên KHÔNG thể chạy tiếp mà thiếu link "
                "gốc (cần nhận diện lại). Hãy chạy lại link đó ở tab 'Tạo kịch bản'.")
            return
        tts_settings = self._collect_tts_settings()   # main thread (đọc tk.Var)
        if tts_settings is None:
            return   # cấu hình giọng sai → đã cảnh báo
        # Đăng YouTube (nếu đang tick): kiểm tra đăng nhập ngay trên main thread.
        upload = bool(self.var_upload.get())
        if upload and not self._upload_check_ready():
            return
        self._pipe_set_busy(True)
        self.pipe_progress.set(0)
        self.pipe_link_status.set(f"▶ Chuẩn bị chạy tiếp tập {episode}...")
        self._show_view("script")   # chuyển sang tab có thanh tiến trình + nhật ký
        threading.Thread(target=self._resume_folder_worker,
                         args=(folder, episode, tts_settings, upload), daemon=True).start()

    def _resume_folder_worker(self, folder, episode, tts_settings, upload=False):
        """Chạy tiếp 1 thư mục tập đã có sẵn (bỏ qua bước đã xong, làm phần còn thiếu).
        Dùng lại toàn bộ logic resume/bỏ-qua như batch nhưng cho ĐÚNG 1 tập."""
        import datetime as _dt
        driver = None
        file_handler = None
        try:
            import nhandien_giongnoi as recog
            import dich_gemini as g
            youtube_dir = str(YOUTUBE_DIR)
            if youtube_dir not in sys.path:
                sys.path.insert(0, youtube_dir)
            import seo_youtube_gemini as seo
            prefix = load_prefix()
            try:
                file_handler = logging.FileHandler(SCRIPT_DIR / "batch_log.txt", encoding="utf-8")
                file_handler.setFormatter(
                    logging.Formatter("%(asctime)s  %(message)s", "%Y-%m-%d %H:%M:%S"))
                logging.getLogger().addHandler(file_handler)
                logging.info("\n" + "═" * 12 + f" CHẠY TIẾP TẬP {episode} "
                             f"({_dt.datetime.now():%Y-%m-%d %H:%M:%S}) " + "═" * 12)
            except Exception as e:
                # Không mở được batch_log.txt → chạy xong không có nhật ký để dò lỗi.
                logging.warning(f"⚠️ Không ghi được batch_log.txt (vẫn chạy tiếp): {e}")

            # Nguồn (lấy từ manifest nếu có) — chỉ để ghi log/manifest, KHÔNG cần để chạy.
            m = load_manifest()
            src = next((s for s, e in m.items()
                        if str(e.get("episode", "")).zfill(2) == episode), str(folder))

            gemini_docx = folder / "gemini_result.docx"
            input_txt = folder / "input.txt"
            seo_docx = folder / "seoYoutube.docx"

            # 1+2) Nhận diện: dùng lại bản đã có (tiengTrung.docx hoặc *_zh.docx cũ).
            existing_zh = find_zh_docx(folder)
            chunks = read_zh_docx_chunks(existing_zh) if existing_zh else []
            if not chunks:
                logging.error(f"❌ Tập {episode}: không đọc được đoạn từ *_zh.docx.")
                self.pipe_status.set(f"❌ Tập {episode}: lỗi đọc nhận diện.")
                return
            logging.info(f"♻ Dùng lại nhận diện ({existing_zh.name}, {len(chunks)} đoạn).")

            # 3) Dịch Gemini — đủ thì bỏ qua; thiếu thì tiếp tục (xem _dich_gemini_cho_tap).
            driver, translated_now, translation_ok = self._dich_gemini_cho_tap(
                gemini_docx, chunks, prefix, episode, driver)

            if not translation_ok:
                self.pipe_status.set(f"⛔ Tập {episode}: dịch chưa xong — dừng.")
                self._manifest_update(src, episode, folder, done=False)
                return

            # 4) input.txt — tạo lại nếu vừa dịch (bản cũ có thể dở) hoặc chưa có.
            if translated_now or not (input_txt.exists() and input_txt.stat().st_size > 0):
                self._batch_prepare_input(gemini_docx, input_txt)

            # 5) SEO
            if not self._seo_docx_valid(seo_docx):
                if driver is None:
                    driver = g.init_firefox()
                logging.info("🔎 Tạo SEO YouTube...")
                seo.run(str(gemini_docx), str(seo_docx),
                        keep_open=True, log=logging.info, driver=driver)
            else:
                logging.info("♻ Bỏ qua SEO (đã có).")
            self._save_youtube_seo_copy(seo_docx, folder / "youtube_seo.txt", episode)

            # 6) Thumbnail (ngang + dọc) + cập nhật số tập
            if not (folder / f"thumbnail{episode}.png").exists() \
                    or not (folder / f"thumbnail{episode}_dọc.png").exists():
                self._make_thumbnail_for_folder(folder, episode)
            save_episode_number(max(load_episode_number(), int(episode)))
            # 7) KHÔNG tự động tải input.txt lên Drive nữa (theo yêu cầu). Muốn tải
            # thủ công thì dùng ô tick "Tải kịch bản lên Drive" ở tab Thumbnail.

            # Đóng Firefox trước khi render video (nhả RAM).
            if driver is not None:
                try:
                    driver.quit()
                except Exception:
                    pass
                driver = None

            # 8) Tạo giọng + video (reuse=True → chỉ render phần còn thiếu, vd video dọc).
            if tts_settings:
                try:
                    recog.free_model()
                except Exception:
                    pass
                self.pipe_status.set(f"🎧 Tập {episode}: đang tạo giọng + video...")
                self._batch_run_tts(folder, tts_settings, episode)

            self._don_rac_audio(folder)     # 🧹 xoá audio trung gian (giữ output.wav)
            self._manifest_update(src, episode, folder, done=True)
            if upload:
                self._upload_enqueue(folder, episode)
            self.pipe_progress.set(100)
            self.pipe_link_status.set(f"✅ Tập {episode} đã chạy tiếp xong.")
            self.pipe_status.set(f"✅ Tập {episode} hoàn tất.")
            logging.info(f"🎉 Tập {episode} chạy tiếp xong.")
        except Exception as e:
            import traceback
            logging.error(f"❌ Lỗi chạy tiếp tập {episode}: {e}")
            logging.error(traceback.format_exc())
            self.pipe_link_status.set(f"❌ Lỗi chạy tiếp tập {episode}.")
            self.pipe_status.set(f"Lỗi: {e}")
        finally:
            if driver is not None:
                try:
                    driver.quit()
                except Exception:
                    pass
            try:
                import nhandien_giongnoi as recog
                recog.free_model()
            except Exception:
                pass
            self._upload_wait_drain()   # đăng nốt trước khi hạ cờ bận (có thể hẹn tắt máy)
            if file_handler is not None:
                try:
                    logging.getLogger().removeHandler(file_handler)
                    file_handler.close()
                except Exception:
                    pass
            self._pipe_set_busy(False)

    # ── TAB COPY SEO: chọn 1 tập (thư mục số) → copy tiêu đề/mô tả/thẻ tag ───────
    def _build_copyseo_panel(self, parent):
        wrap = ttk.Frame(parent, padding=4)
        wrap.grid(row=0, column=0, sticky="nsew")
        wrap.columnconfigure(0, weight=1)
        wrap.rowconfigure(2, weight=1)

        hdr = ttk.Frame(wrap)
        hdr.grid(row=0, column=0, sticky="ew")
        ttk.Label(hdr, text="📑  Copy SEO theo tập", style="Header.TLabel").pack(side="left")
        ttk.Label(wrap, text="Chọn 1 tập (thư mục số trong kịch_bản) rồi copy Tiêu đề · Mô tả · Thẻ tag. "
                  "Tiêu đề mở đầu [FULL]; mô tả có #truyenfull #full; thẻ tag < 499 ký tự.",
                  style="Sub.TLabel").grid(row=1, column=0, sticky="w", pady=(2, 8))

        body = ttk.Frame(wrap)
        body.grid(row=2, column=0, sticky="nsew")
        body.columnconfigure(1, weight=1)
        body.rowconfigure(0, weight=1)

        # Cột trái: danh sách tập có SEO.
        left = ttk.Frame(body)
        left.grid(row=0, column=0, sticky="ns", padx=(0, 12))
        left.rowconfigure(0, weight=1)
        lst = tk.Listbox(left, width=12, exportselection=False, activestyle="dotbox",
                         font=("Segoe UI", 10))
        lst.grid(row=0, column=0, sticky="ns")
        lsb = ttk.Scrollbar(left, orient="vertical", command=lst.yview)
        lsb.grid(row=0, column=1, sticky="ns")
        lst.configure(yscrollcommand=lsb.set)
        lst.bind("<<ListboxSelect>>", lambda e: self._copyseo_load_selected())
        self._copyseo_list = lst

        # Cột phải: 4 nút copy + xem trước nội dung.
        right = ttk.Frame(body)
        right.grid(row=0, column=1, sticky="nsew")
        right.columnconfigure(0, weight=1)
        right.rowconfigure(1, weight=1)

        btns = ttk.Frame(right)
        btns.grid(row=0, column=0, sticky="ew", pady=(0, 8))
        ttk.Button(btns, text="📋  Tiêu đề youtube",
                   command=lambda: self._copyseo_copy("title")).pack(side="left")
        ttk.Button(btns, text="📋  Tiêu đề tiktok",
                   command=lambda: self._copyseo_copy("title_tiktok")).pack(side="left", padx=(8, 0))
        ttk.Button(btns, text="📋  Mô tả", command=lambda: self._copyseo_copy("desc")).pack(side="left", padx=(8, 0))
        ttk.Button(btns, text="📋  Thẻ tag", command=lambda: self._copyseo_copy("tags")).pack(side="left", padx=(8, 0))
        ttk.Button(btns, text="📋  Cả 3", style="Accent.TButton",
                   command=lambda: self._copyseo_copy("all")).pack(side="left", padx=(8, 0))

        txt = tk.Text(right, wrap="word", height=18, font=("Consolas", 10),
                      bg="white", relief="solid", borderwidth=1)
        txt.grid(row=1, column=0, sticky="nsew")
        tsb = ttk.Scrollbar(right, orient="vertical", command=txt.yview)
        tsb.grid(row=1, column=1, sticky="ns")
        txt.configure(yscrollcommand=tsb.set, state="disabled")
        self._copyseo_text = txt

        foot = ttk.Frame(wrap)
        foot.grid(row=3, column=0, sticky="ew", pady=(8, 0))
        ttk.Button(foot, text="🔄  Làm mới", command=self._refresh_copyseo).pack(side="left")
        self._copyseo_status = tk.StringVar(value="")
        ttk.Label(foot, textvariable=self._copyseo_status,
                  style="Sub.TLabel").pack(side="left", padx=(12, 0))

        self._copyseo_episodes = []       # các số tập đang hiển thị (khớp index listbox)
        self._copyseo_blocks = {}         # cache: episode -> {'title','desc','tags'}

    def _copyseo_selected_episode(self):
        lst = getattr(self, "_copyseo_list", None)
        sel = lst.curselection() if lst else ()
        eps = getattr(self, "_copyseo_episodes", [])
        if sel and sel[0] < len(eps):
            return eps[sel[0]]
        return None

    def _refresh_copyseo(self):
        """Quét các thư mục số trong kịch_bản CÓ seoYoutube.docx → đổ vào danh sách."""
        lst = getattr(self, "_copyseo_list", None)
        if lst is None:
            return
        prev = self._copyseo_selected_episode()
        self._copyseo_blocks = {}        # SEO có thể đã đổi → đọc lại khi chọn
        # Giữ TÊN thư mục (có kèm tên nguồn) để hiện ra danh sách + dựng đường dẫn;
        # số tập thật lấy lại bằng episode_of() khi cần dựng tiêu đề SEO.
        eps = [p.name for p in episode_dirs() if (p / "seoYoutube.docx").exists()]
        self._copyseo_episodes = eps
        lst.delete(0, tk.END)
        for ep in eps:
            lst.insert(tk.END, f"  Tập {ep}")
        self._copyseo_status.set(f"{len(eps)} tập có SEO trong kịch_bản")
        if prev in eps:                  # giữ nguyên tập đang chọn nếu còn
            i = eps.index(prev)
            lst.selection_set(i)
            lst.see(i)
            self._copyseo_load_selected()
        else:
            self._set_copyseo_preview("← Chọn một tập ở cột trái để xem và copy.")

    def _set_copyseo_preview(self, text: str):
        txt = getattr(self, "_copyseo_text", None)
        if txt is None:
            return
        txt.configure(state="normal")
        txt.delete("1.0", tk.END)
        txt.insert("1.0", text or "")
        txt.configure(state="disabled")

    def _copyseo_load_selected(self):
        """Đọc SEO của tập đang chọn (có cache) rồi hiện xem trước 3 phần."""
        ep = self._copyseo_selected_episode()
        if not ep:
            return
        blocks = self._copyseo_blocks.get(ep)
        if blocks is None:
            blocks = self._seo_copy_blocks(SCRIPT_DIR / ep / "seoYoutube.docx",
                                           episode_of(ep) or "")
            self._copyseo_blocks[ep] = blocks
        if not blocks:
            self._set_copyseo_preview(f"(Không đọc được SEO của tập {ep}.)")
            self._copyseo_status.set(f"Tập {ep}: lỗi đọc SEO")
            return
        self._set_copyseo_preview(
            "===== TIÊU ĐỀ YOUTUBE =====\n" + blocks["title"] + "\n\n"
            "===== TIÊU ĐỀ TIKTOK =====\n" + blocks.get("title_tiktok", "") + "\n\n"
            "===== MÔ TẢ =====\n" + blocks["desc"] + "\n\n"
            "===== THẺ TAG =====\n" + blocks["tags"])
        self._copyseo_status.set(f"Tập {ep} • thẻ tag {len(blocks['tags'])} ký tự")

    def _copyseo_copy(self, which: str):
        """Copy 1 phần (title/title_tiktok/desc/tags) hoặc 'all' của tập đang chọn."""
        ep = self._copyseo_selected_episode()
        if not ep:
            messagebox.showinfo("Chọn tập", "Hãy chọn 1 tập trong danh sách bên trái.")
            return
        blocks = self._copyseo_blocks.get(ep) or \
            self._seo_copy_blocks(SCRIPT_DIR / ep / "seoYoutube.docx",
                                  episode_of(ep) or "")
        if not blocks:
            messagebox.showwarning("Không có SEO", f"Tập {ep} chưa đọc được nội dung SEO.")
            return
        self._copyseo_blocks[ep] = blocks
        label = {"title": "tiêu đề youtube", "title_tiktok": "tiêu đề tiktok",
                 "desc": "mô tả", "tags": "thẻ tag", "all": "cả 3 phần"}[which]
        if which == "all":
            text = blocks["title"] + "\n\n" + blocks["desc"] + "\n\n" + blocks["tags"]
        else:
            text = blocks.get(which, "")
        if not text.strip():
            self._copyseo_status.set(f"Tập {ep}: không có {label} để copy.")
            return
        self.clipboard_clear()
        self.clipboard_append(text)
        self._copyseo_status.set(f"✓ Đã copy {label} tập {ep} ({len(text)} ký tự)")

    # ── TAB '🎧 NHẬN DIỆN': nhận diện link + dịch + tạo giọng/video HÀNG LOẠT ──────
    def _build_recog_panel(self, parent):
        """Quy trình 5 bước cho HÀNG LOẠT tập trong kịch_bản/:
          ①  Nhận diện các link → tiengTrung.docx (mỗi link 1 thư mục tập).
          ②  Dịch Gemini + tạo input.txt cho các tập đã nhận diện.
          ③  Gửi SEO lên Gemini → seoYoutube.docx + youtube_seo.txt (cần Firefox).
          ④  Tạo thumbnail ngang/dọc từ tiêu đề SEO (chỉ render ảnh, KHÔNG cần Firefox).
          ⑤  Tạo giọng clone + video (y hệt tab Home) cho các tập đã có input.txt.
        ③ và ④ tách riêng: render lại ảnh khỏi phải đụng tới Gemini/Firefox.
        Có BẢNG trạng thái các tập (tick chọn tập để chạy riêng) + nút Làm mới; bảng
        tự cập nhật khi mở tab và sau mỗi bước."""
        self._recog_checked = set()   # số tập được tick để chạy riêng (rỗng = chạy hết)
        self._recog_seo_cache = {}    # (file, mtime, size) -> SEO hợp lệ? (xem _recog_seo_ok)
        wrap = ttk.Frame(parent, padding=6)
        wrap.grid(row=0, column=0, sticky="nsew")
        # Xếp 2 CỘT: trái = điều khiển (①②③ + tiến trình), phải = bảng tập + nhật ký.
        # Xếp dọc 1 cột thì tổng chiều cao (~965px) vượt cửa sổ (~684px) → grid bóp
        # hàng có weight → BẢNG TẬP bị co còn 1px (không thấy tập nào để tick).
        wrap.columnconfigure(0, weight=0, minsize=460)   # cột trái: rộng cố định
        wrap.columnconfigure(1, weight=1)                # cột phải: giãn
        wrap.rowconfigure(2, weight=1)

        ttk.Label(wrap, text="🎧  Nhận diện & tạo kịch bản/video hàng loạt",
                  style="Header.TLabel").grid(row=0, column=0, columnspan=2, sticky="w")
        ttk.Label(wrap, text="Quy trình 5 bước chạy cho các thư mục tập trong kịch_bản/. Tick tập "
                  "trong bảng để chỉ chạy tập đó; KHÔNG tick = chạy hết tập đủ điều kiện.",
                  style="Sub.TLabel").grid(row=1, column=0, columnspan=2, sticky="w", pady=(2, 8))

        left = ttk.Frame(wrap)
        left.grid(row=2, column=0, sticky="nsew", padx=(0, 10))
        left.columnconfigure(0, weight=1)
        right = ttk.Frame(wrap)
        right.grid(row=2, column=1, sticky="nsew")
        right.columnconfigure(0, weight=1)
        right.rowconfigure(0, weight=3, minsize=240)   # bảng tập: KHÔNG co dưới 240px
        right.rowconfigure(1, weight=1)                # nhật ký: nhường chỗ cho bảng

        # ── ①  Nhận diện link → tiếng Trung ─────────────────────────────────────
        s1 = ttk.LabelFrame(left, text="  ①  Nhận diện link → tiếng Trung  ")
        s1.grid(row=0, column=0, sticky="ew")
        s1.columnconfigure(0, weight=1)
        self.recog_txt_sources = tk.Text(s1, height=3, wrap="none",
                                         font=("Segoe UI", 10), relief="solid", bd=1)
        self.recog_txt_sources.grid(row=0, column=0, sticky="ew", padx=6, pady=(6, 4))
        ttk.Button(s1, text="Chọn…", width=8, command=self._recog_pick_file).grid(
            row=0, column=1, padx=(0, 6), pady=(6, 4), sticky="n")
        orow = ttk.Frame(s1)
        orow.grid(row=1, column=0, columnspan=2, sticky="w", padx=6, pady=(0, 4))
        ttk.Label(orow, text="Mỗi dòng 1 link/file   •   Model:").pack(side="left")
        ttk.Combobox(orow, textvariable=self.pipe_var_model, width=9, state="readonly",
                     values=["tiny", "base", "small", "medium", "large-v3"]).pack(side="left", padx=(4, 12))
        ttk.Label(orow, text="Tốc độ:").pack(side="left")
        ttk.Combobox(orow, textvariable=self.pipe_var_speed, width=5, state="readonly",
                     values=["0.6", "0.7", "0.8", "0.9", "1.0"]).pack(side="left", padx=(4, 0))
        self.recog_tab_btn = ttk.Button(
            s1, text="🎧  ①  Nhận diện các link rồi ngưng",
            style="Accent.TButton", command=self._recog_tab_run)
        self.recog_tab_btn.grid(row=2, column=0, columnspan=2, sticky="ew", padx=6, pady=(0, 8))

        # ── Bảng trạng thái các tập trong kịch_bản/ ─────────────────────────────
        tbl = ttk.LabelFrame(right, text="  Các tập trong kịch_bản/  ")
        tbl.grid(row=0, column=0, sticky="nsew")
        tbl.columnconfigure(0, weight=1)
        tbl.rowconfigure(1, weight=1)
        bar = ttk.Frame(tbl)
        bar.grid(row=0, column=0, columnspan=2, sticky="ew", padx=6, pady=(6, 4))
        ttk.Button(bar, text="🔄 Làm mới", command=self._recog_refresh_table).pack(side="left")
        ttk.Button(bar, text="Chọn tất cả", command=lambda: self._recog_check_all(True)).pack(side="left", padx=(8, 0))
        ttk.Button(bar, text="Bỏ chọn", command=lambda: self._recog_check_all(False)).pack(side="left", padx=(6, 0))
        ttk.Label(bar, text="(bấm vào DÒNG để tick tập cần chạy riêng • ✅ = đã có · — = chưa "
                            "· ½ = thumbnail mới có 1 bản)",
                  style="Sub.TLabel").pack(side="left", padx=(10, 0))

        cols = ("sel", "ep", "zh", "input", "seo", "thumb", "audio", "video", "up")
        self._recog_tree = tree = ttk.Treeview(tbl, columns=cols, show="headings",
                                               height=8, selectmode="none")
        heads = {"sel": ("", 34), "ep": ("Tập", 54), "zh": ("Tiếng Trung", 100),
                 "input": ("input.txt", 88), "seo": ("SEO", 62), "thumb": ("Thumbnail", 78),
                 "audio": ("Giọng", 72), "video": ("Video", 72), "up": ("Đăng", 62)}
        for c, (txt, w) in heads.items():
            tree.heading(c, text=txt)
            tree.column(c, width=w, anchor="center", stretch=(c in ("zh", "input")))
        tree.grid(row=1, column=0, sticky="nsew", padx=(6, 0), pady=(0, 6))
        vsb = ttk.Scrollbar(tbl, orient="vertical", command=tree.yview)
        vsb.grid(row=1, column=1, sticky="ns", pady=(0, 6), padx=(0, 6))
        tree.configure(yscrollcommand=vsb.set)
        tree.bind("<Button-1>", self._recog_on_tree_click)
        # Tag "sel" đặt TRƯỚC need/done trong danh sách tag của dòng → màu tick thắng
        # (Tk lấy option ở tag đứng đầu), nhờ vậy tập đang tick luôn nhìn ra ngay.
        tree.tag_configure("sel", background="#CFE3FF")     # tập ĐANG TICK → xanh dương nhạt
        tree.tag_configure("need", background="#FFF3CD")   # tập CÒN VIỆC → tô vàng nhạt
        tree.tag_configure("done", background="#E7F6E7")   # tập đã đủ video → xanh nhạt
        # Dòng tóm tắt: tổng số tập + số tập cần bước ②/③ (đếm ở _recog_refresh_table).
        self._recog_count_var = tk.StringVar(value="")
        ttk.Label(tbl, textvariable=self._recog_count_var, style="Sub.TLabel").grid(
            row=2, column=0, columnspan=2, sticky="w", padx=8, pady=(0, 6))

        # ── ② + ③ + ④ + ⑤  (chạy cho tập được tick, hoặc hết tập đủ điều kiện) ───
        # Cột trái hẹp → xếp DỌC cho đủ chỗ chữ (nút xếp ngang sẽ bị cắt tên).
        btnrow = ttk.Frame(left)
        btnrow.grid(row=1, column=0, sticky="ew", pady=(10, 0))
        btnrow.columnconfigure(0, weight=1)
        self.recog_translate_btn = ttk.Button(
            btnrow, text="🌐  ②  Dịch + tạo input.txt",
            style="Accent.TButton", command=self._recog_translate_all)
        self.recog_translate_btn.grid(row=0, column=0, sticky="ew")
        # ③ SEO và ④ thumbnail đứng CÙNG 1 HÀNG, chia đôi bề ngang (uniform) — 2 bước
        # này đi liền nhau nên để cạnh nhau. Nửa hàng ~342px, tên nút cần ~205/281px
        # nên vẫn đủ chỗ, không bị cắt chữ như các nút dài ở hàng riêng.
        seorow = ttk.Frame(btnrow)
        seorow.grid(row=1, column=0, sticky="ew", pady=(8, 0))
        seorow.columnconfigure(0, weight=1, uniform="seo")
        seorow.columnconfigure(1, weight=1, uniform="seo")
        self.recog_seo_btn = ttk.Button(
            seorow, text="🔎  ③  Gửi SEO (Gemini)",
            style="Accent.TButton", command=self._recog_seo_all)
        self.recog_seo_btn.grid(row=0, column=0, sticky="ew", padx=(0, 4))
        self.recog_thumb_btn = ttk.Button(
            seorow, text="🖼  ④  Tạo thumbnail (ngang + dọc)",
            style="Accent.TButton", command=self._recog_thumb_all)
        self.recog_thumb_btn.grid(row=0, column=1, sticky="ew", padx=(4, 0))
        self.recog_tts_btn = ttk.Button(
            btnrow, text="🎬  ⑤  Tạo giọng + video (clone như Home)",
            style="Accent.TButton", command=self._recog_make_video_all)
        self.recog_tts_btn.grid(row=2, column=0, sticky="ew", pady=(8, 0))
        self.recog_upload_btn = ttk.Button(
            btnrow, text=f"⬆  ⑥  Đăng YouTube (hẹn {upload_slots_text()})",
            style="Accent.TButton", command=self._recog_upload_all)
        self.recog_upload_btn.grid(row=3, column=0, sticky="ew", pady=(8, 0))
        # ⚡ 1 nút chạy TRỌN quy trình: tập thiếu bước nào thì tự chạy bước đó ②→⑤
        # (hỏi trước có kèm ⑥), khỏi phải bấm lần lượt từng nút.
        self.recog_chain_btn = ttk.Button(
            btnrow, text="⚡  Chạy TIẾP các bước còn thiếu (②→⑥)",
            style="Accent.TButton", command=self._recog_run_missing)
        self.recog_chain_btn.grid(row=4, column=0, sticky="ew", pady=(14, 0))

        # ── Tiến trình + Tạm dừng/Dừng (dùng chung biến với batch) ──────────────
        pf = ttk.Frame(left)
        pf.grid(row=2, column=0, sticky="ew", pady=(10, 0))
        pf.columnconfigure(0, weight=1)
        ttk.Progressbar(pf, variable=self.pipe_progress, maximum=100).grid(row=0, column=0, sticky="ew")
        # wraplength BẮT BUỘC: label không giới hạn bề ngang sẽ ĐÒI đúng bề ngang của
        # chữ; cột trái để weight=0 nên nó phình theo → bảng bên phải co lại, cả giao
        # diện "nhảy" mỗi lần tick tập (dòng trạng thái liệt kê tên các tập đang tick).
        # 430 < minsize 460 của cột trái → chữ dài mấy cũng không đẩy được cột.
        ttk.Label(pf, textvariable=self.pipe_link_status,
                  font=("Segoe UI", 10, "bold"), foreground=UI["accent"],
                  wraplength=430, justify="left").grid(
            row=1, column=0, sticky="w", pady=(4, 0))
        ttk.Label(pf, textvariable=self.pipe_status, style="Sub.TLabel",
                  wraplength=430, justify="left").grid(
            row=2, column=0, sticky="w", pady=(2, 0))
        bctl = ttk.Frame(pf)
        bctl.grid(row=3, column=0, sticky="ew", pady=(8, 0))
        b_pause = ttk.Button(bctl, text="⏸  Tạm dừng", width=15, state="disabled",
                             command=self._batch_toggle_pause)
        b_pause.pack(side="left")
        b_stop = ttk.Button(bctl, text="⏹  Xong tập này rồi dừng", width=24, state="disabled",
                            command=self._batch_request_stop)
        b_stop.pack(side="left", padx=(8, 0))
        self._batch_pause_widgets.append(b_pause)
        self._batch_stop_widgets.append(b_stop)
        ttk.Label(pf, textvariable=self.upload_status, style="Sub.TLabel",
                  foreground=UI["accent"], wraplength=430, justify="left").grid(
            row=4, column=0, sticky="w", pady=(2, 0))

        # ── Nhật ký ─────────────────────────────────────────────────────────────
        logf = ttk.LabelFrame(right, text="  Nhật ký  ")
        logf.grid(row=1, column=0, sticky="nsew", pady=(10, 0))
        logf.columnconfigure(0, weight=1)
        logf.rowconfigure(0, weight=1)
        self._make_log_box(logf).configure(height=6)   # thấp hơn mặc định: nhường chỗ cho bảng

        self._recog_refresh_table()

    # ── BẢNG trạng thái các tập + tick chọn ──────────────────────────────────────
    def _all_episode_folders(self) -> list:
        """Mọi thư mục tập kịch_bản/NN (hoặc 'NN - tên nguồn'), theo số tập."""
        return episode_dirs()

    def _recog_seo_ok(self, folder) -> bool:
        """SEO của 1 tập đã có tiêu đề thật chưa (dùng chung logic với Home).

        Có CACHE theo (đường dẫn, mtime, cỡ file): mỗi lần làm mới bảng mà mở lại
        seoYoutube.docx của mọi tập thì bảng giật; file đổi thì key đổi → tự đọc lại.
        """
        p = Path(folder) / "seoYoutube.docx"
        try:
            st = p.stat()
        except OSError:
            return False
        key = (str(p), st.st_mtime_ns, st.st_size)
        cache = self._recog_seo_cache
        if key not in cache:
            if len(cache) > 400:
                cache.clear()
            cache[key] = self._seo_docx_valid(p)
        return cache[key]

    def _thumb_count(self, folder, episode: str) -> int:
        """Số bản thumbnail đã có của 1 tập: 0 · 1 (thiếu ngang hoặc dọc) · 2 (đủ)."""
        folder = Path(folder)
        return int((folder / f"thumbnail{episode}.png").is_file()) \
            + int((folder / f"thumbnail{episode}_dọc.png").is_file())

    def _recog_refresh_table(self):
        """Đổ lại bảng: mỗi tập 1 dòng + trạng thái (tiếng Trung / input.txt / SEO /
        thumbnail / giọng / video)."""
        tree = getattr(self, "_recog_tree", None)
        if tree is None:
            return
        try:
            tree.delete(*tree.get_children())
        except Exception:
            return
        yes, no = "✅", "—"
        total = need2 = need_seo = need_thumb = need5 = 0
        for folder in self._all_episode_folders():
            total += 1
            # ep = TÊN thư mục (khóa tick + hiển thị, có thể là "01 - 95");
            # epnum = SỐ TẬP thật, dùng khi cần đánh số (vd đếm thumbnail).
            ep = folder.name
            epnum = episode_of(ep)
            has_zh = find_zh_docx(folder) is not None
            inp = folder / "input.txt"
            try:
                has_inp = inp.is_file() and inp.stat().st_size > 0
            except OSError:
                has_inp = False
            gem = folder / "gemini_result.docx"
            try:
                has_gem = gem.is_file() and gem.stat().st_size > 0
            except OSError:
                has_gem = False
            has_seo = self._recog_seo_ok(folder)
            n_thumb = self._thumb_count(folder, epnum)
            has_aud = (folder / "output.wav").is_file()
            has_vid = any((folder / n).is_file()
                          for n in ("YOUTUBE.mp4", "facebook.mp4", "tiktok.mp4"))
            has_up = (folder / "youtube_upload.json").is_file()
            # "Cần làm": có tiếng Trung mà chưa có input.txt (②); đã dịch mà chưa có
            # SEO (③); đã có SEO mà thiếu thumbnail (④ — thumbnail lấy tiêu đề từ SEO
            # nên chưa SEO thì chưa tính là cần); có input.txt mà chưa có video (⑤).
            need_step2 = has_zh and not has_inp
            need_step3 = has_gem and not has_seo
            need_step4 = has_seo and n_thumb < 2
            need_step5 = has_inp and not has_vid
            need2 += int(need_step2)
            need_seo += int(need_step3)
            need_thumb += int(need_step4)
            need5 += int(need_step5)
            tag = "need" if (need_step2 or need_step3 or need_step4 or need_step5) \
                else ("done" if has_vid else "")
            tags = (("sel",) if ep in self._recog_checked else ()) \
                + ((tag,) if tag else ())
            tree.insert("", "end", tags=tags, values=(
                "☑" if ep in self._recog_checked else "☐",
                ep,
                yes if has_zh else no,
                yes if has_inp else no,
                yes if has_seo else no,
                yes if n_thumb == 2 else ("½" if n_thumb else no),
                yes if has_aud else no,
                yes if has_vid else no,
                yes if has_up else no,
            ))
        var = getattr(self, "_recog_count_var", None)
        if var is not None:
            if total == 0:
                var.set("⚠️ Chưa có thư mục tập nào trong kịch_bản/ — chạy ① để tạo, hoặc bấm 🔄 Làm mới.")
            else:
                var.set(f"Tổng {total} tập   •   cần ② (dịch+input.txt): {need2}   •   "
                        f"cần ③ (SEO): {need_seo}   •   cần ④ (thumbnail): {need_thumb}   •   "
                        f"cần ⑤ (giọng+video): {need5}   •   dòng tô vàng = còn việc")
        self._recog_selection_hint()

    @staticmethod
    def _ep_list_short(names, max_chars: int = 22) -> str:
        """Danh sách tên tập RÚT GỌN cho dòng trạng thái (nhật ký vẫn ghi đủ tên).

        Liệt kê hết tên tập thì dòng trạng thái dài vô tận → label đòi bề ngang đúng
        bằng chữ, kéo cả cột/cửa sổ rộng ra mỗi lần tick (giao diện nhảy). Cắt theo
        SỐ KÝ TỰ (tên tập có thể là '01 - 95') để dòng luôn gói gọn trong 1 hàng."""
        names = list(names)
        out, used = [], 0
        for n in names:
            add = len(n) + (2 if out else 0)
            if out and used + add > max_chars:
                break
            out.append(n)
            used += add
        rest = len(names) - len(out)
        return ", ".join(out) + (f" …+{rest}" if rest > 0 else "")

    def _recog_selection_hint(self):
        """Ghi rõ ĐANG tick tập nào ra dòng trạng thái — để biết chắc bấm ②→⑤ sẽ
        chạy tập nào TRƯỚC khi bấm (không tick = chạy hết). Đang chạy thì không ghi
        đè trạng thái của tác vụ."""
        if self._pipe_busy:
            return
        eps = sorted(getattr(self, "_recog_checked", set()))
        if eps:
            self.pipe_status.set(f"☑ Đang tick {len(eps)} tập: {self._ep_list_short(eps)} → "
                                 "②→⑤ CHỈ chạy các tập này.")
        else:
            self.pipe_status.set("Chưa tick tập nào → ②→⑤ sẽ chạy HẾT tập đủ điều kiện.")

    def _recog_on_tree_click(self, event):
        """Bấm vào BẤT KỲ ô nào của 1 dòng = tick/bỏ tick tập đó.

        Trước đây chỉ ô ☐ (rộng 34px) mới ăn → bấm vào số tập tưởng đã chọn nhưng
        thật ra chưa tick, bấm ②→⑤ là chạy HẾT tập. Nay bấm đâu trên dòng cũng được.
        """
        tree = self._recog_tree
        if tree.identify_region(event.x, event.y) != "cell":
            return
        row = tree.identify_row(event.y)
        if not row:
            return
        ep = tree.set(row, "ep")
        if ep in self._recog_checked:
            self._recog_checked.discard(ep)
        else:
            self._recog_checked.add(ep)
        on = ep in self._recog_checked
        tree.set(row, "sel", "☑" if on else "☐")
        # Đổi màu dòng ngay (khỏi dựng lại cả bảng): "sel" phải đứng đầu để thắng need/done.
        tags = [t for t in tree.item(row, "tags") if t != "sel"]
        tree.item(row, tags=(["sel"] + tags) if on else tags)
        self._recog_selection_hint()
        return "break"

    def _recog_check_all(self, on: bool):
        """Tick tất cả / bỏ tick tất cả các tập trong bảng."""
        self._recog_checked = {f.name for f in self._all_episode_folders()} if on else set()
        self._recog_refresh_table()

    def _recog_apply_selection(self, folders: list) -> list:
        """Có tick tập nào → chỉ giữ tập được tick; không tick → giữ nguyên (chạy hết)."""
        sel = getattr(self, "_recog_checked", set())
        return [f for f in folders if f.name in sel] if sel else folders

    def _recog_log_preview(self, label: str, folders: list):
        """Xem trước: ghi danh sách tập SẼ chạy ra nhật ký + dòng trạng thái."""
        names = [f.name for f in folders]
        logging.info(f"▶ {label}: sẽ chạy {len(folders)} tập → {', '.join(names)}")
        # Dòng trạng thái chỉ ghi vài tên đầu (danh sách đủ đã có trong nhật ký).
        self.pipe_status.set(f"▶ {label}: {len(folders)} tập → {self._ep_list_short(names)}")

    def _recog_schedule_table_refresh(self):
        """Refresh bảng an toàn từ thread nền (đẩy về main thread)."""
        try:
            self.after(0, self._recog_refresh_table)
        except Exception:
            pass

    def _recog_tab_sources(self) -> list:
        """Link/file hợp lệ từ ô nhập của tab Nhận diện (mỗi dòng 1 mục), giữ thứ tự."""
        out = []
        for line in self.recog_txt_sources.get("1.0", "end").splitlines():
            s = line.strip().strip('"').strip("'")
            if s and (os.path.isfile(s) or s.lower().startswith(("http://", "https://"))):
                out.append(s)
        return out

    def _recog_pick_file(self):
        paths = filedialog.askopenfilenames(
            title="Chọn file audio/video tiếng Trung (có thể chọn nhiều)",
            filetypes=[("Audio/Video", "*.mp3 *.wav *.m4a *.aac *.flac *.ogg *.opus *.wma "
                                       "*.mp4 *.mkv *.mov *.avi *.webm *.flv"),
                       ("Tất cả", "*.*")])
        if paths:
            cur = self.recog_txt_sources.get("1.0", "end").strip()
            self.recog_txt_sources.insert("end", ("\n" if cur else "") + "\n".join(paths))

    def _recog_tab_run(self):
        """Nút tab Nhận diện: lấy link từ ô RIÊNG của tab rồi chạy luồng chỉ-nhận-diện."""
        self._pipe_recognize_only(self._recog_tab_sources())

    # ── Gộp bước ②+③: DỊCH Gemini + tạo input.txt cho MỌI tập đã nhận diện ────────
    def _recognized_folders(self) -> list:
        """Thư mục tập ĐÃ nhận diện (có tiengTrung.docx/*_zh.docx), theo số tập."""
        return [p for p in episode_dirs() if find_zh_docx(p) is not None]

    def _recog_translate_all(self):
        """Nút '🌐 Dịch + tạo input.txt': gộp bước ② (dịch Gemini) + ③ (input.txt) của
        tab Tạo kịch bản, chạy LẦN LƯỢT cho hết các tập ĐÃ NHẬN DIỆN (do nút 'Nhận diện
        các link rồi ngưng' tạo ra). KHÔNG nhận diện lại, KHÔNG SEO/thumbnail/video."""
        if self._pipe_busy:
            return
        folders = self._recog_apply_selection(self._recognized_folders())
        if not folders:
            self.pipe_status.set("⚠️ Không có tập đã nhận diện phù hợp (kiểm tra tick / tiengTrung.docx).")
            logging.warning("Không có thư mục tập đã nhận diện để dịch.")
            return
        self._recog_log_preview("② Dịch + input.txt", folders)   # xem trước danh sách sẽ chạy
        self._save_pipe_settings()
        self._pipe_set_busy(True)
        self._batch_pause_evt.clear()
        self._batch_stop_evt.clear()
        self._batch_running = True
        self._set_batch_pause_btns(state="normal", text="⏸  Tạm dừng")
        self._set_batch_stop_btns(state="normal")
        self.pipe_progress.set(0)
        self.pipe_link_status.set(f"⏳ Dịch + input.txt cho {len(folders)} tập...")
        threading.Thread(target=self._recog_translate_all_worker,
                         args=(folders,), daemon=True).start()

    def _recog_translate_all_worker(self, folders, chained=False):
        """Với mỗi tập đã nhận diện: dịch Gemini (bỏ qua nếu đủ) rồi tạo input.txt.
        Dùng chung Firefox + hỗ trợ Tạm dừng/Dừng như batch. KHÔNG SEO/thumbnail/video.

        chained=True → đang chạy trong chuỗi ⚡ (nút 'Chạy TIẾP các bước còn thiếu'):
        KHÔNG nhả cờ busy/nút điều khiển ở finally — chuỗi ⚡ tự nhả khi xong hết."""
        driver = None
        total = len(folders)
        ok_count = 0

        # (Gemini treo → Firefox mở lại → driver mới) nay do _dich_gemini_cho_tap trả về
        # và được gán lại ngay tại chỗ gọi, khỏi cần closure _on_driver riêng.

        # map số tập → nguồn gốc (để cập nhật manifest ĐÚNG dòng, không tạo dòng trùng).
        m = load_manifest()
        ep2src = {str(v.get("episode", "")).zfill(2): k for k, v in m.items()
                  if str(v.get("episode", "")).isdecimal()}
        try:
            import dich_gemini as g
            prefix = load_prefix()
            logging.info("\n" + "═" * 10 +
                         f" DỊCH + INPUT.TXT cho {total} tập đã nhận diện " + "═" * 10)

            for i, folder in enumerate(folders, 1):
                self._batch_pause_wait()
                if self._batch_stop_evt.is_set():
                    logging.info(f"⏹ ĐÃ DỪNG — xong {ok_count}/{total} tập.")
                    self.pipe_link_status.set(f"⏹ Đã dừng — xong {ok_count}/{total} tập.")
                    break

                episode = episode_of(folder.name)   # SỐ TẬP (tên thư mục có thể kèm tên nguồn)
                src = ep2src.get(episode)
                gemini_docx = folder / "gemini_result.docx"
                input_txt = folder / "input.txt"
                self.pipe_link_status.set(f"🌐 Dịch: Tập {episode} ({i}/{total})")
                self.pipe_status.set(f"🌐 Tập {episode} ({i}/{total})")

                try:
                    zh = find_zh_docx(folder)
                    chunks = read_zh_docx_chunks(zh) if zh else []
                    if not chunks:
                        logging.warning(f"⚠️ Tập {episode}: không đọc được đoạn — bỏ qua.")
                        continue

                    # ── ②) DỊCH GEMINI — đủ thì bỏ qua; thiếu thì dịch tiếp (giống batch) ──
                    driver, translated_now, translation_ok = self._dich_gemini_cho_tap(
                        gemini_docx, chunks, prefix, episode, driver,
                        on_status=lambda i2, t2: self.pipe_status.set(
                            f"🌐 Tập {episode} • Gemini {i2 + 1}/{t2}"))
                    if src:
                        self._manifest_update(src, episode, folder)

                    if not translation_ok:
                        self.pipe_status.set(f"⛔ Tập {episode}: dịch chưa xong — bỏ qua input.txt.")
                        continue

                    # ── ③) input.txt — tạo lại nếu vừa dịch, hoặc chưa có (giống batch) ──
                    if not translated_now and input_txt.exists() and input_txt.stat().st_size > 0:
                        logging.info(f"♻ Tập {episode}: đã có input.txt — bỏ qua.")
                    elif self._batch_prepare_input(gemini_docx, input_txt):
                        logging.info(f"💾 Đã tạo: {input_txt}")
                    ok_count += 1
                except Exception as e:
                    import traceback
                    logging.error(f"❌ Lỗi tập {episode}: {e}")
                    logging.error(traceback.format_exc())
                    continue

            self.pipe_progress.set(100)
            self.pipe_link_status.set(f"✅ Xong dịch + input.txt: {ok_count}/{total} tập.")
            self.pipe_status.set(f"✅ Đã dịch + tạo input.txt {ok_count}/{total} tập.")
            logging.info(f"🎉 XONG: {ok_count}/{total} tập có bản dịch + input.txt.")
        except Exception as e:
            logging.error(f"Lỗi dịch hàng loạt: {e}")
            self.pipe_status.set(f"Lỗi: {e}")
        finally:
            if driver is not None:
                try:
                    driver.quit()
                except Exception:
                    pass
            if not chained:
                self._pipe_set_busy(False)
                self._batch_controls_reset()
            self._recog_schedule_table_refresh()   # cập nhật bảng trạng thái tab Nhận diện

    # ── Bước ③: GỬI SEO (Gemini) cho MỌI tập đã dịch ─────────────────────────────
    def _folders_with_gemini(self) -> list:
        """Thư mục tập (kịch_bản/NN) đã có gemini_result.docx KHÔNG rỗng — SEO lấy
        đoạn đầu của bản dịch này làm nguồn, theo số tập."""
        out = []
        for p in episode_dirs():
            gem = p / "gemini_result.docx"
            try:
                if gem.is_file() and gem.stat().st_size > 0:
                    out.append(p)
            except OSError:
                pass
        return out

    def _recog_seo_all(self):
        """Nút '🔎 Gửi SEO (Gemini)': chạy bước SEO của tab Home LẦN LƯỢT cho các tập
        ĐÃ DỊCH (có gemini_result.docx) → seoYoutube.docx + youtube_seo.txt. Tập nào đã
        có SEO hợp lệ thì bỏ qua. KHÔNG nhận diện/dịch lại, KHÔNG thumbnail/video
        (thumbnail tách sang nút ④)."""
        if self._pipe_busy:
            return
        folders = self._recog_apply_selection(self._folders_with_gemini())
        if not folders:
            self.pipe_status.set("⚠️ Không có tập đã dịch phù hợp (kiểm tra tick / gemini_result.docx).")
            logging.warning("Không có thư mục tập nào có gemini_result.docx để gửi SEO.")
            return
        self._recog_log_preview("③ Gửi SEO", folders)   # xem trước danh sách sẽ chạy
        self._save_pipe_settings()
        self._pipe_set_busy(True)
        self._batch_pause_evt.clear()
        self._batch_stop_evt.clear()
        self._batch_running = True
        self._set_batch_pause_btns(state="normal", text="⏸  Tạm dừng")
        self._set_batch_stop_btns(state="normal")
        self.pipe_progress.set(0)
        self.pipe_link_status.set(f"⏳ Gửi SEO cho {len(folders)} tập...")
        threading.Thread(target=self._recog_seo_all_worker,
                         args=(folders,), daemon=True).start()

    def _recog_seo_all_worker(self, folders, chained=False):
        """Với mỗi tập đã dịch: gửi SEO lên Gemini (bỏ qua nếu seoYoutube.docx đã có
        tiêu đề) → lưu youtube_seo.txt. Dùng CHUNG 1 Firefox cho mọi tập + hỗ trợ
        Tạm dừng/Dừng như batch. Thumbnail do nút ④ lo.
        chained=True → chạy trong chuỗi ⚡, không nhả busy ở finally."""
        driver = None
        total = len(folders)
        ok_count = 0
        try:
            import dich_gemini as g
            youtube_dir = str(YOUTUBE_DIR)
            if youtube_dir not in sys.path:
                sys.path.insert(0, youtube_dir)
            import seo_youtube_gemini as seo
            logging.info("\n" + "═" * 10 +
                         f" GỬI SEO cho {total} tập đã dịch " + "═" * 10)

            for i, folder in enumerate(folders, 1):
                self._batch_pause_wait()
                if self._batch_stop_evt.is_set():
                    logging.info(f"⏹ ĐÃ DỪNG — xong {ok_count}/{total} tập.")
                    self.pipe_link_status.set(f"⏹ Đã dừng — xong {ok_count}/{total} tập.")
                    break

                episode = episode_of(folder.name)   # SỐ TẬP (tên thư mục có thể kèm tên nguồn)
                gemini_docx = folder / "gemini_result.docx"
                seo_docx = folder / "seoYoutube.docx"
                self.pipe_link_status.set(f"🔎 SEO: Tập {episode} ({i}/{total})")
                self.pipe_status.set(f"🔎 Tập {episode} ({i}/{total})")

                try:
                    # ── SEO YouTube — bỏ qua nếu seoYoutube.docx đã có tiêu đề ──
                    if self._seo_docx_valid(seo_docx):
                        logging.info(f"♻ Tập {episode}: đã có seoYoutube.docx hợp lệ — bỏ qua SEO.")
                    else:
                        if driver is None:
                            logging.info("🌐 Mở Firefox cho SEO...")
                            driver = g.init_firefox()
                        logging.info(f"🔎 Tập {episode}: tạo SEO YouTube...")
                        # seo.run mở cuộc trò chuyện MỚI cho mỗi tập nhưng vẫn dùng lại
                        # đúng Firefox này; keep_open=True → worker đóng ở cuối.
                        seo.run(str(gemini_docx), str(seo_docx),
                                keep_open=True, log=logging.info, driver=driver)
                        if not self._seo_docx_valid(seo_docx):
                            logging.error(f"⛔ Tập {episode}: không lấy được SEO hợp lệ "
                                          "— bỏ qua, chạy lại để làm SEO mới.")
                            continue
                        logging.info(f"💾 Đã tạo: {seo_docx}")

                    # ── Nội dung 3 nút Copy (tiêu đề/mô tả/thẻ tag) ra .txt — LUÔN
                    # tạo lại (nhẹ) để áp dụng logic mới nhất, như tab Home. ──
                    self._save_youtube_seo_copy(
                        seo_docx, folder / "youtube_seo.txt", episode)

                    # Tiêu đề trùng tập khác = SEO lấy nhầm kết quả cũ → cất bản đó đi
                    # để lần chạy sau làm lại, không để tập này đi tiếp với tiêu đề sai.
                    blocks_seo = self._seo_copy_blocks(seo_docx, episode) or {}
                    trung = self._seo_title_duplicate(
                        folder, episode, blocks_seo.get("title") or "")
                    if trung:
                        logging.error(f"⛔ Tập {episode}: tiêu đề SEO trùng — {trung}. "
                                      "Bỏ qua tập này (chạy lại để làm SEO mới).")
                        try:
                            parked = seo.park_docx(seo_docx)
                            if parked:
                                logging.info(f"📦 Đã cất bản SEO nghi sai → {parked.name}")
                        except Exception as e:
                            logging.warning(f"⚠️ Không cất được {seo_docx.name}: {e}")
                        continue
                    self._manifest_update_if_known(folder, episode)
                    ok_count += 1
                except Exception as e:
                    import traceback
                    logging.error(f"❌ Lỗi SEO tập {episode}: {e}")
                    logging.error(traceback.format_exc())
                    continue

            self.pipe_progress.set(100)
            self.pipe_link_status.set(f"✅ Xong SEO: {ok_count}/{total} tập.")
            self.pipe_status.set(f"✅ Đã gửi SEO {ok_count}/{total} tập — bấm ④ để tạo thumbnail.")
            logging.info(f"🎉 XONG: {ok_count}/{total} tập có SEO.")
        except Exception as e:
            logging.error(f"Lỗi SEO hàng loạt: {e}")
            self.pipe_status.set(f"Lỗi: {e}")
        finally:
            if driver is not None:
                try:
                    driver.quit()
                except Exception:
                    pass
            if not chained:
                self._pipe_set_busy(False)
                self._batch_controls_reset()
            self._recog_schedule_table_refresh()   # cập nhật bảng trạng thái tab Nhận diện

    # ── Bước ④: TẠO THUMBNAIL (ngang + dọc) cho MỌI tập đã có SEO ────────────────
    def _folders_with_seo(self) -> list:
        """Thư mục tập (kịch_bản/NN) đã có seoYoutube.docx CÓ tiêu đề — thumbnail lấy
        tiêu đề từ đó nên chưa SEO thì chưa render được, theo số tập."""
        return [p for p in self._all_episode_folders() if self._recog_seo_ok(p)]

    def _recog_thumb_all(self):
        """Nút '🖼 Tạo thumbnail': render thumbnail ngang + dọc (y hệt tab Home) cho các
        tập ĐÃ CÓ SEO. Tập nào đủ 2 bản thì bỏ qua. Chỉ render ảnh — KHÔNG mở Firefox,
        KHÔNG gửi gì lên Gemini, nên chạy lại thoải mái."""
        if self._pipe_busy:
            return
        folders = self._recog_apply_selection(self._folders_with_seo())
        if not folders:
            self.pipe_status.set("⚠️ Không có tập nào có SEO hợp lệ (kiểm tra tick / chạy ③ trước).")
            logging.warning("Không có thư mục tập nào có seoYoutube.docx hợp lệ để tạo thumbnail.")
            return
        self._recog_log_preview("④ Tạo thumbnail", folders)   # xem trước danh sách sẽ chạy
        self._save_pipe_settings()
        self._pipe_set_busy(True)
        self._batch_pause_evt.clear()
        self._batch_stop_evt.clear()
        self._batch_running = True
        self._set_batch_pause_btns(state="normal", text="⏸  Tạm dừng")
        self._set_batch_stop_btns(state="normal")
        self.pipe_progress.set(0)
        self.pipe_link_status.set(f"⏳ Tạo thumbnail cho {len(folders)} tập...")
        threading.Thread(target=self._recog_thumb_all_worker,
                         args=(folders,), daemon=True).start()

    def _recog_thumb_all_worker(self, folders, chained=False):
        """Với mỗi tập đã có SEO: render thumbnail ngang + dọc (bỏ qua nếu đã đủ 2 bản)
        → cập nhật số tập + manifest. Hỗ trợ Tạm dừng/Dừng.
        chained=True → chạy trong chuỗi ⚡, không nhả busy ở finally."""
        total = len(folders)
        ok_count = 0
        try:
            logging.info("\n" + "═" * 10 +
                         f" TẠO THUMBNAIL cho {total} tập đã có SEO " + "═" * 10)
            for i, folder in enumerate(folders, 1):
                self._batch_pause_wait()
                if self._batch_stop_evt.is_set():
                    logging.info(f"⏹ ĐÃ DỪNG — xong {ok_count}/{total} tập.")
                    self.pipe_link_status.set(f"⏹ Đã dừng — xong {ok_count}/{total} tập.")
                    break

                episode = episode_of(folder.name)   # SỐ TẬP (tên thư mục có thể kèm tên nguồn)
                self.pipe_link_status.set(f"🖼 Thumbnail: Tập {episode} ({i}/{total})")
                self.pipe_status.set(f"🖼 Tập {episode} ({i}/{total})")
                try:
                    if self._thumb_count(folder, episode) == 2:
                        logging.info(f"♻ Tập {episode}: đã có thumbnail ngang & dọc — bỏ qua.")
                    elif not self._make_thumbnail_for_folder(folder, episode):
                        # Thiếu tiêu đề SEO / ảnh mèo / tiêu đề bất thường — đã ghi log
                        # lý do trong _make_thumbnail_for_folder.
                        logging.warning(f"⚠️ Tập {episode}: chưa tạo được thumbnail.")
                        continue
                    # Xong thumbnail → cập nhật SỐ TẬP (không lùi) + ghi tiến độ manifest.
                    save_episode_number(max(load_episode_number(), int(episode)))
                    self._manifest_update_if_known(folder, episode)
                    ok_count += 1
                except Exception as e:
                    import traceback
                    logging.error(f"❌ Lỗi thumbnail tập {episode}: {e}")
                    logging.error(traceback.format_exc())
                    continue

            self.pipe_progress.set(100)
            self.pipe_link_status.set(f"✅ Xong thumbnail: {ok_count}/{total} tập.")
            self.pipe_status.set(f"✅ Đã tạo thumbnail {ok_count}/{total} tập.")
            logging.info(f"🎉 XONG: {ok_count}/{total} tập có thumbnail.")
        except Exception as e:
            logging.error(f"Lỗi thumbnail hàng loạt: {e}")
            self.pipe_status.set(f"Lỗi: {e}")
        finally:
            if not chained:
                self._pipe_set_busy(False)
                self._batch_controls_reset()
            self._recog_schedule_table_refresh()   # cập nhật bảng trạng thái tab Nhận diện

    # ── Bước ⑤: TẠO GIỌNG (clone) + VIDEO cho MỌI tập đã có input.txt ─────────────
    def _folders_with_input(self) -> list:
        """Thư mục tập (kịch_bản/NN) đã có input.txt KHÔNG rỗng, theo số tập."""
        out = []
        for p in episode_dirs():
            inp = p / "input.txt"
            try:
                if inp.is_file() and inp.stat().st_size > 0:
                    out.append(p)
            except OSError:
                pass
        return out

    def _recog_make_video_all(self):
        """Nút '🎬 Tạo giọng + video': với MỌI tập đã có input.txt, tạo giọng clone +
        video Y HỆT tab Home (cùng cài đặt giọng/cắt/video), lưu output.wav + video vào
        TỪNG thư mục tập. KHÔNG nhận diện/dịch lại."""
        if self._pipe_busy:
            return
        folders = self._recog_apply_selection(self._folders_with_input())
        if not folders:
            self.pipe_status.set("⚠️ Không có tập có input.txt phù hợp (kiểm tra tick / input.txt).")
            logging.warning("Không có thư mục tập nào có input.txt để tạo giọng/video.")
            return
        # Đọc cài đặt TTS (clone/cắt/video) trên MAIN THREAD — y như chế độ nhiều link.
        tts_settings = self._collect_tts_settings()
        if tts_settings is None:
            return   # cấu hình sai (đã hiện cảnh báo)
        self._recog_log_preview("③ Tạo giọng + video", folders)   # xem trước danh sách sẽ chạy
        self._save_pipe_settings()
        self._pipe_set_busy(True)
        self._batch_pause_evt.clear()
        self._batch_stop_evt.clear()
        self._batch_running = True
        self._set_batch_pause_btns(state="normal", text="⏸  Tạm dừng")
        self._set_batch_stop_btns(state="normal")
        self.pipe_progress.set(0)
        self.pipe_link_status.set(f"⏳ Tạo giọng + video cho {len(folders)} tập...")
        threading.Thread(target=self._recog_make_video_all_worker,
                         args=(folders, tts_settings), daemon=True).start()

    def _recog_make_video_all_worker(self, folders, tts_settings, chained=False):
        """Với mỗi tập: đọc input.txt → tạo giọng OmniVoice + dựng video vào thư mục tập
        (tái dùng _batch_run_tts như chế độ nhiều link). Hỗ trợ Tạm dừng/Dừng.
        chained=True → chạy trong chuỗi ⚡, không nhả busy ở finally."""
        total = len(folders)
        ok_count = 0
        try:
            logging.info("\n" + "═" * 10 +
                         f" TẠO GIỌNG + VIDEO cho {total} tập (clone như Home) " + "═" * 10)
            for i, folder in enumerate(folders, 1):
                self._batch_pause_wait()
                if self._batch_stop_evt.is_set():
                    logging.info(f"⏹ ĐÃ DỪNG — xong {ok_count}/{total} tập.")
                    self.pipe_link_status.set(f"⏹ Đã dừng — xong {ok_count}/{total} tập.")
                    break
                episode = episode_of(folder.name)   # SỐ TẬP (tên thư mục có thể kèm tên nguồn)
                self.pipe_link_status.set(f"🎬 Tạo giọng + video: Tập {episode} ({i}/{total})")
                self.pipe_status.set(f"🎬 Tập {episode} ({i}/{total})")
                try:
                    if self._batch_run_tts(folder, tts_settings, episode):
                        ok_count += 1
                        self._manifest_update_if_known(folder, episode)
                except Exception as e:
                    import traceback
                    logging.error(f"❌ Lỗi tạo giọng/video tập {episode}: {e}")
                    logging.error(traceback.format_exc())
                    continue

            self.pipe_progress.set(100)
            self.pipe_link_status.set(f"✅ Xong giọng + video: {ok_count}/{total} tập.")
            self.pipe_status.set(f"✅ Đã tạo giọng + video {ok_count}/{total} tập.")
            logging.info(f"🎉 XONG: {ok_count}/{total} tập có giọng + video.")
        except Exception as e:
            logging.error(f"Lỗi tạo giọng/video hàng loạt: {e}")
            self.pipe_status.set(f"Lỗi: {e}")
        finally:
            if not chained:
                self._pipe_set_busy(False)
                self._batch_controls_reset()
            self._recog_schedule_table_refresh()   # cập nhật bảng trạng thái tab Nhận diện

    # ── ⑥ ĐĂNG YOUTUBE hàng loạt cho các tập đã dựng xong video ─────────────────
    def _folders_ready_to_upload(self) -> list:
        """Tập đã có YOUTUBE.mp4 + SEO hợp lệ và CHƯA đăng.

        Loại cả tập mà KÊNH đã có video cùng số (đăng tay từ trước) — xem
        _episodes_on_channel. Cần cache kênh còn mới, nên gọi sau _upload_check_ready.
        """
        _ensure_youtube_path()
        import dang_tap_youtube as up
        on_channel = self._episodes_on_channel()

        def _new(p):
            ep = episode_of(p.name)
            return ep is not None and int(ep) not in on_channel

        return [p for p in episode_dirs()
                if up.find_video(p) is not None
                and up.already_uploaded(p) is None
                and _new(p)
                and self._recog_seo_ok(p)]

    def _recog_upload_all(self):
        """Nút '⑥ Đăng YouTube': đăng các tập đã dựng xong mà chưa đăng, mỗi tập hẹn
        vào một khung giờ trống (08:00 / 18:00). Tập đã đăng rồi thì bỏ qua."""
        if self._pipe_busy:
            return
        # Kiểm tra đăng nhập TRƯỚC: vừa chặn sớm, vừa nạp lại cache kênh để danh sách
        # dưới đây loại đúng những tập đã có trên kênh.
        if not self._upload_check_ready():
            return
        folders = self._recog_apply_selection(self._folders_ready_to_upload())
        if not folders:
            self.pipe_status.set("⚠️ Không có tập nào cần đăng (thiếu video/SEO, "
                                 "hoặc kênh đã có hết).")
            logging.warning("Không có tập nào đủ điều kiện đăng YouTube.")
            return
        names = ", ".join(f.name for f in folders)
        if not messagebox.askyesno(
                "Đăng YouTube",
                f"Sẽ đăng {len(folders)} tập lên kênh đang đăng nhập:\n\n{names}\n\n"
                f"Mỗi tập hẹn vào một khung giờ trống ({upload_slots_text()}), để chế "
                "độ riêng tư tới giờ mới tự công khai.\n\nTiếp tục?"):
            return
        self._recog_log_preview("⑥ Đăng YouTube", folders)
        self._pipe_set_busy(True)
        self.pipe_link_status.set(f"⏳ Đăng {len(folders)} tập lên YouTube...")
        threading.Thread(target=self._recog_upload_all_worker,
                         args=(folders,), daemon=True).start()

    def _recog_upload_all_worker(self, folders, chained=False):
        """Xếp hết vào hàng đợi đăng rồi chờ chạy xong (luồng đăng lo phần tải lên).
        chained=True → chạy trong chuỗi ⚡, không nhả busy ở finally."""
        self._upload_done = 0
        try:
            logging.info("\n" + "═" * 10 +
                         f" ĐĂNG YOUTUBE cho {len(folders)} tập " + "═" * 10)
            for folder in folders:
                self._upload_enqueue(folder, episode_of(folder.name))
            self._upload_wait_drain()
            self.pipe_link_status.set(
                f"✅ Đã đăng {self._upload_done}/{len(folders)} tập.")
            self.pipe_status.set(f"✅ Đăng xong {self._upload_done}/{len(folders)} tập.")
        except Exception as e:
            logging.error(f"Lỗi đăng YouTube hàng loạt: {e}")
            self.pipe_status.set(f"Lỗi đăng YouTube: {e}")
        finally:
            if not chained:
                self._pipe_set_busy(False)
                self._batch_controls_reset()
            self._recog_schedule_table_refresh()

    # ── ⚡ CHẠY TIẾP các bước còn thiếu (②→⑥) — 1 nút chạy trọn quy trình ─────────
    # Tiêu chí "còn thiếu" của từng bước Y HỆT cột "cần ..." trong bảng trạng thái
    # (_recog_refresh_table): dòng tô vàng cần gì thì nút ⚡ chạy đúng cái đó.
    def _recog_need_translate(self) -> list:
        """Tập cần ②: đã nhận diện (có *_zh.docx) mà CHƯA có input.txt."""
        out = []
        for p in self._recognized_folders():
            inp = p / "input.txt"
            try:
                if not (inp.is_file() and inp.stat().st_size > 0):
                    out.append(p)
            except OSError:
                out.append(p)
        return out

    def _recog_need_seo(self) -> list:
        """Tập cần ③: đã dịch (gemini_result.docx) mà SEO chưa có tiêu đề thật."""
        return [p for p in self._folders_with_gemini() if not self._recog_seo_ok(p)]

    def _recog_need_thumb(self) -> list:
        """Tập cần ④: đã có SEO mà chưa đủ 2 bản thumbnail (ngang + dọc)."""
        return [p for p in self._folders_with_seo()
                if self._thumb_count(p, episode_of(p.name)) < 2]

    def _recog_need_video(self) -> list:
        """Tập cần ⑤: đã có input.txt mà chưa có video nào (ngang/dọc/facebook)."""
        return [p for p in self._folders_with_input()
                if not any((p / n).is_file()
                           for n in ("YOUTUBE.mp4", "facebook.mp4", "tiktok.mp4"))]

    def _recog_run_missing(self):
        """Nút '⚡ Chạy TIẾP các bước còn thiếu': nhìn từng tập thiếu bước nào thì chạy
        đúng bước đó, LẦN LƯỢT ②→⑤ (hỏi trước có kèm ⑥ đăng YouTube không). Danh sách
        tập của mỗi bước TÍNH LẠI sau khi bước trước xong — tập vừa được ② dịch xong sẽ
        được ③④⑤ làm nốt ngay trong cùng lượt bấm. Tập đã đủ thì bước đó tự bỏ qua."""
        if self._pipe_busy:
            return
        # Chụp lại danh sách tick LÚC BẤM: đổi tick giữa chừng không ảnh hưởng lượt đang chạy.
        sel = set(getattr(self, "_recog_checked", set()))

        def _pick(folders):
            return [f for f in folders if f.name in sel] if sel else list(folders)

        n2 = len(_pick(self._recog_need_translate()))
        n3 = len(_pick(self._recog_need_seo()))
        n4 = len(_pick(self._recog_need_thumb()))
        n5 = len(_pick(self._recog_need_video()))
        scope = f"{len(sel)} tập đang tick" if sel else "TẤT CẢ tập trong kịch_bản/"
        ans = messagebox.askyesnocancel(
            "⚡ Chạy tiếp các bước còn thiếu",
            f"Phạm vi: {scope}.\n"
            f"Đang thiếu:  ② dịch+input: {n2}   •   ③ SEO: {n3}   •   "
            f"④ thumbnail: {n4}   •   ⑤ giọng+video: {n5}\n"
            "(tập vừa dịch xong ở ② sẽ được làm tiếp ③④⑤ trong cùng lượt này)\n\n"
            f"Có kèm bước ⑥ ĐĂNG YOUTUBE khi xong không?\n"
            f"• Yes  = chạy ②→⑤ rồi TỰ ĐĂNG các tập đủ điều kiện (hẹn {upload_slots_text()}, "
            "KHÔNG hỏi lại từng tập)\n"
            "• No   = chỉ chạy ②→⑤, không đăng\n"
            "• Cancel = không chạy gì cả")
        if ans is None:
            return
        include_upload = bool(ans)
        # Kiểm tra đăng nhập YouTube NGAY (main thread) — token hỏng thì hỏi lúc còn
        # ngồi đây, đừng để nửa đêm chuỗi ⚡ đứng im chờ đăng nhập.
        if include_upload and not self._upload_check_ready():
            self.pipe_status.set("⛔ Chưa sẵn sàng đăng YouTube — bấm ⚡ lại, hoặc chọn No để bỏ ⑥.")
            return
        # Cài đặt giọng/video đọc từ tk.Var → phải lấy trên MAIN THREAD ngay bây giờ.
        # ② có thể tạo input.txt mới nên dù n5=0 vẫn có thể tới ⑤ → cứ đòi hợp lệ khi
        # có khả năng chạy ⑤; chỉ ③/④ thì cho qua (⑤ sẽ bị bỏ nếu thiếu cài đặt).
        tts_settings = self._collect_tts_settings() if (n2 or n5) else None
        if (n2 or n5) and tts_settings is None:
            return   # cấu hình sai (đã hiện cảnh báo trong _collect_tts_settings)

        self._save_pipe_settings()
        self._pipe_set_busy(True)
        self._batch_pause_evt.clear()
        self._batch_stop_evt.clear()
        self._batch_running = True
        self._set_batch_pause_btns(state="normal", text="⏸  Tạm dừng")
        self._set_batch_stop_btns(state="normal")
        self.pipe_progress.set(0)
        self.pipe_link_status.set("⚡ Bắt đầu chạy tiếp các bước còn thiếu...")
        threading.Thread(target=self._recog_run_missing_worker,
                         args=(sel, tts_settings, include_upload), daemon=True).start()

    def _recog_run_missing_worker(self, sel, tts_settings, include_upload):
        """Chạy LẦN LƯỢT ②→⑤(→⑥) bằng chính các worker của từng nút (chained=True →
        chúng không nhả busy; chuỗi này nhả 1 lần ở cuối). Trước mỗi bước tính lại
        danh sách tập còn thiếu; Tạm dừng/Dừng tác dụng ở ranh giới tập VÀ bước."""
        def _pick(folders):
            return [f for f in folders if f.name in sel] if sel else list(folders)

        ran = []
        try:
            logging.info("\n" + "═" * 10 + " ⚡ CHẠY TIẾP CÁC BƯỚC CÒN THIẾU " + "═" * 10)
            steps = [
                ("② Dịch + input.txt", self._recog_need_translate,
                 lambda fs: self._recog_translate_all_worker(fs, chained=True)),
                ("③ Gửi SEO", self._recog_need_seo,
                 lambda fs: self._recog_seo_all_worker(fs, chained=True)),
                ("④ Tạo thumbnail", self._recog_need_thumb,
                 lambda fs: self._recog_thumb_all_worker(fs, chained=True)),
                ("⑤ Tạo giọng + video", self._recog_need_video,
                 lambda fs: self._recog_make_video_all_worker(fs, tts_settings, chained=True)),
            ]
            if include_upload:
                # Đã _upload_check_ready ở nút bấm (cache kênh còn mới) + người dùng
                # đã chọn Yes → không hỏi lại từng tập như nút ⑥ rời.
                steps.append(("⑥ Đăng YouTube", self._folders_ready_to_upload,
                              lambda fs: self._recog_upload_all_worker(fs, chained=True)))

            stopped = False
            for label, compute, run in steps:
                self._batch_pause_wait()
                if self._batch_stop_evt.is_set():
                    stopped = True
                    break
                if label.startswith("⑤") and tts_settings is None:
                    # Chỉ xảy ra khi lúc bấm không có gì cần ②/⑤ nhưng giữa chừng lại
                    # phát sinh (hiếm) — thiếu cài đặt giọng thì đành bỏ bước này.
                    logging.warning("⚠️ Bỏ qua ⑤: chưa lấy cài đặt giọng/video ở lúc bấm nút.")
                    continue
                folders = _pick(compute())
                if not folders:
                    logging.info(f"♻ {label}: không tập nào thiếu — bỏ qua.")
                    continue
                self.pipe_progress.set(0)
                self._recog_log_preview(f"⚡ {label}", folders)
                run(folders)   # worker tự lo Tạm dừng/Dừng + log từng tập
                ran.append(f"{label}: {len(folders)} tập")

            done = " · ".join(ran) if ran else "không có bước nào thiếu"
            if stopped:
                logging.info(f"⏹ ĐÃ DỪNG chuỗi ⚡ — đã chạy: {done}.")
                self.pipe_link_status.set("⏹ Đã dừng chuỗi ⚡ theo yêu cầu.")
                self.pipe_status.set(f"⏹ Chuỗi ⚡ dừng — đã chạy: {done}.")
            else:
                logging.info(f"🎉 ⚡ XONG chuỗi — đã chạy: {done}.")
                self.pipe_link_status.set("✅ ⚡ Xong: đã chạy hết các bước còn thiếu.")
                self.pipe_status.set(f"✅ ⚡ Đã chạy: {done}.")
        except Exception as e:
            import traceback
            logging.error(f"Lỗi chuỗi ⚡: {e}")
            logging.error(traceback.format_exc())
            self.pipe_status.set(f"Lỗi chuỗi ⚡: {e}")
        finally:
            self._pipe_set_busy(False)
            self._batch_controls_reset()
            self._recog_schedule_table_refresh()

    def _manifest_update_if_known(self, folder, episode):
        """Cập nhật tiến độ manifest cho tập (tìm nguồn gốc theo số tập; không có → bỏ)."""
        try:
            m = load_manifest()
            src = next((k for k, v in m.items()
                        if str(v.get("episode", "")).zfill(2) == str(episode).zfill(2)), None)
            if src:
                self._manifest_update(src, episode, folder)
        except Exception:
            pass

    def _build_thumbnail_panel(self, parent):
        """Nhúng GUI tạo thumbnail (YOUTUBE/thumbnail_gui.py) vào 1 panel.
        Lỗi (thiếu thư viện/ảnh nguồn) chỉ hiện thông báo, không làm hỏng app."""
        try:
            youtube_dir = str(YOUTUBE_DIR)
            if youtube_dir not in sys.path:
                sys.path.insert(0, youtube_dir)
            import thumbnail_gui as tg
            host = tk.Frame(parent, bg="#F4F6FB")
            host.grid(row=0, column=0, sticky="nsew")
            self._thumb_gui = tg.ThumbnailGUI(
                host, embed=True,
                on_upload_scripts=self._upload_all_scripts_to_drive)
            # Ô 'Số tập' TikTok trên Home bám theo số Thumbnail (+1): đổi số ở tab
            # Thumbnail → ô tự cập nhật. (Panel này build TRƯỚC ô nên _sync bỏ qua an
            # toàn lúc đầu; ô lấy giá trị khởi tạo từ _default_tiktok_episode() khi dựng.)
            try:
                self._thumb_gui.number_var.trace_add("write", self._sync_tiktok_episode)
                self._sync_tiktok_episode()
            except Exception:
                pass
        except Exception as e:
            logging.error(f"Không tải được tab Thumbnail: {e}")
            ttk.Label(parent, text=f"Không mở được Thumbnail Studio:\n{e}",
                      style="Sub.TLabel").grid(row=0, column=0, padx=24, pady=24)

    @staticmethod
    def _drive_script_name(number: str) -> str:
        number = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", str(number or "").strip())
        return f"{number or 'input'}.txt"

    @staticmethod
    def _drive_log(msg, level="info"):
        if level == "err":
            logging.error(msg)
        elif level == "warn":
            logging.warning(msg)
        else:
            logging.info(msg)

    def _set_thumbnail_upload_status(self, text: str, ok: bool = True):
        def _apply():
            gui = getattr(self, "_thumb_gui", None)
            if not gui:
                return
            try:
                gui.status_var.set(text)
                gui.status_label.configure(fg="#15803D" if ok else "#DC2626")
            except Exception:
                pass
        self.after(0, _apply)

    def _collect_scripts_to_upload(self) -> list[tuple[Path, str]]:
        """Gom (đường_dẫn_input.txt, số_tập) của MỌI kịch bản đang có để tải lên Drive.

        • Nhiều link (batch): mỗi thư mục kịch_bản/<số tập>/input.txt → số tập = tên thư mục.
        • Một link: kịch_bản/input.txt → số tập lấy từ tab Thumbnail (>0 mới tính).
        Bỏ file rỗng; nếu trùng số tập thì ưu tiên thư mục batch (đã thêm trước)."""
        jobs: list[tuple[Path, str]] = []
        seen: set[str] = set()

        def _add(path: Path, episode: str) -> None:
            episode = str(episode).strip()
            if not episode or episode in seen:
                return
            try:
                if path.is_file() and path.stat().st_size > 0:
                    jobs.append((path, episode))
                    seen.add(episode)
            except OSError:
                pass

        # 1) Nhiều link: thư mục con tập ("01" hoặc "01 - tên nguồn").
        for p in episode_dirs():
            _add(p / "input.txt", episode_of(p.name))

        # 2) Một link: kịch_bản/input.txt theo số tập ở tab Thumbnail.
        ep = self._current_episode_number()
        if ep > 0:
            _add(SCRIPT_DIR / "input.txt", f"{ep:02d}")
        return jobs

    def _set_upload_button_state(self, state: str) -> None:
        """Bật/tắt nút 'Tải kịch bản lên Drive' ở tab Thumbnail (an toàn đa luồng)."""
        def _apply():
            gui = getattr(self, "_thumb_gui", None)
            btn = getattr(gui, "upload_drive_button", None) if gui else None
            if btn is not None:
                try:
                    btn.configure(state=state)
                except Exception:
                    pass
        self.after(0, _apply)

    def _upload_all_scripts_to_drive(self):
        """Nút 'Tải kịch bản lên Drive' (chủ động): tải MỌI kịch bản (nhiều link +
        1 link) lên Drive trong 1 luồng nền. Tự bỏ qua file đã có trên Drive."""
        threading.Thread(target=self._upload_all_scripts_worker, daemon=True).start()

    def _upload_all_scripts_worker(self):
        self._set_upload_button_state("disabled")
        try:
            jobs = self._collect_scripts_to_upload()
            if not jobs:
                self._set_thumbnail_upload_status(
                    "Không có kịch bản nào để tải (kiểm tra số tập / input.txt).", ok=False)
                return
            logging.info(f"⬆ Bắt đầu tải {len(jobs)} kịch bản lên Drive: "
                         f"{', '.join(ep for _, ep in jobs)}")
            done = 0
            for input_path, episode in jobs:
                if self._upload_input_script_to_drive(input_path, episode):
                    done += 1
            self._set_thumbnail_upload_status(
                f"Đã tải {done}/{len(jobs)} kịch bản lên Drive.", ok=(done == len(jobs)))
            logging.info(f"✅ Tải kịch bản lên Drive xong: {done}/{len(jobs)} thành công.")
        except Exception as e:
            logging.error(f"Lỗi tải kịch bản lên Drive: {e}")
            self._set_thumbnail_upload_status("Lỗi tải kịch bản lên Drive.", ok=False)
        finally:
            self._set_upload_button_state("normal")

    def _upload_input_script_to_drive(self, input_path: Path, episode: str) -> bool:
        """Tải 1 file input.txt lên Drive/kịch bản với tên <số tập>.txt. Trả về True
        nếu tải thành công HOẶC Drive đã có (bỏ qua); False nếu lỗi/thiếu file."""
        drive_name = self._drive_script_name(episode)
        try:
            if not input_path.exists():
                raise FileNotFoundError(f"Không tìm thấy {input_path}")
            if input_path.stat().st_size == 0:
                raise RuntimeError(f"{input_path.name} đang rỗng, chưa tải lên Drive.")

            import taive_drive

            missing = taive_drive._check_deps()
            if missing:
                logging.info(f"Thiếu thư viện Google API ({missing}). Đang cài...")
                taive_drive.install_deps(self._drive_log)
                if taive_drive._check_deps():
                    raise RuntimeError("Không cài được thư viện Google API.")

            self._set_thumbnail_upload_status(f"Đang kiểm tra {drive_name} trên Drive...", ok=True)
            creds = taive_drive.get_credentials(self._drive_log)
            existing = taive_drive.find_drive_file(
                drive_name,
                folder_id=DRIVE_SCRIPT_FOLDER_ID,
                log=self._drive_log,
                creds=creds,
            )
            if existing:
                link = existing.get("webViewLink") or existing.get("id", "")
                logging.info(f"↪ Drive đã có {drive_name}, bỏ qua upload: {link}")
                self._set_thumbnail_upload_status(f"Drive đã có {drive_name}, bỏ qua", ok=True)
                return True

            self._set_thumbnail_upload_status(f"Đang tải {drive_name} lên Drive...", ok=True)
            logging.info(f"⬆ Tải {input_path.name} lên Drive/kịch bản với tên {drive_name}...")
            result = taive_drive.upload_to_drive(
                input_path,
                folder_id=DRIVE_SCRIPT_FOLDER_ID,
                log=self._drive_log,
                creds=creds,
                drive_name=drive_name,
                mimetype="text/plain",
            )
            link = result.get("webViewLink") or result.get("id", "")
            logging.info(f"✅ Đã tải kịch bản lên Drive: {drive_name} → {link}")
            self._set_thumbnail_upload_status(f"Đã tải {drive_name} lên Drive", ok=True)
            return True
        except Exception as e:
            logging.error(f"Lỗi tải input.txt lên Drive: {e}")
            self._set_thumbnail_upload_status("Lỗi tải input.txt lên Drive", ok=False)
            return False

    def _make_log_box(self, parent):
        """Tạo 1 ô nhật ký (ScrolledText) đã set màu/tag và đăng ký vào danh sách
        _log_boxes để _poll_log ghi log đồng thời ra mọi ô (panel video + tab kịch bản)."""
        C = UI
        box = scrolledtext.ScrolledText(
            parent, width=46, height=10, state="disabled",
            font=("Consolas", 9), relief="flat", borderwidth=0,
            background=C["log_bg"], foreground=C["log_info"],
            insertbackground=C["fg"], selectbackground=C["accent_soft"],
            padx=10, pady=8, wrap="word",
        )
        box.grid(row=0, column=0, sticky="nsew")
        box.tag_config("info", foreground=C["log_info"])
        box.tag_config("warn", foreground=C["log_warn"])
        box.tag_config("err", foreground=C["log_err"])
        self._log_boxes.append(box)
        return box

    def _build_log_panel(self, parent, row):
        """Ô nhật ký — đặt DƯỚI dòng trạng thái 'Sẵn sàng' trong panel video."""
        log_frame = ttk.LabelFrame(parent, text="  Nhật ký  ")
        log_frame.grid(row=row, column=0, sticky="nsew", pady=(10, 0))
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        self.log_box = self._make_log_box(log_frame)

    def _build_design_dropdowns(self):
        for w in self.design_attr_frame.winfo_children():
            w.destroy()
        self._design_vars.clear()
        self._design_maps: list[dict] = []

        NONE = "—"

        if self.var_lang.get() == "en":
            self._design_sep = ", "
            # (label_vi, [(hiển_thị_vi, giá_trị_model), ...])
            groups = [
                [
                    ("Giới tính", [(NONE,""), ("Nữ","female"), ("Nam","male")]),
                    ("Tuổi",      [(NONE,""), ("Trẻ em","child"), ("Thiếu niên","teenager"),
                                   ("Thanh niên","young adult"), ("Trung niên","middle-aged"),
                                   ("Cao tuổi","elderly")]),
                    ("Âm điệu",  [(NONE,""), ("Rất thấp","very low pitch"), ("Thấp","low pitch"),
                                   ("Vừa","moderate pitch"), ("Cao","high pitch"),
                                   ("Rất cao","very high pitch")]),
                    ("Đặc biệt", [(NONE,""), ("Thì thầm","whisper")]),
                ],
                [
                    ("Giọng vùng", [(NONE,""), ("Mỹ","american accent"), ("Úc","australian accent"),
                                    ("Anh","british accent"), ("Canada","canadian accent"),
                                    ("Trung Quốc","chinese accent"), ("Ấn Độ","indian accent"),
                                    ("Nhật Bản","japanese accent"), ("Hàn Quốc","korean accent"),
                                    ("Bồ Đào Nha","portuguese accent"), ("Nga","russian accent")]),
                ],
            ]
        else:
            self._design_sep = "，"
            groups = [
                [
                    ("Giới tính", [(NONE,""), ("Nữ","女"), ("Nam","男")]),
                    ("Tuổi",      [(NONE,""), ("Trẻ em","儿童"), ("Thiếu niên","少年"),
                                   ("Thanh niên","青年"), ("Trung niên","中年"), ("Cao tuổi","老年")]),
                    ("Âm điệu",  [(NONE,""), ("Rất thấp","极低音调"), ("Thấp","低音调"),
                                   ("Vừa","中音调"), ("Cao","高音调"), ("Rất cao","极高音调")]),
                    ("Đặc biệt", [(NONE,""), ("Thì thầm","耳语")]),
                ],
                [
                    ("Phương ngữ", [(NONE,""), ("Đông Bắc","东北话"), ("Vân Nam","云南话"),
                                    ("Tứ Xuyên","四川话"), ("Ninh Hạ","宁夏话"), ("Cam Túc","甘肃话"),
                                    ("Quế Lâm","桂林话"), ("Hà Nam","河南话"), ("Tế Nam","济南话"),
                                    ("Thiểm Tây","陕西话"), ("Thạch Gia Trang","石家庄话"),
                                    ("Quý Châu","贵州话"), ("Thanh Đảo","青岛话")]),
                ],
            ]

        for group in groups:
            row_f = ttk.Frame(self.design_attr_frame)
            row_f.pack(anchor="w", pady=3)
            for label, options in group:
                displays = [d for d, _ in options]
                mapping  = {d: v for d, v in options}
                cell = ttk.Frame(row_f)
                cell.pack(side="left", padx=(0, 14))
                ttk.Label(cell, text=label, font=("Segoe UI", 8)).pack(anchor="w")
                var = tk.StringVar(value=NONE)
                self._design_vars.append(var)
                self._design_maps.append(mapping)
                ttk.Combobox(cell, textvariable=var, values=displays,
                             width=17, state="readonly").pack()
                var.trace_add("write", lambda *_: self._update_instruct())

    def _on_lang_change(self):
        self._build_design_dropdowns()
        self._update_instruct()

    def _update_instruct(self):
        parts = []
        for var, mapping in zip(self._design_vars, self._design_maps):
            actual = mapping.get(var.get(), "")
            if actual:
                parts.append(actual)
        self.var_instruct.set(self._design_sep.join(parts))

    def _on_mode_change(self):
        for frm in (self.frm_clone, self.frm_design, self.frm_default):
            frm.pack_forget()
        {"clone":   self.frm_clone,
         "design":  self.frm_design,
         "default": self.frm_default}[self.var_mode.get()].pack(anchor="w")

    def _current_voice(self) -> str:
        """Tên file giọng mẫu thật đang chọn (đã bỏ tiền tố ★)."""
        return strip_star(self.var_ref.get())

    def _reload_voice_combo(self, keep: str | None = None):
        """Dựng lại danh sách giọng: yêu thích (★) lên đầu, còn lại theo a-z.

        keep = tên file thật muốn giữ chọn; mặc định giữ mục đang chọn.
        """
        files = list_voice_files()
        favs = [f for f in files if f in self._favorites]
        rest = [f for f in files if f not in self._favorites]
        ordered = favs + rest
        display = [(STAR + f if f in self._favorites else f) for f in ordered]
        self.cb_ref["values"] = display

        want = keep if keep is not None else self._current_voice()
        if want in ordered:
            self.var_ref.set(STAR + want if want in self._favorites else want)
        elif display:
            self.var_ref.set(display[0])
        else:
            self.var_ref.set("")
        self._update_fav_button()

    def _update_fav_button(self):
        fav = self._current_voice() in self._favorites
        self.btn_fav.config(text="★" if fav else "☆")

    def _toggle_favorite(self):
        name = self._current_voice()
        if not name:
            return
        if name in self._favorites:
            self._favorites.discard(name)
            logging.info(f"☆ Bỏ yêu thích: {name}")
        else:
            self._favorites.add(name)
            logging.info(f"★ Đã thêm yêu thích: {name}")
        save_favorites(self._favorites)
        self._reload_voice_combo(keep=name)

    def _refresh_voices(self):
        self._reload_voice_combo()
        logging.info(f"Tìm thấy {len(list_voice_files())} file giọng trong {VOICE_DIR}")

    def _current_effect(self) -> str:
        """Tên file hiệu ứng thật đang chọn (đã bỏ tiền tố ★); EFFECT_NONE nếu không chọn."""
        return strip_star(self.var_effect.get())

    def _reload_effect_combo(self, keep: str | None = None):
        """Dựng lại danh sách hiệu ứng: yêu thích (★) lên đầu, kèm mục 'Không'.

        keep = tên file thật muốn giữ chọn; mặc định giữ mục đang chọn.
        """
        files = list_effect_files()
        favs = [f for f in files if f in self._effect_favorites]
        rest = [f for f in files if f not in self._effect_favorites]
        ordered = favs + rest
        display = [EFFECT_NONE] + [
            (STAR + f if f in self._effect_favorites else f) for f in ordered]
        self.cb_effect["values"] = display

        want = keep if keep is not None else self._current_effect()
        if want in ordered:
            self.var_effect.set(STAR + want if want in self._effect_favorites else want)
        else:
            self.var_effect.set(EFFECT_NONE)
        self._update_effect_fav_button()

    def _update_effect_fav_button(self):
        cur = self._current_effect()
        is_file = cur in list_effect_files()
        fav = cur in self._effect_favorites
        self.btn_effect_fav.config(text="★" if fav else "☆",
                                   state="normal" if is_file else "disabled")

    def _toggle_effect_favorite(self):
        name = self._current_effect()
        if not name or name == EFFECT_NONE or name not in list_effect_files():
            return
        if name in self._effect_favorites:
            self._effect_favorites.discard(name)
            logging.info(f"☆ Bỏ yêu thích hiệu ứng: {name}")
        else:
            self._effect_favorites.add(name)
            logging.info(f"★ Đã thêm yêu thích hiệu ứng: {name}")
        save_effect_favorites(self._effect_favorites)
        self._reload_effect_combo(keep=name)

    def _refresh_effects(self):
        """Nạp lại danh sách hiệu ứng trong scripts/hieuung/ (giữ mục đang chọn)."""
        self._reload_effect_combo(keep=self._current_effect())
        logging.info(f"Tìm thấy {len(list_effect_files())} hiệu ứng trong {EFFECTS_DIR}")

    def _o_mau_sub(self, hang, nhan: str, var, tieude: str):
        """Dựng một ô chọn màu phụ đề (nhãn + nút mã màu + nút × xoá) → nút màu.

        Dùng chung cho màu CHỮ và màu VIỀN: cùng bảng chọn màu Windows, cùng
        cách hiện '(kiểu)' khi để trống = giữ màu gốc của kiểu.
        """
        from tkinter import ttk as _ttk
        _ttk.Label(hang, text=nhan).pack(side="left", padx=(10, 2))
        btn = _ttk.Button(hang, width=9)
        btn.config(command=lambda: self._chon_mau_sub(var, btn, tieude))
        btn.pack(side="left")
        _ttk.Button(hang, text="×", width=2,
                    command=lambda: (var.set(""),
                                     self._ve_nut_mau_sub(var, btn))).pack(side="left")
        self._ve_nut_mau_sub(var, btn)
        return btn

    def _chon_mau_sub(self, var, btn, tieude: str):
        """Mở bảng chọn màu Windows cho một ô màu của phụ đề (chữ / viền)."""
        from tkinter import colorchooser
        hien = var.get()
        c = colorchooser.askcolor(color=("#" + hien) if hien else None,
                                  title=tieude)
        if c and c[1]:
            var.set(c[1].lstrip("#").upper())
        self._ve_nut_mau_sub(var, btn)

    @staticmethod
    def _ve_nut_mau_sub(var, btn):
        """Nút hiện mã màu đang chọn; để trống thì ghi '(kiểu)'."""
        btn.config(text=("#" + var.get()) if var.get() else "(kiểu)")

    def _save_opt_settings(self):
        """Lưu cài đặt mục 'Cài đặt' hiện tại để lần sau mở lại dùng làm mặc định."""
        try:
            save_opt_settings(dict(
                from_gemini=self.var_from_gemini.get(),
                chunk=int(self.var_chunk.get()),
                make_video=self.var_make_video.get(),
                ngang_speed=self.var_ngang_speed.get(),
                ngang_source=self.var_ngang_source.get(),
                effect=self._current_effect(),
                make_video_doc=self.var_make_video_doc.get(),
                doc_speed=self.var_doc_speed.get(),
                doc_percent=self._parse_percent(self.var_doc_percent, 100),
                doc_from_ngang=self.var_doc_from_ngang.get(),
                doc_from_subfolder=self.var_doc_from_subfolder.get(),
                doc_no_effect=self.var_doc_no_effect.get(),
                make_tiktok=self.var_make_tiktok.get(),
                make_short=self.var_make_short.get(),
                tiktok_speed=self.var_tiktok_speed.get(),
                tiktok_percent=self._parse_percent(self.var_tiktok_percent),
                tiktok_no_effect=self.var_tiktok_no_effect.get(),
                tiktok_caption_pos=int(self.var_tiktok_caption_pos.get()),
                tiktok_music=self.var_tiktok_music.get(),
                tiktok_music_db=int(self.var_tiktok_music_db.get()),
                bring_front=self.var_bring_front.get(),
                make_sub=self.var_make_sub.get(),
                sub_mode=self.var_sub_mode.get(),
                sub_model=self.var_sub_model.get(),
                sub_max_chars=int(self.var_sub_max_chars.get()),
                sub_kieu=self.var_sub_kieu.get(),
                sub_font=self.var_sub_font.get(),
                sub_mau=self.var_sub_mau.get(),
                sub_mau_vien=self.var_sub_mau_vien.get(),
                sub_vitri=self.var_sub_vitri.get().strip(),
                sub_cochu=self._parse_sub_cochu(),
                sub_bengang=self._parse_sub_bengang(),
                sub_dong=self._parse_sub_dong(),
                make_sub_doc=self.var_make_sub_doc.get(),
            ))
        except Exception as e:
            logging.warning(f"Không lưu được cài đặt: {e}")

    @staticmethod
    def _send_to_recycle_bin(paths):
        """Đưa danh sách file/thư mục vào THÙNG RÁC Windows (khôi phục được).

        Phần cài đặt nằm ở module dùng chung xoa_antoan.py — để các script khác
        (doiten_video.py...) cũng xoá an toàn được, không phải chép lại code.
        Trả về (số thành công, số lỗi).
        """
        import xoa_antoan
        return xoa_antoan.send_to_recycle_bin(paths, on_log=logging.warning)

    def _clear_output(self):
        """Đưa output/, các thư mục tập và nội dung file kịch_bản/ vào THÙNG RÁC.

        Khác bản cũ (xóa thẳng, không cứu được): giờ mọi thứ vào Thùng rác Windows nên
        LỠ TAY vẫn khôi phục được, và có 1 bước XÁC NHẬN trước khi làm.
        - output/ (kịch_bản/output): đưa mọi file + thư mục con vào Thùng rác.
        - Thư mục tập (tên toàn chữ số 01, 02, 17...): đưa cả thư mục vào Thùng rác.
        - kịch_bản/: các file trực tiếp (input.txt, *.docx...) được đưa vào Thùng rác
          (giữ nội dung cũ để cứu được) rồi tạo lại bản RỖNG cùng tên cho pipeline chạy
          tiếp không lỗi. Nếu file KHÔNG vào được Thùng rác thì GIỮ NGUYÊN, không ghi đè.
        """
        out_items = list(OUTPUT_DIR.iterdir()) if OUTPUT_DIR.exists() else []
        kb_files = [p for p in SCRIPT_DIR.iterdir() if p.is_file()] if SCRIPT_DIR.exists() else []
        # Thư mục tập = thư mục con của kịch_bản bắt đầu bằng SỐ TẬP
        # ("01", "02", hoặc kèm tên nguồn "01 - 95"). output/ không nằm trong này.
        ep_dirs = episode_dirs()

        if not out_items and not kb_files and not ep_dirs:
            self.status.set("Không có gì để xóa hay làm rỗng.")
            return

        # Xác nhận trước khi xóa (phòng lỡ tay). Chỉ 1 hộp Yes/No nên vẫn nhanh.
        from tkinter import messagebox
        if not messagebox.askyesno(
                "Xác nhận xóa output",
                "Đưa vào THÙNG RÁC (có thể khôi phục lại):\n"
                f"  • {len(out_items)} mục trong output/\n"
                f"  • {len(ep_dirs)} thư mục tập (01, 02...)\n"
                f"  • {len(kb_files)} file trong kịch_bản/ (bản cũ vào Thùng rác, "
                "tạo lại bản rỗng)\n\nTiếp tục?"):
            self.status.set("Đã hủy — không xóa gì.")
            return

        self._stop_preview()   # nhả file đang nghe (nếu có) để xóa được

        # 1) output/ + 2) thư mục tập -> Thùng rác
        out_ok, out_fail = self._send_to_recycle_bin(out_items)
        ep_ok, ep_fail = self._send_to_recycle_bin(ep_dirs)

        # 3) File trong kịch_bản/: đưa bản cũ vào Thùng rác rồi tạo lại bản RỖNG.
        kb_ok, kb_fail = self._send_to_recycle_bin(kb_files)
        emptied = 0
        for p in kb_files:
            if p.exists():
                # Chưa vào được Thùng rác -> KHÔNG ghi đè để khỏi mất nội dung.
                logging.warning(f"Giữ nguyên (chưa vào Thùng rác được): {p.name}")
                continue
            try:
                if p.suffix.lower() == ".docx":
                    from docx import Document
                    Document().save(str(p))          # docx rỗng nhưng hợp lệ
                else:
                    p.write_text("", encoding="utf-8")
                emptied += 1
            except Exception as e:
                logging.warning(f"Không tạo lại được {p.name}: {e}")

        # Đặt lại tên kết quả về output.wav để đánh số lại từ đầu
        self.var_out.set(str(OUTPUT_DIR / "output.wav"))
        self._last_output = None
        self.btn_preview.config(state="disabled")
        total_fail = out_fail + ep_fail + kb_fail
        logging.info(f"Vào Thùng rác: {out_ok} mục output, {ep_ok} thư mục tập, "
                     f"{kb_ok} file kịch_bản; tạo lại {emptied} bản rỗng; lỗi {total_fail}")
        msg = "Đã đưa output + thư mục tập + file kịch_bản vào Thùng rác (khôi phục được)."
        if total_fail:
            msg += f"  ({total_fail} mục lỗi — xem nhật ký)"
        self.status.set(msg)

    def _pick_file(self, var, filetypes):
        path = filedialog.askopenfilename(filetypes=filetypes)
        if path:
            var.set(path)

    def _pick_save(self, var, filetypes):
        path = filedialog.asksaveasfilename(defaultextension=".wav", filetypes=filetypes)
        if path:
            var.set(path)

    def _open_nhan_dien(self):
        """Mở GUI nhận diện giọng nói tiếng Trung trong cửa sổ/tiến trình riêng."""
        import subprocess
        gui = Path(__file__).resolve().parent / "nhandien_gui.py"
        if not gui.exists():
            messagebox.showerror("Thiếu file", f"Không thấy:\n{gui}")
            return
        try:
            subprocess.Popen([sys.executable, str(gui)])
            logging.info("🎙  Đã mở cửa sổ Nhận diện giọng nói.")
        except Exception as e:
            messagebox.showerror("Lỗi mở Nhận diện", str(e))

    # ── BẢNG ĐIỀU KHIỂN WEB (myvoice/web) ─────────────────────────────────────
    @staticmethod
    def _web_port() -> int:
        try:
            return int(os.environ.get("MYVOICE_WEB_PORT", "8765"))
        except ValueError:
            return 8765

    @staticmethod
    def _web_alive(port: int, timeout: float = 0.4) -> bool:
        """Đã có server nghe ở cổng đó chưa (mở 2 lần cùng cổng là lỗi bind)."""
        import socket
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.settimeout(timeout)
            return s.connect_ex(("127.0.0.1", port)) == 0

    def _open_web_panel(self):
        """Bật bảng điều khiển web rồi mở trình duyệt.

        Server chạy trong CỬA SỔ CONSOLE RIÊNG (không phải nền ẩn) để đóng nó là
        tắt hẳn — và để lỡ có lỗi khởi động thì còn nhìn thấy. Đang chạy sẵn thì
        chỉ mở lại trình duyệt, không bật thêm cái thứ hai.

        Bảng web CHỈ nghe trên 127.0.0.1 — máy khác trong mạng không vào được.
        """
        import subprocess
        import time
        import webbrowser

        port = self._web_port()
        web_dir = BASE_DIR / "web"
        server = web_dir / "server.py"
        if not server.exists():
            messagebox.showerror("Thiếu bảng web", f"Không thấy:\n{server}")
            return

        self.btn_web.config(state="disabled")
        try:
            if not self._web_alive(port):
                logging.info("🌐 Đang bật bảng điều khiển web...")
                # python.exe (KHÔNG phải pythonw) để có cửa sổ console tắt được.
                exe = _VENV_PYTHON if os.path.exists(_VENV_PYTHON) else sys.executable
                CREATE_NEW_CONSOLE = 0x00000010
                # Server tự mở trình duyệt khi chạy một mình; ở đây TẮT đi vì
                # hàm này mở lấy (còn lo cả trường hợp server đã chạy sẵn) —
                # không thì ra hai tab.
                env = {**os.environ, "MYVOICE_WEB_NO_OPEN": "1"}
                subprocess.Popen([exe, "-m", "myvoice.web.server"],
                                 cwd=str(BASE_DIR.parent), env=env,
                                 creationflags=CREATE_NEW_CONSOLE)
                for _ in range(40):            # chờ tối đa ~20 giây
                    self.update()
                    if self._web_alive(port, timeout=0.2):
                        break
                    time.sleep(0.5)
                else:
                    messagebox.showerror(
                        "Không bật được bảng web",
                        "Server chưa phản hồi sau 20 giây.\n"
                        "Xem cửa sổ console vừa mở để biết lỗi.")
                    return
            else:
                logging.info("🌐 Bảng điều khiển web đang chạy sẵn.")

            token = ""
            try:
                token = (web_dir / "token.txt").read_text(encoding="utf-8").strip()
            except OSError:
                pass
            query = f"?token={token}" if token else ""
            webbrowser.open(f"http://127.0.0.1:{port}/{query}")
            logging.info(f"🌐 Bảng web: http://127.0.0.1:{port}/{query}")
        except Exception as e:
            messagebox.showerror("Lỗi mở bảng web", str(e))
        finally:
            self.btn_web.config(state="normal")

    def _start_web_log_mirror(self):
        """Chiếu nhật ký hàng đợi ĐĂNG của bảng web sang ô Nhật ký của GUI.

        Bảng web là TIẾN TRÌNH RIÊNG (cửa sổ console riêng), nên bấm “⬆ Đăng ngay”
        bên đó thì cửa sổ này không hay biết gì — nhìn vào chỉ thấy im lìm, không rõ
        có chạy hay không. Luồng nền dưới đây hỏi server 2 giây/lần “có dòng nào mới
        hơn số này chưa”; server chưa bật thì chỉ tốn một lần thử kết nối tới
        127.0.0.1 rồi ngủ tiếp, không kêu ca gì.

        Chỉ chiếu hàng đợi ĐĂNG, không chiếu hàng đợi chính: bên đó là hàng nghìn
        dòng ffmpeg/Whisper, đổ hết sang đây thì ô nhật ký thành bãi rác.
        """
        def loop():
            import json
            import time
            import urllib.request

            since = -1          # -1 = chỉ lấy từ giờ trở đi, bỏ qua mẻ đăng cũ
            port = self._web_port()
            token_file = BASE_DIR / "web" / "token.txt"
            while True:
                if not self._web_alive(port, timeout=0.2):
                    since = -1          # server tắt rồi bật lại → xin lại từ đầu
                    time.sleep(5)
                    continue
                try:
                    token = token_file.read_text(encoding="utf-8").strip()
                    with urllib.request.urlopen(
                            f"http://127.0.0.1:{port}/api/nhatky-dang"
                            f"?since={since}&token={token}", timeout=5) as resp:
                        data = json.loads(resp.read().decode("utf-8"))
                except Exception:
                    time.sleep(5)       # server đang bận/khởi động lại → thử lại sau
                    continue
                since = int(data.get("next", since))
                for line in data.get("lines") or []:
                    level = (logging.ERROR if line.startswith(("❌", "⛔"))
                             else logging.WARNING if line.startswith("⚠") else logging.INFO)
                    logging.log(level, f"🌐 {line}")    # 🌐 = việc do bảng web chạy
                time.sleep(2)

        threading.Thread(target=loop, daemon=True).start()

    def _prepare_input_from_gemini(self) -> bool:
        """Quy trình trước khi tạo audio: lấy nội dung từ gemini_result.docx.

        1) KIỂM TRA câu dẫn nhập/thừa (dich_kiemtra). Có lỗi → báo + DỪNG,
           không tạo audio (để sửa docx trước).
        2) Bỏ cấu trúc 'Kết quả dịch từ Gemini' / 'Đoạn k', ghép thành 1 nội dung.
        3) Ghi vào file 'Văn bản' (input.txt) đang cấu hình.

        CÓ SẴN INPUT: nếu gemini_result.docx THIẾU hoặc RỖNG (0 đoạn) mà file 'Văn bản'
        đã có sẵn nội dung → DÙNG LUÔN file đó (bỏ qua bước lấy từ Gemini) thay vì báo
        lỗi/dừng, để chạy được với input soạn sẵn.

        Trả về True nếu sẵn sàng tạo audio; False thì DỪNG.
        """
        # Fallback dùng input soạn sẵn khi Gemini thiếu/rỗng.
        def _use_existing_input(reason: str) -> bool:
            try:
                txt = Path(self.var_txt.get())
                existing = (clean_text(txt.read_text(encoding="utf-8")).strip()
                            if txt.exists() else "")
            except Exception:
                txt, existing = None, ""
            if existing:
                logging.info(f"ℹ️ {reason} → dùng sẵn nội dung trong '{txt.name}' "
                             f"({len(existing)} ký tự), bỏ qua bước lấy từ Gemini.")
                return True
            return False

        if not GEMINI_DOCX.exists():
            if _use_existing_input("Không thấy gemini_result.docx"):
                return True
            messagebox.showerror(
                "Thiếu file Gemini",
                f"Không thấy:\n{GEMINI_DOCX}\n\nHãy dịch Gemini trước, hoặc bỏ tick "
                "'Lấy nội dung từ Gemini' để dùng input.txt thủ công.")
            return False
        try:
            import dich_kiemtra as cg
        except Exception as e:
            messagebox.showerror("Thiếu dich_kiemtra", str(e))
            return False

        # 1) KIỂM TRA
        logging.info("🔎 Kiểm tra gemini_result.docx trước khi tạo audio...")
        findings = cg.check_docx(GEMINI_DOCX, on_log=logging.info)
        if findings:
            lines = []
            for label, hits in findings:
                phrases = ", ".join(f'"{p}"' for p, _ in hits)
                lines.append(f"• {label}: {phrases}")
            messagebox.showwarning(
                "Gemini còn câu dẫn nhập/thừa",
                "gemini_result.docx còn câu dẫn nhập/thừa:\n\n"
                + "\n".join(lines)
                + "\n\nHãy sửa lại docx rồi thử lại. (CHƯA tạo audio.)")
            self.status.set("⛔ Gemini còn câu thừa — chưa tạo audio.")
            return False

        # 2) BỎ CẤU TRÚC + GHÉP NỘI DUNG
        chunks = cg.read_docx_chunks(GEMINI_DOCX)
        content = "\n".join(t for _, t in chunks).strip()
        if not content:
            if _use_existing_input("gemini_result.docx rỗng (0 đoạn)"):
                return True
            messagebox.showerror("Trống", f"Không lấy được nội dung từ:\n{GEMINI_DOCX}")
            return False

        # 2b) DỌN câu quảng bá kênh rải rác → CHÈN LẠI (mở đầu / thân bài / kết bài)
        content, n_promo, n_add = replace_channel_promo(content)
        logging.info(f"🔁 Quảng bá kênh: xóa {n_promo} câu rải rác trong bản dịch, "
                     f"chèn lại {n_add} câu (mở đầu/thân/kết).")

        # 2c) SỬA từ TIẾNG ANH Gemini sót → tiếng Việt (but→nhưng, If→Nếu,
        #     Twenty→Hai mươi…) + cảnh báo từ lạ còn lại.
        content, _n_en, _sus_en = replace_leaked_english(content)

        # 2d) XỬ LÝ chữ Hán Gemini bỏ sót: câu dài → dịch NGHĨA (MT offline),
        #     chữ ngắn / tên / thành ngữ → phiên âm Hán-Việt.
        try:
            import dich_hanviet as hv
            content, n_mt, n_am = hv.translate_han(content, on_log=logging.info)
            if n_mt or n_am:
                logging.info(f"🈶 Chữ Hán sót: {n_mt} đoạn dịch nghĩa (MT), "
                             f"{n_am} chữ phiên âm Hán-Việt.")
        except Exception as e:
            logging.warning(f"⚠️ Bỏ qua xử lý chữ Hán: {e}")

        # 2e) SỬA từ/cụm cố định cho TTS: giết→giớt, máu→máo, tỳ→tì… (_WORD_FIXES)
        try:
            import dich_chuanbi_input as prep
            content = prep.apply_word_fixes(content)
        except Exception as e:
            logging.warning(f"⚠️ Bỏ qua sửa từ cố định (giết→giớt, tỳ→tì): {e}")

        # 3) GHI VÀO input.txt (đường dẫn ở ô 'Văn bản')
        try:
            out = Path(self.var_txt.get())
            out.parent.mkdir(parents=True, exist_ok=True)
            out.write_text(content, encoding="utf-8")
        except Exception as e:
            messagebox.showerror("Lỗi ghi input.txt", str(e))
            return False
        logging.info(f"✅ Đã lấy {len(content)} ký tự từ Gemini → {out.name} (đã qua kiểm tra)")
        return True

    # ── QUY TRÌNH TẠO KỊCH BẢN (cột trái) ─────────────────────────────────────
    def _build_pipeline_column(self, parent, col):
        """Cột trái: nhận diện giọng nói → dịch Gemini → chuẩn bị input.txt."""
        wrap = ttk.Frame(parent)
        wrap.grid(row=0, column=col, sticky="nsew", padx=(0, 16))
        wrap.columnconfigure(0, weight=1)
        self._pipe_wrap = wrap

        hdr = ttk.Frame(wrap)
        hdr.grid(row=0, column=0, sticky="ew", pady=(0, 14))
        self._pipe_hdr = hdr
        ttk.Label(hdr, text="🛠  Tạo kịch bản", style="Header.TLabel").pack(anchor="w")
        ttk.Label(hdr, text="Audio/Video → 中文 → Gemini → input.txt",
                  style="Sub.TLabel").pack(anchor="w", pady=(2, 0))

        # ① Nhận diện giọng nói
        s1 = ttk.LabelFrame(wrap, text="  ①  Nhận diện giọng nói  ")
        s1.grid(row=1, column=0, sticky="ew", pady=(0, 12))
        s1.columnconfigure(0, weight=1)
        frow = ttk.Frame(s1)
        frow.grid(row=0, column=0, sticky="ew")
        frow.columnconfigure(0, weight=1)
        # Nhiều dòng: mỗi dòng 1 link hoặc 1 file. 1 dòng → xử lý như cũ; ≥2 dòng
        # → mỗi link tạo 1 thư mục kịch bản (01, 02, ...) chạy full pipeline.
        # width nhỏ (mặc định Text là 80 ký tự → phình ngang); sticky="ew" vẫn cho
        # ô giãn vừa theo cột nên không cần để rộng.
        self.pipe_txt_sources = tk.Text(frow, height=2, width=20, wrap="none",
                                        font=("Segoe UI", 10), relief="solid", bd=1)
        self.pipe_txt_sources.grid(row=0, column=0, sticky="ew")
        ttk.Button(frow, text="Chọn…", width=8,
                   command=self._pipe_pick_file).grid(row=0, column=1, padx=(8, 0), sticky="n")
        orow = ttk.Frame(s1)
        orow.grid(row=1, column=0, sticky="w", pady=(8, 0))
        ttk.Label(orow, text="Model:").pack(side="left")
        self.pipe_var_model = tk.StringVar(value=self._pipe_settings["model"])
        ttk.Combobox(orow, textvariable=self.pipe_var_model, width=9, state="readonly",
                     values=["tiny", "base", "small", "medium", "large-v3"]).pack(side="left", padx=(4, 12))
        ttk.Label(orow, text="Tốc độ:").pack(side="left")
        self.pipe_var_speed = tk.StringVar(value=self._pipe_settings["speed"])
        ttk.Combobox(orow, textvariable=self.pipe_var_speed, width=5, state="readonly",
                     values=["0.6", "0.7", "0.8", "0.9", "1.0"]).pack(side="left", padx=(4, 0))
        # Luồng "chỉ nhận diện rồi ngưng" đã CHUYỂN sang tab riêng '🎧 Nhận diện'
        # (nav bên trái). Ở đây chỉ còn nút nhận diện chạy full pipeline như cũ.
        self.btn_recog = ttk.Button(s1, text="🎙  Nhận diện → tiếng Trung",
                                    style="Accent.TButton", command=self._pipe_recognize)
        self.btn_recog.grid(row=2, column=0, sticky="ew", pady=(10, 0))
        ttk.Label(s1, text="(mỗi dòng 1 link/file)",
                  style="Hint.TLabel").grid(row=3, column=0, sticky="w", pady=(4, 0))
        self.var_auto2 = tk.BooleanVar(value=self._pipe_settings["auto2"])
        ttk.Checkbutton(s1, text="⛓  Tự động chạy bước ② sau khi xong",
                        variable=self.var_auto2).grid(row=4, column=0, sticky="w", pady=(6, 0))

        # ② Dịch Gemini
        s2 = ttk.LabelFrame(wrap, text="  ②  Dịch qua Gemini  ")
        s2.grid(row=2, column=0, sticky="ew", pady=(0, 12))
        s2.columnconfigure(0, weight=1)
        self.btn_gemini = ttk.Button(s2, text="🌐  Gửi Gemini (Firefox)",
                                     style="Accent.TButton", command=self._pipe_send_gemini)
        self.btn_gemini.grid(row=0, column=0, sticky="ew")
        ttk.Label(s2, text="(tiengTrung.docx → gemini_result.docx)",
                  style="Hint.TLabel").grid(row=1, column=0, sticky="w", pady=(4, 0))
        self.var_seo = tk.BooleanVar(value=self._pipe_settings["seo"])
        ttk.Checkbutton(s2, text="🔎  Tạo SEO YouTube (Gemini) sau khi xong",
                        variable=self.var_seo).grid(row=2, column=0, sticky="w", pady=(6, 0))
        ttk.Label(s2, text="(gemini_result.docx → seoYoutube.docx)",
                  style="Hint.TLabel").grid(row=3, column=0, sticky="w", pady=(2, 0))
        self.var_auto3 = tk.BooleanVar(value=self._pipe_settings["auto3"])
        ttk.Checkbutton(s2, text="⛓  Tự động chạy bước ③ sau khi xong",
                        variable=self.var_auto3).grid(row=4, column=0, sticky="w", pady=(6, 0))

        # ③ Chuẩn bị input.txt
        s3 = ttk.LabelFrame(wrap, text="  ③  Chuẩn bị input.txt  ")
        s3.grid(row=3, column=0, sticky="ew", pady=(0, 12))
        s3.columnconfigure(0, weight=1)
        self.btn_prep = ttk.Button(s3, text="📝  Tạo input.txt",
                                   style="Accent.TButton", command=self._pipe_prepare_input)
        self.btn_prep.grid(row=0, column=0, sticky="ew")
        ttk.Label(s3, text="(kiểm tra + gemini_result.docx → input.txt)",
                  style="Hint.TLabel").grid(row=1, column=0, sticky="w", pady=(4, 0))
        self.var_auto_tts = tk.BooleanVar(value=self._pipe_settings["auto_tts"])
        ttk.Checkbutton(s3, text="⛓  Chạy tiếp tạo giọng (OmniVoice) sau khi xong",
                        variable=self.var_auto_tts).grid(row=2, column=0, sticky="w", pady=(6, 0))

        # ⬆ ĐĂNG YOUTUBE — dựng xong video tập nào là xếp hàng đăng tập đó NGAY, chạy
        # song song với việc dựng tập kế (xem _upload_enqueue). Video để riêng tư tới
        # giờ hẹn mới tự công khai nên còn kịp sửa/xoá.
        self.var_upload = tk.BooleanVar(value=self._pipe_settings.get("upload", False))
        ttk.Checkbutton(s3, text="⬆  Đăng YouTube sau khi dựng xong video",
                        variable=self.var_upload).grid(row=7, column=0, sticky="w", pady=(6, 0))
        ttk.Label(s3, text=f"(tự hẹn giờ vào khung {upload_slots_text()} còn trống — "
                           "mỗi ngày 2 tập)",
                  style="Hint.TLabel").grid(row=8, column=0, sticky="w", pady=(2, 0))

        # TẮT MÁY / CHO MÁY NGỦ khi chạy xong — để chạy batch qua đêm. Một lần duy
        # nhất: xong là tự bỏ tick, khỏi lỡ tắt/ngủ ở lần chạy sau. Biến dùng chung
        # với ô tick bên tab Giọng nói (tạo ở __init__), và hai ô loại trừ nhau —
        # hẹn tắt máy rồi lại cho ngủ thì lệnh tắt nổ ngay lúc máy dậy.
        ttk.Checkbutton(
            s3, text=f"⏻  Xong hết thì TẮT MÁY (sau {SHUTDOWN_DELAY_MIN} phút)",
            variable=self.var_shutdown,
            command=lambda: self._pick_power_mode("shutdown")
        ).grid(row=3, column=0, sticky="w", pady=(6, 0))
        ttk.Label(s3, text="Huỷ bất cứ lúc nào: mở CMD gõ  shutdown /a",
                  style="Hint.TLabel").grid(row=4, column=0, sticky="w")
        ttk.Checkbutton(
            s3, text=f"🌙  Xong hết thì CHO MÁY NGỦ (sau {SLEEP_DELAY_MIN} phút)",
            variable=self.var_sleep,
            command=lambda: self._pick_power_mode("sleep")
        ).grid(row=5, column=0, sticky="w", pady=(6, 0))
        ttk.Label(s3, text="Huỷ: bỏ tick ô này. Ngủ rồi thì chạm chuột là dậy, "
                           "mọi thứ còn nguyên.",
                  style="Hint.TLabel").grid(row=6, column=0, sticky="w")

        self._pipe_steps = (s1, s2, s3)

        # Tiến trình + trạng thái của quy trình
        pf = ttk.Frame(wrap)
        pf.grid(row=4, column=0, sticky="ew", pady=(2, 0))
        pf.columnconfigure(0, weight=1)
        self._pipe_pf = pf
        self.pipe_progress = tk.IntVar(value=0)
        ttk.Progressbar(pf, variable=self.pipe_progress, maximum=100).grid(
            row=0, column=0, sticky="ew")
        # Dòng ĐANG CHẠY LINK MẤY — luôn hiển thị (không bị các thông báo bước con
        # ghi đè) ngay dưới thanh tiến trình.
        self.pipe_link_status = tk.StringVar(value="")
        lb_link = ttk.Label(pf, textvariable=self.pipe_link_status,
                            font=("Segoe UI", 10, "bold"), foreground=UI["accent"],
                            justify="left")
        lb_link.grid(row=1, column=0, sticky="w", pady=(4, 0))
        self.pipe_status = tk.StringVar(value="Sẵn sàng.")
        lb_stt = ttk.Label(pf, textvariable=self.pipe_status, style="Sub.TLabel",
                           justify="left")
        lb_stt.grid(row=2, column=0, sticky="w", pady=(2, 0))
        # Dòng RIÊNG cho việc đăng YouTube: nó chạy song song với dựng video nên
        # không được dùng chung pipe_status (hai bên sẽ ghi đè lẫn nhau).
        self.upload_status = tk.StringVar(value="")
        lb_up = ttk.Label(pf, textvariable=self.upload_status, style="Sub.TLabel",
                          foreground=UI["accent"], justify="left")
        lb_up.grid(row=4, column=0, sticky="w", pady=(2, 0))
        # Ba dòng trạng thái này dùng CHUNG biến với tab '🎧 Nhận diện' (②→⑤ ghi vào
        # đây). Label không wraplength thì đòi bề ngang đúng bằng chữ → cột quy trình
        # nở ra theo câu trạng thái, nhìn như giao diện bị nhảy. wraplength đặt lại
        # theo kiểu dàn trang ở _pipeline_set_layout (dọc hẹp ở Home / ngang rộng ở
        # tab Tạo kịch bản).
        self._pipe_status_labels = (lb_link, lb_stt, lb_up)

        # ── Điều khiển batch NHIỀU LINK: Tạm dừng/Tiếp tục + Xong link này rồi dừng ──
        # CHỈ hiện ở tab "Tạo kịch bản" (view script) — ẩn ở Home cho đỡ chật (do
        # _pipeline_set_layout ẩn/hiện). Chỉ BẬT khi đang chạy batch. Tác dụng ở ĐIỂM
        # AN TOÀN (ranh giới bước/link), KHÔNG cắt ngang thao tác đang chạy (Whisper/
        # Gemini/tạo giọng) — nên có thể trễ tới khi bước hiện tại xong.
        self._batch_pause_evt = threading.Event()   # set = đang TẠM DỪNG
        self._batch_stop_evt = threading.Event()    # set = DỪNG sau khi xong link hiện tại
        self._batch_running = False
        self._batch_ctrl_frame = bctl = ttk.Frame(pf)
        bctl.grid(row=3, column=0, sticky="ew", pady=(8, 0))
        self.btn_batch_pause = ttk.Button(
            bctl, text="⏸  Tạm dừng", width=15, state="disabled",
            command=self._batch_toggle_pause)
        self.btn_batch_pause.pack(side="left")
        self.btn_batch_stop = ttk.Button(
            bctl, text="⏹  Xong link này rồi dừng", width=24, state="disabled",
            command=self._batch_request_stop)
        self.btn_batch_stop.pack(side="left", padx=(8, 0))
        # Đăng ký để đồng bộ với nút cùng loại ở tab '🎧 Nhận diện'.
        self._batch_pause_widgets.append(self.btn_batch_pause)
        self._batch_stop_widgets.append(self.btn_batch_stop)

        self._pipe_btn_open = ttk.Button(wrap, text="↗  Mở cửa sổ nhận diện đầy đủ",
                                         command=self._open_nhan_dien)
        self._pipe_btn_open.grid(row=5, column=0, sticky="w", pady=(8, 0))
        self._pipe_btn_reset = ttk.Button(wrap, text="↺  Reset cài đặt quy trình về gốc",
                                          command=self._reset_pipe_settings)
        self._pipe_btn_reset.grid(row=6, column=0, sticky="w", pady=(6, 0))

        # Nhật ký riêng của tab "Tạo kịch bản" — chỉ hiện khi xem ngang (dàn 3 bước).
        # Cùng nhận log với ô nhật ký ở panel video (qua _log_boxes).
        self._pipe_log_frame = ttk.LabelFrame(wrap, text="  Nhật ký  ")
        self._pipe_log_frame.columnconfigure(0, weight=1)
        self._pipe_log_frame.rowconfigure(0, weight=1)
        self._make_log_box(self._pipe_log_frame)
        self._pipe_log_frame.grid(row=1, column=3, rowspan=4, sticky="nsew", padx=(12, 0))
        self._pipe_log_frame.grid_remove()   # mặc định ẩn (dọc); _show_view bật khi cần

    def _pipe_pick_file(self):
        # askopenfilenames (số nhiều) → chọn được NHIỀU file 1 lần (Ctrl/Shift để
        # quét khối). Mỗi file thêm thành 1 dòng trong ô nhập.
        paths = filedialog.askopenfilenames(
            title="Chọn file audio/video tiếng Trung (có thể chọn nhiều)",
            filetypes=[("Audio/Video", "*.mp3 *.wav *.m4a *.aac *.flac *.ogg *.opus *.wma "
                                       "*.mp4 *.mkv *.mov *.avi *.webm *.flv"),
                       ("Tất cả", "*.*")])
        if paths:
            cur = self.pipe_txt_sources.get("1.0", "end").strip()
            block = "\n".join(paths)
            self.pipe_txt_sources.insert("end", ("\n" if cur else "") + block)

    def _run_tts_then_shutdown(self, *args, **kwargs):
        """Bọc run_tts cho các luồng ở tab Giọng nói (không đi qua _pipe_set_busy):
        chạy xong/lỗi thì xét ô tick ⏻/🌙 để hẹn tắt máy hoặc cho máy ngủ. Gọi lại trên luồng Tk
        bằng after() vì BooleanVar không an toàn khi đụng từ thread khác."""
        try:
            run_tts(*args, **kwargs)
        finally:
            try:
                self.after(0, self._maybe_schedule_shutdown)
            except Exception:
                pass

    def _pick_power_mode(self, which: str):
        """⏻ và 🌙 loại trừ nhau — tick ô này thì bỏ ô kia.

        Hẹn tắt máy rồi lại cho máy ngủ là kiểu hỏng nhất: lệnh `shutdown /s` vẫn
        nằm đó và nổ ngay khi máy vừa dậy, giữa lúc bạn đang ngồi làm việc.
        """
        if which == "sleep" and self.var_sleep.get():
            self.var_shutdown.set(False)
        elif which == "shutdown" and self.var_shutdown.get():
            self.var_sleep.set(False)
        if not self.var_sleep.get():        # vừa bỏ tick 🌙 → huỷ luôn lệnh đang đếm
            self._cancel_sleep("✖ Đã huỷ hẹn cho máy ngủ.")
        try:
            self._save_pipe_settings()
        except Exception as e:
            logging.warning(f"⚠️ Không lưu được cài đặt pipeline: {e}")

    def _sleep_status(self, msg: str):
        """Ghi dòng trạng thái, bỏ qua nếu cửa sổ đã đóng — mấy hàm dưới đây chạy ở
        luồng Timer, không được phép chết vì một dòng chữ."""
        try:
            self.pipe_status.set(msg)
        except Exception:
            pass

    def _cancel_sleep(self, msg: str = "") -> bool:
        """Huỷ đếm ngược giờ ngủ nếu đang có. True = vừa huỷ được một lệnh."""
        timer, self._sleep_timer = self._sleep_timer, None
        if timer is None:
            return False
        try:
            timer.cancel()
        except Exception:
            pass
        if msg:
            logging.warning(msg)
            self._sleep_status(msg)
        return True

    def _schedule_sleep(self):
        """Đếm ngược SLEEP_DELAY_MIN phút rồi cho máy ngủ.

        KHÔNG bỏ tick ngay như ô ⏻: lệnh tắt máy còn huỷ được bằng `shutdown /a`,
        còn lệnh ngủ này là của chính app — giữ tick lại chính là nút huỷ. Tick chỉ
        tự bỏ khi máy đã thật sự ngủ (một lần duy nhất).
        """
        if self._sleep_timer is not None:      # đang đếm rồi, đừng chồng lệnh
            return
        timer = threading.Timer(SLEEP_DELAY_MIN * 60, self._do_sleep)
        timer.daemon = True                    # đóng app là lệnh ngủ cũng tan theo
        self._sleep_timer = timer
        timer.start()
        logging.warning(f"🌙 ĐÃ HẸN CHO MÁY NGỦ sau {SLEEP_DELAY_MIN} phút — "
                        "muốn huỷ thì bỏ tick ô 🌙.")
        self._sleep_status(f"🌙 Máy ngủ sau {SLEEP_DELAY_MIN} phút — huỷ: bỏ tick 🌙")

    def _do_sleep(self):
        """Tới giờ ngủ (chạy ở luồng của Timer nên không treo cửa sổ lúc máy ngủ)."""
        self._sleep_timer = None
        try:
            if not self.var_sleep.get():       # vừa bỏ tick sát giờ
                return
            if self._pipe_busy:                # có việc mới chen vào → để lượt sau
                logging.warning("🌙 Có việc đang chạy → hoãn cho máy ngủ.")
                return
            self.var_sleep.set(False)          # một lần duy nhất
            self._save_pipe_settings()
        except Exception as e:
            # Cửa sổ vừa đóng / Tk không trả lời → KHÔNG ngủ. Đọc không ra ý người
            # dùng mà vẫn úp máy xuống thì phiền hơn nhiều so với bỏ qua lượt này.
            logging.warning(f"⚠️ Bỏ qua lệnh cho máy ngủ: {e}")
            return
        logging.warning("🌙 Cho máy ngủ… (chạm chuột/bàn phím là dậy, mọi thứ còn nguyên)")
        self._sleep_status("🌙 Đang cho máy ngủ…")
        err = suspend_computer()               # chỉ trả về khi máy đã DẬY
        if err:
            logging.error(f"Không cho máy ngủ được: {err}")
            self._sleep_status("❌ Không cho máy ngủ được — xem nhật ký.")
            return
        logging.warning("☀️ Máy đã dậy — app chạy tiếp bình thường.")
        self._sleep_status("☀️ Máy đã dậy — sẵn sàng chạy tiếp.")

    def _maybe_schedule_shutdown(self):
        """Xong hết thì TẮT MÁY (ô ⏻) hoặc CHO MÁY NGỦ (ô 🌙) — để chạy batch qua đêm.
        Chỉ hẹn khi KHÔNG còn tiến trình nào chạy.

        MỘT LẦN DUY NHẤT: hẹn xong tự bỏ tick, tránh lần chạy sau bị tắt/ngủ ngoài ý
        muốn. Vẫn hẹn khi luồng kết thúc do LỖI — máy đằng nào cũng rảnh, log đã ghi
        lại lỗi. Huỷ: `shutdown /a` (tắt máy) hoặc bỏ tick 🌙 (ngủ).
        """
        if self._pipe_busy:
            return
        if getattr(self, "var_sleep", None) is not None and self.var_sleep.get():
            self._schedule_sleep()
            return
        var = getattr(self, "var_shutdown", None)
        if var is None or not var.get():
            return
        var.set(False)                 # một lần duy nhất
        try:
            self._save_pipe_settings()
        except Exception as e:
            logging.warning(f"⚠️ Không lưu được cài đặt pipeline trước khi tắt máy: {e}")
        try:
            import subprocess
            secs = int(SHUTDOWN_DELAY_MIN * 60)
            subprocess.run(
                ["shutdown", "/s", "/t", str(secs), "/c",
                 f"OmniVoice da chay xong - tat may sau {SHUTDOWN_DELAY_MIN} phut. "
                 f"Huy: shutdown /a"],
                creationflags=CREATE_NO_WINDOW, check=False)
            logging.warning(f"⏻ ĐÃ HẸN TẮT MÁY sau {SHUTDOWN_DELAY_MIN} phút — "
                            f"muốn huỷ thì mở CMD gõ:  shutdown /a")
            self.pipe_status.set(
                f"⏻ Tắt máy sau {SHUTDOWN_DELAY_MIN} phút — huỷ: shutdown /a")
        except Exception as e:
            logging.error(f"Không hẹn được tắt máy: {e}")

    def _pipe_set_busy(self, busy: bool):
        self._pipe_busy = busy
        # Việc mới chen vào giữa lúc đang đếm ngược giờ ngủ → huỷ đếm, chờ xong hết
        # đã. Ô tick 🌙 vẫn giữ nguyên nên xong việc này là hẹn lại từ đầu.
        if busy:
            self._cancel_sleep("🌙 Có việc mới → huỷ đếm ngược ngủ, chờ xong hết đã.")
        state = "disabled" if busy else "normal"
        for b in (self.btn_recog, self.btn_gemini, self.btn_prep,
                  getattr(self, "recog_tab_btn", None),
                  getattr(self, "recog_translate_btn", None),
                  getattr(self, "recog_seo_btn", None),
                  getattr(self, "recog_thumb_btn", None),
                  getattr(self, "recog_tts_btn", None),
                  getattr(self, "recog_upload_btn", None),
                  getattr(self, "recog_chain_btn", None)):
            if b is not None:
                b.config(state=state)
        # Mọi luồng dài (batch nhiều link, các bước ①②③④⑤) đều đi qua đây khi kết
        # thúc → chỗ duy nhất cần móc lệnh hẹn tắt máy.
        if not busy:
            self._maybe_schedule_shutdown()

    def _save_pipe_settings(self):
        """Lưu cài đặt quy trình hiện tại (auto + model/tốc độ) cho lần sau."""
        save_pipe_settings(dict(
            auto2=self.var_auto2.get(), auto3=self.var_auto3.get(),
            auto_tts=self.var_auto_tts.get(), seo=self.var_seo.get(),
            model=self.pipe_var_model.get(), speed=self.pipe_var_speed.get(),
            shutdown=getattr(self, "var_shutdown", tk.BooleanVar()).get(),
            sleep=getattr(self, "var_sleep", tk.BooleanVar()).get(),
            upload=getattr(self, "var_upload", tk.BooleanVar()).get(),
        ))

    def _reset_pipe_settings(self):
        """Đưa các tùy chọn quy trình về mặc định gốc và lưu lại."""
        self.var_auto2.set(PIPE_DEFAULTS["auto2"])
        self.var_auto3.set(PIPE_DEFAULTS["auto3"])
        self.var_auto_tts.set(PIPE_DEFAULTS["auto_tts"])
        self.var_seo.set(PIPE_DEFAULTS["seo"])
        if hasattr(self, "var_shutdown"):
            self.var_shutdown.set(PIPE_DEFAULTS["shutdown"])
        if hasattr(self, "var_sleep"):
            self.var_sleep.set(PIPE_DEFAULTS["sleep"])
            self._cancel_sleep("✖ Đã huỷ hẹn cho máy ngủ.")
        if hasattr(self, "var_upload"):
            self.var_upload.set(PIPE_DEFAULTS["upload"])
        self.pipe_var_model.set(PIPE_DEFAULTS["model"])
        self.pipe_var_speed.set(PIPE_DEFAULTS["speed"])
        self.pipe_txt_sources.delete("1.0", "end")
        self._save_pipe_settings()
        self.pipe_status.set("↺ Đã reset cài đặt quy trình về mặc định.")

    def _pipe_sources(self) -> list:
        """Danh sách link/file hợp lệ từ ô nhập ① (mỗi dòng 1 mục), giữ thứ tự."""
        out = []
        for line in self.pipe_txt_sources.get("1.0", "end").splitlines():
            s = line.strip().strip('"').strip("'")
            if s and (os.path.isfile(s) or s.lower().startswith(("http://", "https://"))):
                out.append(s)
        return out

    def _pipe_recognize(self):
        if self._pipe_busy:
            return
        sources = self._pipe_sources()
        if not sources:
            self.pipe_status.set("⚠️ Chưa có link/file — hãy nhập đầu vào (mỗi dòng 1 mục).")
            logging.warning("Chưa có link video / file để nhận diện.")
            return
        self._save_pipe_settings()   # ấn chạy → nhớ cài đặt cho lần sau

        # ≥2 link → chế độ NHIỀU LINK: mỗi link 1 thư mục + full pipeline.
        if len(sources) >= 2:
            self._pipe_start_batch(sources)
            return

        self._pipe_set_busy(True)
        self.pipe_progress.set(0)
        self.pipe_link_status.set("")   # 1 link đơn → không hiển thị "đang chạy link mấy"
        self.pipe_status.set("🎙  Đang chuẩn bị...")
        threading.Thread(
            target=self._pipe_recognize_worker,
            args=(sources[0], self.pipe_var_model.get(), self.pipe_var_speed.get()),
            daemon=True).start()

    def _pipe_recognize_worker(self, media, model, speed):
        ok = False
        try:
            import nhandien_giongnoi as recog
            # Đầu vào là LINK → tải MP3 trước; file có sẵn → dùng trực tiếp.
            if not os.path.isfile(media) and str(media).lower().startswith(("http://", "https://")):
                logging.info(f"🌐 Tải audio từ link: {media}")
                self.pipe_status.set("🌐  Đang tải audio từ link...")
                media = download_audio_mp3(media, DOWNLOAD_DIR)
                if not media:
                    logging.error("❌ Không tải được audio từ link.")
                    self.pipe_status.set("❌ Tải link thất bại.")
                    return
            self.pipe_status.set("🎙  Đang nhận diện...")
            logging.info(f"🎙  Nhận diện: {Path(media).name} (model={model}, tốc độ={speed})")
            transcript = recog.transcribe_chinese(
                media, model_name=model, speed=float(speed),
                on_progress=lambda f: self.pipe_progress.set(int(f * 100)))
            if not transcript:
                logging.error("❌ Nhận diện thất bại / không có nội dung.")
                self.pipe_status.set("❌ Nhận diện thất bại.")
                return
            recog.save_docx(transcript, str(CHINESE_DOCX), title=Path(media).name)
            n = len(recog.split_into_chunks(transcript))
            self.pipe_progress.set(100)
            logging.info(f"✅ Nhận diện xong: {len(transcript)} ký tự, {n} đoạn → {CHINESE_DOCX.name}")
            self.pipe_status.set(f"✅ Xong ({n} đoạn) → {CHINESE_DOCX.name}")
            ok = True
        except Exception as e:
            logging.error(f"Lỗi nhận diện: {e}")
            self.pipe_status.set(f"Lỗi: {e}")
        finally:
            # Nhận diện xong → KHÔNG còn dùng model Whisper nữa, giải phóng khỏi
            # VRAM ngay để nhường chỗ cho OmniVoice ở bước tạo giọng (GPU 8GB dễ
            # nghẽn nếu large-v3 ~3GB vẫn nằm lại trong VRAM khi nạp OmniVoice).
            try:
                import nhandien_giongnoi as recog
                recog.free_model()
                logging.info("🧹 Đã giải phóng model nhận diện khỏi VRAM.")
            except Exception as e:
                logging.warning(f"Không giải phóng được model nhận diện: {e}")
            self._pipe_set_busy(False)
            if ok and self.var_auto2.get():   # ⛓ tự động sang bước ②
                self.after(600, lambda: self._pipe_send_gemini(auto=True))

    # ── LUỒNG RIÊNG: CHỈ NHẬN DIỆN các link rồi NGƯNG (tab '🎧 Nhận diện') ──────
    def _pipe_recognize_only(self, sources=None):
        """Nhận diện MỌI link/file rồi NGƯNG — KHÔNG dịch Gemini, KHÔNG SEO, KHÔNG
        input.txt, KHÔNG video. Kết quả *_zh.docx lưu vào từng thư mục tập. Không
        truyền sources → lấy từ ô ① pipeline; tab Nhận diện truyền ô riêng của nó."""
        if self._pipe_busy:
            return
        if sources is None:                 # gọi từ đâu không truyền → lấy ô ① pipeline
            sources = self._pipe_sources()
        if not sources:
            self.pipe_status.set("⚠️ Chưa có link/file — hãy nhập đầu vào (mỗi dòng 1 mục).")
            logging.warning("Chưa có link/file để nhận diện.")
            return
        self._save_pipe_settings()   # nhớ model/tốc độ cho lần sau
        self._pipe_set_busy(True)
        # Cho phép Tạm dừng / Dừng-sau-link giống batch.
        self._batch_pause_evt.clear()
        self._batch_stop_evt.clear()
        self._batch_running = True
        self._set_batch_pause_btns(state="normal", text="⏸  Tạm dừng")
        self._set_batch_stop_btns(state="normal")
        self.pipe_progress.set(0)
        self.pipe_link_status.set(f"⏳ Chỉ nhận diện {len(sources)} link (không chạy bước khác)...")
        self.pipe_status.set(f"⏳ Bắt đầu nhận diện {len(sources)} link...")
        threading.Thread(
            target=self._pipe_recognize_only_worker,
            args=(sources, self.pipe_var_model.get(), self.pipe_var_speed.get()),
            daemon=True).start()

    def _pipe_recognize_only_worker(self, sources, model, speed):
        """Nhận diện lần lượt từng nguồn (mỗi nguồn 1 thư mục tập) rồi DỪNG. Bỏ qua
        nguồn đã có *_zh.docx. Không chạm số tập / không dịch / không video."""
        total = len(sources)
        ok_count = 0
        try:
            import nhandien_giongnoi as recog
            logging.info("\n" + "═" * 10 +
                         f" CHỈ NHẬN DIỆN {total} link (không dịch/SEO/video) " + "═" * 10)
            logging.info(f"📦 Model: {model}  •  Tốc độ: {speed}x")

            for i, src in enumerate(sources, 1):
                self._batch_pause_wait()
                if self._batch_stop_evt.is_set():
                    logging.info(f"⏹ ĐÃ DỪNG — đã nhận diện {ok_count}/{total} link.")
                    self.pipe_link_status.set(f"⏹ Đã dừng — nhận diện {ok_count}/{total} link.")
                    break

                episode, is_new = self._allocate_episode(src)
                # Tên thư mục kèm TÊN NGUỒN ("01 - 95") để biết tập làm từ đâu;
                # tập đã có thư mục thì dùng lại thư mục cũ, không đổi tên.
                folder = episode_dir_for(episode, src)
                folder.mkdir(parents=True, exist_ok=True)
                # Ghi nguồn↔tập vào manifest (chưa done) để sau chạy tiếp/batch đúng tập.
                self._manifest_update(src, episode, folder, done=False)
                logging.info(f"🔖 {'Tập MỚI' if is_new else 'Tập'} {episode} "
                             f"← {norm_source(src)[:70]}")
                self.pipe_link_status.set(f"🎧 Nhận diện: Link {i}/{total} — tập {episode}")
                self.pipe_status.set(f"🎧 Link {i}/{total} → tập {episode}/")

                try:
                    s = src.strip().strip('"').strip("'")
                    # Đã có bản nhận diện hợp lệ → bỏ qua (không tải + không chạy lại).
                    existing_zh = find_zh_docx(folder)
                    if existing_zh and read_zh_docx_chunks(existing_zh):
                        logging.info(f"♻ Bỏ qua (đã có {existing_zh.name}).")
                        ok_count += 1
                        continue

                    if os.path.isfile(s):
                        media = s
                        logging.info(f"📁 File local: {media}")
                    elif s.lower().startswith(("http://", "https://")):
                        logging.info(f"🌐 Tải từ link: {s}")
                        media = download_audio_mp3(s, DOWNLOAD_DIR)
                        if not media:
                            logging.error(f"❌ Link {i}: không tải được audio — bỏ qua.")
                            continue
                    else:
                        logging.error(f"❌ Link {i}: không hợp lệ — bỏ qua.")
                        continue

                    self.pipe_progress.set(0)
                    transcript = recog.transcribe_chinese(
                        media, model_name=model, speed=float(speed),
                        on_progress=lambda f: self.pipe_progress.set(int(f * 100)))
                    if not transcript:
                        logging.error(f"❌ Link {i}: không nhận diện được — bỏ qua.")
                        continue
                    zh_docx = folder / ZH_DOCX_NAME   # tên cố định, không theo tên video
                    recog.save_docx(transcript, str(zh_docx), title=Path(media).name)
                    logging.info(f"💾 Đã lưu bản nhận diện: {zh_docx}")
                    ok_count += 1
                except Exception as e:
                    import traceback
                    logging.error(f"❌ Lỗi nhận diện link {i}/{total}: {e}")
                    logging.error(traceback.format_exc())
                    continue

            self.pipe_progress.set(100)
            self.pipe_link_status.set(f"✅ Đã nhận diện {ok_count}/{total} link — đã ngưng.")
            self.pipe_status.set(f"✅ Xong nhận diện {ok_count}/{total} link (chưa dịch/SEO/video).")
            logging.info(f"🎉 XONG NHẬN DIỆN: {ok_count}/{total} link. Dùng 'Chạy tiếp' "
                         f"hoặc chạy nhiều link để làm các bước còn lại.")
        except Exception as e:
            logging.error(f"Lỗi luồng chỉ-nhận-diện: {e}")
            self.pipe_status.set(f"Lỗi: {e}")
        finally:
            try:
                import nhandien_giongnoi as recog
                recog.free_model()
                logging.info("🧹 Đã giải phóng model nhận diện khỏi VRAM.")
            except Exception:
                pass
            self._pipe_set_busy(False)
            self._batch_controls_reset()
            self._recog_schedule_table_refresh()   # cập nhật bảng trạng thái tab Nhận diện

    def _collect_tts_settings(self):
        """Thu thập cài đặt TTS hiện tại (phải gọi trên MAIN THREAD vì đọc các tk.Var)
        để chế độ NHIỀU LINK tạo giọng cho mỗi tập y như khi chạy 1 link.

        Trả về dict cài đặt, hoặc None nếu cấu hình sai (đã hiện cảnh báo)."""
        mode = self.var_mode.get()
        if mode == "clone":
            voice_name = self._current_voice()
            if not voice_name:
                messagebox.showwarning("Thiếu giọng mẫu",
                                       f"Không tìm thấy file audio trong:\n{VOICE_DIR}")
                return None
            voice_param = str(VOICE_DIR / voice_name)
        elif mode == "design":
            voice_param = self.var_instruct.get().strip()
            if not voice_param:
                messagebox.showwarning("Thiếu mô tả", "Vui lòng nhập mô tả giọng đọc.")
                return None
        else:
            voice_param = None

        try:
            doc_speed = float(str(self.var_doc_speed.get()).replace(",", ".").strip())
        except (TypeError, ValueError):
            doc_speed = 1.0
        doc_speed = max(0.5, min(doc_speed, 2.0))
        try:
            ngang_speed = float(str(self.var_ngang_speed.get()).replace(",", ".").strip())
        except (TypeError, ValueError):
            ngang_speed = 1.0
        ngang_speed = max(0.5, min(ngang_speed, 2.0))
        try:
            tiktok_speed = float(str(self.var_tiktok_speed.get()).replace(",", ".").strip())
        except (TypeError, ValueError):
            tiktok_speed = 1.0
        tiktok_speed = max(0.5, min(tiktok_speed, 2.0))
        try:
            tiktok_caption_pos = int(self.var_tiktok_caption_pos.get())
        except Exception:
            tiktok_caption_pos = 40
        tiktok_caption_pos = max(0, min(tiktok_caption_pos, 100))
        try:
            tiktok_music_db = int(self.var_tiktok_music_db.get())
        except Exception:
            tiktok_music_db = -12
        tiktok_music_db = max(-40, min(tiktok_music_db, 0))

        effect_name = self._current_effect()
        effect_path = None
        if effect_name and effect_name != EFFECT_NONE:
            p = EFFECTS_DIR / effect_name
            effect_path = str(p) if p.exists() else None

        return dict(
            mode=mode, voice_param=voice_param, chunk=self.var_chunk.get(),
            make_video=self.var_make_video.get(), effect=effect_path,
            make_video_doc=self.var_make_video_doc.get(),
            doc_speed=doc_speed,
            doc_percent=self._parse_percent(self.var_doc_percent, 100),
            ngang_speed=ngang_speed, ngang_source=self.var_ngang_source.get(),
            doc_from_ngang=self.var_doc_from_ngang.get(),
            doc_from_subfolder=self.var_doc_from_subfolder.get(),
            doc_no_effect=self.var_doc_no_effect.get(),
            make_tiktok=self.var_make_tiktok.get(), tiktok_speed=tiktok_speed,
            make_short=self.var_make_short.get(),
            tiktok_percent=self._parse_percent(self.var_tiktok_percent),
            tiktok_no_effect=self.var_tiktok_no_effect.get(),
            tiktok_caption_pos=tiktok_caption_pos,
            tiktok_music=self.var_tiktok_music.get(), tiktok_music_db=tiktok_music_db,
            bring_front=self.var_bring_front.get(),
            make_sub=self.var_make_sub.get(), sub_mode=self.var_sub_mode.get(),
            sub_model=self.var_sub_model.get(),
            sub_max_chars=self._parse_sub_max_chars(),
            sub_kieu=self.var_sub_kieu.get(),
            sub_font=self.var_sub_font.get(),
            sub_mau=self.var_sub_mau.get(),
            sub_mau_vien=self.var_sub_mau_vien.get(),
            sub_vitri=self.var_sub_vitri.get().strip(),
            sub_cochu=self._parse_sub_cochu(),
            sub_bengang=self._parse_sub_bengang(),
            sub_dong=self._parse_sub_dong(),
            make_sub_doc=self.var_make_sub_doc.get(),
        )

    # ── CHẾ ĐỘ NHIỀU LINK: mỗi link 1 thư mục kịch_bản/NN, full pipeline ───────
    def _pipe_start_batch(self, sources):
        """Xử lý nhiều link lần lượt: mỗi link 1 thư mục (01, 02, ...) trong
        kịch_bản, chạy full pipeline (nhận diện → Gemini → input.txt → SEO → giọng)."""
        # Nếu bật "⛓ Chạy tiếp tạo giọng" thì thu thập cài đặt TTS NGAY BÂY GIỜ trên
        # main thread (worker chạy thread khác, không nên đọc tk.Var). Cấu hình sai
        # (vd clone mà chưa chọn giọng) → dừng trước khi bắt đầu cả batch.
        tts_settings = None
        if self.var_auto_tts.get():
            tts_settings = self._collect_tts_settings()
            if tts_settings is None:
                return
        # Đăng YouTube: kiểm tra thư viện + đăng nhập NGAY BÂY GIỜ (main thread), vì
        # giữa mẻ mà phải mở trình duyệt đăng nhập là treo cả mẻ chạy qua đêm.
        upload = bool(self.var_upload.get())
        if upload and not self._upload_check_ready():
            return
        # Chạy ngay, KHÔNG hỏi xác nhận. Lưu ý: bước dịch & SEO dùng Firefox — hãy
        # ĐÓNG Firefox đang mở (profile bị khoá khi đang chạy) và đảm bảo đã đăng nhập.
        logging.info(f"⛓ Xử lý {len(sources)} link theo thứ tự (mỗi link 1 tập)"
                     + (" + tạo giọng." if tts_settings else ".")
                     + (" Đăng YouTube: BẬT." if upload else ""))
        self._pipe_set_busy(True)
        # Bật điều khiển batch: xoá cờ cũ + cho phép 2 nút Tạm dừng / Dừng-sau-link.
        self._batch_pause_evt.clear()
        self._batch_stop_evt.clear()
        self._batch_running = True
        self._set_batch_pause_btns(state="normal", text="⏸  Tạm dừng")
        self._set_batch_stop_btns(state="normal")
        self.pipe_progress.set(0)
        self.pipe_link_status.set(f"⏳ Chuẩn bị xử lý {len(sources)} link...")
        self.pipe_status.set(f"⏳ Bắt đầu xử lý {len(sources)} link...")
        threading.Thread(
            target=self._pipe_batch_worker,
            args=(sources, self.pipe_var_model.get(), self.pipe_var_speed.get(),
                  tts_settings, upload),
            daemon=True).start()

    # ── Điều khiển batch nhiều link: tạm dừng / dừng-sau-link ──────────────────
    def _set_batch_pause_btns(self, **kw):
        """Đổi trạng thái/chữ cho MỌI nút Tạm dừng (tab Tạo kịch bản + tab Nhận diện)."""
        for b in self._batch_pause_widgets:
            try:
                b.config(**kw)
            except Exception:
                pass

    def _set_batch_stop_btns(self, **kw):
        """Đổi trạng thái cho MỌI nút Dừng-sau-link (đồng bộ giữa các tab)."""
        for b in self._batch_stop_widgets:
            try:
                b.config(**kw)
            except Exception:
                pass

    def _batch_toggle_pause(self):
        """Tạm dừng ↔ cho chạy tiếp batch. Áp dụng ở điểm an toàn kế tiếp (giữa bước/link),
        không cắt ngang thao tác đang chạy."""
        if not self._batch_running:
            return
        if self._batch_pause_evt.is_set():
            self._batch_pause_evt.clear()
            self._set_batch_pause_btns(text="⏸  Tạm dừng")
            self.pipe_link_status.set("▶  Cho chạy tiếp...")
            logging.info("▶ Người dùng cho CHẠY TIẾP batch.")
        else:
            self._batch_pause_evt.set()
            self._set_batch_pause_btns(text="▶  Tiếp tục")
            self.pipe_link_status.set("⏸  Sẽ tạm dừng ở điểm an toàn kế tiếp (giữa bước/link)...")
            logging.info("⏸ Người dùng yêu cầu TẠM DỪNG batch.")

    def _batch_request_stop(self):
        """Đánh dấu DỪNG: link đang chạy sẽ hoàn tất RỒI batch dừng (không bắt đầu link kế)."""
        if not self._batch_running:
            return
        self._batch_stop_evt.set()
        # Nếu đang tạm dừng thì bỏ cờ tạm dừng để link hiện tại chạy nốt rồi mới dừng.
        self._batch_pause_evt.clear()
        self._set_batch_pause_btns(state="disabled", text="⏸  Tạm dừng")
        self._set_batch_stop_btns(state="disabled")
        self.pipe_link_status.set("⏹  Sẽ dừng sau khi xong link đang chạy...")
        logging.info("⏹ Người dùng yêu cầu DỪNG sau khi xong link hiện tại.")

    def _batch_pause_wait(self):
        """Chặn (chờ) tại điểm an toàn khi đang TẠM DỪNG; nhả khi bấm Tiếp tục HOẶC khi
        đã bấm Dừng (để link hiện tại chạy nốt rồi dừng ở đầu vòng kế)."""
        if not self._batch_pause_evt.is_set():
            return
        import time
        self.pipe_status.set("⏸ Đã tạm dừng — bấm Tiếp tục để chạy tiếp.")
        while self._batch_pause_evt.is_set() and not self._batch_stop_evt.is_set():
            time.sleep(0.3)

    def _batch_controls_reset(self):
        """Về trạng thái KHÔNG chạy: xoá cờ + tắt 2 nút điều khiển batch."""
        self._batch_running = False
        self._batch_pause_evt.clear()
        self._batch_stop_evt.clear()
        self._set_batch_pause_btns(state="disabled", text="⏸  Tạm dừng")
        self._set_batch_stop_btns(state="disabled")

    def _batch_prepare_input(self, gemini_docx, out_txt) -> bool:
        """Tạo input.txt cho 1 link: bỏ cấu trúc + thay câu quảng bá kênh + bỏ chú
        thích () []. (Bản batch không hỏi/không dừng như bước ③ thủ công.)"""
        # ⛔ CHỐT ĐOẠN HỎNG: BẤT KỲ đoạn nào chưa dịch / bị Gemini từ chối / dịch
        # cụt → KHÔNG ghi input.txt, bỏ cả tập. Kiểm ở đây (ngoài chốt dịch) vì
        # bước tạo input còn được gọi thẳng không qua _translation_complete.
        bad = kiem_ban_dich_folder(Path(gemini_docx).parent)
        if bad:
            mota = ", ".join(f"đoạn {j} {ly_do}" for j, ly_do in bad)
            logging.error(f"⛔ {Path(gemini_docx).parent.name}: bản dịch có đoạn hỏng "
                          f"({mota}) → KHÔNG ghi input.txt, BỎ CẢ TẬP. Chạy lại bước "
                          "dịch Gemini để dịch lại các đoạn đó.")
            return False
        try:
            import dich_kiemtra as cg
            findings = cg.check_docx(gemini_docx, on_log=logging.info)
            if findings:
                logging.warning(f"⚠️ {gemini_docx.parent.name}: gemini_result.docx còn "
                                f"{len(findings)} đoạn có câu dẫn nhập/thừa — vẫn ghi input.txt.")
            chunks = cg.read_docx_chunks(gemini_docx)
            content = "\n".join(t for _, t in chunks).strip()
            if not content:
                return False
            # ⛔ CHẶN: còn đoạn "(chưa dịch)"/"(trống)" thì KHÔNG ghi input.txt. Phải
            # kiểm ở ĐÂY — sau đó remove_annotations sẽ xoá chúng như chú thích trong
            # ngoặc, mất nguyên đoạn mà không còn dấu vết (xem tập 42).
            try:
                import dich_chuanbi_input as prep
                marks = prep.find_untranslated(content)
            except Exception:
                marks = []
            if marks:
                logging.error(
                    f"⛔ {gemini_docx.parent.name}: gemini_result.docx còn đoạn chưa dịch "
                    f"{marks} → KHÔNG ghi input.txt (tránh ra audio/video thiếu thời "
                    "lượng). Chạy lại bước dịch Gemini cho các đoạn còn thiếu.")
                return False
            content, n_promo, n_add = replace_channel_promo(content)
            logging.info(f"🔁 Quảng bá kênh: xóa {n_promo} câu rải rác, "
                         f"chèn lại {n_add} câu.")
            content, _n_en, _sus_en = replace_leaked_english(content)
            try:
                import dich_hanviet as hv
                content, n_mt, n_am = hv.translate_han(content, on_log=logging.info)
                if n_mt or n_am:
                    logging.info(f"🈶 Chữ Hán sót: {n_mt} đoạn MT, {n_am} chữ phiên âm.")
            except Exception as e:
                logging.warning(f"⚠️ Bỏ qua xử lý chữ Hán: {e}")
            try:
                import dich_chuanbi_input as prep
                content = prep.remove_annotations(content)  # bỏ chú thích () []
                # Sửa từ cố định: giết→giớt, máu→máo, tỳ→tì… (_WORD_FIXES)
                content = prep.apply_word_fixes(content)
            except Exception as e:
                logging.warning(f"⚠️ Không dọn được chú thích () [] / sửa từ cố định: {e}")
            # ⛔ CHỐT ĐỘ DÀI: so với bản nhận diện tiếng Trung của cùng tập. Bắt các ca
            # mất đoạn mà marker "(chưa dịch)" KHÔNG phủ được (vd Gemini trả về lời xin
            # lỗi ngắn thay vì bản dịch, hoặc docx bị sửa tay hụt).
            try:
                import kiemtra_daura as kt
                ok_len, msg = kt.check_ban_dich(content, Path(out_txt).parent)
                if not ok_len:
                    logging.error(f"⛔ {gemini_docx.parent.name}: {msg}\n"
                                  "   → KHÔNG ghi input.txt (tránh ra audio/video thiếu).")
                    return False
            except Exception as e:
                logging.warning(f"⚠️ Bỏ qua chốt kiểm độ dài bản dịch: {e}")
            Path(out_txt).write_text(content, encoding="utf-8")
            return True
        except Exception as e:
            logging.error(f"⚠️ Lỗi tạo input.txt: {e}")
            return False

    def _batch_run_tts(self, folder, ts, episode=None) -> bool:
        """Tạo giọng OmniVoice cho 1 tập: đọc folder/input.txt → folder/output.wav
        (kèm cắt/dựng video theo cài đặt ts). Chạy ĐỒNG BỘ trong thread batch.

        episode: số tập (để ghi chữ 'Mimi audio Số <episode>' lên video TikTok, khớp
                 số trên thumbnail). None → không ghi chữ."""
        # ⛔ CHỐT ĐOẠN HỎNG: bản dịch còn đoạn chưa dịch / bị từ chối / dịch cụt →
        # KHÔNG tạo giọng, KHÔNG dựng video. Phải kiểm lại ở đây dù input.txt đã
        # tồn tại: input cũ có thể được tạo TRƯỚC khi có chốt (ca tập 85/87 —
        # render + đăng video thiếu 14-25% nội dung).
        bad = kiem_ban_dich_folder(folder)
        if bad:
            mota = ", ".join(f"đoạn {j} {ly_do}" for j, ly_do in bad)
            logging.error(f"⛔ {folder.name}: bản dịch có đoạn hỏng ({mota}) → KHÔNG "
                          "tạo giọng/video, BỎ CẢ TẬP. Chạy lại bước dịch Gemini "
                          "trước (input.txt hiện có được tạo từ bản dịch hỏng).")
            return False
        input_txt = folder / "input.txt"
        if not input_txt.exists():
            logging.warning(f"⚠️ {folder.name}: chưa có input.txt → bỏ qua tạo giọng.")
            return False
        try:
            full_text = clean_text(input_txt.read_text(encoding="utf-8"))
        except Exception as e:
            logging.error(f"⚠️ {folder.name}: không đọc được input.txt: {e}")
            return False
        chunks = split_chunks(full_text.lower(), ts["chunk"])
        if not chunks:
            logging.warning(f"⚠️ {folder.name}: input.txt trống → bỏ qua tạo giọng.")
            return False

        # Whisper đã được giải phóng ở lượt gọi (sau khi đóng Firefox) nên ở đây
        # chỉ cần nạp OmniVoice để tạo giọng.
        output = folder / "output.wav"
        stub = _NullWidget()
        pause_event = threading.Event()
        pause_event.set()
        # Tới bước clone giọng cho link này → (tùy chọn) bật cửa sổ GUI lên trên cùng.
        if ts.get("bring_front"):
            self._bring_to_front()
        logging.info(f"🎧 Tạo giọng OmniVoice cho tập {folder.name} ({len(chunks)} đoạn)...")
        # Tiến trình chi tiết (tạo giọng + dựng video) hiện ở THANH TTS (self.progress/
        # self.status); thanh "Tạo kịch bản" để hiển thị tiến độ tổng theo số tập.
        run_tts(
            ts["mode"], ts["voice_param"], chunks, str(output),
            self.progress, self.status,
            stub, stub, stub, pause_event,
            ts["make_video"], ts["effect"],
            ts["make_video_doc"], ts["doc_speed"],
            ts.get("doc_percent", 100),                # % cắt audio video dọc (100 = cả bài)
            ts["ngang_speed"], True,                   # reuse=True → TIẾP TỤC: dùng
            ts["doc_from_ngang"], ts["doc_no_effect"],  # lại audio/video đã có, chỉ
            doc_from_subfolder=ts.get("doc_from_subfolder", False),
            ngang_source=ts.get("ngang_source"),   # nguồn clip video ngang theo chủ đề
            # Bản tự động: đặt tên riêng cho các video kết quả trong thư mục tập.
            ngang_out=folder / "YOUTUBE.mp4",      # video NGANG (đăng YouTube)
            doc_out=folder / "facebook.mp4",       # video DỌC  (đăng Facebook)
            make_tiktok=ts.get("make_tiktok", False),
            tiktok_out=folder / "tiktok.mp4",      # video TIKTOK (cắt theo %)
            make_short=ts.get("make_short", False),
            short_out=folder / "short.mp4",        # YOUTUBE SHORT (≤2:50, đăng tự động)
            tiktok_speed=ts.get("tiktok_speed", 1.0),
            tiktok_percent=ts.get("tiktok_percent", 50),
            tiktok_no_effect=ts.get("tiktok_no_effect", False),
            # Chữ trên TikTok = 'Mimi audio Số <số ở thumbnail>' (khớp số tập).
            tiktok_caption=(f"Mimi audio Số {episode}" if episode else None),
            tiktok_caption_pos=ts.get("tiktok_caption_pos", 40),
            tiktok_music=ts.get("tiktok_music", False),
            tiktok_music_db=ts.get("tiktok_music_db", -12),
            # Phụ đề cho YOUTUBE.mp4 (tắt mặc định — xem hàng "📝 Phụ đề" ở Cài đặt).
            make_sub=ts.get("make_sub", False),
            sub_mode=ts.get("sub_mode", SUB_MODE_SRT),
            sub_model=ts.get("sub_model", "large-v3-turbo"),
            sub_max_chars=ts.get("sub_max_chars", 50),
            sub_kieu=ts.get("sub_kieu", "hopbo"),
            sub_font=ts.get("sub_font", ""),
            sub_mau=ts.get("sub_mau", ""),
            sub_mau_vien=ts.get("sub_mau_vien", ""),
            sub_vitri=ts.get("sub_vitri", ""),
            sub_cochu=ts.get("sub_cochu", ""),
            sub_bengang=ts.get("sub_bengang", ""),
            sub_dong=ts.get("sub_dong", 2),
            # Phụ đề cho facebook.mp4 (khung dọc) — chạy SAU TikTok, xem run_tts.
            make_sub_doc=ts.get("make_sub_doc", False),
        )                                          # render phần còn thiếu (vd video dọc).
        logging.info(f"🎧 Xong tạo giọng tập {folder.name} → {output.name}")
        # ⚠️ CHỐT ĐỘ DÀI AUDIO: so thời lượng output.wav với số ký tự input.txt. Bắt ca
        # TTS hụt chunk / file bị cắt cụt — audio đã tạo xong nên chỉ CẢNH BÁO (kèm beep)
        # chứ không chặn, để bạn tự quyết định render lại.
        try:
            import kiemtra_daura as kt
            ok_len, msg = kt.check_audio(output, input_txt)
            if not ok_len:
                logging.error(f"⚠️ {folder.name}: {msg}")
                try:
                    import winsound
                    winsound.MessageBeep(winsound.MB_ICONHAND)
                except Exception:
                    pass
        except Exception as e:
            logging.warning(f"⚠️ Bỏ qua chốt kiểm độ dài audio: {e}")
        return True

    def _don_rac_audio(self, folder) -> None:
        """🧹 Xoá audio TRUNG GIAN của 1 tập sau khi video đã xong (~215 MB/tập).

        GIỮ output.wav — mọi file bị xoá (output_chunks/, *_sped*.wav, *_tiktok*.wav,
        tiktok_bgm.wav) đều dựng lại được từ nó bằng ffmpeg trong vài giây, còn
        output.wav thì phải chạy lại TTS hàng chục phút.

        Xoá HẲN, không qua Thùng rác — file vào Thùng rác vẫn chiếm ổ đĩa nên sẽ làm
        tính năng dọn rác thành vô nghĩa. Các chốt an toàn (phải có video, phải còn
        output.wav) nằm trong don_rac.don_rac_audio.

        Lỗi ở bước dọn KHÔNG được làm hỏng tập đã chạy xong → nuốt lỗi nhưng vẫn log.
        """
        try:
            import don_rac
            don_rac.don_rac_audio(folder, on_log=logging.info)
        except Exception as e:
            logging.warning(f"⚠️ Không dọn được audio trung gian của {Path(folder).name}: {e}")

    def _make_thumbnail_for_folder(self, folder, episode: str) -> bool:
        """Render thumbnail cho 1 tập — CẢ 2 bản, dùng CHUNG ảnh mèo + tiêu đề + số tập:
          • NGANG (1920×1080): thumbnail<episode>.png
          • DỌC  (1080×1920): thumbnail<episode>_dọc.png  (chuẩn YouTube Shorts)
        Bản nào đã có thì bỏ qua (tạo tiếp phần còn thiếu). Lỗi bản DỌC KHÔNG chặn
        bản NGANG. Tiêu đề lấy từ folder/seoYoutube.docx."""
        try:
            import random
            youtube_dir = str(YOUTUBE_DIR)
            if youtube_dir not in sys.path:
                sys.path.insert(0, youtube_dir)
            import dien_tieu_de_thumbnail as renderer
            from seo_docx_parser import parse_seo_docx

            out_png = folder / f"thumbnail{episode}.png"        # bản NGANG
            out_doc = folder / f"thumbnail{episode}_dọc.png"    # bản DỌC
            need_ngang = not out_png.exists()
            need_doc = not out_doc.exists()
            if not need_ngang and not need_doc:
                return True   # cả 2 bản đã có → khỏi làm lại

            seo_docx = folder / "seoYoutube.docx"
            title = ""
            if seo_docx.exists():
                try:
                    title = (parse_seo_docx(seo_docx).get("title") or "").strip()
                except Exception as e:
                    logging.warning(f"Không đọc được tiêu đề SEO: {e}")
            if not title:
                logging.warning(f"⚠️ {folder.name}: không có tiêu đề SEO — bỏ qua thumbnail.")
                return False
            # Tiêu đề hợp lệ luôn NGẮN (1 câu đã chọn). Nếu parse SEO lấy nhầm cả đoạn
            # (vd câu mở đầu Gemini "Dưới đây là 5 tiêu đề...") thì title rất dài → BỎ QUA
            # thumbnail thay vì nhồi vào renderer (tránh treo CPU + thumbnail xấu).
            if len(title) > 120 or len(title.split()) > 18:
                logging.warning(
                    f"⚠️ {folder.name}: tiêu đề SEO BẤT THƯỜNG ({len(title.split())} từ, "
                    f"{len(title)} ký tự) — có thể parse nhầm câu mở đầu. BỎ QUA thumbnail. "
                    f"Tiêu đề: {title[:80]}…")
                return False

            photos = renderer.list_photo_files(renderer.CAT_IMAGE_DIR)
            if not photos:
                logging.warning(f"⚠️ Không có ảnh mèo trong {renderer.CAT_IMAGE_DIR} — bỏ qua thumbnail.")
                return False

            photo = random.choice(photos)   # dùng CHUNG cho cả bản ngang & dọc
            made = False
            if need_ngang:
                renderer.add_title(
                    renderer.SOURCE_IMAGE, out_png, title, photo,
                    renderer.FRAME_IMAGE, episode, renderer.NUMBER_FRAME_IMAGE)
                logging.info(f"🖼  Đã tạo thumbnail ngang: {out_png.name}")
                made = True
            # Bản DỌC 1080×1920 — bọc riêng để lỗi bản dọc KHÔNG làm hỏng bản ngang.
            if need_doc:
                try:
                    renderer.add_title_vertical(out_doc, title, photo, episode)
                    logging.info(f"🖼  Đã tạo thumbnail dọc: {out_doc.name}")
                    made = True
                except Exception as e:
                    logging.warning(f"⚠️ {folder.name}: lỗi tạo thumbnail dọc (giữ bản ngang): {e}")
            return made
        except Exception as e:
            logging.error(f"Lỗi tạo thumbnail {folder.name}: {e}")
            return False

    def _seo_copy_blocks(self, seo_docx, episode: str):
        """Đọc seoYoutube.docx → {'title','title_tiktok','desc','tags'} đã chuẩn hóa
        cho 1 tập (None nếu lỗi). Khớp tuyệt đối với các nút Copy của tab Thumbnail:
        tiêu đề YouTube 'Mimi audio Số <tập> | <tên truyện>', bản TikTok thêm 'Full ở'
        phía trước; mô tả thêm hashtag #truyenfull #full; thẻ tag gắn tag tập rồi cắt
        cho tổng < 499 ký tự."""
        try:
            youtube_dir = str(YOUTUBE_DIR)
            if youtube_dir not in sys.path:
                sys.path.insert(0, youtube_dir)
            from seo_docx_parser import parse_seo_docx
            import thumbnail_gui as tg   # dùng đúng hàm của 3 nút để khớp tuyệt đối

            seo = parse_seo_docx(str(seo_docx))
            ep = episode if str(episode).strip().isdecimal() else ""
            # Tiêu đề YouTube 'Mimi audio Số <tập> | <tên truyện>'; bản TikTok thêm
            # 'Full ở' phía trước. Mô tả thêm hashtag #truyenfull #full.
            title = tg.compose_youtube_title(seo.get("title", ""), ep)
            title_tiktok = tg.compose_tiktok_title(seo.get("title", ""), ep)
            desc = tg.add_episode_to_description(seo.get("description", ""), ep)
            desc = tg.add_episode_hashtag_top(desc, ep)   # '#MimiAudioSo<ep>' lên đầu
            desc = tg.add_full_hashtags(desc)

            # Thẻ tag: gắn thẻ tập 'mimi audio số <ep>' + thẻ TÊN TRUYỆN lên đầu rồi cắt
            # cho tổng (đếm kiểu YouTube) < 500 ký tự — LUÔN GIỮ 2 thẻ này; ưu tiên bỏ
            # TRƯỚC các thẻ KHÔNG phải tiếng Việt (chữ Hán / latin không dấu, trừ 'mimi audio').
            raw_tags = seo.get("tags", [])
            raw_title = seo.get("title", "")
            full_list = tg.add_episode_tag(tg.add_title_tag(raw_tags, raw_title), ep)
            tags = tg.cap_tags(raw_tags, ep, raw_title)
            dropped = len(full_list) - len([t for t in tags.split(", ") if t])
            if dropped:
                ep_tag = f"mimi audio số {ep}" if ep else None
                logging.info(f"✂ Thẻ tag ≥{tg.MAX_TAGS_LEN} ký tự → bỏ {dropped} tag cuối "
                             + (f"(giữ '{ep_tag}' + tên truyện)." if ep_tag else "(giữ tên truyện)."))
            return {"title": title or "", "title_tiktok": title_tiktok or "",
                    "desc": desc or "", "tags": tags or ""}
        except Exception as e:
            logging.warning(f"Không đọc được nội dung SEO copy: {e}")
            return None

    def _save_youtube_seo_copy(self, seo_docx, out_path, episode: str) -> bool:
        """Lưu sẵn nội dung 3 nút Copy của tab Thumbnail (tiêu đề · mô tả · thẻ tag)
        ra 1 file .txt để dán nhanh khi đăng YouTube.

        Tab Thumbnail chỉ đọc seoYoutube.docx CHUNG (kịch_bản/) nên khi chạy NHIỀU
        LINK không copy được nội dung từng tập. File này ghi đúng nội dung 3 nút đó
        (tiêu đề 'Full ở Mimi audio Số <tập> | <tên truyện>') cho SEO của riêng tập.
        """
        blocks = self._seo_copy_blocks(seo_docx, episode)
        if not blocks:
            return False
        try:
            content = (
                "===== TIÊU ĐỀ YOUTUBE =====\n" + blocks["title"] + "\n\n"
                "===== TIÊU ĐỀ TIKTOK =====\n" + blocks.get("title_tiktok", "") + "\n\n"
                "===== MÔ TẢ =====\n" + blocks["desc"] + "\n\n"
                "===== THẺ TAG =====\n" + blocks["tags"] + "\n"
            )
            Path(out_path).write_text(content, encoding="utf-8")
            logging.info(f"💾 Đã lưu nội dung copy YouTube (tiêu đề/mô tả/thẻ tag) "
                         f"→ {Path(out_path).name}")
            return True
        except Exception as e:
            logging.warning(f"Không lưu được file copy YouTube SEO: {e}")
            return False

    # ── ĐĂNG YOUTUBE: kiểm tra trước + hàng đợi chạy SONG SONG với dựng video ────
    @staticmethod
    def _upload_log(msg, level="info"):
        """Bộ log cho phần đăng YouTube: đổi mức của dang_video_youtube sang logging."""
        fn = {"err": logging.error, "warn": logging.warning}.get(level, logging.info)
        fn(msg)

    def _upload_check_ready(self) -> bool:
        """Kiểm tra TRƯỚC KHI chạy: đủ thư viện Google API + đã đăng nhập đúng kênh.

        BẮT BUỘC gọi trên MAIN THREAD. Lý do: token hỏng thì get_credentials sẽ MỞ
        TRÌNH DUYỆT đòi đăng nhập — nếu việc đó rơi vào giữa mẻ chạy lúc 2 giờ sáng
        thì cả mẻ đứng im chờ người bấm. Thà hỏi ngay lúc bạn còn ngồi đây.

        Tiện thể đọc luôn danh sách video của kênh: vừa xác nhận token dùng được,
        vừa cho biết ĐĂNG LÊN KÊNH NÀO (1 Gmail có thể có nhiều kênh), vừa nạp cache
        các giờ đã hẹn để xếp lịch 08:00/18:00 không bị trùng.
        """
        try:
            _ensure_youtube_path()
            import dang_video_youtube as yt
        except Exception as e:
            messagebox.showerror("Không nạp được phần đăng YouTube", str(e))
            return False

        missing = yt._check_deps()
        if missing:
            messagebox.showerror(
                "Thiếu thư viện Google API",
                f"Chưa cài thư viện để đăng YouTube:\n{missing}\n\nCài bằng lệnh:\n"
                "pip install google-api-python-client google-auth-oauthlib "
                "google-auth-httplib2")
            return False
        if not yt.CLIENT_SECRET_FILE.exists():
            messagebox.showerror(
                "Thiếu client_secret.json",
                f"Chưa có file:\n{yt.CLIENT_SECRET_FILE}\n\n"
                "Mở app 'Đăng video YouTube' và xem nút 'Hướng dẫn' để lấy file này.")
            return False

        try:
            creds = yt.get_credentials(self._upload_log, interactive=False)
        except Exception as e:
            logging.warning(f"Không đọc được token YouTube: {e}")
            creds = None
        if creds is None:
            if not messagebox.askyesno(
                    "Chưa đăng nhập YouTube",
                    "Chưa có đăng nhập YouTube dùng được (token thiếu/hết hạn).\n\n"
                    "Bấm Yes để đăng nhập NGAY bây giờ — trình duyệt sẽ mở ra và cửa "
                    "sổ này đứng yên tới khi bạn đăng nhập xong.\n\n"
                    "⚠ Một Gmail có thể có NHIỀU KÊNH: hãy chọn đúng kênh muốn đăng."):
                return False
            try:
                yt.get_credentials(self._upload_log)
            except Exception as e:
                messagebox.showerror("Đăng nhập thất bại", str(e))
                return False

        # Đọc kênh + video đã hẹn giờ (nạp cache cho việc xếp lịch, và để báo rõ kênh).
        try:
            chan, _videos = yt.fetch_channel_videos(self._upload_log)
        except Exception as e:
            messagebox.showerror(
                "Không đọc được kênh YouTube",
                f"Đăng nhập được nhưng không đọc được kênh:\n{e}\n\n"
                "Chưa chạy tiếp để tránh đăng nhầm chỗ.")
            return False
        n_sched = len(yt.cached_scheduled_videos())
        logging.info(f"⬆ Sẽ đăng lên kênh: {chan['title']} ({chan['id']}) — "
                     f"{chan['video_count']} video, {n_sched} video đang hẹn giờ.")
        try:
            slot = yt.next_publish_slot()
            logging.info(f"⬆ Tập đầu tiên của mẻ này sẽ hẹn: {slot:%d/%m/%Y %H:%M}")
        except Exception as e:
            logging.warning(f"Không tính được khung giờ đăng: {e}")
        return True

    @staticmethod
    def _episodes_on_channel() -> set:
        """Số tập ĐÃ CÓ trên kênh, tách từ tiêu đề video trong cache ('... Số 47 ...').

        Chốt chặn thứ hai chống đăng trùng: tập đã đăng TAY bằng app 'Đăng video
        YouTube' thì thư mục không có youtube_upload.json, chỉ nhìn file là không
        biết. Cache giữ ~50 video gần nhất — thừa sức phủ các tập đang làm.
        """
        try:
            _ensure_youtube_path()
            import dang_video_youtube as yt
            data = yt.load_video_cache()
            entry = data["channels"].get(data.get("current"))
            return yt.episode_numbers(entry["videos"]) if entry else set()
        except Exception:
            return set()

    def _upload_enqueue(self, folder, episode):
        """Xếp 1 tập vào hàng đợi đăng rồi TRẢ VỀ NGAY.

        Không chờ tải xong: quy trình chạy tiếp tập kế (dựng video dùng GPU, tải lên
        dùng mạng — chạy chồng nhau mới hết thời gian chết). Luồng đăng tự khởi động
        khi có việc và tự kết thúc khi hết hàng đợi."""
        with self._upload_lock:
            self._upload_q.put((Path(folder), str(episode)))
            if self._upload_thread is None:
                self._upload_thread = threading.Thread(target=self._upload_worker,
                                                       daemon=True)
                self._upload_thread.start()
        logging.info(f"⬆ Đã xếp tập {episode} vào hàng đợi đăng "
                     f"({self._upload_q.qsize()} tập đang chờ).")

    def _upload_worker(self):
        """Luồng đăng: xử lý lần lượt từng tập, MỘT video một lúc (tải song song chỉ
        chia nhỏ băng thông chứ không nhanh hơn). Hết hàng đợi thì tự kết thúc —
        việc gỡ cờ _upload_thread nằm trong lock để không bỏ sót tập vừa được xếp
        vào đúng lúc luồng đang thoát."""
        import time
        try:
            while True:
                try:
                    folder, episode = self._upload_q.get(timeout=2)
                except queue.Empty:
                    with self._upload_lock:
                        if self._upload_q.empty():
                            self._upload_thread = None
                            return
                    continue
                try:
                    if self._upload_one(folder, episode):
                        self._upload_done += 1
                except Exception as e:
                    import traceback
                    logging.error(f"❌ Lỗi đăng tập {episode}: {e}")
                    logging.error(traceback.format_exc())
                    self.upload_status.set(f"❌ Lỗi đăng tập {episode} — xem nhật ký.")
                finally:
                    self._upload_q.task_done()
                    self._recog_schedule_table_refresh()
                    time.sleep(0.5)   # nghỉ nhẹ giữa 2 lượt gọi API
        finally:
            # Luồng chết vì lỗi NGOÀI DỰ KIẾN vẫn phải gỡ cờ. Không gỡ thì
            # _upload_enqueue tưởng còn luồng nên không dựng luồng mới, hàng đợi nằm
            # đó mãi và _upload_wait_drain chờ vô tận → treo cả mẻ.
            with self._upload_lock:
                if self._upload_thread is threading.current_thread():
                    self._upload_thread = None

    def _upload_one(self, folder, episode) -> bool:
        """Đăng YOUTUBE.mp4 của 1 tập. Trả về True nếu ĐÃ đăng lượt này."""
        _ensure_youtube_path()
        import dang_tap_youtube as up

        folder = Path(folder)
        if up.already_uploaded(folder):
            logging.info(f"♻ Bỏ qua đăng tập {episode} (đã có "
                         f"{up.RECORD_NAME} — không đăng trùng).")
            return False
        if str(episode).isdecimal() and int(episode) in self._episodes_on_channel():
            logging.warning(f"⚠ Bỏ qua đăng tập {episode}: trên kênh ĐÃ CÓ video "
                            f"'Số {int(episode)}' (nhiều khả năng đã đăng tay) — "
                            "tránh đăng trùng. Muốn đăng lại thì xoá video cũ trên kênh.")
            return False
        # ⛔ CHỐT CUỐI trước khi đăng: bản dịch còn đoạn hỏng (chưa dịch / bị Gemini
        # từ chối / dịch cụt) → TUYỆT ĐỐI KHÔNG ĐĂNG, dù video đã render xong.
        # Đây là lưới cuối cùng: tập 85/87 từng render từ bản dịch thiếu 14-25%
        # nội dung rồi lên thẳng YouTube vì không cửa nào kiểm lại.
        bad = kiem_ban_dich_folder(folder)
        if bad:
            mota = ", ".join(f"đoạn {j} {ly_do}" for j, ly_do in bad)
            logging.error(f"⛔ KHÔNG ĐĂNG tập {episode}: bản dịch có đoạn hỏng "
                          f"({mota}). Video hiện có được dựng từ bản dịch thiếu — "
                          "dịch lại các đoạn hỏng, render lại rồi mới đăng.")
            self.upload_status.set(f"⛔ Tập {episode}: bản dịch hỏng — KHÔNG đăng.")
            return False
        blocks = self._seo_copy_blocks(folder / "seoYoutube.docx", str(episode))
        if not blocks:
            logging.error(f"❌ Tập {episode}: không đọc được SEO → không đăng.")
            return False

        self.upload_status.set(f"⬆ Đang đăng tập {episode}...")
        rec = up.upload_episode(
            folder, episode, blocks, self._upload_log,
            progress_cb=lambda p: self.upload_status.set(
                f"⬆ Đang đăng tập {episode} — {p}%"))
        if rec is None:
            self.upload_status.set(f"⚠ Tập {episode}: chưa đăng (xem nhật ký).")
            return False
        self.upload_status.set(
            f"✅ Tập {episode} đã đăng — công khai {rec['publish_at_text']}")
        return True

    def _upload_wait_drain(self):
        """Chờ hàng đợi đăng chạy hết. Gọi ở CUỐI mẻ, TRƯỚC khi báo xong/hẹn tắt máy —
        tắt máy giữa lúc còn video đang tải lên là mất công cả mẻ.

        ⚠ CHỈ gọi từ LUỒNG NỀN, tuyệt đối không từ luồng Tk. Hàm này chặn luồng gọi
        nó; mà luồng đăng lại phải cập nhật StringVar — lời gọi Tcl từ luồng khác chỉ
        được phục vụ khi luồng Tk đang chạy mainloop. Chặn luồng Tk ở đây là kẹt cả
        hai bên: luồng đăng đứng chờ Tcl, còn hàm này đứng chờ luồng đăng.
        """
        import time
        if self._upload_q.empty() and self._upload_thread is None:
            return
        logging.info("⬆ Chờ tải nốt các video còn trong hàng đợi đăng...")
        while True:
            with self._upload_lock:
                left = self._upload_q.qsize()
                t = self._upload_thread
                alive = t is not None and t.is_alive()
                if left == 0 and not alive:
                    break
                if left and not alive:
                    # Luồng đăng chết bất thường mà hàng đợi còn việc: dựng lại luồng
                    # thay vì chờ mãi một luồng đã chết.
                    logging.warning("⬆ Luồng đăng đã dừng bất thường — chạy lại.")
                    self._upload_thread = threading.Thread(target=self._upload_worker,
                                                           daemon=True)
                    self._upload_thread.start()
            self.pipe_link_status.set(
                f"⬆ Đang đăng nốt {left + 1} video (dựng video đã xong)...")
            time.sleep(2)
        logging.info(f"⬆ Đã đăng xong {self._upload_done} video trong mẻ này.")

    def _seo_docx_valid(self, seo_docx) -> bool:
        """True nếu seoYoutube.docx đã có nội dung SEO thật (tiêu đề khác rỗng).

        Dùng để TIẾP TỤC: file rỗng/chỉ có tiêu đề (do reset) → coi như chưa làm SEO.
        """
        try:
            if not Path(seo_docx).exists():
                return False
            youtube_dir = str(YOUTUBE_DIR)
            if youtube_dir not in sys.path:
                sys.path.insert(0, youtube_dir)
            from seo_docx_parser import parse_seo_docx
            return bool((parse_seo_docx(str(seo_docx)).get("title") or "").strip())
        except Exception:
            return False

    def _seo_title_duplicate(self, folder, episode, title) -> str:
        """Tập KHÁC đang mang đúng TÊN TRUYỆN này → mô tả chỗ trùng; '' nếu sạch.

        Mỗi tập là một truyện khác nhau, nên trùng tên truyện = bước SEO lấy nhầm
        kết quả của tập khác (đã dính: tập 55–58 / 49 / 08 cùng một tiêu đề vì đọc
        trúng câu trả lời cũ trong chat Gemini). Dò cả hai nơi:
          • kênh YouTube (theo cache đọc gần nhất) — bắt tập ĐÃ ĐĂNG,
          • các thư mục tập khác trong kịch_bản/ — bắt tập dựng rồi mà chưa đăng.
        So bằng dang_video_youtube.title_key nên 'Số N', 'Full ở', 'Mimi audio',
        '#Shorts' không ảnh hưởng.
        """
        try:
            youtube_dir = str(YOUTUBE_DIR)
            if youtube_dir not in sys.path:
                sys.path.insert(0, youtube_dir)
            import dang_video_youtube as yt
            from seo_docx_parser import parse_seo_docx
        except Exception as e:
            logging.warning(f"⚠️ Không dò được tiêu đề trùng: {e}")
            return ""

        key = yt.title_key(title)
        if not key:
            return ""
        try:
            dup = yt.duplicate_title_on_channel(title, episode)
            if dup:
                return f"kênh đã có “{dup.get('title')}”"
        except Exception as e:
            logging.warning(f"⚠️ Không dò được tiêu đề trùng trên kênh: {e}")

        ep = str(episode).strip().zfill(2)
        for p in episode_dirs():
            ep_khac = episode_of(p.name)
            if Path(p) == Path(folder) or ep_khac == ep:
                continue
            docx = p / "seoYoutube.docx"
            if not docx.exists():
                continue
            try:
                khac = parse_seo_docx(str(docx)).get("title") or ""
            except Exception:
                continue
            if khac and yt.title_key(khac) == key:
                return f"tập {ep_khac} ({p.name}) cũng mang tên truyện này"
        return ""

    def _dich_gemini_cho_tap(self, gemini_docx, chunks, prefix, episode,
                             driver=None, on_status=None):
        """Dịch 1 tập sang Gemini: bỏ qua đoạn đã dịch, gửi phần thiếu, rồi KIỂM ĐỦ ĐOẠN.

        Gộp 3 khối GIỐNG HỆT nhau trước đây nằm ở _resume_folder_worker /
        _pipe_batch_worker / _batch_worker. Trước đây sửa 1 lỗi phải nhớ sửa cả 3 chỗ —
        đúng cái làm lỗi tập 42 tồn tại được (chốt _translation_complete bị vô hiệu ở
        cả 3 bản sao). Nay chỉ còn 1 nơi để sửa.

        on_status(i, total): callback cập nhật dòng trạng thái GUI (None = không cập nhật).

        Trả về (driver, translated_now, translation_ok):
          • driver         — driver ĐANG dùng; có thể là bản MỚI nếu Firefox phải mở lại
                             giữa chừng, nên bên gọi PHẢI gán lại biến driver của mình.
          • translated_now — lượt này có thực sự gửi Gemini không (→ cần tạo lại input.txt)
          • translation_ok — đã dịch đủ mọi đoạn chưa (False → KHÔNG tạo input/audio/video)
        """
        import time
        import dich_gemini as g

        # Firefox có thể bị đóng/mở lại trong send_chunks_to_gemini → giữ driver trong
        # dict để callback cập nhật được, rồi trả bản mới nhất về cho bên gọi.
        state = {"driver": driver}

        def _on_driver(d):
            state["driver"] = d

        prior = (g.read_results_docx(gemini_docx, len(chunks))
                 if gemini_docx.exists() else [None] * len(chunks))
        # Đoạn HỎNG phải gửi lại (bộ tiêu chí chung g.bad_chunks: chưa dịch / câu
        # TỪ CHỐI của Gemini / dịch CỤT — ca tập 85-87).
        todo = [j for j, _ly_do in g.bad_chunks(chunks, prior)]
        if SKIP_TRANSLATE_DETAIL_CHECK and gemini_docx.exists() and not todo:
            # Đã có gemini_result.docx đủ đoạn lành → KHÔNG gửi lại Gemini (tránh
            # dịch lại đoạn đã xong). Xem cờ SKIP_TRANSLATE_DETAIL_CHECK ở đầu file.
            translated_now = False
            logging.info(f"♻ Bỏ qua gửi Gemini — đã có {gemini_docx.name}.")
        else:
            translated_now = bool(todo)
            if not todo:
                logging.info(f"♻ Bỏ qua dịch Gemini (đã đủ {len(chunks)} đoạn).")
            else:
                if gemini_docx.exists():
                    logging.info(f"⚠️ {gemini_docx.name} có đoạn hỏng {todo} (từ chối/"
                                 "dịch cụt/chưa dịch) → gửi lại các đoạn đó.")
                if state["driver"] is None:
                    logging.info("🌐 Mở Firefox + Gemini...")
                    state["driver"] = g.init_firefox()
                else:
                    state["driver"].get(g.GEMINI_URL)   # chat mới cho tập này
                    time.sleep(8)
                logging.info(f"🌐 Tập {episode}: dịch {n_todo}/{len(chunks)} đoạn còn thiếu...")
                results = g.send_chunks_to_gemini(
                    chunks, prefix=prefix, on_log=logging.info, out_path=gemini_docx,
                    driver=state["driver"], keep_open=True, on_driver=_on_driver,
                    resume=True,
                    on_result=((lambda i2, t2, _a: on_status(i2, t2)) if on_status else None))
                g.save_results_docx(chunks, results, gemini_docx)
                logging.info(f"💾 Đã lưu bản dịch: {gemini_docx}")

        # ⛔ CHẶN: LUÔN kiểm đủ đoạn, kể cả khi đã bỏ qua bước gửi ở trên. Bỏ chốt này
        # là ra audio/video thiếu nội dung mà không ai biết (lỗi tập 42).
        translation_ok = self._translation_complete(gemini_docx, chunks, episode)
        return state["driver"], translated_now, translation_ok

    def _translation_complete(self, gemini_docx, chunks, episode) -> bool:
        """True nếu MỌI đoạn trong gemini_result.docx đã dịch xong (không rỗng, không
        '(chưa dịch)', hết chữ Hán). Thiếu đoạn nào thì ghi log rõ để biết mà sửa.

        Ngoài ra: nếu còn MỘT ĐOẠN HÁN LIÊN TIẾP đủ dài (Gemini bỏ sót nguyên câu) thì
        cũng trả False → BỎ QUA cả tập, sang tập khác (dù tỉ lệ chữ Hán toàn đoạn thấp)."""
        import dich_gemini as g
        n = len(chunks)
        check = g.read_results_docx(gemini_docx, n) if Path(gemini_docx).exists() else [None] * n
        # Bộ tiêu chí chung g.bad_chunks: chưa dịch / câu TỪ CHỐI của Gemini /
        # bản dịch CỤT — cùng bộ với _dich_gemini_cho_tap và các chốt input/tts/đăng.
        bad = g.bad_chunks(chunks, check)
        if bad:
            head = ", ".join(f"đoạn {j} {ly_do}" for j, ly_do in bad[:10]) \
                   + ("..." if len(bad) > 10 else "")
            logging.error(
                f"⛔ Tập {episode}: bản dịch HỎNG {len(bad)}/{n} đoạn ({head}) → DỪNG, "
                "KHÔNG tạo audio/video. Chạy lại để dịch tiếp; nếu Gemini cứ lỗi 1 "
                "đoạn, sửa tay đoạn đó trong gemini_result.docx rồi chạy lại.")
            return False

        # Còn MỘT ĐOẠN HÁN LIÊN TIẾP đủ dài (>= MAX_CHINESE_RUN chữ) → Gemini bỏ sót
        # nguyên câu. BỎ QUA cả tập, sang tập khác (khỏi tạo input/SEO/audio/video).
        worst_j, worst_run = 0, 0
        for j, r in enumerate(check, 1):
            run = g.max_chinese_run(r or "")
            if run > worst_run:
                worst_j, worst_run = j, run
        if worst_run >= g.MAX_CHINESE_RUN:
            logging.error(
                f"⛔ Tập {episode}: đoạn {worst_j} còn CHUỖI HÁN LIÊN TIẾP {worst_run} chữ "
                f"(>= {g.MAX_CHINESE_RUN}) — Gemini bỏ sót nguyên câu → BỎ QUA cả tập, "
                "chuyển sang tập khác.")
            return False

        return True

    # ── MANIFEST: tiến độ + map nguồn↔tập (để chạy tiếp & báo cáo) ──────────────
    def _folder_steps(self, folder, episode: str) -> dict:
        """Bước ĐÃ XONG của 1 tập, suy từ file thực tế trong thư mục."""
        folder = Path(folder)
        zh = find_zh_docx(folder)
        gem = folder / "gemini_result.docx"
        inp = folder / "input.txt"
        translate_done = False
        if gem.exists():
            try:
                import dich_gemini as g
                chunks = read_zh_docx_chunks(zh) if zh else []
                if chunks:
                    prior = g.read_results_docx(gem, len(chunks))
                    # Cùng bộ tiêu chí đoạn hỏng với _translation_complete (từ chối/
                    # dịch cụt cũng tính là CHƯA xong).
                    translate_done = not g.bad_chunks(chunks, prior)
                else:
                    translate_done = True   # không rõ số đoạn → coi gem tồn tại là xong
            except Exception:
                translate_done = True
        return {
            "recognize": bool(zh),
            "translate": translate_done,
            "input": inp.exists() and inp.stat().st_size > 0,
            "seo": self._seo_docx_valid(folder / "seoYoutube.docx"),
            "thumbnail": (folder / f"thumbnail{episode}.png").exists(),
            "audio": (folder / "output.wav").exists(),
            # Đã đăng YouTube chưa — suy từ bản ghi youtube_upload.json của tập.
            "upload": (folder / "youtube_upload.json").exists(),
            # Bản tự động đặt tên YOUTUBE.mp4 / facebook.mp4; sau khi đăng YouTube
            # thì bản dọc thành "facebook <ngày giờ>.mp4" (xem dang_tap_youtube.
            # rename_doc). Vẫn nhận tên cũ (*_videodone.mp4 / *_doc.mp4) cho các
            # tập tạo trước đây.
            "video_ngang": (folder / "YOUTUBE.mp4").exists()
                            or bool(list(folder.glob("*_videodone.mp4"))),
            "video_doc": (folder / "facebook.mp4").exists()
                          or bool(list(folder.glob("facebook *.mp4")))
                          or bool(list(folder.glob("*_doc.mp4"))),
        }

    def _manifest_update(self, source, episode, folder, done=None) -> None:
        """Ghi/cập nhật 1 mục manifest (nguồn→tập + tiến độ) rồi lưu ngay."""
        import datetime as _dt
        m = load_manifest()
        key = norm_source(source)
        entry = m.get(key, {})
        entry["source"] = key
        entry["episode"] = str(episode)
        entry["folder"] = str(folder)
        entry["steps"] = self._folder_steps(folder, episode)
        if done is not None:
            entry["done"] = bool(done)
        entry["updated"] = _dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        m[key] = entry
        save_manifest(m)

    def _allocate_episode(self, source):
        """(episode, is_new). Nguồn đã có trong manifest → ĐÚNG tập cũ; nguồn mới →
        số tập kế tiếp (không trùng manifest / thư mục đã có / số đã lưu, và nhảy
        qua các số trong danh sách 'tập bỏ qua')."""
        m = load_manifest()
        key = norm_source(source)
        entry = m.get(key)
        if entry is None:
            # Manifest CŨ khoá theo chuỗi THÔ (trước khi norm_source chuẩn hoá path).
            # Dò mục nào re-chuẩn-hoá ra CÙNG khoá → coi là cùng nguồn; dời sang khoá
            # mới cho gọn (tự vá dần). Nhờ vậy file local đã làm trước đây, dù gõ khác
            # kiểu, vẫn map về ĐÚNG tập cũ thay vì bị cấp tập mới + làm lại.
            for old_key in list(m):
                if old_key != key and norm_source(old_key) == key:
                    entry = m.pop(old_key)
                    m[key] = entry
                    save_manifest(m)
                    break
        if entry is not None and str(entry.get("episode", "")).isdecimal():
            return str(entry["episode"]).zfill(2), False
        # Nhảy qua các số đã đặt trước trong danh sách "tập bỏ qua".
        return str(next_episode_number(self._last_used_episode())).zfill(2), True

    @staticmethod
    def _last_used_episode() -> int:
        """Số tập LỚN NHẤT đã dùng: manifest + thư mục tập có sẵn + số đã lưu. Chưa
        có tập nào → 0 (nên tập đầu tiên sẽ là 1)."""
        used = {int(v["episode"]) for v in load_manifest().values()
                if str(v.get("episode", "")).isdecimal()}
        for p in episode_dirs():
            used.add(int(episode_of(p.name)))
        used.add(load_episode_number())
        return max(used) if used else 0

    def _pipe_batch_worker(self, sources, model, speed, tts_settings=None, upload=False):
        import time
        driver = None
        total = len(sources)
        ok_count = 0
        self._upload_done = 0

        # Gemini treo → dịch_gemini đóng & mở lại Firefox; tham chiếu driver mới nay do
        # _dich_gemini_cho_tap trả về và được gán lại ngay tại chỗ gọi, nên SEO và link
        # kế (driver.get / driver.quit) vẫn dùng đúng trình duyệt đang mở.

        # ④ Ghi nhật ký batch ra kịch_bản/batch_log.txt (append, giữ lịch sử) để sau
        # sự cố mở file là biết dừng ở link/bước nào.
        file_handler = None
        try:
            import datetime as _dt
            file_handler = logging.FileHandler(SCRIPT_DIR / "batch_log.txt", encoding="utf-8")
            file_handler.setFormatter(
                logging.Formatter("%(asctime)s  %(message)s", "%Y-%m-%d %H:%M:%S"))
            logging.getLogger().addHandler(file_handler)
            logging.info("\n" + "═" * 12 + f" BẮT ĐẦU BATCH {total} link "
                         f"({_dt.datetime.now():%Y-%m-%d %H:%M:%S}) " + "═" * 12)
        except Exception as e:
            logging.warning(f"Không mở được batch_log.txt: {e}")

        try:
            import nhandien_giongnoi as recog
            import dich_gemini as g
            youtube_dir = str(YOUTUBE_DIR)
            if youtube_dir not in sys.path:
                sys.path.insert(0, youtube_dir)
            import seo_youtube_gemini as seo
            prefix = load_prefix()
            logging.info(f"📚 NHIỀU LINK: {total} link — mỗi link 1 tập (theo manifest).")
            logging.info(f"📦 Model: {model}  •  Tốc độ: {speed}x")

            for i, src in enumerate(sources, 1):
                # ⏸/⏹ ĐIỂM AN TOÀN GIỮA CÁC LINK: đang tạm dừng thì CHỜ tại đây; đã bấm
                # dừng thì THOÁT vòng (link trước đã hoàn tất, không bắt đầu link mới).
                self._batch_pause_wait()
                if self._batch_stop_evt.is_set():
                    logging.info(f"⏹ ĐÃ DỪNG theo yêu cầu — xong {ok_count}/{total} link, "
                                 f"còn {total - i + 1} link chưa chạy.")
                    self.pipe_link_status.set(
                        f"⏹ Đã dừng — xong {ok_count}/{total} link (còn {total - i + 1}).")
                    break
                # Số tập theo MANIFEST: nguồn cũ → đúng tập cũ (chạy tiếp); nguồn mới
                # → cấp số kế tiếp. Nhờ vậy nhập khác thứ tự / thiếu link vẫn đúng tập.
                episode, is_new = self._allocate_episode(src)
                episode_num = int(episode)
                # Tên thư mục = "<số tập> - <tên nguồn>" (tập cũ giữ nguyên tên cũ).
                folder = episode_dir_for(episode, src)
                folder.mkdir(parents=True, exist_ok=True)
                # Ghi nguồn↔tập vào manifest NGAY (kể cả lỗi sau đó vẫn biết link→tập).
                self._manifest_update(src, episode, folder, done=False)
                logging.info(f"🔖 {'Tập MỚI' if is_new else 'Tiếp tục tập'} {episode} "
                             f"← {norm_source(src)[:70]}")
                # Dòng ĐANG CHẠY (luôn hiển thị dưới thanh tiến trình) — giữ nguyên
                # suốt link, không bị thông báo bước con (Gemini/giọng/video) ghi đè.
                self.pipe_link_status.set(f"🔗 Đang chạy: Link {i}/{total} — tập {episode}")
                self.pipe_status.set(f"🔗 Link {i}/{total} → tập {episode}/")
                logging.info(f"════════ 🔗 LINK {i}/{total} → TẬP {episode}/ ════════")
                try:
                    s = src.strip().strip('"').strip("'")
                    gemini_docx = folder / "gemini_result.docx"
                    input_txt = folder / "input.txt"
                    seo_docx = folder / "seoYoutube.docx"

                    # ── 1+2) NHẬN DIỆN — bỏ qua nếu đã có *_zh.docx hợp lệ ──────
                    # (chỉ tải MP3 + chạy Whisper khi THỰC SỰ cần nhận diện lại).
                    existing_zh = find_zh_docx(folder)
                    chunks = read_zh_docx_chunks(existing_zh) if existing_zh else []
                    if chunks:
                        logging.info(f"♻ Bỏ qua tải + nhận diện (đã có {existing_zh.name}, "
                                     f"{len(chunks)} đoạn).")
                    else:
                        if os.path.isfile(s):
                            media = s
                            logging.info(f"📁 File local: {media}")
                        elif s.lower().startswith(("http://", "https://")):
                            logging.info(f"🌐 Tải từ link: {s}")
                            media = download_audio_mp3(s, DOWNLOAD_DIR)
                            if not media:
                                logging.error(f"❌ Link {i}: không tải được audio — bỏ qua.")
                                continue
                        else:
                            logging.error(f"❌ Link {i}: không hợp lệ — bỏ qua.")
                            continue
                        self.pipe_progress.set(0)
                        transcript = recog.transcribe_chinese(
                            media, model_name=model, speed=float(speed),
                            on_progress=lambda f: self.pipe_progress.set(int(f * 100)))
                        if not transcript:
                            logging.error(f"❌ Link {i}: không nhận diện được — bỏ qua.")
                            continue
                        zh_docx = folder / ZH_DOCX_NAME   # tên cố định, không theo tên video
                        recog.save_docx(transcript, str(zh_docx), title=Path(media).name)
                        logging.info(f"💾 Đã lưu bản nhận diện: {zh_docx}")
                        chunks = recog.split_into_chunks(transcript)
                    if not chunks:
                        logging.warning(f"⚠️ Link {i}: không có đoạn nào — bỏ qua dịch/SEO.")
                        continue

                    self._batch_pause_wait()          # ⏸ điểm tạm dừng trước khi dịch
                    # ── 3) DỊCH GEMINI — đủ thì bỏ qua; thiếu thì TIẾP TỤC dịch ──
                    # ⛔ CHẶN: DỊCH CHƯA XONG thì KHÔNG tạo input/audio/video (tránh tình
                    # trạng như tập 29 / tập 42: dịch dở vẫn ra audio/video thiếu thời
                    # lượng). Lần sau chạy lại (đặt SKIP_TRANSLATE_DETAIL_CHECK = False)
                    # sẽ dịch tiếp các đoạn còn thiếu.
                    driver, translated_now, translation_ok = self._dich_gemini_cho_tap(
                        gemini_docx, chunks, prefix, episode, driver,
                        on_status=lambda i2, t2: self.pipe_status.set(
                            f"🌐 Link {i}/{total} • Gemini {i2 + 1}/{t2}"))
                    self._manifest_update(src, episode, folder)   # ghi tiến độ sau dịch

                    if not translation_ok:
                        self.pipe_status.set(f"⛔ Tập {episode}: dịch chưa xong — bỏ qua.")
                        continue

                    # ── 4) input.txt — TẠO LẠI nếu vừa dịch (bản cũ có thể dở), hoặc
                    # chưa có. Tạo lại → chữ ký đổi → audio/video tự render lại đúng.
                    if not translated_now and input_txt.exists() and input_txt.stat().st_size > 0:
                        logging.info("♻ Bỏ qua tạo input.txt (đã có).")
                    elif self._batch_prepare_input(gemini_docx, input_txt):
                        logging.info(f"💾 Đã tạo: {input_txt}")

                    self._batch_pause_wait()          # ⏸ điểm tạm dừng trước khi SEO
                    # ── 5) SEO YouTube — bỏ qua nếu seoYoutube.docx đã có tiêu đề ─
                    if self._seo_docx_valid(seo_docx):
                        logging.info("♻ Bỏ qua SEO (đã có seoYoutube.docx hợp lệ).")
                    else:
                        if driver is None:        # dịch đã bỏ qua → mở Firefox cho SEO
                            logging.info("🌐 Mở Firefox cho SEO...")
                            driver = g.init_firefox()
                        logging.info("🔎 Tạo SEO YouTube...")
                        seo.run(str(gemini_docx), str(seo_docx),
                                keep_open=True, log=logging.info, driver=driver)
                        if not self._seo_docx_valid(seo_docx):
                            logging.error(f"⛔ Tập {episode}: không có SEO hợp lệ → BỎ QUA "
                                          "tập này (thumbnail/tên video đều lấy tiêu đề "
                                          "từ SEO). Chạy lại để làm SEO mới.")
                            continue
                        logging.info(f"💾 Đã tạo: {seo_docx}")

                    # 5b) Nội dung 3 nút Copy (tiêu đề/mô tả/thẻ tag) ra .txt — LUÔN
                    # tạo lại (nhẹ, suy ra từ seoYoutube.docx) để áp dụng logic mới nhất
                    # (vd cắt thẻ tag ≤500 ký tự) kể cả khi các bước khác đã bỏ qua.
                    self._save_youtube_seo_copy(
                        seo_docx, folder / "youtube_seo.txt", episode)

                    # 5c) Chốt sớm: tiêu đề trùng tập khác = SEO lấy nhầm kết quả cũ.
                    # Bắt TRƯỚC khi dựng thumbnail/video cho khỏi hỏng cả mẻ.
                    blocks_seo = self._seo_copy_blocks(seo_docx, episode) or {}
                    trung = self._seo_title_duplicate(
                        folder, episode, blocks_seo.get("title") or "")
                    if trung:
                        logging.error(f"⛔ Tập {episode}: tiêu đề SEO trùng — {trung}. "
                                      "BỎ QUA tập này (chạy lại để làm SEO mới).")
                        try:
                            parked = seo.park_docx(seo_docx)
                            if parked:
                                logging.info(f"📦 Đã cất bản SEO nghi sai → {parked.name}")
                        except Exception as e:
                            logging.warning(f"⚠️ Không cất được {seo_docx.name}: {e}")
                        continue

                    # ── 6) Thumbnail (ngang + dọc) — bỏ qua nếu đã có CẢ 2 bản ──
                    if (folder / f"thumbnail{episode}.png").exists() \
                            and (folder / f"thumbnail{episode}_dọc.png").exists():
                        logging.info("♻ Bỏ qua thumbnail (đã có cả ngang & dọc).")
                    else:
                        self._make_thumbnail_for_folder(folder, episode)
                    # Xong thumbnail → cập nhật SỐ TẬP (không lùi) + ghi tiến độ manifest.
                    save_episode_number(max(load_episode_number(), episode_num))
                    self._manifest_update(src, episode, folder)

                    # 7) KHÔNG tự động tải input.txt lên Drive nữa (theo yêu cầu).
                    # Muốn tải thủ công thì dùng ô tick ở tab Thumbnail.

                    # ── 8) TẠO GIỌNG + VIDEO NGAY cho link này (tuần tự từng link) ──
                    # Mỗi link chạy ĐẦY ĐỦ: dịch → SEO → video xong mới sang link kế.
                    # Đóng Firefox trước khi render để nhả RAM (video chỉ dùng GPU);
                    # link sau tự mở lại Firefox cho bước dịch.
                    self._batch_pause_wait()          # ⏸ điểm tạm dừng trước khi tạo giọng/video
                    if tts_settings:
                        if driver is not None:
                            try:
                                driver.quit()
                                logging.info("🦊 Đã đóng Firefox trước khi tạo video.")
                            except Exception:
                                pass
                            driver = None
                        # Nhả Whisper trước khi nạp OmniVoice (GPU 8GB không chứa cả 2).
                        try:
                            recog.free_model()
                            logging.info("🧹 Giải phóng Whisper trước khi tạo giọng.")
                        except Exception:
                            pass
                        self.pipe_status.set(f"🎧 Tập {episode}: đang tạo giọng + video...")
                        self._batch_run_tts(folder, tts_settings, episode)

                    self._don_rac_audio(folder)   # 🧹 xoá audio trung gian (giữ output.wav)
                    self._manifest_update(src, episode, folder, done=True)  # link XONG
                    # ── 9) ĐĂNG YOUTUBE — chỉ XẾP HÀNG rồi đi tiếp: luồng đăng tải
                    # lên trong lúc link kế đang dựng video (mạng ↔ GPU chạy chồng).
                    if upload:
                        self._upload_enqueue(folder, episode)
                    ok_count += 1
                    done_what = "dịch + SEO + video" if tts_settings else "dịch + SEO"
                    logging.info(f"✅ Link {i}/{total} (tập {episode}) HOÀN TẤT ({done_what}).")
                except Exception as e:
                    # Một link lỗi không làm hỏng cả batch — ghi log (kèm traceback để
                    # dễ tìm nguyên nhân) rồi sang link kế.
                    import traceback
                    logging.error(f"❌ Lỗi ở link {i}/{total}: {e}")
                    logging.error(traceback.format_exc())
                    continue

            # Số tập đã cập nhật theo TỪNG link (sau thumbnail). Manifest giữ map
            # nguồn↔tập + tiến độ → chạy lại (chưa xóa output) sẽ TIẾP TỤC đúng tập,
            # kể cả khi nhập khác thứ tự / thiếu link.
            logging.info(f"🔢 Xong vòng batch: {ok_count}/{total} link hoàn tất.")

            # Mọi link đã chạy ĐẦY ĐỦ TUẦN TỰ (dịch → SEO → video) ngay trong vòng lặp.
            self.pipe_progress.set(100)
            self.pipe_link_status.set(f"✅ Hoàn tất {ok_count}/{total} link.")
            self.pipe_status.set(f"✅ Xong {ok_count}/{total} link → thư mục trong kịch_bản.")
            logging.info(f"🎉 XONG: {ok_count}/{total} link hoàn tất.")
        except Exception as e:
            logging.error(f"Lỗi batch nhiều link: {e}")
            self.pipe_link_status.set(f"❌ Lỗi batch (đã xong {ok_count}/{total} link).")
            self.pipe_status.set(f"Lỗi batch: {e}")
        finally:
            if driver is not None:          # đóng Firefox dùng chung sau khi xong/lỗi
                try:
                    driver.quit()
                except Exception:
                    pass
            try:
                import nhandien_giongnoi as recog
                recog.free_model()
                logging.info("🧹 Đã giải phóng model nhận diện khỏi VRAM.")
            except Exception:
                pass
            # Chờ đăng nốt TRƯỚC khi hạ cờ bận: _pipe_set_busy(False) là chỗ móc lệnh
            # hẹn TẮT MÁY — tắt lúc còn video đang tải lên thì mất công cả mẻ.
            self._upload_wait_drain()
            if file_handler is not None:        # ngừng ghi nhật ký batch ra file
                try:
                    logging.info("──────── KẾT THÚC BATCH ────────")
                    logging.getLogger().removeHandler(file_handler)
                    file_handler.close()
                except Exception:
                    pass
            self._pipe_set_busy(False)
            self._batch_controls_reset()   # tắt nút Tạm dừng/Dừng khi batch kết thúc
            self._recog_schedule_table_refresh()   # cập nhật bảng trạng thái tab Nhận diện

    def _pipe_send_gemini(self, auto=False):
        if self._pipe_busy:
            return
        if not auto:
            self._save_pipe_settings()   # ấn chạy → nhớ cài đặt
        if not CHINESE_DOCX.exists():
            if auto:
                logging.error("⛓ Tự động dừng: chưa có tiengTrung.docx.")
                self.pipe_status.set("❌ Chưa có tiengTrung.docx — dừng tự động.")
                return
            messagebox.showwarning(
                "Chưa có nội dung",
                f"Chưa thấy {CHINESE_DOCX.name}.\nHãy bấm ① Nhận diện trước "
                "(hoặc đặt file tiengTrung.docx vào thư mục kịch_bản).")
            return
        # Khi chạy tự động (chuỗi) thì bỏ qua hộp hỏi xác nhận cho liền mạch.
        if not auto and not messagebox.askyesno(
                "Gửi Gemini",
                "Sẽ mở Firefox và gửi nội dung sang Gemini.\n\n"
                "Hãy ĐÓNG Firefox đang mở (nếu có) và đảm bảo profile đã đăng nhập "
                "Google.\n\nTiếp tục?"):
            return
        if auto:
            logging.info("⛓ Tự động: gửi Gemini (bỏ qua hỏi xác nhận).")
        self._pipe_set_busy(True)
        self.pipe_progress.set(0)
        self.pipe_status.set("🌐  Đang gửi Gemini...")
        threading.Thread(target=self._pipe_gemini_worker, daemon=True).start()

    def _pipe_gemini_worker(self):
        ok = False
        seo_on = self.var_seo.get()
        driver = None

        # Khi dịch_gemini phải đóng & mở lại Firefox (Gemini treo) thì cập nhật lại
        # tham chiếu driver ở đây để bước SEO dùng đúng Firefox đang mở.
        def _on_driver(d):
            nonlocal driver
            driver = d

        try:
            chunks = read_chinese_docx_chunks(CHINESE_DOCX)
            if not chunks:
                logging.error("❌ Không đọc được nội dung tiếng Trung để gửi.")
                self.pipe_status.set("❌ Nội dung trống.")
                return
            import dich_gemini as g
            prefix = load_prefix()
            # Nếu bật SEO: tự mở MỘT Firefox và DÙNG CHUNG cho cả dịch lẫn SEO. Tránh
            # đóng Firefox sau khi dịch rồi mở lại cho SEO — lần mở thứ hai hay kẹt
            # khóa profile khiến SEO không chạy được.
            if seo_on:
                logging.info("🌐 Mở Firefox (dùng chung cho dịch + SEO)...")
                driver = g.init_firefox()
            logging.info(f"🌐 Gửi {len(chunks)} đoạn sang Gemini...")
            results = g.send_chunks_to_gemini(
                chunks, prefix=prefix, on_log=logging.info, out_path=GEMINI_DOCX,
                driver=driver,                 # None → tự mở; có driver → tái dùng
                keep_open=(driver is not None),  # còn SEO ở sau thì giữ Firefox mở
                on_driver=_on_driver,          # đóng/mở lại Firefox → cập nhật driver
                on_result=lambda i, total, ans: self.pipe_status.set(
                    f"🌐 Gemini: đoạn {i + 1}/{total}"))
            g.save_results_docx(chunks, results, GEMINI_DOCX)
            n_ok = sum(1 for r in results if r and r.strip())
            logging.info(f"✅ Gemini xong: {n_ok}/{len(results)} đoạn → {GEMINI_DOCX.name}")
            self.pipe_status.set(f"✅ Gemini xong → {GEMINI_DOCX.name}")
            ok = True
            # Ngay sau khi có nội dung Gemini → tạo SEO YouTube (tiêu đề/mô tả/hashtag)
            # nếu người dùng bật. SEO lỗi KHÔNG làm hỏng quy trình (ok đã True nên
            # vẫn chạy tiếp bước ③).
            if seo_on:
                self._run_seo_youtube(driver=driver)
        except Exception as e:
            logging.error(f"Lỗi gửi Gemini: {e}")
            self.pipe_status.set(f"Lỗi: {e}")
        finally:
            if driver is not None:          # đóng Firefox dùng chung sau khi xong/ lỗi
                try:
                    driver.quit()
                except Exception:
                    pass
            self._pipe_set_busy(False)
            if ok and self.var_auto3.get():   # ⛓ tự động sang bước ③
                self.after(600, lambda: self._pipe_prepare_input(auto=True))

    def _run_seo_youtube(self, driver=None):
        """Chạy SEO YouTube (Gemini) ngay sau bước lấy nội dung Gemini.

        Lấy ĐOẠN ĐẦU của gemini_result.docx, gửi lên cuộc trò chuyện Gemini chuyên
        SEO YouTube (đã có sẵn chỉ dẫn) rồi lưu seoYoutube.docx. driver: tái dùng
        Firefox đang mở (do bước dịch mở) để khỏi mở lại. Mọi lỗi đều được nuốt để
        không chặn các bước tiếp theo của quy trình."""
        try:
            youtube_dir = str(YOUTUBE_DIR)
            if youtube_dir not in sys.path:
                sys.path.insert(0, youtube_dir)
            import seo_youtube_gemini as seo
            self.pipe_status.set("🔎  Đang tạo SEO YouTube (Gemini)...")
            logging.info("🔎 Tạo SEO YouTube từ gemini_result.docx...")
            # Có driver → keep_open=True để seo KHÔNG tự đóng (worker đóng sau cùng).
            seo.run(str(GEMINI_DOCX), str(SEO_DOCX),
                    keep_open=(driver is not None), log=logging.info, driver=driver)
            self.pipe_status.set(f"✅ SEO YouTube xong → {SEO_DOCX.name}")

            # Luồng 1 LINK không có bước thumbnail như batch (batch làm ở bước 6) →
            # tạo LUÔN ở đây, ngay sau khi có SEO (nguồn tiêu đề). Số tập lấy từ tab
            # Thumbnail (đã tự lưu vào state → đọc file, an toàn đa luồng); 0 → bỏ qua.
            ep = load_episode_number()
            if ep > 0:
                self.pipe_status.set(f"🖼  Đang tạo thumbnail tập {ep:02d}...")
                if self._make_thumbnail_for_folder(SCRIPT_DIR, f"{ep:02d}"):
                    logging.info(f"🖼  Đã tạo thumbnail tập {ep:02d} (kịch_bản/).")
            else:
                logging.info("ℹ️ Chưa đặt số tập ở tab Thumbnail (0) → chưa tạo thumbnail 1 link.")
        except Exception as e:
            logging.error(f"Lỗi tạo SEO YouTube (bỏ qua, tiếp tục quy trình): {e}")
            self.pipe_status.set(f"⚠️ SEO YouTube lỗi: {e}")

    def _pipe_prepare_input(self, auto=False):
        if self._pipe_busy:
            return
        if not auto:
            self._save_pipe_settings()   # ấn chạy → nhớ cài đặt
        if self._prepare_input_from_gemini():
            self.pipe_status.set("✅ Đã tạo input.txt từ Gemini.")
            if self.var_auto_tts.get():   # ⛓ chạy tiếp tạo giọng (OmniVoice)
                logging.info("⛓ Tự động: chạy tiếp tạo giọng (OmniVoice)...")
                self.after(400, self._start)
            elif not auto:
                messagebox.showinfo(
                    "Xong",
                    "Đã tạo input.txt từ gemini_result.docx.\n"
                    "Giờ có thể bấm '▶ Chạy' để tạo audio.")

    def _start(self):
        self._save_pipe_settings()   # ấn ▶ Chạy → nhớ cài đặt quy trình cho lần sau
        self._save_opt_settings()    # nhớ cả mục "Cài đặt" để lần sau làm mặc định
        mode = self.var_mode.get()
        if mode == "clone":
            voice_name = self._current_voice()
            if not voice_name:
                messagebox.showwarning("Thiếu giọng mẫu",
                                       f"Không tìm thấy file audio trong:\n{VOICE_DIR}")
                return
            voice_param = str(VOICE_DIR / voice_name)
        elif mode == "design":
            voice_param = self.var_instruct.get().strip()
            if not voice_param:
                messagebox.showwarning("Thiếu mô tả", "Vui lòng nhập mô tả giọng đọc.")
                return
        else:
            voice_param = None

        # ── (TÙY CHỌN) Lấy nội dung từ Gemini + KIỂM TRA trước khi tạo audio ──
        if self.var_from_gemini.get():
            if not self._prepare_input_from_gemini():
                return

        # ── Chia text ngay tại đây, trước khi khởi động thread ──────────────
        text_file = Path(self.var_txt.get())
        try:
            full_text = clean_text(text_file.read_text(encoding="utf-8"))
        except Exception as e:
            messagebox.showerror("Lỗi đọc file", str(e))
            return

        chunks = split_chunks(full_text.lower(), self.var_chunk.get())
        if not chunks:
            messagebox.showwarning("File trống", "File văn bản không có nội dung.")
            return

        # Video dọc: bật/tắt + tốc độ
        make_video_doc = self.var_make_video_doc.get()
        doc_from_ngang = self.var_doc_from_ngang.get()   # dùng lại video ngang cho dọc
        doc_from_subfolder = self.var_doc_from_subfolder.get()  # ghép từ thư mục con videodoc
        doc_no_effect = self.var_doc_no_effect.get()     # không phủ hiệu ứng lên video dọc
        try:
            doc_speed = float(str(self.var_doc_speed.get()).replace(",", ".").strip())
        except (TypeError, ValueError):
            doc_speed = 1.0
        doc_speed = max(0.5, min(doc_speed, 2.0))   # atempo chỉ nhận 0.5–2.0
        doc_percent = self._parse_percent(self.var_doc_percent, 100)  # % audio cho video dọc (100 = cả bài)
        if make_video_doc and doc_speed > 1.001:
            logging.info(f"Video dọc sẽ tăng tốc audio x{doc_speed:.2f} (giữ cao độ).")
        if make_video_doc and doc_percent < 99:
            logging.info(f"Video dọc sẽ cắt ~{doc_percent}% audio (cắt cuối câu).")

        # Tốc độ audio cho VIDEO NGANG (audio full) — atempo, giữ cao độ
        try:
            ngang_speed = float(str(self.var_ngang_speed.get()).replace(",", ".").strip())
        except (TypeError, ValueError):
            ngang_speed = 1.0
        ngang_speed = max(0.5, min(ngang_speed, 2.0))
        if self.var_make_video.get() and ngang_speed > 1.001:
            logging.info(f"Video ngang sẽ tăng tốc audio x{ngang_speed:.2f} (giữ cao độ).")
        ngang_source = self.var_ngang_source.get()   # nguồn clip video ngang theo chủ đề

        # Tốc độ audio cho VIDEO TIKTOK (cắt theo %) — atempo, giữ cao độ
        make_tiktok = self.var_make_tiktok.get()
        try:
            tiktok_speed = float(str(self.var_tiktok_speed.get()).replace(",", ".").strip())
        except (TypeError, ValueError):
            tiktok_speed = 1.0
        tiktok_speed = max(0.5, min(tiktok_speed, 2.0))
        tiktok_no_effect = self.var_tiktok_no_effect.get()   # không phủ hiệu ứng lên TikTok
        # Chữ TikTok = 'Mimi audio Số <số tập>' (ô 'Số tập' = thumbnail + 1); LUÔN ghi.
        _ep = self._tiktok_episode_number()
        tiktok_caption = f"Mimi audio Số {_ep:02d}"   # LUÔN ghi chữ (kể cả số 00)
        try:
            tiktok_caption_pos = int(self.var_tiktok_caption_pos.get())
        except Exception:
            tiktok_caption_pos = 40
        tiktok_caption_pos = max(0, min(tiktok_caption_pos, 100))
        tiktok_music = self.var_tiktok_music.get()
        try:
            tiktok_music_db = int(self.var_tiktok_music_db.get())
        except Exception:
            tiktok_music_db = -12
        tiktok_music_db = max(-40, min(tiktok_music_db, 0))
        tiktok_percent = self._parse_percent(self.var_tiktok_percent)  # % thời lượng audio cho TikTok
        if make_tiktok and tiktok_speed > 1.001:
            logging.info(f"Video TikTok sẽ tăng tốc audio x{tiktok_speed:.2f} (giữ cao độ).")

        preview_path = text_file.parent / (text_file.stem + "_preview.txt")
        if preview_path.exists():
            preview_path.unlink()
        preview_path.write_text("\n\n".join(chunks), encoding="utf-8")
        logging.info(f"Chia {len(chunks)} đoạn (chunk={self.var_chunk.get()} ký tự) → {preview_path.name}")

        # ♻ Dùng lại audio/video đã có (chỉ dựng phần còn thiếu).
        reuse = self.var_reuse.get()

        # Bình thường: nếu file kết quả đã có → tự đặt tên mới (output.wav →
        # output1.wav…) để KHÔNG ghi đè bản cũ. Khi DÙNG LẠI thì GIỮ NGUYÊN tên
        # để tái dùng audio/video cũ thay vì tạo bản mới.
        if reuse:
            out_path = Path(self.var_out.get())
            logging.info(f"♻ Dùng lại: {out_path.name} (chỉ dựng phần còn thiếu).")
        else:
            out_path = unique_path(Path(self.var_out.get()))
            if str(out_path) != self.var_out.get():
                logging.info(f"File kết quả đã có → dùng tên mới: {out_path.name}")
                self.status.set(f"Kết quả sẽ lưu thành: {out_path.name}")
        self.var_out.set(str(out_path))
        self._last_output = self.var_out.get()

        # Áp QUY TẮC ĐẶT TÊN như bản tự động: 3 video theo nền tảng (cùng thư mục output).
        _out_dir = out_path.parent
        ngang_out = _out_dir / "YOUTUBE.mp4"     # video ngang → YouTube
        doc_out = _out_dir / "facebook.mp4"      # video dọc  → Facebook
        tiktok_out = _out_dir / "tiktok.mp4"     # video TikTok

        self._stop_preview()                         # dừng audio đang nghe (nếu có)
        self.btn_preview.config(state="disabled")    # khóa tới khi tạo xong lần này
        self._pause_event = threading.Event()
        self._pause_event.set()
        self.btn_run.config(state="disabled")
        self.btn_pause.config(state="normal", text="⏸  Tạm dừng")
        # Hiệu ứng phủ video (nếu chọn) — bỏ tiền tố ★ rồi chuyển thành đường dẫn đầy đủ
        effect_name = self._current_effect()
        effect_path = None
        if effect_name and effect_name != EFFECT_NONE:
            p = EFFECTS_DIR / effect_name
            effect_path = str(p) if p.exists() else None
            if effect_path:
                logging.info(f"Hiệu ứng phủ video: {effect_name}")

        self.progress.set(0)
        self.status.set(f"Đã chia {len(chunks)} đoạn — đang khởi động...")
        threading.Thread(
            target=self._run_tts_then_shutdown,   # bọc để xong thì xét ô tick ⏻
            args=(mode, voice_param, chunks, self.var_out.get(),
                  self.progress, self.status,
                  self.btn_run, self.btn_pause, self.btn_preview, self._pause_event,
                  self.var_make_video.get(), effect_path,
                  make_video_doc, doc_speed, doc_percent,
                  ngang_speed, reuse, doc_from_ngang, doc_no_effect),
            kwargs={"make_tiktok": make_tiktok,   # video TikTok (cắt theo %)
                    "make_short": self.var_make_short.get(),   # Short ≤2:50 cắt từ TikTok
                    "tiktok_speed": tiktok_speed,
                    "tiktok_percent": tiktok_percent,
                    "tiktok_no_effect": tiktok_no_effect,
                    "tiktok_caption": tiktok_caption,
                    "tiktok_caption_pos": tiktok_caption_pos,
                    "tiktok_music": tiktok_music,
                    "tiktok_music_db": tiktok_music_db,
                    "doc_from_subfolder": doc_from_subfolder,   # ghép video dọc từ thư mục con videodoc
                    "ngang_source": ngang_source,               # nguồn clip video ngang theo chủ đề
                    # Quy tắc đặt tên 3 video (giống bản tự động).
                    "ngang_out": ngang_out,
                    "doc_out": doc_out,
                    "tiktok_out": tiktok_out,
                    "short_out": Path(tiktok_out).with_name("short.mp4"),
                    # Phụ đề cho video ngang (YouTube)
                    "make_sub": self.var_make_sub.get(),
                    "sub_mode": self.var_sub_mode.get(),
                    "sub_model": self.var_sub_model.get(),
                    "sub_max_chars": self._parse_sub_max_chars(),
                    "sub_kieu": self.var_sub_kieu.get(),
                    "sub_font": self.var_sub_font.get(),
                    "sub_mau": self.var_sub_mau.get(),
                    "sub_mau_vien": self.var_sub_mau_vien.get(),
                    "sub_vitri": self.var_sub_vitri.get().strip(),
                    "sub_cochu": self._parse_sub_cochu(),
                    "sub_bengang": self._parse_sub_bengang(),
                    "sub_dong": self._parse_sub_dong(),
                    "make_sub_doc": self.var_make_sub_doc.get()},
            daemon=True,
        ).start()

    @staticmethod
    def _parse_percent(var, default: int = 50) -> int:
        """Đọc % (combobox TikTok) → int, kẹp trong [10, 100]; lỗi thì dùng mặc định."""
        try:
            v = int(round(float(str(var.get()).replace(",", ".").strip())))
        except (TypeError, ValueError):
            v = default
        return max(10, min(v, 100))

    def _parse_sub_max_chars(self, default: int = 50) -> int:
        """Đọc ô 'Dài mỗi dòng' phụ đề → int, kẹp trong [20, 90]."""
        try:
            v = int(str(self.var_sub_max_chars.get()).strip())
        except (TypeError, ValueError, AttributeError):
            v = default
        return max(20, min(v, 90))

    def _parse_sub_cochu(self) -> str:
        """Đọc ô 'Cỡ chữ %' → chuỗi số kẹp trong [50, 200]; rỗng/100 = giữ cỡ gốc."""
        raw = str(self.var_sub_cochu.get()).replace(",", ".").strip()
        if not raw:
            return ""
        try:
            v = int(round(float(raw)))
        except (TypeError, ValueError):
            return ""
        v = max(50, min(v, 200))
        return "" if v == 100 else str(v)

    def _parse_sub_bengang(self) -> str:
        """Đọc ô 'Ngang %' → chuỗi số kẹp trong [20, 150]; rỗng/100 = như cũ."""
        raw = str(self.var_sub_bengang.get()).replace(",", ".").strip()
        if not raw:
            return ""
        try:
            v = int(round(float(raw)))
        except (TypeError, ValueError):
            return ""
        v = max(20, min(v, 150))
        return "" if v == 100 else str(v)

    def _parse_sub_dong(self) -> int:
        """Đọc ô 'Số dòng' phụ đề → 1 hoặc 2 (mặc định 2)."""
        try:
            return 1 if int(str(self.var_sub_dong.get()).strip()) == 1 else 2
        except (TypeError, ValueError, AttributeError):
            return 2

    @staticmethod
    def _parse_speed(var) -> float:
        """Đọc tốc độ (combobox) → float, kẹp trong [0.5, 2.0] (giới hạn atempo)."""
        try:
            v = float(str(var.get()).replace(",", ".").strip())
        except (TypeError, ValueError):
            v = 1.0
        return max(0.5, min(v, 2.0))

    def _current_episode_number(self) -> int:
        """Số tập ở tab Thumbnail: ƯU TIÊN số ĐANG nhập (khớp đúng số người dùng vừa
        đặt, kể cả khi chưa lưu), không có thì lấy số đã lưu gần nhất."""
        tg = getattr(self, "_thumb_gui", None)
        if tg is not None:
            try:
                n = str(tg.number_var.get()).strip()
                if n.isdecimal():
                    return int(n)
            except Exception:
                pass
        return load_episode_number()

    def _default_tiktok_episode(self) -> str:
        """Số tập MẶC ĐỊNH cho chữ TikTok = số tab Thumbnail + 1 (2 chữ số). TikTok của
        tập ĐANG dựng đi TRƯỚC thumbnail 1 số, nên +1 để khớp đúng tập mới.
        Số nằm trong danh sách "tập bỏ qua" thì nhảy tiếp cho tới số dùng được."""
        return f"{next_episode_number(self._current_episode_number()):02d}"

    def _sync_tiktok_episode(self, *_):
        """Cập nhật ô 'Số tập' TikTok theo số Thumbnail + 1 (gọi khi số thumbnail đổi).
        Bỏ qua nếu ô chưa được dựng (thứ tự build)."""
        var = getattr(self, "var_tiktok_episode", None)
        if var is not None:
            try:
                var.set(self._default_tiktok_episode())
            except Exception:
                pass

    def _tiktok_episode_number(self) -> int:
        """Số tập DÙNG cho chữ TikTok: lấy từ ô 'Số tập' trên Home (người dùng có thể
        sửa); trống/không hợp lệ → quay về mặc định (số thumbnail + 1)."""
        var = getattr(self, "var_tiktok_episode", None)
        if var is not None:
            try:
                n = str(var.get()).strip()
                if n.isdecimal():
                    return int(n)
            except Exception:
                pass
        return next_episode_number(self._current_episode_number())

    # ── TẬP BỎ QUA (đặt trước cho tương lai) ────────────────────────────────
    @staticmethod
    def _parse_episode_list(s: str) -> set:
        """'33, 34' / '33 34' / '33-35' → {33, 34, 35}. Bỏ qua phần không hợp lệ."""
        out = set()
        for token in re.split(r'[,\s]+', (s or "").strip()):
            if not token:
                continue
            m = re.fullmatch(r'(\d+)\s*-\s*(\d+)', token)
            if m:
                a, b = int(m.group(1)), int(m.group(2))
                out.update(range(min(a, b), max(a, b) + 1))
            elif token.isdecimal():
                out.add(int(token))
        return {n for n in out if n > 0}

    def _refresh_skip_episodes_label(self):
        """Cập nhật chữ gợi ý cạnh nút: các tập đang bỏ qua + tập mới kế tiếp."""
        lbl = getattr(self, "lbl_skip_eps", None)
        if lbl is None:
            return
        skip = load_skip_episodes()
        nxt = next_episode_number(self._last_used_episode(), skip)
        lbl.config(text=(f"bỏ qua {', '.join(str(n) for n in sorted(skip))} · "
                         f"tập mới → {nxt:02d}") if skip else f"tập mới → {nxt:02d}")

    def _edit_skip_episodes(self):
        """Hộp thoại xem/sửa danh sách số tập sẽ BỎ QUA khi cấp số tập mới."""
        win = tk.Toplevel(self)
        win.title("Tập bỏ qua")
        win.configure(bg=UI["bg"])
        win.transient(self)
        win.resizable(False, False)

        wrap = ttk.Frame(win, padding=14)
        wrap.pack(fill="both", expand=True)
        ttk.Label(wrap, text="Số tập sẽ BỎ QUA khi cấp số cho tập mới").pack(anchor="w")
        ttk.Label(wrap, text="Ví dụ: đang ở tập 31, thêm 33 và 34 → sau 32 nhảy thẳng "
                             "sang 35.\nTập ĐÃ tạo vẫn giữ nguyên số của nó.",
                  style="Hint.TLabel", justify="left").pack(anchor="w", pady=(2, 10))

        mid = ttk.Frame(wrap)
        mid.pack(fill="both", expand=True)
        lst = tk.Listbox(mid, width=14, height=8, selectmode="extended",
                         exportselection=False, activestyle="dotbox",
                         font=("Segoe UI", 10))
        lst.pack(side="left", fill="y")
        sb = ttk.Scrollbar(mid, orient="vertical", command=lst.yview)
        sb.pack(side="left", fill="y")
        lst.configure(yscrollcommand=sb.set)

        side = ttk.Frame(mid)
        side.pack(side="left", fill="both", expand=True, padx=(12, 0))
        ttk.Label(side, text="Thêm số tập:").pack(anchor="w")
        var_add = tk.StringVar()
        ent = ttk.Entry(side, textvariable=var_add, width=16)
        ent.pack(anchor="w", pady=(2, 2))
        ttk.Label(side, text="cách nhau bởi dấu phẩy,\nhoặc khoảng: 33-35",
                  style="Hint.TLabel", justify="left").pack(anchor="w")
        info = ttk.Label(side, text="", style="Hint.TLabel", justify="left")
        info.pack(anchor="w", pady=(10, 0))

        def refresh():
            skip = sorted(load_skip_episodes())
            lst.delete(0, "end")
            for n in skip:
                lst.insert("end", f"  Tập {n:02d}")
            nxt = next_episode_number(self._last_used_episode(), set(skip))
            info.config(text=f"Tập lớn nhất đã có: {self._last_used_episode():02d}\n"
                             f"Tập mới kế tiếp: {nxt:02d}")
            self._refresh_skip_episodes_label()
            return skip

        def add():
            nums = self._parse_episode_list(var_add.get())
            if not nums:
                messagebox.showwarning("Chưa hợp lệ",
                                       "Nhập số tập, ví dụ: 33, 34 hoặc 33-35.",
                                       parent=win)
                return
            save_skip_episodes(load_skip_episodes() | nums)
            var_add.set("")
            refresh()

        def remove():
            skip = sorted(load_skip_episodes())
            chosen = {skip[i] for i in lst.curselection()}
            if not chosen:
                messagebox.showinfo("Chưa chọn", "Chọn dòng cần xóa trong danh sách.",
                                    parent=win)
                return
            save_skip_episodes(set(skip) - chosen)
            refresh()

        ent.bind("<Return>", lambda e: add())
        btns = ttk.Frame(wrap)
        btns.pack(fill="x", pady=(12, 0))
        ttk.Button(btns, text="➕  Thêm", command=add).pack(side="left")
        ttk.Button(btns, text="🗑  Xóa mục chọn", command=remove).pack(side="left", padx=(8, 0))
        ttk.Button(btns, text="Đóng", command=win.destroy).pack(side="right")

        refresh()
        win.update_idletasks()
        win.geometry(f"+{self.winfo_rootx() + 90}+{self.winfo_rooty() + 110}")
        ent.focus_set()
        win.grab_set()

    def _exclusive_doc_source(self, chosen: str):
        """2 nguồn hình video dọc loại trừ nhau: '♻ dùng lại video ngang' và
        '📁 ghép từ thư mục con videodoc'. Vừa bật cái này thì tắt cái kia.
        (Cả hai cùng tắt vẫn hợp lệ → ghép random cả kho videodoc/.)"""
        if chosen == "ngang" and self.var_doc_from_ngang.get():
            self.var_doc_from_subfolder.set(False)
        elif chosen == "subfolder" and self.var_doc_from_subfolder.get():
            self.var_doc_from_ngang.set(False)

    def _rebuild_video(self, kind: str):
        """Nút 'Dựng lại' của từng mục: dựng LẠI đúng 1 loại video (ngang/dọc/tiktok)
        từ AUDIO ĐÃ CÓ trong kịch_bản, KHÔNG chạy lại TTS. Tái dùng run_tts(video_only)."""
        out_path = Path(self.var_out.get())
        if not (out_path.exists() and out_path.stat().st_size > 4096):
            messagebox.showwarning(
                "Chưa có audio",
                f"Không thấy audio đã tạo:\n{out_path}\n\n"
                "Hãy tạo giọng trước, hoặc chọn đúng file ở ô 'Kết quả'.")
            return

        # Bấm "Dựng lại" cũng NHỚ cài đặt lần chạy gần nhất (tốc độ, chữ TikTok, nhạc
        # nền...) như nút "Chạy" chính, để lần sau mở app dùng lại làm mặc định.
        self._save_opt_settings()

        # 3 video theo nền tảng (cùng thư mục với audio) — như bản tự động.
        _out_dir = out_path.parent
        ngang_out = _out_dir / "YOUTUBE.mp4"
        doc_out = _out_dir / "facebook.mp4"
        tiktok_out = _out_dir / "tiktok.mp4"

        # Hiệu ứng phủ (nếu chọn) — bỏ tiền tố ★, chuyển thành đường dẫn đầy đủ.
        effect_name = self._current_effect()
        effect_path = None
        if effect_name and effect_name != EFFECT_NONE:
            p = EFFECTS_DIR / effect_name
            effect_path = str(p) if p.exists() else None

        # Chữ TikTok = 'Mimi audio Số <số tập>' (ô 'Số tập' = thumbnail + 1); LUÔN ghi.
        _ep = self._tiktok_episode_number()
        tiktok_caption = f"Mimi audio Số {_ep:02d}"   # LUÔN ghi chữ (kể cả số 00)
        try:
            tiktok_caption_pos = max(0, min(int(self.var_tiktok_caption_pos.get()), 100))
        except Exception:
            tiktok_caption_pos = 40
        try:
            tiktok_music_db = max(-40, min(int(self.var_tiktok_music_db.get()), 0))
        except Exception:
            tiktok_music_db = -12

        btn = {"ngang": self.btn_run_ngang, "doc": self.btn_run_doc,
               "tiktok": self.btn_run_tiktok}[kind]
        btn.config(state="disabled")
        self.progress.set(0)
        self.status.set(f"🎬 Dựng lại video {kind} từ audio có sẵn...")
        logging.info(f"🎬 Dựng lại video {kind} từ {out_path.name} (không chạy lại TTS).")

        ev = threading.Event()
        ev.set()   # không tạm dừng
        threading.Thread(
            target=self._run_tts_then_shutdown,   # bọc để xong thì xét ô tick ⏻
            args=("", None, [], str(out_path), self.progress, self.status,
                  btn, _NullWidget(), _NullWidget(), ev),
            kwargs=dict(
                video_only=True, reuse=True, effect=effect_path,
                make_video=(kind == "ngang"), ngang_speed=self._parse_speed(self.var_ngang_speed),
                ngang_source=self.var_ngang_source.get(),
                ngang_out=ngang_out,
                make_video_doc=(kind == "doc"),
                doc_speed=self._parse_speed(self.var_doc_speed),
                doc_percent=self._parse_percent(self.var_doc_percent, 100),
                doc_from_ngang=self.var_doc_from_ngang.get(),
                doc_from_subfolder=self.var_doc_from_subfolder.get(),
                doc_no_effect=self.var_doc_no_effect.get(), doc_out=doc_out,
                make_tiktok=(kind == "tiktok"), tiktok_out=tiktok_out,
                # Dựng lại TikTok là short cũ hết hạn theo → cắt lại luôn cho khớp.
                make_short=(kind == "tiktok" and self.var_make_short.get()),
                short_out=_out_dir / "short.mp4",
                tiktok_speed=self._parse_speed(self.var_tiktok_speed),
                tiktok_percent=self._parse_percent(self.var_tiktok_percent),
                tiktok_no_effect=self.var_tiktok_no_effect.get(),
                tiktok_caption=tiktok_caption, tiktok_caption_pos=tiktok_caption_pos,
                tiktok_music=self.var_tiktok_music.get(), tiktok_music_db=tiktok_music_db,
                # Chỉ nút "Dựng lại" của video NGANG mới làm phụ đề (2 video kia không có).
                make_sub=(kind == "ngang" and self.var_make_sub.get()),
                # Nút "Dựng lại" video DỌC thì làm phụ đề khung dọc cho chính nó.
                make_sub_doc=(kind == "doc" and self.var_make_sub_doc.get()),
                sub_mode=self.var_sub_mode.get(), sub_model=self.var_sub_model.get(),
                sub_max_chars=self._parse_sub_max_chars(),
                sub_kieu=self.var_sub_kieu.get(),
                sub_font=self.var_sub_font.get(),
                sub_mau=self.var_sub_mau.get(),
                sub_mau_vien=self.var_sub_mau_vien.get(),
                sub_vitri=self.var_sub_vitri.get().strip(),
                sub_cochu=self._parse_sub_cochu(),
                sub_bengang=self._parse_sub_bengang(),
                sub_dong=self._parse_sub_dong(),
            ),
            daemon=True,
        ).start()

    def _toggle_pause(self):
        if self._pause_event.is_set():
            self._pause_event.clear()
            self.btn_pause.config(text="▶  Tiếp tục")
        else:
            self._pause_event.set()
            self.btn_pause.config(text="⏸  Tạm dừng")

    # ── NGHE THỬ KẾT QUẢ ──────────────────────────────────────────────────────
    def _toggle_preview(self):
        if self._playing:
            self._stop_preview()
        else:
            self._play_preview()

    def _play_preview(self):
        path = Path(self._last_output or self.var_out.get())
        if not path.exists():
            messagebox.showinfo("Chưa có audio",
                                "Chưa có file kết quả để nghe. Hãy chạy tạo giọng trước.")
            return
        try:
            import winsound
            winsound.PlaySound(str(path), winsound.SND_FILENAME | winsound.SND_ASYNC)
        except Exception:
            # Không phát nội bộ được → mở bằng trình phát mặc định của hệ thống
            try:
                os.startfile(str(path))  # type: ignore[attr-defined]
            except Exception as e:
                messagebox.showerror("Lỗi phát audio", str(e))
            return
        self._playing = True
        self.btn_preview.config(text="⏹  Dừng nghe")
        # Tự nhả nút khi nghe hết (winsound không báo kết thúc)
        try:
            dur = sf.info(str(path)).duration
            self._preview_after = self.after(int(dur * 1000) + 300, self._stop_preview)
        except Exception:
            self._preview_after = None

    def _stop_preview(self):
        try:
            import winsound
            winsound.PlaySound(None, winsound.SND_PURGE)
        except Exception:
            pass
        if self._preview_after is not None:
            try:
                self.after_cancel(self._preview_after)
            except Exception:
                pass
            self._preview_after = None
        self._playing = False
        self.btn_preview.config(text="🔊  Nghe thử")

    def _poll_log(self):
        while not log_queue.empty():
            levelno, msg = log_queue.get_nowait()
            tag = ("err" if levelno >= logging.ERROR
                   else "warn" if levelno >= logging.WARNING else "info")
            for box in self._log_boxes:   # ghi ra mọi ô nhật ký (video + tab kịch bản)
                box.config(state="normal")
                box.insert("end", msg + "\n", tag)
                box.see("end")
                box.config(state="disabled")
        self.after(200, self._poll_log)


if __name__ == "__main__":
    App().mainloop()
