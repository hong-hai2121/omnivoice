# -*- coding: utf-8 -*-
"""
seo_youtube_gemini.py — Lấy ĐOẠN ĐẦU của nội dung đã dịch (gemini_result.docx),
CHÈN TRỌN BỘ YÊU CẦU SEO (SEO_PROMPT) lên trước rồi gửi kèm đoạn truyện lên
Gemini, rồi lưu kết quả ra seoYoutube.docx.

Ý tưởng:
• Trước đây script dựa vào 1 cuộc trò chuyện Gemini "đã ra đề sẵn"; khi Gemini bỏ
  id rồi mở CHAT MỚI thì mất ngữ cảnh → SEO lỗi hoặc không gửi được. Nay TIN NHẮN
  TỰ CHỨA đầy đủ yêu cầu (SEO_PROMPT) nên gửi vào chat nào — kể cả chat mới trống —
  cũng ra đúng định dạng.
• Vì thế mặc định MỖI TẬP gửi vào một cuộc trò chuyện MỚI (SEO_NEW_CHAT): chat cũ
  chứa sẵn SEO của mọi tập trước, chỉ cần đọc nhầm một lần là cả loạt tập ra chung
  một tiêu đề (đã dính: tập 55–58, 49, 08).
• Kết quả phải ĐÚNG DẠNG SEO mới được lưu (_looks_like_seo); không đạt thì gửi lại,
  hết lượt vẫn hỏng thì để file .docx RỖNG chứ không lưu bừa.

Chạy:
    python seo_youtube_gemini.py                          # dùng gemini_result.docx mặc định
    python seo_youtube_gemini.py "gemini_result.docx"     # đổi file nguồn
    python seo_youtube_gemini.py -o "seoYoutube.docx"      # đổi file kết quả
    python seo_youtube_gemini.py --chars 1500             # cắt còn N ký tự đầu (0 = cả đoạn)
    python seo_youtube_gemini.py --no-keep-open           # đóng Firefox sau khi xong
"""

import sys
import os

# ── Tự chuyển sang python của venv (giống các script khác trong thư mục này) ──
# CHỈ làm khi CHẠY THẲNG file này (python seo_youtube_gemini.py). Khi được GUI
# import (import seo_youtube_gemini), TUYỆT ĐỐI không tự chuyển: lúc đó sys.argv[0]
# là amain_taogiong_gui.py nên subprocess.run sẽ MỞ LẠI GUI (GUI bật 2 lần) rồi
# sys.exit() giết luôn thread SEO → SEO YouTube không chạy. Dùng normcase để tránh
# lệch hoa/thường ổ đĩa ("d:" vs "D:") gây tự chuyển nhầm.
_REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, os.pardir))
_VENV_PYTHON = os.path.join(_REPO_ROOT, "venv", "Scripts", "python.exe")
if __name__ == "__main__" and os.path.exists(_VENV_PYTHON) and \
        os.path.normcase(os.path.abspath(sys.executable)) != \
        os.path.normcase(os.path.abspath(_VENV_PYTHON)):
    import subprocess
    subprocess.run([_VENV_PYTHON] + sys.argv)
    sys.exit()

# Thư mục hiện tại (myvoice/YOUTUBE) + thư mục scripts (nơi có dich_gemini.py,
# dich_docx.py). Thêm cả hai vào sys.path để import được.
_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
_SCRIPTS_DIR = os.path.join(os.path.dirname(_THIS_DIR), "scripts")
for _p in (_THIS_DIR, _SCRIPTS_DIR):
    if _p not in sys.path:
        sys.path.insert(0, _p)

import time
import argparse
from pathlib import Path

import dich_gemini as g
from dich_docx import read_chunks

# ── Link cuộc trò chuyện Gemini chuyên SEO YouTube (đã có sẵn chỉ dẫn) ─────────
# CHỈ giữ phần /app/<id> — KHÔNG kèm tham số quảng cáo (gclid, utm_*, campaign_id,
# is_sa...). Link 1 cuộc trò chuyện CỤ THỂ mà còn dính tham số chiến dịch thì Gemini
# hay BỎ id rồi mở CHAT MỚI → SEO gửi nhầm vào trò chuyện trống (mất ngữ cảnh đã
# huấn luyện). Vì vậy luôn cắt bỏ phần query (kể cả khi đặt qua biến môi trường).
SEO_GEMINI_URL = os.environ.get(
    "OMNI_SEO_GEMINI_URL",
    "https://gemini.google.com/app/e1a7ace4a71c48b2",
).split("?", 1)[0].strip()

# ── Mặc định: MỖI TẬP một cuộc trò chuyện MỚI ─────────────────────────────────
# Chat SEO cố định ở trên chứa sẵn câu trả lời của mọi tập trước. Khi Gemini chậm
# hoặc không trả lời, phần đọc kết quả từng lấy nhầm CÂU TRẢ LỜI CŨ đang nằm cuối
# chat → nhiều tập ra CÙNG MỘT bản SEO (tập 55–58, 49, 08 cùng một tiêu đề). Chat
# TRỐNG thì không có gì cũ để lấy nhầm, mà tin nhắn đã tự chứa trọn SEO_PROMPT nên
# vẫn ra đúng định dạng. Đặt OMNI_SEO_NEW_CHAT=0 để quay lại dùng chat cố định.
NEW_CHAT_URL = "https://gemini.google.com/app"
SEO_NEW_CHAT = os.environ.get("OMNI_SEO_NEW_CHAT", "1").strip().lower() not in (
    "0", "false", "no")
# Số lần gửi lại khi Gemini không trả lời / trả lời sai định dạng (mỗi lần một chat mới).
SEO_ATTEMPTS = max(1, int(os.environ.get("OMNI_SEO_ATTEMPTS", "2")))


# ── Trọn bộ yêu cầu SEO chèn LÊN TRƯỚC đoạn truyện khi gửi Gemini ──────────────
# Nhờ tin nhắn tự chứa đủ yêu cầu này, không còn phụ thuộc cuộc trò chuyện "đã ra
# đề sẵn": gửi vào chat mới trống vẫn ra đúng định dạng. Đặt qua biến môi trường
# OMNI_SEO_PROMPT nếu muốn đổi mà không sửa code.
SEO_PROMPT = os.environ.get("OMNI_SEO_PROMPT", "").strip() or """\
Tôi là kênh "Mimi audio" (chuyên về truyện audio, truyện full, audio dài kỳ, nội dung kịch tính, nghẹt thở, KHÔNG PHẢI truyện dễ ngủ). Hãy đóng vai một chuyên gia nội dung và SEO YouTube để giúp tôi xử lý các đoạn trích từ truyện thành các tài nguyên đăng tải theo đúng quy trình và định dạng bắt buộc sau:

---

### 🎯 BƯỚC 2: CHỌN TIÊU ĐỀ VIDEO
Chọn 5 tiêu đề hay, độ dài từ 60 – 75 ký tự (bao gồm cả phần "| Mimi audio" ở cuối). Sau đó chọn ra 1 tiêu đề tốt nhất làm quán quân.
*Yêu cầu nghiêm ngặt:*
- Tiêu đề số 1: Luôn bắt đầu bằng cụm từ "Sự Thật Đằng Sau...".
- Tiêu đề số 2: Luôn bắt đầu bằng cụm từ "Màn...".
- Tiêu đề số 3: Luôn viết theo mẫu "Câu Chuyện Đời [Tên Nhân Vật Chính] Và Sự Thật Đau Đớn Sau Nhiều Năm...".
- Tiêu đề số 4: Luôn bắt đầu bằng cụm từ "Truyện Tâm Lý Xã Hội Về Luật Nhân Quả...".
- Tiêu đề số 5: Luôn bắt đầu bằng cụm từ "Bí Mật Sau Cuộc Đời...".
- TUYỆT ĐỐI KHÔNG dùng các từ liên quan đến tiền bạc như: tiền, nghèo, giàu, gia tài, phú hộ, triệu phú, tỷ phú, vàng, bạc... (Nếu cốt truyện có tiền bạc, phải dùng từ thay thế né tránh như "phú quý", "hào môn", "sự nghiệp").

---

### 🏷️ BƯỚC 3: VIẾT THẺ TAG YOUTUBE
Tạo một đoạn văn bản chứa các thẻ tag cách nhau bằng dấu phẩy, tổng độ dài khoảng 400 ký tự. Phải phân bổ đủ 4 nhóm từ khóa sau:
- Nhóm 1 (Mặc định cố định): "truyện audio, truyện đêm khuya, truyện tâm lý xã hội, truyện đời sống, truyện Mimi audio, truyện gia đình, truyện cảm động, truyện người vợ, truyện mẹ chồng nàng dâu, truyện nhân quả, Mimi audio, nghe truyện Mimi audio, truyện audio Mimi, kênh Mimi audio, nghe truyện hay, truyện hay mỗi ngày, truyện kể đêm khuya, truyện Việt Nam, audio truyện hay,"
- Nhóm 2: Từ khóa liên quan đến thể loại cụ thể của truyện vừa đưa (ví dụ: trùng sinh, xuyên không, cổ đại...).
- Nhóm 3: Tên các nhân vật chính xuất hiện trong đoạn truyện.
- Nhóm 4: Các cụm từ nổi bật, ấn tượng lấy từ nội dung cốt truyện.

---

### 📝 BƯỚC 4: VIẾT MÔ TẢ VIDEO CHUẨN SEO
Viết mô tả siêu rút gọn theo đúng cấu trúc 4 phần sau (không thêm bớt dòng, không viết tràn lan):
- Dòng 1 (Bộ hashtag): #MimiAudio #TruyenAudio #TruyenTamLyXaHoi #TruyenKichTinh #Audio
- Dòng 2: [Tiêu đề tốt nhất đã chọn ở Bước 2]
- Dòng 3 & 4 (Tóm tắt nội dung): Viết 2 câu ngắn gọn, súc tích, mô tả tình huống kịch tính, nghẹt thở của truyện (Tuyệt đối KHÔNG nhắc đến các từ "ngủ ngon", "giấc ngủ", "dễ ngủ").
- Dòng 5 (Kêu gọi hành động cố định): 👉 Đăng ký Mimi audio và bật chuông để nghe truyện mới mỗi ngày.

---

### ⚠️ QUY TẮC ĐỊNH DẠNG ĐẦU RA (OUTPUT):
- Luôn sử dụng các thanh ngăn cách (---) và tiêu đề lớn bằng định dạng Markdown (## 🎯 BƯỚC 2..., ## 🏷️ BƯỚC 3..., ## 📝 BƯỚC 4...) để tạo cấu trúc rõ ràng, scannable, dễ đọc lướt.
- Phần Thẻ Tag ở Bước 3 phải được đặt riêng trong một khối mã (code block) để dễ dàng sao chép.
- Sử dụng trích dẫn (>) để làm nổi bật Tiêu đề được chọn ở Bước 2.
- Luôn dùng tiếng Việt, giọng điệu adaptive, sắc sảo và tuân thủ tuyệt đối các quy định trên.

Bây giờ, tôi sẽ gửi đoạn trích truyện đầu tiên, hãy thực hiện ngay lập tức từ BƯỚC 2 đến BƯỚC 4 cho tôi."""


def _seo_conversation_id():
    """ID cuộc trò chuyện trong SEO_GEMINI_URL (phần sau '/app/'); '' nếu không có."""
    if "/app/" not in SEO_GEMINI_URL:
        return ""
    return SEO_GEMINI_URL.split("/app/", 1)[-1].split("/", 1)[0].strip()


def _open_seo_chat(driver, log, wait=8):
    """Điều hướng Firefox sang ĐÚNG cuộc trò chuyện SEO và XÁC NHẬN Gemini không
    đẩy sang chat mới. Thử lại 1 lần nếu bị mở chat mới. True nếu đang ở đúng chat."""
    conv_id = _seo_conversation_id()
    last = ""
    for attempt in (1, 2):
        driver.get(SEO_GEMINI_URL)
        time.sleep(wait)
        last = driver.current_url or ""
        if not conv_id or conv_id in last:
            return True
        log(f"⚠️ Lần {attempt}: Gemini mở CHAT MỚI ({last}) thay vì cuộc trò chuyện "
            f"SEO ({conv_id}) — thử lại...")
    log(f"❗ KHÔNG vào được cuộc trò chuyện SEO (đang ở {last}). KHÔNG gửi để tránh "
        "tạo SEO sai. Kiểm tra lại link OMNI_SEO_GEMINI_URL và đăng nhập Gemini.")
    return False


def _open_new_chat(driver, log, wait=6):
    """Mở một cuộc trò chuyện MỚI (trống) để gửi SEO — xem SEO_NEW_CHAT."""
    driver.get(NEW_CHAT_URL)
    time.sleep(wait)
    try:
        con = len(g._get_responses(driver))
    except Exception:
        con = 0
    if con:
        log(f"⚠️ Cuộc trò chuyện mới vẫn còn {con} khối nội dung cũ trên màn hình — "
            "vẫn gửi (chốt chống lấy nhầm câu cũ nằm trong send_to_gemini).")
    else:
        log("🆕 Đã mở cuộc trò chuyện Gemini MỚI (trống) cho SEO.")
    return True


def _open_chat(driver, log):
    """Đưa Firefox tới nơi sẽ gửi SEO: chat MỚI (mặc định) hoặc chat SEO cố định."""
    if SEO_NEW_CHAT:
        return _open_new_chat(driver, log)
    if _open_seo_chat(driver, log):
        log("✅ Đã mở cuộc trò chuyện Gemini SEO. Đang gửi nội dung...")
    else:
        log("↪️ Không vào được chat SEO — vẫn gửi KÈM TRỌN YÊU CẦU vào chat hiện tại.")
    return True


# Kết quả SEO thật luôn có phần TIÊU ĐỀ và phần MÔ TẢ (xem SEO_PROMPT). Câu trả lời
# thiếu hai mốc này là câu vu vơ/báo lỗi của Gemini — KHÔNG lưu, gửi lại.
_SEO_MARKS = ("TIÊU ĐỀ", "MÔ TẢ")
_SEO_MIN_CHARS = 200


def _looks_like_seo(ans: str) -> bool:
    """True nếu câu trả lời đúng dạng kết quả SEO (đủ mốc TIÊU ĐỀ + MÔ TẢ)."""
    text = " ".join((ans or "").split()).upper()
    return len(text) >= _SEO_MIN_CHARS and all(m in text for m in _SEO_MARKS)


# ── Đường dẫn mặc định ───────────────────────────────────────────────────────
KICHBAN_DIR = Path(_SCRIPTS_DIR).parent / "kịch_bản"
DEFAULT_INPUT = KICHBAN_DIR / "gemini_result.docx"
DEFAULT_OUTPUT = KICHBAN_DIR / "seoYoutube.docx"


def first_chunk(path):
    """Lấy đoạn nội dung ĐẦU TIÊN (bỏ Heading/Title) trong file .docx đã dịch."""
    chunks = read_chunks(path)
    return chunks[0] if chunks else ""


def save_seo_docx(content, out_path):
    """Lưu phần SEO YouTube ra .docx (1 mục, giữ nguyên xuống dòng từ Gemini)."""
    from docx import Document
    doc = Document()
    doc.add_heading("SEO YouTube (Gemini)", level=1)
    for line in (content or "(trống)").splitlines() or ["(trống)"]:
        doc.add_paragraph(line)
    doc.save(str(out_path))
    return out_path


PARKED_STEM = "seoYoutube_trung"    # nơi cất bản SEO nghi lấy nhầm của tập khác


def park_docx(out_path):
    """Cất file SEO nghi LẤY NHẦM của tập khác sang 'seoYoutube_trung*.docx'.

    KHÔNG xoá gì: bản cũ vẫn còn đó để đối chiếu, mà thư mục tập không còn
    seoYoutube.docx nên lần chạy sau tự làm SEO mới. Trả về đường dẫn mới / None.
    """
    src = Path(out_path)
    if not src.is_file():
        return None
    dst, i = src.with_name(f"{PARKED_STEM}.docx"), 2
    while dst.exists():
        dst = src.with_name(f"{PARKED_STEM}_{i}.docx")
        i += 1
    src.rename(dst)
    return dst


def reset_docx(out_path):
    """Làm RỖNG file kết quả ngay đầu mỗi lần chạy: ghi đè bằng .docx hợp lệ nhưng
    chưa có nội dung (chỉ tiêu đề). Nhờ vậy nếu lần chạy này lỗi giữa chừng thì
    file cũng không còn dính kết quả của lần chạy trước."""
    from docx import Document
    doc = Document()
    doc.add_heading("SEO YouTube (Gemini)", level=1)
    doc.save(str(out_path))
    return out_path


def run(input_path, output_path, max_chars=0, keep_open=True, log=print, driver=None,
        attempts=SEO_ATTEMPTS):
    """driver: truyền 1 Firefox/driver đang mở để TÁI DÙNG (chỉ điều hướng sang cuộc
    trò chuyện dùng cho SEO). Nhờ vậy không phải đóng rồi mở lại Firefox — tránh kẹt
    khóa profile khiến lần mở thứ hai thất bại. None thì tự mở Firefox mới.

    attempts: số lần gửi (mỗi lần một cuộc trò chuyện mới) khi Gemini không trả lời
    hoặc trả lời sai định dạng. Hết lượt vẫn không có kết quả thì file .docx được
    GIỮ RỖNG — bước sau (thumbnail/đăng) nhìn file rỗng là biết SEO chưa làm, thay
    vì dùng nhầm nội dung của tập khác."""
    text = first_chunk(input_path)
    if not text:
        log(f"❌ Không đọc được đoạn đầu từ: {input_path}")
        return ""
    if max_chars and max_chars > 0:
        text = text[:max_chars]

    log(f"📄 Nguồn : {input_path}")
    log(f"💾 Kết quả: {output_path}")

    # Làm RỖNG file kết quả trước, rồi mới gửi & ghi nội dung mới của lần chạy này.
    reset_docx(output_path)
    log("🧹 Đã làm rỗng file kết quả (sẽ ghi nội dung mới sau khi Gemini trả lời).")

    log(f"📤 Gửi đoạn đầu ({len(text)} ký tự) — KÈM TRỌN YÊU CẦU SEO — lên Gemini...")

    own_driver = driver is None
    try:
        if driver is None:
            driver = g.init_firefox(url=NEW_CHAT_URL if SEO_NEW_CHAT else SEO_GEMINI_URL)

        ans = ""
        for lan in range(1, max(1, attempts) + 1):
            if lan > 1:
                log(f"🔁 Gửi lại SEO lần {lan}/{attempts} (cuộc trò chuyện khác)...")
            _open_chat(driver, log)
            # Chèn trọn yêu cầu SEO (SEO_PROMPT) lên trước, rồi tới đoạn truyện.
            ans = (g.send_to_gemini(driver, text, prefix=SEO_PROMPT, on_log=log) or "").strip()
            if not ans:
                log("⚠️ Gemini không trả về kết quả SEO.")
                continue
            if not _looks_like_seo(ans):
                log("⚠️ Nội dung nhận được KHÔNG đúng dạng kết quả SEO (thiếu phần "
                    "TIÊU ĐỀ/MÔ TẢ) → bỏ, không lưu.")
                ans = ""
                continue
            break

        if not ans:
            # File đã được reset ở trên → GIỮ RỖNG. Đừng lưu chuỗi rỗng vào .docx
            # kiểu "(trống)" nữa: bước sau chỉ coi là chưa làm SEO khi file rỗng.
            log("❌ Không lấy được SEO cho tập này — giữ file kết quả RỖNG để bước "
                "sau biết là CHƯA làm SEO (không dùng nhầm nội dung của tập khác).")
            return ""

        save_seo_docx(ans, output_path)
        log(f"\n========== KẾT QUẢ SEO YOUTUBE ==========\n{ans}\n")
        log(f"✅ XONG: đã lưu SEO → {output_path}")
        return ans
    finally:
        if own_driver and not keep_open and driver is not None:
            try:
                driver.quit()
            except Exception:
                pass


def main(argv=None):
    parser = argparse.ArgumentParser(
        description="Gửi đoạn đầu của truyện (đã dịch) lên Gemini SEO YouTube và lưu kết quả."
    )
    parser.add_argument("input", nargs="?", default=str(DEFAULT_INPUT),
                        help="File .docx nguồn (mặc định: gemini_result.docx).")
    parser.add_argument("-o", "--output", default=str(DEFAULT_OUTPUT),
                        help="File .docx lưu kết quả (mặc định: seoYoutube.docx).")
    parser.add_argument("--chars", type=int, default=0,
                        help="Chỉ gửi N ký tự đầu của đoạn (0 = cả đoạn).")
    parser.add_argument("--no-keep-open", dest="keep_open", action="store_false",
                        help="Đóng Firefox sau khi xong.")
    args = parser.parse_args(argv)

    run(args.input, args.output, max_chars=args.chars, keep_open=args.keep_open)


if __name__ == "__main__":
    main()
